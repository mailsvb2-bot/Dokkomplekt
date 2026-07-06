from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Iterable, Sequence
import re

from docx import Document

from diagnostic_logging import record_soft_exception
from diary_dates import parse_full_date, parse_full_datetime, parse_optional_discharge_date
from diary_gender import adapt_text_to_patient_gender, detect_gender_from_patient_name
from diary_models import DiaryBatchResult
from diary_paths import available_path, make_diary_output_name, safe_filename_part
from diary_schedule import DEFAULT_CALENDAR_DIARY_DAY_OFFSETS, expand_day_offsets, expand_hour_intervals, expand_minute_intervals
from diary_text_parser import clean_status_text, extract_statuses_from_docx, is_signature_paragraph_text, remove_examinee_words
from diary_writer_apply import NEUTRAL_FINAL_DIARY_TEXT
from medical_calendar import next_working_day
from medical_docx_xml_fragments import ensure_docx_compatible, existing_word_file
from medical_formatting import redact_technical_text, safe_filename, technical_ref, technical_report_path

MAX_INTRADAY_TEXT_DIARY_ENTRIES = 20000
TEXT_DIARY_SIGNATURE_LOCK_VERSION = "v1.0"
TEXT_DIARY_SIGNATURE_LINES = (
    "Лечащий врач ____________________",
    "Зав. отделением ____________________",
)


@dataclass(frozen=True)
class DynamicEpicrisisInput:
    patient_name: str = ""
    birth_date: str = ""
    sick_leave_from: str = ""
    complaints: str = ""
    treatment: str = ""
    profile_status: str = ""
    treatment_correction: str = ""


def default_observation_diary_dates(admission: date, *, limit: int = 20, discharge_date: date | None = None) -> tuple[date, ...]:
    if limit <= 0:
        return ()
    offsets = [0, 1, 2, 7]
    current = 10
    while len(offsets) < max(limit * 3, 12):
        offsets.append(current)
        current += 3 if len(offsets) % 2 else 4
    result: list[date] = []
    for offset in offsets:
        planned = admission + timedelta(days=max(0, int(offset)))
        if discharge_date is not None and planned > discharge_date:
            break
        result.append(planned)
        if len(result) >= limit:
            break
    return tuple(result)


def dynamic_epicrisis_dates(admission: date, *, discharge_date: date | None = None, limit: int = 12) -> tuple[date, ...]:
    result: list[date] = []
    current = admission + timedelta(days=10)
    while len(result) < limit:
        if discharge_date is not None and current >= discharge_date:
            break
        adjusted = next_working_day(current, used=result)
        if discharge_date is not None and adjusted >= discharge_date:
            break
        result.append(adjusted)
        current += timedelta(days=10)
    return tuple(result)


def build_dynamic_epicrisis_text(data: DynamicEpicrisisInput) -> str:
    correction = str(data.treatment_correction or "").strip() or "Лекарства принимает согласно назначениям."
    return "\n".join([
        "Динамический эпикриз.",
        f"ФИО: {data.patient_name or 'не указано'}.",
        f"Дата рождения: {data.birth_date or 'не указана'}.",
        f"Лечится с: {data.sick_leave_from or 'не указано'}.",
        f"Жалобы: {data.complaints or 'без существенной динамики'}.",
        f"Принимает: {data.treatment or 'согласно листу назначений'}.",
        f"Профильный статус: {data.profile_status or 'без существенной динамики'}.",
        correction,
        "Продолжение лечения по листу нетрудоспособности.",
        "Заведующий отделением ____________________",
        "Лечащий врач ____________________",
    ])


def _existing_docx_files(paths: Iterable[str | Path], label: str) -> list[Path]:
    result: list[Path] = []
    seen: set[Path] = set()
    for raw_path in paths:
        if raw_path is None or str(raw_path).strip() == "":
            raise ValueError(f"Пустой путь к файлу ({label}).")
        source = existing_word_file(raw_path, label)
        path = ensure_docx_compatible(source, label=label)
        key = path.resolve()
        if key not in seen:
            seen.add(key)
            result.append(path)
    return result


def _resolve_output_dir(output_dir: str | Path | None, fallback_dir: Path) -> Path:
    result = fallback_dir if output_dir is None or str(output_dir).strip() == "" else Path(output_dir).expanduser()
    if result.exists() and not result.is_dir():
        raise ValueError(f"Папка результата указывает на файл, а не на папку: {result}")
    result.mkdir(parents=True, exist_ok=True)
    return result


def _fallback_statuses_from_docx(path: str | Path) -> list[str]:
    doc = Document(str(path))
    result: list[str] = []
    seen: set[str] = set()

    def add(text: str) -> None:
        cleaned = clean_status_text(text)
        low = cleaned.lower().replace("ё", "е")
        if not cleaned or len(cleaned) < 3 or is_signature_paragraph_text(cleaned):
            return
        if re.fullmatch(r"[\d\s./-]+", cleaned):
            return
        if low in seen:
            return
        seen.add(low)
        result.append(cleaned)

    for paragraph in doc.paragraphs:
        add(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            seen_cells: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_cells:
                    continue
                seen_cells.add(tc_id)
                for paragraph in cell.paragraphs:
                    add(paragraph.text)
    return result


def read_statuses_from_files(paths: Iterable[str | Path]) -> list[str]:
    statuses: list[str] = []
    seen: set[str] = set()
    for path in _existing_docx_files(paths, "тексты дневников"):
        path_statuses = extract_statuses_from_docx(path) or _fallback_statuses_from_docx(path)
        for status in path_statuses:
            key = " ".join(status.strip().lower().replace("ё", "е").split())
            if key not in seen:
                statuses.append(status.strip())
                seen.add(key)
    return statuses


def open_folder(path: str | Path) -> bool:
    folder = Path(path).expanduser()
    try:
        from printer_platform import open_desktop_path
        return bool(open_desktop_path(folder, require_dir=True))
    except Exception as exc:
        record_soft_exception("diary_batch.open_folder", exc, detail=str(folder))
        return False


def _optional_full_date(value: str) -> date | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return parse_full_date(text)
    except ValueError:
        return None


def _dynamic_epicrisis_base_date(admission_date_value: date, sick_leave_from: str) -> date:
    sick_leave_date = _optional_full_date(sick_leave_from)
    return max(admission_date_value, sick_leave_date) if sick_leave_date is not None else admission_date_value


def _calendar_text_diary_dates(admission_date_value: date, discharge_date_value: date | None, *, limit: int, day_offsets: Sequence[int] = ()) -> tuple[date, ...]:
    if limit <= 0:
        return ()
    offsets = tuple(int(item) for item in (day_offsets or DEFAULT_CALENDAR_DIARY_DAY_OFFSETS))
    expanded_offsets = expand_day_offsets(offsets, max(limit * 3, len(offsets), 10))
    result: list[date] = []
    for offset in expanded_offsets:
        planned = admission_date_value + timedelta(days=max(1, int(offset)))
        if discharge_date_value is not None and planned > discharge_date_value:
            break
        result.append(planned)
        if len(result) >= limit:
            break
    return tuple(result)


def _text_diary_dates(admission_date_value: date, discharge_date_value: date | None, *, limit: int, day_offsets: Sequence[int] = ()) -> tuple[date, ...]:
    explicit_offsets = tuple(int(item) for item in day_offsets)
    return _calendar_text_diary_dates(admission_date_value, discharge_date_value, limit=limit, day_offsets=explicit_offsets or DEFAULT_CALENDAR_DIARY_DAY_OFFSETS)


def _build_dated_entries(statuses: Sequence[str], dates: Sequence[date], patient_gender: str | None, repeat_statuses: bool) -> tuple[tuple[date, str], ...]:
    entries: list[tuple[date, str]] = []
    status_index = 0
    for item_date in dates:
        if not statuses:
            text = ""
        else:
            if status_index >= len(statuses):
                if repeat_statuses:
                    status_index = 0
                else:
                    break
            text = statuses[status_index]
            status_index += 1
        adapted, _changed = adapt_text_to_patient_gender(text, patient_gender)
        entries.append((item_date, f"{item_date:%d.%m.%y} {clean_status_text(adapted)}".rstrip()))
    return tuple(entries)


def _intraday_minutes_for_one_day(intervals: Sequence[int]) -> tuple[int, ...]:
    values = tuple(max(1, int(item)) for item in intervals if int(item) > 0)
    if not values:
        return ()
    result = [0]
    for offset in expand_minute_intervals(values, 1440):
        if offset >= 1440:
            break
        result.append(offset)
    return tuple(dict.fromkeys(result))


def _intraday_text_diary_datetimes(
    admission_value: str,
    admission_date_value: date,
    discharge_date_value: date | None,
    *,
    day_offsets: Sequence[int],
    minute_offsets: Sequence[int],
    limit: int,
) -> tuple[datetime, ...]:
    if limit <= 0:
        return ()
    try:
        base = parse_full_datetime(admission_value)
    except ValueError:
        base = datetime.combine(admission_date_value, datetime.min.time())
    minutes = _intraday_minutes_for_one_day(minute_offsets)
    if not minutes:
        return ()
    if day_offsets:
        day_count = max(1, min(370, limit))
        days = _text_diary_dates(admission_date_value, discharge_date_value, limit=day_count, day_offsets=day_offsets)
        result: list[datetime] = []
        for item_date in days:
            start = datetime.combine(item_date, base.time())
            for minute in minutes:
                moment = start + timedelta(minutes=minute)
                if moment.date() != item_date:
                    break
                if discharge_date_value is not None and moment.date() > discharge_date_value:
                    break
                result.append(moment)
                if len(result) >= MAX_INTRADAY_TEXT_DIARY_ENTRIES:
                    return tuple(result)
        return tuple(result)
    result: list[datetime] = []
    step_values = tuple(max(1, int(item)) for item in minute_offsets)
    expanded = expand_minute_intervals(step_values, MAX_INTRADAY_TEXT_DIARY_ENTRIES)
    for minute in expanded:
        moment = base + timedelta(minutes=minute)
        if discharge_date_value is not None and moment.date() > discharge_date_value:
            break
        result.append(moment)
        if len(result) >= MAX_INTRADAY_TEXT_DIARY_ENTRIES:
            break
    return tuple(result)


def _hourly_text_diary_datetimes(admission_value: str, discharge_date_value: date | None, *, limit: int, hour_offsets: Sequence[int]) -> tuple[datetime, ...]:
    if limit <= 0:
        return ()
    try:
        base = parse_full_datetime(admission_value)
    except ValueError:
        return ()
    offsets = expand_hour_intervals(tuple(int(item) for item in hour_offsets) or (1,), limit)
    result: list[datetime] = []
    for offset in offsets:
        planned = base + timedelta(hours=int(offset))
        # Hourly diaries are multiple observations within the same treatment day.
        # The old uniqueness-by-date guard moved the second hourly record to the
        # next workday (15:00, then tomorrow 16:00), breaking the user-visible
        # hourly route. Only non-working days are shifted; same-date observations
        # stay on the same date.
        if is_non_working_day(planned.date()):
            planned = datetime.combine(next_working_day(planned.date()), planned.time())
        if discharge_date_value is not None and planned.date() > discharge_date_value:
            break
        result.append(planned)
        if len(result) >= limit:
            break
    return tuple(result)


def _build_hourly_entries(statuses: Sequence[str], moments: Sequence[datetime], patient_gender: str | None, repeat_statuses: bool) -> tuple[tuple[date, str], ...]:
    entries: list[tuple[date, str]] = []
    status_index = 0
    for moment in moments:
        if not statuses:
            text = ""
        else:
            if status_index >= len(statuses):
                if repeat_statuses:
                    status_index = 0
                else:
                    break
            text = statuses[status_index]
            status_index += 1
        adapted, _changed = adapt_text_to_patient_gender(text, patient_gender)
        entries.append((moment.date(), f"{moment:%d.%m.%y %H:%M} {clean_status_text(adapted)}".rstrip()))
    return tuple(entries)


def _split_regular_and_final_text_diary_dates(dates: Sequence[date], *, discharge_date_value: date | None, force_final_diary: bool) -> tuple[tuple[date, ...], date | None]:
    normalized_dates = tuple(d for d in dates if isinstance(d, date))
    if not force_final_diary:
        return normalized_dates, None
    if discharge_date_value is not None:
        final_date = discharge_date_value
        return tuple(d for d in normalized_dates if d < discharge_date_value), final_date
    if not normalized_dates:
        return normalized_dates, None
    return normalized_dates[:-1], normalized_dates[-1]


def _neutral_final_diary_entry(date_value: date, patient_gender: str | None) -> tuple[date, str]:
    adapted_final_text, _changed = adapt_text_to_patient_gender(NEUTRAL_FINAL_DIARY_TEXT, patient_gender)
    adapted_final_text = remove_examinee_words(adapted_final_text)
    return date_value, f"{date_value:%d.%m.%y} {adapted_final_text}".rstrip()


def _add_text_diary_signature_block(doc: Document) -> None:
    """Append the legacy doctor/head signature block after one diary entry."""
    for line in TEXT_DIARY_SIGNATURE_LINES:
        doc.add_paragraph(line)


def _create_text_diary_document(output_dir: Path, patient_name: str, entries: Sequence[tuple[date, str]], epicrisis_entries: Sequence[tuple[date, str]]) -> Path:
    target = available_path(output_dir / safe_filename(make_diary_output_name(safe_filename_part(patient_name), file_index=1, total_files=1)))
    doc = Document()
    blocks: list[tuple[date, int, tuple[str, ...]]] = []
    for item_date, entry in entries:
        blocks.append((item_date, 0, (str(entry or "").strip(),)))
    for item_date, text in epicrisis_entries:
        lines = str(text or "").splitlines()
        first_line = f"{item_date:%d.%m.%y} {lines[0] if lines else ''}".rstrip()
        blocks.append((item_date, 1, tuple([first_line, *lines[1:]])))
    for _item_date, block_kind, lines in sorted(blocks, key=lambda block: (block[0], block[1])):
        if doc.paragraphs:
            doc.add_paragraph("")
        for line in lines:
            doc.add_paragraph(line)
        if block_kind == 0:
            _add_text_diary_signature_block(doc)
    doc.save(str(target))
    return target


def _fill_text_diary_batch(
    *,
    statuses: Sequence[str],
    result_dir: Path,
    patient_name: str,
    admission_value: str,
    admission_date_value,
    discharge_date_value,
    gender_source_name: str,
    repeat_statuses: bool,
    patient_gender: str | None,
    sick_leave_dynamic_epicrisis: bool,
    treatment_correction: str,
    birth_date: str,
    complaints: str,
    treatment: str,
    profile_status: str,
    sick_leave_from: str,
    write_report: bool,
    diary_day_offsets: Sequence[int],
    force_final_diary: bool,
    diary_hour_offsets: Sequence[int] = (),
    diary_minute_offsets: Sequence[int] = (),
    diary_frequency_mode: str = "daily",
    removed_after_discharge_rows: int = 0,
) -> DiaryBatchResult:
    """Build a text diary document and optional dynamic epicrisis entries."""
    if admission_date_value is None:
        admission_date_value = parse_full_date(admission_value)
    rough_limit = max(10, min(370, (discharge_date_value - admission_date_value).days + 10)) if discharge_date_value else max(10, len(statuses) or 10)
    hourly_mode = str(diary_frequency_mode or "daily").strip().lower() == "hourly"
    if hourly_mode:
        if diary_minute_offsets:
            moments = _intraday_text_diary_datetimes(admission_value, admission_date_value, discharge_date_value, day_offsets=diary_day_offsets, minute_offsets=diary_minute_offsets, limit=rough_limit)
        else:
            hourly_limit = max(rough_limit, min(240, rough_limit * 12))
            moments = _hourly_text_diary_datetimes(admission_value, discharge_date_value, limit=hourly_limit, hour_offsets=diary_hour_offsets)
        dates = tuple(moment.date() for moment in moments)
        final_date = None
        entries = list(_build_hourly_entries(statuses, moments, patient_gender, repeat_statuses))
    else:
        dates = _text_diary_dates(admission_date_value, discharge_date_value, limit=rough_limit, day_offsets=diary_day_offsets)
        regular_dates, final_date = _split_regular_and_final_text_diary_dates(dates, discharge_date_value=discharge_date_value, force_final_diary=force_final_diary)
        entries = list(_build_dated_entries(statuses, regular_dates, patient_gender, repeat_statuses))
    final_rows_filled = 0
    if final_date is not None:
        entries.append(_neutral_final_diary_entry(final_date, patient_gender))
        final_rows_filled = 1
    epicrisis_entries: list[tuple[date, str]] = []
    if sick_leave_dynamic_epicrisis:
        epicrisis_base_date = _dynamic_epicrisis_base_date(admission_date_value, sick_leave_from)
        data = DynamicEpicrisisInput(patient_name, birth_date, f"{epicrisis_base_date:%d.%m.%Y}", complaints, treatment, profile_status, treatment_correction)
        epicrisis_entries = [(d, build_dynamic_epicrisis_text(data)) for d in dynamic_epicrisis_dates(epicrisis_base_date, discharge_date=discharge_date_value, limit=12)]
    created = _create_text_diary_document(result_dir, patient_name, entries, epicrisis_entries)
    report_path: Path | None = None
    if write_report:
        patient_filename = safe_filename_part(patient_name)
        lines = [
            "ОТЧЁТ: текстовые дневники",
            f"Дата запуска: {datetime.now():%d.%m.%Y %H:%M:%S}",
            "Карточка пациента: обезличена",
            "Технический идентификатор: " + technical_ref(patient_filename, gender_source_name, admission_value),
            "Технический контекст: " + redact_technical_text(f"{patient_filename} {gender_source_name} {admission_value}"),
            f"Дневниковых дат: {len(dates)}",
            f"Финальных записей: {final_rows_filled}",
            f"Динамических эпикризов: {len(epicrisis_entries)}",
        ]
        report_path = technical_report_path(result_dir, "ОТЧЁТ_дневники.txt")
        report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return DiaryBatchResult([created], report_path, 1, len(entries), len(entries), 0, final_rows_filled, 0, 0, removed_after_discharge_rows)


def fill_diary_batch(
    *,
    status_files: Sequence[str | Path],
    diary_files: Sequence[str | Path],
    output_dir: str | Path | None,
    patient_name: str,
    admission_value: str,
    gender_source_name: str | None = None,
    discharge_value: str = "",
    repeat_statuses: bool = True,
    reset_each_file: bool = True,
    keep_signature: bool = True,
    fill_months: bool = True,
    force_final_diary: bool = True,
    remove_holiday_rows: bool = True,
    open_result_folder: bool = False,
    write_report: bool = False,
    diary_day_offsets: Sequence[int] = (),
    diary_hour_offsets: Sequence[int] = (),
    diary_minute_offsets: Sequence[int] = (),
    diary_frequency_mode: str = "daily",
    allow_empty_statuses: bool = False,
    text_output: bool = False,
    sick_leave_dynamic_epicrisis: bool = False,
    treatment_correction: str = "",
    birth_date: str = "",
    complaints: str = "",
    treatment: str = "",
    profile_status: str = "",
    sick_leave_from: str = "",
) -> DiaryBatchResult:
    """Validate diary inputs and create text-route diary output for one patient."""
    _ = (diary_files, reset_each_file, keep_signature, fill_months, remove_holiday_rows, text_output)
    status_file_paths = _existing_docx_files(status_files, "тексты дневников") if status_files else []
    if not status_file_paths and not allow_empty_statuses:
        raise ValueError("Сначала выберите тексты дневников. Даты берутся из календаря программы, а не из таблицы дневников.")
    try:
        admission_date_value = parse_full_datetime(admission_value).date()
    except ValueError:
        admission_date_value = None
    discharge_date_value = parse_optional_discharge_date(discharge_value)
    if admission_date_value is not None and discharge_date_value is not None and discharge_date_value < admission_date_value:
        raise ValueError("Дата выписки не может быть раньше даты поступления.")
    gender_name = safe_filename_part(gender_source_name or patient_name)
    patient_gender = detect_gender_from_patient_name(gender_name)
    if patient_gender is None:
        raise ValueError("Введите ФИО так, чтобы первым словом была фамилия пациента. Например: Иванов И.И. или Петрова А.А.")
    statuses = read_statuses_from_files(status_file_paths)
    if status_file_paths and not statuses:
        raise ValueError("В выбранных файлах с текстами дневников не найдено подходящих текстов.")
    first_dir = status_file_paths[0].parent if status_file_paths else Path.cwd()
    result_dir = _resolve_output_dir(output_dir, first_dir)
    result = _fill_text_diary_batch(
        statuses=statuses,
        result_dir=result_dir,
        patient_name=patient_name,
        admission_value=admission_value,
        admission_date_value=admission_date_value,
        discharge_date_value=discharge_date_value,
        gender_source_name=gender_name,
        repeat_statuses=repeat_statuses,
        patient_gender=patient_gender,
        sick_leave_dynamic_epicrisis=sick_leave_dynamic_epicrisis,
        treatment_correction=treatment_correction,
        birth_date=birth_date,
        complaints=complaints,
        treatment=treatment,
        profile_status=profile_status,
        sick_leave_from=sick_leave_from,
        write_report=write_report,
        diary_day_offsets=tuple(int(x) for x in diary_day_offsets),
        diary_hour_offsets=tuple(int(x) for x in diary_hour_offsets),
        diary_minute_offsets=tuple(int(x) for x in diary_minute_offsets),
        diary_frequency_mode=str(diary_frequency_mode or "daily").strip().lower(),
        force_final_diary=force_final_diary,
        removed_after_discharge_rows=0,
    )
    if open_result_folder:
        open_folder(result_dir)
    return result
