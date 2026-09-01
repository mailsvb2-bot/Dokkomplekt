from __future__ import annotations

import re
from pathlib import Path

from docx import Document

_NAMES = (
    "cell_int", "is_holiday_skip_date", "should_remove_holiday",
    "is_data_row", "find_column_by_header", "find_diary_column", "find_day_column",
    "find_hospitalization_day_column", "find_month_year_column", "clear_paragraph_keep_properties",
    "reset_cell_to_one_paragraph", "is_structural_diary_prefix", "first_signature_paragraph_index",
    "add_run_with_size", "fill_text_cell", "write_diary_text_into_existing_paragraph",
    "fill_diary_text_cell", "remove_row", "collect_dated_entries",
)


def detect_first_month_year_from_docx(path: str | Path):
    """Read-only classifier helper for legacy 01–31 date templates.

    The old table writer stays removed, but drag-and-drop still needs a safe way
    to recognise a date template.  Return ``(month, year)`` for the first explicit
    date/month-year found, otherwise ``None``.
    """
    from medical_docx_xml_fragments import ensure_docx_compatible
    readable = ensure_docx_compatible(path, label="diary date template")
    doc = Document(str(readable))
    chunks: list[str] = [p.text for p in doc.paragraphs]
    seen_cells: set[int] = set()

    def walk_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen_cells:
                    continue
                seen_cells.add(key)
                chunks.extend(p.text for p in cell.paragraphs)
                for nested in cell.tables:
                    walk_table(nested)

    for table in doc.tables:
        walk_table(table)
    text = "\n".join(chunks)
    for match in re.finditer(r"(?<!\d)(?:0?[1-9]|[12]\d|3[01])[./-](0?[1-9]|1[0-2])[./-](20\d{2}|\d{2})(?!\d)", text):
        month = int(match.group(1))
        year = int(match.group(2))
        if year < 100:
            year += 2000
        return month, year
    for match in re.finditer(r"(?<!\d)(0?[1-9]|1[0-2])[./-](20\d{2})(?!\d)", text):
        return int(match.group(1)), int(match.group(2))
    return None


def _removed(*args, **kwargs):
    _ = (args, kwargs)
    raise NotImplementedError("Legacy DOCX row writer is removed")


def __getattr__(name: str):
    if name in _NAMES:
        return _removed
    raise AttributeError(name)


__all__ = ["detect_first_month_year_from_docx", *_NAMES]
