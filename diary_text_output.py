"""Text-mode diary document generation.

This is the new diary output style requested by the doctor: entries are written
as dated text paragraphs instead of forcing every workflow through a Word table.
The old table filler remains available for existing templates.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path
from typing import Sequence

from docx import Document

from diary_gender import adapt_text_to_patient_gender
from diary_paths import available_path, make_diary_output_name, safe_filename_part
from diary_text_parser import clean_status_text
from medical_formatting import safe_filename

DIARY_TEXT_OUTPUT_LOCK_VERSION = "v1.0"
DIARY_TEXT_OUTPUT_USES_DATED_PARAGRAPHS = True
DIARY_TEXT_OUTPUT_SUPPORTS_DYNAMIC_EPICRISIS = True


@dataclass(frozen=True)
class DynamicEpicrisisInput:
    patient_name: str = ""
    birth_date: str = ""
    sick_leave_from: str = ""
    complaints: str = ""
    treatment: str = ""
    profile_status: str = ""
    treatment_correction: str = ""


def diary_text_output_name(patient_name: str) -> str:
    return safe_filename(make_diary_output_name(safe_filename_part(patient_name), file_index=1, total_files=1))


def build_dated_diary_entries(
    *,
    statuses: Sequence[str],
    dates: Sequence[date],
    patient_gender: str | None = None,
    repeat_statuses: bool = True,
) -> tuple[str, ...]:
    entries: list[str] = []
    status_index = 0
    for item_date in dates:
        if not statuses:
            text = ""
        else:
            if status_index >= len(statuses):
                if repeat_statuses:
                    status_index = 0
                else:
                    break
            text = statuses[status_index]
            status_index += 1
        adapted, _changed = adapt_text_to_patient_gender(text, patient_gender)
        cleaned = clean_status_text(adapted)
        entries.append(f"{item_date:%d.%m.%y} {cleaned}".rstrip())
    return tuple(entries)


def dynamic_epicrisis_dates(admission: date, *, discharge_date: date | None = None, limit: int = 12) -> tuple[date, ...]:
    result: list[date] = []
    current = admission + timedelta(days=10)
    while len(result) < limit:
        if discharge_date is not None and current > discharge_date:
            break
        result.append(current)
        current += timedelta(days=10)
    return tuple(result)


def build_dynamic_epicrisis_text(data: DynamicEpicrisisInput) -> str:
    correction = str(data.treatment_correction or "").strip() or "Лекарства принимает согласно назначениям."
    lines = [
        "Динамический эпикриз.",
        f"ФИО: {data.patient_name or 'не указано'}.",
        f"Дата рождения: {data.birth_date or 'не указана'}.",
        f"Лечится с: {data.sick_leave_from or 'не указано'}.",
        f"Жалобы: {data.complaints or 'без существенной динамики'}.",
        f"Принимает: {data.treatment or 'согласно листу назначений'}.",
        f"Профильный статус: {data.profile_status or 'без существенной динамики'}.",
        correction,
        "Продолжение лечения по листу нетрудоспособности.",
        "Заведующий отделением ____________________",
        "Лечащий врач ____________________",
    ]
    return "\n".join(lines)


def create_text_diary_document(
    *,
    output_dir: str | Path,
    patient_name: str,
    entries: Sequence[str],
    epicrisis_entries: Sequence[tuple[date, str]] = (),
) -> Path:
    result_dir = Path(output_dir).expanduser()
    result_dir.mkdir(parents=True, exist_ok=True)
    target = available_path(result_dir / diary_text_output_name(patient_name))
    doc = Document()
    for entry in entries:
        doc.add_paragraph(str(entry or "").strip())
    for item_date, text in epicrisis_entries:
        if doc.paragraphs:
            doc.add_paragraph("")
        doc.add_paragraph(f"{item_date:%d.%m.%y} {text.splitlines()[0]}")
        for line in text.splitlines()[1:]:
            doc.add_paragraph(line)
    doc.save(str(target))
    return target


def assert_diary_text_output_lock() -> None:
    if DIARY_TEXT_OUTPUT_LOCK_VERSION != "v1.0":
        raise AssertionError("Diary text output lock changed unexpectedly")
    if not DIARY_TEXT_OUTPUT_USES_DATED_PARAGRAPHS or not DIARY_TEXT_OUTPUT_SUPPORTS_DYNAMIC_EPICRISIS:
        raise AssertionError("Text diary output must keep dated paragraphs and dynamic epicrisis support")
    entries = build_dated_diary_entries(statuses=["Пациент был спокоен, инструкции выполнял."], dates=[date(2026, 1, 10)], patient_gender="female")
    if not entries or not entries[0].startswith("10.01.26"):
        raise AssertionError("Text diary entry date formatting is broken")
