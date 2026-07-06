from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable, Mapping, Sequence

from docx import Document

from diagnostic_logging import record_soft_exception
from medical_docx_xml_fragments import ensure_docx_compatible

from .models import DocumentBlueprint, DocumentSource, FieldSpec, PopupField, SectionSpec, SignatureSpec
from .pdf_reader import read_pdf_text
from .text_utils import custom_field_id, field_id_from_placeholder, normalize, value_kind

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
LABEL_BLANK_RE = re.compile(r"^\s*(?P<label>[^:：\n]{2,90}?)\s*[:：]\s*[_—–\-.\s]{3,}\s*$")
LABEL_UNDERLINE_RE = re.compile(r"^\s*(?P<label>[^_—–\n]{2,90}?)\s+[_—–-]{3,}\s*$")
BLANK_RE = re.compile(r"[_—–-]{3,}|\.{4,}")
SIGNATURE_RE = re.compile(r"(?i)(подпись|директор|бухгалтер|врач|зав\.?\s*отдел|исполнитель|составил|утверждаю|signature|accountant|director|approved)")
ACCOUNTING_RE = re.compile(r"(?i)(счет|счёт|инн|кпп|ндс|банк|бухгалтер|контрагент|invoice|vat|accountant)")
MEDICAL_RE = re.compile(r"(?i)(пациент|диагноз|лечение|анамнез|выпис|эпикриз|patient|diagnosis)")
LEGAL_RE = re.compile(r"(?i)(договор|иск|претензи|истец|ответчик|суд|contract|claim|court)")
HR_RE = re.compile(r"(?i)(приказ|работник|сотрудник|должность|employee|hr|position)")
TABLE_RE = re.compile(r"(?i)(сумма|количество|цена|услуг|товар|работ|итого|amount|qty|price|total)")


def _read_docx(source: DocumentSource) -> tuple[tuple[str, ...], tuple[tuple[str, ...], ...]]:
    path = ensure_docx_compatible(source.path, label="document intelligence source")
    doc = Document(str(path))
    lines: list[str] = []
    rows: list[tuple[str, ...]] = []
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells: list[str] = []
            seen: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen:
                    continue
                seen.add(tc_id)
                text = normalize(" ".join(paragraph.text for paragraph in cell.paragraphs))
                cells.append(text)
                if text:
                    lines.append(text)
            if any(cells):
                rows.append(tuple(cells))
    return tuple(lines), tuple(rows)


def _read_pdf(source: DocumentSource) -> tuple[tuple[str, ...], tuple[tuple[str, ...], ...]]:
    result = read_pdf_text(source.path)
    if not result.has_text:
        warnings = "; ".join(result.warnings) or "PDF has no extractable text"
        raise ValueError(warnings)
    lines = tuple(normalize(line) for line in result.text.splitlines() if normalize(line))
    return lines, ()


def _read_source(source: DocumentSource) -> tuple[tuple[str, ...], tuple[tuple[str, ...], ...]]:
    suffix = Path(source.path).suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(source)
    return _read_docx(source)


def _signature_role(label: str) -> str:
    low = normalize(label).lower().replace("ё", "е")
    if "бухгалтер" in low:
        return "accountant"
    if "директор" in low or "руковод" in low:
        return "director"
    if "зав" in low and "отдел" in low:
        return "head"
    if "врач" in low:
        return "author"
    if "исполн" in low or "состав" in low:
        return "executor"
    return "custom"


def _infer_fields(lines: Sequence[str], rows: Sequence[Sequence[str]]) -> tuple[FieldSpec, ...]:
    found: dict[str, FieldSpec] = {}

    def add(label: str, source: str, confidence: float = 0.75, field_id: str | None = None) -> None:
        label = normalize(label).strip(" :：\t")
        if not label or SIGNATURE_RE.search(label):
            return
        fid = field_id or custom_field_id(label)
        spec = FieldSpec(fid, label, value_kind(label, fid), True, source, confidence)
        old = found.get(fid)
        if old is None or spec.confidence > old.confidence:
            found[fid] = spec

    for line in lines:
        for match in PLACEHOLDER_RE.finditer(line):
            raw = match.group(1)
            add(raw, "placeholder", 0.98, field_id_from_placeholder(raw))
        match = LABEL_BLANK_RE.match(line) or LABEL_UNDERLINE_RE.match(line)
        if match:
            add(match.group("label"), "visible_blank", 0.82)
    for row in rows:
        for idx, cell in enumerate(row[:-1]):
            if normalize(cell) and (not normalize(row[idx + 1]) or BLANK_RE.search(normalize(row[idx + 1]))):
                add(cell, "table_neighbor_blank", 0.86)
    return tuple(found.values())


def _infer_signatures(lines: Sequence[str]) -> tuple[SignatureSpec, ...]:
    result: dict[str, SignatureSpec] = {}
    for line in lines:
        if SIGNATURE_RE.search(line):
            label = normalize(line).strip(" _—–-.")
            if label:
                result.setdefault(label.lower().replace("ё", "е"), SignatureSpec(label, _signature_role(label)))
    return tuple(result.values())


def _classify_domain(lines: Sequence[str], fields: Sequence[FieldSpec], source: DocumentSource) -> str:
    hint = normalize(source.profession_hint).lower()
    if hint and hint != "auto":
        return hint
    text = "\n".join([*lines, *(field.label for field in fields)])
    scores = {"medical": len(MEDICAL_RE.findall(text)), "accounting": len(ACCOUNTING_RE.findall(text)), "legal": len(LEGAL_RE.findall(text)), "hr": len(HR_RE.findall(text))}
    best, value = max(scores.items(), key=lambda item: item[1])
    return best if value else "custom"


def _infer_shape(lines: Sequence[str], rows: Sequence[Sequence[str]], fields: Sequence[FieldSpec]) -> str:
    has_blanks = any(field.source in {"visible_blank", "table_neighbor_blank"} for field in fields)
    has_tables = bool(rows)
    long_text = sum(len(line) for line in lines) > 700 and len(lines) > 4
    if has_tables and has_blanks:
        return "mixed"
    if has_tables:
        return "table_form"
    if has_blanks:
        return "blank_form"
    if long_text:
        return "free_text"
    return "letter"


def _infer_sections(lines: Sequence[str], rows: Sequence[Sequence[str]]) -> tuple[SectionSpec, ...]:
    result: list[SectionSpec] = []
    if rows and any(TABLE_RE.search(" ".join(row)) for row in rows):
        result.append(SectionSpec("table", "table", 0.78))
    if any(SIGNATURE_RE.search(line) for line in lines):
        result.append(SectionSpec("signature_block", "signature", 0.82))
    if sum(len(line) for line in lines) > 700 and not rows:
        result.append(SectionSpec("free_text_body", "body", 0.72))
    return tuple(result)


class DocumentIntelligenceCore:
    """Analyze arbitrary user-owned documents without hard-coding a profession."""

    def analyze_source(self, source: DocumentSource | str | Path) -> DocumentBlueprint:
        if not isinstance(source, DocumentSource):
            source = DocumentSource(str(source))
        try:
            lines, rows = _read_source(source)
        except Exception as exc:
            record_soft_exception("document_intelligence_core.read", exc, detail=str(source.path))
            return DocumentBlueprint(source.user_label or Path(source.path).stem, "custom", "letter", str(source.path), confidence=0.0)
        fields = _infer_fields(lines, rows)
        title = source.user_label or next((line for line in lines if 2 <= len(line) <= 120), Path(source.path).stem)
        return DocumentBlueprint(title, _classify_domain(lines, fields, source), _infer_shape(lines, rows, fields), str(source.path), fields, _infer_signatures(lines), _infer_sections(lines, rows))

    def analyze_sources(self, sources: Iterable[DocumentSource | str | Path]) -> tuple[DocumentBlueprint, ...]:
        return tuple(self.analyze_source(source) for source in sources)

    def build_popup_plan(self, blueprint: DocumentBlueprint, values: Mapping[str, object] | None = None) -> tuple[PopupField, ...]:
        values = values or {}
        result: list[PopupField] = []
        for field in blueprint.fields:
            if field.required and not str(values.get(field.field_id, "") or "").strip():
                result.append(PopupField(field.field_id, field.label, field.value_kind, True, placeholder="{{" + field.field_id + "}}"))
        return tuple(result)
