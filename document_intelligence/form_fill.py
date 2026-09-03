from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable, Mapping

from docx import Document

from .text_utils import custom_field_id, normalize
BLANK_RE = re.compile(r"[_—–-]{3,}|\.{4,}")
SIGNATURE_RE = re.compile(r"(?i)(подпись|директор|бухгалтер|врач|зав\.?\s*отдел|исполнитель|составил|утверждаю|signature|accountant|director|approved)")


@dataclass(frozen=True)
class VisibleBlankSlot:
    label: str
    start: int
    end: int


def visible_blank_slot(text: object) -> VisibleBlankSlot | None:
    """Find one human-labelled blank while preserving all fixed surrounding text."""
    raw = str(text or "")
    matches = list(BLANK_RE.finditer(raw))
    if not matches:
        return None
    first, last = matches[0], matches[-1]
    prefix = raw[: first.start()].rstrip()
    if not prefix:
        return None
    # Prefer an explicit colon (ASCII or full-width).  Otherwise the historical
    # ``ФИО ______`` form uses everything before the blank as the label.
    colon_at = max(prefix.rfind(":"), prefix.rfind("："))
    if colon_at >= 0:
        label = prefix[:colon_at].strip()
    else:
        label = prefix.strip()
    label = label.strip(" :：\t")
    if len(label) < 2 or len(label) > 90 or SIGNATURE_RE.search(label):
        return None
    return VisibleBlankSlot(label, first.start(), last.end())


def blank_region(text: object) -> tuple[int, int] | None:
    raw = str(text or "")
    matches = list(BLANK_RE.finditer(raw))
    if not matches:
        return None
    return matches[0].start(), matches[-1].end()


def unique_cells(table) -> Iterable:
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            yield cell


def iter_tables_recursive(table) -> Iterable:
    yield table
    for cell in unique_cells(table):
        for nested in cell.tables:
            yield from iter_tables_recursive(nested)


def iter_document_tables(document) -> Iterable:
    seen: set[int] = set()

    def emit(table):
        key = id(table._tbl)
        if key in seen:
            return
        seen.add(key)
        yield from iter_tables_recursive(table)

    for table in document.tables:
        yield from emit(table)
    for section in document.sections:
        for area in (section.header, section.footer):
            for table in area.tables:
                yield from emit(table)


def iter_direct_story_paragraphs(document) -> Iterable:
    """Paragraphs not inside tables: body plus each unique header/footer."""
    seen: set[int] = set()
    for paragraph in document.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for section in document.sections:
        for area in (section.header, section.footer):
            for paragraph in area.paragraphs:
                key = id(paragraph._p)
                if key not in seen:
                    seen.add(key)
                    yield paragraph


def iter_all_story_paragraphs(document) -> Iterable:
    """Yield every Word paragraph once, including tables and text-box XML."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    seen: set[object] = set()

    def emit(paragraph):
        key = paragraph._p
        if key in seen:
            return
        seen.add(key)
        yield paragraph

    for paragraph in iter_direct_story_paragraphs(document):
        yield from emit(paragraph)
    for table in iter_document_tables(document):
        for cell in unique_cells(table):
            for paragraph in cell.paragraphs:
                yield from emit(paragraph)
    # python-docx high-level collections omit text boxes/shapes. Walk raw story
    # XML after structured paragraphs so analyzer and renderer see the same slots.
    roots = [(document.element.body, document)]
    for section in document.sections:
        roots.extend(((section.header._element, section.header), (section.footer._element, section.footer)))
    for root, parent in roots:
        for element in root.iter(qn("w:p")):
            yield from emit(Paragraph(element, parent))

def visible_field_id(label: str, *, role_id: str = "", category: str = "", button_label: str = "") -> str:
    """Resolve a visible human label through the canonical field registry first."""
    try:
        from universal_fields import default_field_registry, normalize_field_id_for_context

        normalized = normalize_field_id_for_context(
            label,
            role_id=role_id,
            category=category,
            document_label=button_label,
        )
        if normalized.startswith("custom.") or normalized in default_field_registry():
            return normalized
        return custom_field_id(label)
    except ValueError:
        return custom_field_id(label)


def _replace_paragraph_span(paragraph, start: int, end: int, value: str) -> bool:
    """Replace a character span without flattening the paragraph's run formatting."""
    if end <= start:
        return False
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(value)
        return True
    cursor = 0
    touched = False
    inserted = False
    for run in runs:
        text = run.text or ""
        run_start, run_end = cursor, cursor + len(text)
        cursor = run_end
        if run_end <= start or run_start >= end:
            continue
        local_start = max(0, start - run_start)
        local_end = min(len(text), end - run_start)
        prefix = text[:local_start]
        suffix = text[local_end:]
        if not inserted:
            run.text = prefix + value + suffix
            inserted = True
        else:
            run.text = prefix + suffix
        touched = True
    return touched


def _fill_labelled_paragraphs(
    document: Document,
    values: Mapping[str, str],
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> list[str]:
    filled: list[str] = []
    for paragraph in iter_all_story_paragraphs(document):
        slot = visible_blank_slot(paragraph.text)
        if slot is None:
            continue
        field_id = visible_field_id(slot.label, role_id=role_id, category=category, button_label=button_label)
        value = str(values.get(field_id, "") or "").strip()
        if value and _replace_paragraph_span(paragraph, slot.start, slot.end, value):
            filled.append(field_id)
    return filled


def _cell_blank_paragraph(cell):
    for paragraph in cell.paragraphs:
        region = blank_region(paragraph.text)
        if region is not None:
            return paragraph, region
    # A genuinely empty cell is also a fill target.
    if not normalize(cell.text):
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        return paragraph, None
    return None, None


def _fill_tables(
    document: Document,
    values: Mapping[str, str],
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> list[str]:
    filled: list[str] = []
    for table in iter_document_tables(document):
        for row in table.rows:
            cells = []
            seen: set[int] = set()
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                cells.append(cell)
            for index, cell in enumerate(cells[:-1]):
                label = normalize(cell.text).strip(" :：")
                if not label:
                    continue
                target = cells[index + 1]
                paragraph, region = _cell_blank_paragraph(target)
                if paragraph is None:
                    continue
                field_id = visible_field_id(label, role_id=role_id, category=category, button_label=button_label)
                value = str(values.get(field_id, "") or "").strip()
                if not value:
                    continue
                if region is None:
                    paragraph.add_run(value)
                else:
                    _replace_paragraph_span(paragraph, region[0], region[1], value)
                filled.append(field_id)
    return filled


def fill_docx_visible_fields(
    path: str | Path,
    values: Mapping[str, str],
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> tuple[str, ...]:
    document = Document(str(path))
    filled = [
        *_fill_labelled_paragraphs(
            document,
            values,
            role_id=role_id,
            category=category,
            button_label=button_label,
        ),
        *_fill_tables(
            document,
            values,
            role_id=role_id,
            category=category,
            button_label=button_label,
        ),
    ]
    if filled:
        document.save(str(path))
    return tuple(dict.fromkeys(filled))


def visible_fill_field_ids(
    path: str | Path,
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> tuple[str, ...]:
    """Return semantic ids for every ordinary Word blank this renderer can fill."""
    document = Document(str(Path(path).expanduser()))
    field_ids: list[str] = []
    for paragraph in iter_all_story_paragraphs(document):
        slot = visible_blank_slot(paragraph.text)
        if slot is not None:
            field_ids.append(visible_field_id(slot.label, role_id=role_id, category=category, button_label=button_label))
    for table in iter_document_tables(document):
        for row in table.rows:
            cells: list = []
            seen: set[int] = set()
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key); cells.append(cell)
            for index, cell in enumerate(cells[:-1]):
                label = normalize(cell.text).strip(" :：")
                if not label:
                    continue
                paragraph, _region = _cell_blank_paragraph(cells[index + 1])
                if paragraph is not None:
                    field_ids.append(visible_field_id(label, role_id=role_id, category=category, button_label=button_label))
    return tuple(dict.fromkeys(field_ids))
