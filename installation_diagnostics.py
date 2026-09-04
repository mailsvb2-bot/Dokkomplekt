"""Doctor-facing installation self-check for Windows/portable builds.

The diagnostic intentionally stays stdlib-only and keeps patient data out of the
report.  It answers the practical support question: why does nothing happen when
an intake DOCX is dropped into ``Выписанные пациенты``?
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import os
import shutil
import sys
from typing import Iterable

from diagnostic_logging import record_soft_exception

INSTALLATION_DIAGNOSTICS_LOCK_VERSION = "v1.0"


@dataclass(frozen=True)
class DiagnosticRow:
    """One doctor-readable self-check row."""

    name: str
    ok: bool
    value: str
    advice: str = ""

    def line(self) -> str:
        icon = "✅" if self.ok else "⚠"
        suffix = f" — {self.advice}" if self.advice else ""
        return f"{icon} {self.name}: {self.value}{suffix}"


def _app_root() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def _data_root() -> Path:
    if getattr(sys, "frozen", False):
        base = os.environ.get("APPDATA")
        root = Path(base) if base else Path.home()
        return root / "MedicalDiaryAutofill"
    return _app_root() / ".medical_diary_autofill_data"


def _startup_dir() -> Path | None:
    if os.name != "nt":
        return None
    startup = os.environ.get("APPDATA")
    if not startup:
        return None
    return Path(startup) / "Microsoft" / "Windows" / "Start Menu" / "Programs" / "Startup"


def _startup_autostart_paths() -> tuple[Path, Path] | tuple[()]:
    folder = _startup_dir()
    if folder is None:
        return ()
    return (
        folder / "MedicalDiaryAutofill Intake Agent.vbs",
        folder / "MedicalDiaryAutofill Intake Agent.lnk",
    )


def _safe_exists(path: Path | None) -> bool:
    try:
        return bool(path and path.exists())
    except Exception as exc:
        record_soft_exception("installation_diagnostics.exists", exc, detail=str(path))
        return False


def _safe_text_tail(path: Path, *, limit: int = 1200) -> str:
    try:
        if not path.exists() or not path.is_file():
            return ""
        text = path.read_text(encoding="utf-8", errors="replace")
        return text[-limit:].strip()
    except Exception as exc:
        record_soft_exception("installation_diagnostics.tail", exc, detail=str(path))
        return ""


def _has_word_or_docx_opener() -> bool:
    if os.name == "nt":
        return True
    return bool(shutil.which("libreoffice") or shutil.which("soffice") or shutil.which("open") or shutil.which("xdg-open"))


def _doctor_button_count(app: object | None) -> int:
    try:
        if app is None:
            return 0
        from layout_checklist import _doctor_buttons_setup_completed
        from universal_main_documents import custom_documents_for_main_ui
        pack = app._load_or_create_universal_pack() if hasattr(app, "_load_or_create_universal_pack") else None
        if pack is None or not _doctor_buttons_setup_completed(pack):
            return 0
        return len(custom_documents_for_main_ui(pack, base_dir=app._universal_profile_path().parent if hasattr(app, "_universal_profile_path") else None))
    except Exception as exc:
        record_soft_exception("installation_diagnostics.button_count", exc)
        return 0


def collect_installation_diagnostics(app: object | None = None) -> list[DiagnosticRow]:
    """Collect actionable installation diagnostics without importing Tk."""
    rows: list[DiagnosticRow] = []
    root = _app_root()
    data = _data_root()
    rows.append(DiagnosticRow("Папка программы", root.exists(), str(root), "проверьте распаковку архива" if not root.exists() else ""))
    rows.append(DiagnosticRow("Папка настроек", data.exists(), str(data), "запустите программу один раз" if not data.exists() else ""))
    try:
        from desktop_intake import default_intake_folder, is_desktop_intake_folder_path
        intake = default_intake_folder()
        rows.append(DiagnosticRow("Папка «Выписанные пациенты»", intake.exists() and intake.is_dir() and is_desktop_intake_folder_path(intake), str(intake), "создайте папку через первый запуск программы" if not intake.exists() else ""))
    except Exception as exc:
        record_soft_exception("installation_diagnostics.intake_folder", exc)
        rows.append(DiagnosticRow("Папка «Выписанные пациенты»", False, "не удалось проверить", "откройте настройки первого запуска"))
    autostart_paths = _startup_autostart_paths()
    if os.name == "nt":
        existing_autostart = [path for path in autostart_paths if _safe_exists(path)]
        value = ", ".join(str(path) for path in existing_autostart) if existing_autostart else "не найден"
        rows.append(DiagnosticRow("Фоновое наблюдение: автозагрузка", bool(existing_autostart), value, "повторно включите фоновое наблюдение из программы" if not existing_autostart else ""))
    else:
        rows.append(DiagnosticRow("Фоновое наблюдение: автозагрузка", True, "не Windows-среда", "боевой тест нужен на Windows"))
    lock = data / "desktop_intake_agent.lock"
    log = data / "desktop_intake_agent.log"
    lock_ok = lock.exists()
    lock_value = str(lock) if lock_ok else "нет активного lock"
    try:
        if lock_ok:
            age = max(0, int(__import__("time").time() - lock.stat().st_mtime))
            lock_value = f"{lock} (обновлён {age} сек назад)"
            if age > 180:
                lock_ok = False
                lock_value += "; похоже, stale-lock"
    except Exception as exc:
        record_soft_exception("installation_diagnostics.lock_age", exc, detail=str(lock))
    rows.append(DiagnosticRow("Защита фонового наблюдения", lock_ok, lock_value, "перезапустите программу, чтобы фоновое наблюдение настроилось заново" if not lock_ok else ""))
    rows.append(DiagnosticRow("Журнал фонового наблюдения", log.exists(), str(log) if log.exists() else "лог пока не создан", "перенесите тестовый DOCX в папку" if not log.exists() else ""))
    tail = _safe_text_tail(log, limit=500)
    if tail:
        rows.append(DiagnosticRow("Последние сообщения фонового наблюдения", True, tail.replace("\n", " | ")))
    rows.append(DiagnosticRow("Открытие DOCX", _has_word_or_docx_opener(), "системная ассоциация/Word/LibreOffice", "установите Word или LibreOffice" if not _has_word_or_docx_opener() else ""))
    count = _doctor_button_count(app)
    rows.append(DiagnosticRow("Кнопки документов блока 03", count > 0, f"{count} активных", "создайте кнопки документов в центре шаблонов" if count <= 0 else ""))
    try:
        diary_texts = bool(getattr(app, "status_files", None) or getattr(app, "diary_texts_dir", "")) if app is not None else False
        diary_dates = bool(getattr(app, "diary_files", None) or getattr(app, "diary_template_dir", "")) if app is not None else False
        rows.append(DiagnosticRow("Дневники — тексты", diary_texts, "выбраны/папка задана" if diary_texts else "не выбраны", "нажмите «Тексты» во втором блоке" if not diary_texts else ""))
        rows.append(DiagnosticRow("Дневники — даты", diary_dates, "выбраны/папка задана" if diary_dates else "не выбраны", "нажмите «Даты» во втором блоке" if not diary_dates else ""))
    except Exception as exc:
        record_soft_exception("installation_diagnostics.diary_state", exc)
    return rows


def render_installation_diagnostics(rows: Iterable[DiagnosticRow]) -> str:
    lines = ["ПРОВЕРКА УСТАНОВКИ MEDICALDIARYAUTOFILL", ""]
    for row in rows:
        lines.append(row.line())
    lines.extend(["", "Если есть ⚠ — исправьте верхний пункт и повторите проверку.", "Боевой тест Windows: закрыть программу → перенести DOCX в «Выписанные пациенты» → программа должна открыться сама."])
    return "\n".join(lines)


def show_installation_diagnostics(app: object) -> None:
    try:
        from tkinter import messagebox
        text = render_installation_diagnostics(collect_installation_diagnostics(app))
        messagebox.showinfo("Проверить программу", text)
        try:
            if hasattr(app, "_set_status"):
                app._set_status("Проверка программы завершена")
        except Exception as status_exc:
            record_soft_exception("installation_diagnostics.status", status_exc)
    except Exception as exc:
        record_soft_exception("installation_diagnostics.show", exc)


def assert_installation_diagnostics_lock() -> None:
    if INSTALLATION_DIAGNOSTICS_LOCK_VERSION != "v1.0":
        raise AssertionError("Installation diagnostics lock changed unexpectedly")
    rows = collect_installation_diagnostics(None)
    if len(rows) < 6:
        raise AssertionError("Installation diagnostics must return actionable rows")


def run_user_journey_check() -> dict[str, object]:
    import tempfile
    import time
    from dataclasses import replace
    from datetime import datetime, timezone
    result: dict[str, object] = {"check": "user_journey", "ok": False, "checks": {}, "error": None}
    checks: dict[str, bool] = result["checks"]  # type: ignore[assignment]
    try:
        from docx import Document
        from desktop_intake import prepare_patient_work_folder, scan_primary_candidates
        from diary_schedule import DiaryScheduleSpec
        from medical_docx_reader import extract_docx_text
        from medical_service import MedicalDocumentService
        from output_transaction import OutputTransaction
        from product_access import apply_watermark_to_files
        from product_access.native import NativeProductAccessManager
        from universal_case_adapter import patient_data_to_case
        from universal_diary_generation import render_diary_documents_from_pack
        from universal_generation import render_documents_from_pack
        from universal_profiles import default_document_pack
        from universal_template_engine import attach_template_to_pack
        def write_docx(path: Path, paragraphs: list[str]) -> Path:
            document = Document()
            for text in paragraphs:
                document.add_paragraph(text)
            document.save(path)
            return path
        with tempfile.TemporaryDirectory(prefix="dokkomplekt-user-journey-") as raw:
            root = Path(raw)
            intake = root / "Выписанные пациенты"
            intake.mkdir()
            source = write_docx(intake / "01.09.2026 Первичный осмотр.docx", ["01.09.2026 Первичный осмотр", "Ф.И.О.: ИВАНОВ ИВАН ИВАНОВИЧ", "Дата поступления: 01.09.2026", "Жалобы: тревога", "Диагноз: K35.8 Острый аппендицит"])
            old = time.time() - 5
            os.utime(source, (old, old))
            candidates = scan_primary_candidates(intake, set())
            checks["desktop_intake_detects_primary"] = len(candidates) == 1 and candidates[0].path == source
            patient_dir, primary = prepare_patient_work_folder(intake, source)
            checks["patient_folder_created"] = patient_dir.name == "ИВАНОВ И.И. сентябрь 2026"
            checks["primary_moved_into_patient_folder"] = primary.exists() and primary.parent == patient_dir and not source.exists()
            data = MedicalDocumentService().parse_primary_document(primary)
            checks["primary_parsed"] = data.fio == "ИВАНОВ ИВАН ИВАНОВИЧ" and data.admission_date == "01.09.2026" and "K35.8" in data.diagnosis
            data.case_number = "К-777"
            data.treatment_plan = "терапия из пользовательского popup"
            data.discharge_date = "05.09.2026"
            case = patient_data_to_case(data, source_document=str(primary))
            checks["popup_values_reach_case"] = case.get("case.number") == "К-777" and case.get("treatment.plan") == "терапия из пользовательского popup" and case.get("discharge.date") == "05.09.2026"
            profile_dir = root / "profile"
            profile_dir.mkdir()
            pack = default_document_pack()
            primary_template = write_docx(root / "Первичный шаблон.docx", ["Пациент {{patient.fio}}", "История болезни № {{case.number}}", "Диагноз {{diagnosis.main}}", "План лечения {{treatment.plan}}"])
            attach_template_to_pack(pack, primary_template, profile_dir, button_label="Первичный осмотр", document_id="doctor_primary", category="medical", role_id="primary_exam")
            discharge_template = write_docx(root / "Выписной шаблон.docx", ["Выписной эпикриз {{patient.fio}}", "История болезни № {{case.number}}", "Диагноз {{diagnosis.main}}", "Лечение {{treatment.plan}}", "Дата выписки {{discharge.date}}"])
            attach_template_to_pack(pack, discharge_template, profile_dir, button_label="Выписной эпикриз", document_id="doctor_discharge", category="medical", role_id="discharge")
            diary_template = write_docx(root / "Дневники шаблон.docx", ["Состояние стабильное, жалоб активно не предъявляет.", "Контактен, ориентирован, назначения выполняет."])
            diary_spec, _ = attach_template_to_pack(pack, diary_template, profile_dir, button_label="Дневники наблюдения", document_id="doctor_diary", category="diaries", role_id="daily_diary")
            diary_spec = replace(diary_spec, category="diaries", role_id="daily_diary", diary_schedule=DiaryScheduleSpec("daily", (1, 2, 3, 4), (), 1.0, "packaged_user_journey").to_dict())
            pack.add_document(diary_spec)
            checks["doctor_owned_buttons_ready"] = {doc.id for doc in pack.documents} == {"doctor_primary", "doctor_discharge", "doctor_diary"}
            transaction = OutputTransaction(final_dir=patient_dir)
            stage_dir = transaction.begin()
            regular = render_documents_from_pack(pack=pack, case=case, document_ids=["doctor_primary", "doctor_discharge"], output_dir=stage_dir, base_dir=profile_dir, strict=True)
            diaries = render_diary_documents_from_pack(pack=pack, case=case, document_ids=["doctor_diary"], output_dir=stage_dir, base_dir=profile_dir, status_files=[], patient_name=case.get("patient.fio"), admission_value=case.get("admission.date"), discharge_value=case.get("discharge.date"), diary_day_offsets=(1, 2, 3, 4), remove_holiday_rows=False, force_final_diary=False)
            checks["regular_documents_created"] = regular.ok and len(regular.created_files) == 2
            checks["diaries_created"] = not diaries.skipped and len(diaries.created_files) == 1

            staged_files = [Path(path) for path in [*regular.created_files, *diaries.created_files]]
            access = NativeProductAccessManager(
                storage_dir=root / "product_access",
                now=datetime(2026, 9, 4, 12, 0, tzinfo=timezone.utc),
            )
            access_decision = access.check_document_creation(len(staged_files))
            checks["trial_access_allows_real_creation"] = access_decision.allowed and access_decision.state.plan == "trial"
            watermark = access.current_watermark_text()
            watermark_result = apply_watermark_to_files(staged_files, watermark)
            checks["trial_watermark_applied"] = bool(watermark) and not watermark_result.errors and watermark_result.changed_count == len(staged_files)

            reservation = access.reserve_created_documents(len(staged_files))
            mapping = transaction.commit(expected_files=staged_files)
            access.finalize_created_documents(reservation)
            committed_regular = [mapping[Path(path)] for path in regular.created_files]
            committed_diaries = [mapping[Path(path)] for path in diaries.created_files]
            committed_files = [*committed_regular, *committed_diaries]
            checks["trial_usage_committed"] = access.current_state().documents_used_total_trial == len(committed_files)
            checks["transaction_published_outputs"] = len(committed_files) == 3 and all(path.exists() for path in committed_files)

            footer_text = "\n".join(
                paragraph.text
                for path in committed_files
                for section in Document(str(path)).sections
                for paragraph in section.footer.paragraphs
            )
            checks["trial_watermark_survives_commit"] = watermark in footer_text
            regular_texts = [extract_docx_text(path) for path in committed_regular]
            diary_text = extract_docx_text(committed_diaries[0]) if committed_diaries else ""
            checks["regular_content_correct"] = all("К-777" in text and "K35.8" in text for text in regular_texts) and any("терапия из пользовательского popup" in text for text in regular_texts) and any("05.09.2026" in text for text in regular_texts)
            checks["diary_calendar_correct"] = "02.09.26" in diary_text and "05.09.26" in diary_text and "01.09.26" not in diary_text and "06.09.26" not in diary_text
            checks["diary_signatures_present"] = "Лечащий врач" in diary_text and "Зав. отделением" in diary_text
            checks["outputs_stay_in_patient_folder"] = all(path.parent == patient_dir for path in committed_files)
            checks["intake_does_not_retrigger"] = scan_primary_candidates(intake, set()) == ()
        result["ok"] = bool(checks) and all(checks.values())
    except Exception as exc:
        result["error"] = repr(exc)
    return result
