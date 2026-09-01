"""Entry point for MedicalDiaryAutofill.

The large Tkinter controller is intentionally split into focused modules:
configuration, reusable UI components, settings persistence, dialogs, file input,
numbered diary-template discovery, drag-and-drop, and creation actions.
``main.py`` stays small so the executable entry point remains stable.
"""

from __future__ import annotations

import json
import sys
import tempfile
import traceback
from pathlib import Path


def __getattr__(name: str):
    """Expose legacy smoke imports lazily without heavy startup imports."""
    if name == "CombinedMedicalDiaryApp":
        from app import CombinedMedicalDiaryApp

        return CombinedMedicalDiaryApp
    try:
        import app_config

        return getattr(app_config, name)
    except AttributeError as exc:
        raise AttributeError(name) from exc


def _check_native_license_core() -> int:
    result = {"check": "native_license_core", "module": "dokkomplekt_license_native", "ok": False, "version": None, "functions": {}, "error": None}
    try:
        import dokkomplekt_license_native as native
        version = native.native_core_version()
        functions = {"native_core_version": callable(getattr(native, "native_core_version", None)), "license_plan": callable(getattr(native, "license_plan", None)), "proof_ok": callable(getattr(native, "proof_ok", None)), "access_decision": callable(getattr(native, "access_decision", None))}
        result["version"] = str(version)
        result["functions"] = functions
        result["ok"] = version == "0.1.0" and all(functions.values())
    except Exception as exc:
        result["error"] = repr(exc)
    print(json.dumps(result, ensure_ascii=False, sort_keys=True))
    return 0 if result["ok"] else 1


def _check_runtime_bundle() -> int:
    result = {"check": "runtime_bundle", "ok": False, "checks": {}, "error": None}
    checks = result["checks"]
    try:
        import lxml  # noqa: F401
        import tkinterdnd2  # noqa: F401
        import win32api  # type: ignore # noqa: F401
        import win32print  # type: ignore # noqa: F401
        from docx import Document
        from medical_docx_blocks import extract_docx_text
        from medical_service import discover_primary_documents
        checks["imports"] = True
        with tempfile.TemporaryDirectory(prefix="dokkomplekt-runtime-smoke-") as temp_dir:
            root = Path(temp_dir)
            primary = root / "01.09.2026 Первичный осмотр.docx"
            document = Document()
            document.add_paragraph("01.09.2026 Первичный осмотр")
            document.add_paragraph("ФИО: ИВАНОВ ИВАН ИВАНОВИЧ")
            document.add_paragraph("Жалобы")
            document.add_paragraph("Анамнез")
            document.add_paragraph("Диагноз")
            document.save(primary)
            text = extract_docx_text(primary)
            checks["docx_read"] = "Первичный осмотр" in text and "ИВАНОВ" in text
            checks["batch_discovery"] = discover_primary_documents(root) == (primary,)
        result["ok"] = bool(checks) and all(bool(value) for value in checks.values())
    except Exception as exc:
        result["error"] = repr(exc)
    print(json.dumps(result, ensure_ascii=False, sort_keys=True))
    return 0 if result["ok"] else 1


def _check_user_journey() -> int:
    from installation_diagnostics import run_user_journey_check
    result = run_user_journey_check()
    print(json.dumps(result, ensure_ascii=False, sort_keys=True))
    return 0 if result.get("ok") else 1


def main() -> None:
    if "--check-native-license-core" in sys.argv:
        raise SystemExit(_check_native_license_core())
    if "--check-runtime-bundle" in sys.argv:
        raise SystemExit(_check_runtime_bundle())
    if "--check-user-journey" in sys.argv:
        raise SystemExit(_check_user_journey())
    if "--install-intake-agent" in sys.argv:
        from desktop_intake_agent import install_agent_autostart
        ok, message = install_agent_autostart(start_now=True)
        print(message)
        raise SystemExit(0 if ok else 1)
    if "--intake-agent" in sys.argv:
        from desktop_intake_agent import run_forever
        run_forever()
        return
    from diagnostic_logging import record_soft_exception
    from tkinter import messagebox
    from app import CombinedMedicalDiaryApp
    from startup import _create_root, _startup_log_path, _write_startup_error
    try:
        root = _create_root()
        CombinedMedicalDiaryApp(root)
        root.mainloop()
    except Exception as exc:
        details = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__))
        _write_startup_error(details)
        try:
            messagebox.showerror("Ошибка запуска", f"Программа не запустилась. Подробности записаны в файл:\n{_startup_log_path()}\n\n{exc}")
        except Exception as dialog_exc:
            record_soft_exception("main.startup_error_dialog", dialog_exc)
        raise


if __name__ == "__main__":
    main()
