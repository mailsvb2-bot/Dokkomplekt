from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from diagnostic_logging import record_soft_exception
from medical_docx_xml_fragments import ensure_docx_compatible
from medical_text_utils import normalize_text


def iter_block_items(parent) -> Iterable[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def extract_docx_text(path: str | Path) -> str:
    """Extract visible Word text from body, tables, headers, footers and XML.

    The primary-document and diary flows rely on this as a broad local scanner:
    it keeps table order, avoids duplicate merged-cell text, includes header and
    footer stories, then supplements python-docx output with raw XML fragments so
    Word fields and less-common text containers are still discoverable.
    """
    compatible_path = ensure_docx_compatible(path, label="Word document")
    doc = Document(str(compatible_path))
    lines: List[str] = []

    def walk(parent):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                lines.append(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    seen_cells: set[int] = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_cells:
                            continue
                        seen_cells.add(tc_id)
                        walk(cell)

    walk(doc)
    for section in doc.sections:
        for area in (section.header, section.footer):
            for paragraph in area.paragraphs:
                lines.append(paragraph.text)
            for table in area.tables:
                for row in table.rows:
                    seen_cells: set[int] = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_cells:
                            continue
                        seen_cells.add(tc_id)
                        walk(cell)
    # Preserve repeated structural labels from tables/body. Only supplementary
    # raw XML fragments are deduplicated against already collected content.
    normalized_lines: list[str] = []
    seen_structured: set[str] = set()
    for line in lines:
        normalized = normalize_text(str(line or ""))
        if not normalized:
            continue
        normalized_lines.append(normalized)
        key = normalized.casefold().replace("ё", "е")
        seen_structured.add(key)

    xml_lines: list[str] = []
    try:
        from medical_docx_xml_fragments import _docx_xml_text_fragments
        xml_lines = _docx_xml_text_fragments(compatible_path)
    except Exception as exc:
        record_soft_exception("medical_docx_blocks.xml_fragments", exc, detail=str(compatible_path))

    for line in xml_lines:
        normalized = normalize_text(str(line or ""))
        if not normalized:
            continue
        key = normalized.casefold().replace("ё", "е")
        if key in seen_structured:
            continue
        seen_structured.add(key)
        normalized_lines.append(normalized)
    deduped = normalized_lines
    return normalize_text("\n".join(deduped))
