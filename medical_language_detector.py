"""Deterministic language detection for uploaded medical DOCX text.

The detector is intentionally lightweight and local-only.  It does not claim
legal certainty; it selects the best language hint for UI defaults, extraction
aliases and orthography rules.
"""

from __future__ import annotations

from dataclasses import dataclass
import re
from pathlib import Path

from docx import Document

from medical_language_catalog import normalize_language_id

LANGUAGE_DETECTION_LOCK_VERSION = "v1.0"

_ARMENIAN_RE = re.compile(r"[\u0531-\u058F]")
_GEORGIAN_RE = re.compile(r"[\u10A0-\u10FF]")
_CYRILLIC_RE = re.compile(r"[\u0400-\u04FF]")
_LATIN_RE = re.compile(r"[A-Za-zĄąĆćĘęŁłŃńÓóŚśŹźŻżƏəĞğİıÖöŞşÜüÇçʻʼ’]")

_MARKERS: dict[str, tuple[str, ...]] = {
    "ru": ("история болезни", "медицинская карта", "жалобы", "анамнез", "диагноз", "лечение", "рекомендации"),
    "uk": ("медична карта", "історія хвороби", "скарги", "анамнез", "діагноз", "лікування", "рекомендації"),
    "be": ("медыцынская карта", "гісторыя хваробы", "скаргі", "дыягназ", "лячэнне"),
    "kk": ("медициналық карта", "ауру тарихы", "шағым", "диагноз", "емдеу", "ұсыным"),
    "tg": ("таърихи беморӣ", "варақаи тиббӣ", "шикоят", "ташхис", "табобат"),
    "hy": ("հիվանդության պատմություն", "բժշկական քարտ", "գանգատ", "ախտորոշում", "բուժում"),
    "ka": ("სამედიცინო ბარათი", "ავადმყოფობის ისტორია", "ჩივილ", "დიაგნოზ", "მკურნალ"),
    "az": ("xəstəlik tarixi", "tibbi kart", "şikayət", "diaqnoz", "müalicə", "tövsiy"),
    "uz": ("tibbiy karta", "kasallik tarixi", "shikoyat", "tashxis", "davolash"),
    "tk": ("lukmançylyk kart", "kesel taryhy", "arz", "diagnoz", "bejergi"),
    "pl": (
        "historia choroby", "karta informacyjna", "karta leczenia", "pacjent",
        "rozpoznanie", "leczenie", "zalecenia", "data przyjęcia", "data przyjecia",
        "data wypisu", "dziennik obserwacji", "badanie", "skierowanie"
    ),
    "en": ("medical record", "case history", "complaints", "diagnosis", "treatment", "recommendations"),
}


@dataclass(frozen=True)
class LanguageDetection:
    language_id: str
    confidence: float
    script: str
    matched_markers: tuple[str, ...]

    @property
    def is_confident(self) -> bool:
        return self.confidence >= 0.62 and self.language_id != "auto"

    def to_dict(self) -> dict:
        return {
            "language_id": self.language_id,
            "confidence": self.confidence,
            "script": self.script,
            "matched_markers": list(self.matched_markers),
        }

    def human_label(self) -> str:
        percent = int(round(self.confidence * 100))
        return f"{self.language_id} ({percent}%, {self.script})"


def detect_text_language(text: str) -> LanguageDetection:
    # Placeholders are technical field ids ({{patient.fio}}, {{procedure.name}})
    # and must not decide the human language of a medical template.
    sample = re.sub(r"\{\{[^}]+\}\}", " ", str(text or ""))
    low = sample.lower()
    if not sample.strip():
        return LanguageDetection("auto", 0.0, "unknown", ())

    script_scores = {
        "hy": len(_ARMENIAN_RE.findall(sample)),
        "ka": len(_GEORGIAN_RE.findall(sample)),
        "cyrillic": len(_CYRILLIC_RE.findall(sample)),
        "latin": len(_LATIN_RE.findall(sample)),
    }
    marker_scores: dict[str, list[str]] = {}
    for lang, markers in _MARKERS.items():
        found = [marker for marker in markers if marker in low]
        if found:
            marker_scores[lang] = found

    if script_scores["hy"]:
        return _with_markers("hy", script_scores["hy"], len(sample), marker_scores.get("hy", ()), "armenian")
    if script_scores["ka"]:
        return _with_markers("ka", script_scores["ka"], len(sample), marker_scores.get("ka", ()), "georgian")

    if marker_scores:
        best_lang, found = max(marker_scores.items(), key=lambda item: (len(item[1]), sum(len(m) for m in item[1])))
        script = "cyrillic" if best_lang in {"ru", "uk", "be", "kk", "tg"} else "latin"
        confidence = min(0.96, 0.48 + len(found) * 0.13)
        return LanguageDetection(normalize_language_id(best_lang), confidence, script, tuple(found))

    cyr = script_scores["cyrillic"]
    lat = script_scores["latin"]
    if cyr > lat and cyr >= 8:
        # Without markers, Russian is the safest Cyrillic default for legacy templates.
        return LanguageDetection("ru", min(0.55, cyr / max(1, len(sample)) * 4), "cyrillic", ())
    if lat >= 8:
        # Latin script alone is not enough: Uzbek/Turkmen/Azerbaijani/English
        # all use Latin-ish alphabets.  Without language markers return auto,
        # so profile button labels can fall back to UI/document preference.
        return LanguageDetection("auto", min(0.45, lat / max(1, len(sample)) * 2), "latin", ())
    return LanguageDetection("auto", 0.0, "unknown", ())


def detect_docx_language(path: str | Path) -> LanguageDetection:
    doc = Document(str(Path(path).expanduser()))
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells if cell.text)
    for section in doc.sections:
        for area in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in area.paragraphs if paragraph.text)
            for table in area.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells if cell.text)
    return detect_text_language("\n".join(chunks))


def _with_markers(lang: str, char_count: int, total: int, markers, script: str) -> LanguageDetection:
    confidence = min(0.98, 0.55 + min(0.25, char_count / max(1, total)) + len(tuple(markers)) * 0.08)
    return LanguageDetection(lang, confidence, script, tuple(markers))


def assert_language_detection_lock() -> None:
    if LANGUAGE_DETECTION_LOCK_VERSION != "v1.0":
        raise AssertionError("Language detection lock changed unexpectedly")
    checks = (
        ("Հիվանդության պատմություն. Ախտորոշում", "hy"),
        ("სამედიცინო ბარათი. დიაგნოზი", "ka"),
        ("Xəstəlik tarixi. Diaqnoz və müalicə", "az"),
        ("Historia choroby. Rozpoznanie i leczenie", "pl"),
        # Russian text with Latin placeholders must stay Russian, not English.
        ("Протокол операции {{patient.fio}} {{procedure.name}}", "ru"),
    )
    for sample, expected in checks:
        actual = detect_text_language(sample).language_id
        if actual != expected:
            raise AssertionError(f"Language detector expected {expected!r}, got {actual!r} for {sample!r}")
