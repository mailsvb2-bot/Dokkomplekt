"""Facade рендера медицинских документов.

Документ-специфичные render_* методы вынесены в отдельные mixin-модули,
а публичный класс MedicalDocumentRenderer сохранён.
"""

from __future__ import annotations

from pathlib import Path

from medical_gender import adapt_patient_data_to_gender
from medical_models import PatientData
from medical_renderer_commission import MedicalRendererCommissionMixin
from medical_renderer_labs import MedicalRendererLabsMixin
from medical_renderer_primary import MedicalRendererPrimaryMixin
from medical_renderer_special import MedicalRendererSpecialMixin


def _append_additional_info_to_docx(output_path: str | Path, data: PatientData) -> None:
    text = str(getattr(data, "additional_info_text", "") or "").strip()
    if not text:
        return
    from docx import Document
    doc = Document(str(output_path))
    doc.add_paragraph("")
    doc.add_paragraph("Дополнительная информация:")
    for line in text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)
    doc.save(str(output_path))


class MedicalDocumentRenderer(
    MedicalRendererPrimaryMixin,
    MedicalRendererCommissionMixin,
    MedicalRendererSpecialMixin,
    MedicalRendererLabsMixin,
):
    def render(self, kind: str, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        methods = {
            "primary": self.render_primary,
            "discharge": self.render_discharge,
            "commission": self.render_commission,
            "vk_mse": self.render_vk_mse,
            "admission_doctor_referral": self.render_admission_doctor_referral,
            "sick_leave_vk": self.render_sick_leave_vk,
            "rvk": self.render_rvk,
        }
        method = methods.get(kind)
        if method is None:
            raise ValueError(f"Неизвестный тип медицинского документа: {kind}")
        gender_adapted_data = adapt_patient_data_to_gender(data)
        method(template_path, output_path, gender_adapted_data)
        _append_additional_info_to_docx(output_path, gender_adapted_data)
