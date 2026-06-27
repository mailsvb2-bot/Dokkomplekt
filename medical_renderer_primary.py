from __future__ import annotations

from pathlib import Path

from docx import Document

from medical_constants import TARGET_MEDICAL_FACILITY
from medical_docx_editor import (
    DocxBlockEditor,
    iter_all_paragraphs,
    remove_exact_paragraphs,
    set_paragraph_text,
)
from medical_expert import put_expert_anamnesis
from medical_formatting import (
    format_birth_for_person_line,
    format_date_with_russian_year_suffix,
    format_military_commissariat_area,
    format_military_commissariat_referral,
    treatment_period_text,
)
from medical_gender import finalize_medical_document
from medical_markers import (
    COMMISSION_MARKERS,
    DISCHARGE_MARKERS,
    PRIMARY_MARKERS,
    RVK_MARKERS,
    SICK_LEAVE_VK_MARKERS,
    VK_MSE_MARKERS,
)
from medical_models import PatientData
from medical_parser_sanitize import sanitize_diagnosis
from medical_text_utils import normalize_match


class MedicalRendererPrimaryMixin:
    def render_primary(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Implement the render_primary workflow with validation, UI state updates and diagnostics."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        header_date = data.admission_date or "Дата"
        header = f"{header_date} 10:00      Первичный осмотр"
        editor.replace_first_matching_paragraph(["Дата, время"], header)
        editor.replace_block(["История болезни №"], "История болезни №", data.case_number, PRIMARY_MARKERS, preserve_when_empty=False, allow_empty=True)
        editor.replace_block(["Ф.И.О.", "ФИО"], "Ф.И.О.:", data.fio, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Год рождения", "Дата рождения"], "Год рождения:", data.birth, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Зарегистрирован"], "Зарегистрирован:", data.registered, PRIMARY_MARKERS, allow_empty=True)
        # В итоговом первичном осмотре не оставляем шаблонные варианты
        # «состоит/не состоит», «нужен/не нужен», «да/нет». Если источник не
        # содержит значения, строка остаётся с чистой подписью без мусорной подсказки.
        editor.replace_block(["На учёте у психиатров", "На учете у психиатров"], "Профильное наблюдение:", data.psych_account, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Работает в организации"], "Работает в организации:", data.work_org, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Должность"], "Должность:", data.position, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Больничный лист"], "Больничный лист:", data.sick_leave, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Оформление инвалидности"], "Оформление инвалидности:", data.disability, PRIMARY_MARKERS, allow_empty=True)
        rvk_referral = data.rvk_referral
        if not rvk_referral and data.rvk_military_commissariat:
            rvk_referral = format_military_commissariat_referral(data.rvk_military_commissariat)
        rvk_replaced = editor.replace_block(["Направление от РВК"], "Направление от РВК:", rvk_referral, PRIMARY_MARKERS, allow_empty=True)
        if rvk_referral and not rvk_replaced:
            normalized_referral = normalize_match(rvk_referral)
            already_present = any(normalized_referral and normalized_referral in normalize_match(paragraph.text or "") for paragraph in iter_all_paragraphs(doc))
            if not already_present:
                inserted = editor.insert_before_first_matching_paragraph(["В 3 отделение КДП поступает", "Жалобы на момент осмотра", "Жалобы"], rvk_referral)
                if not inserted:
                    doc.add_paragraph(rvk_referral)
        editor.remove_all_matching_paragraphs(["Экспертный анамнез"])
        editor.replace_block(["В 3 отделение КДП поступает", "Поступает"], "Поступает:", data.admission, PRIMARY_MARKERS)
        editor.replace_block(["Жалобы на момент осмотра", "Жалобы"], "Жалобы на момент осмотра:", data.complaints, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Психический статус"], "Профильный статус:", data.mental_status, PRIMARY_MARKERS)
        editor.replace_block(["Соматический статус"], "Соматический статус:", data.somatic_status, PRIMARY_MARKERS)
        editor.replace_block(["План обследования"], "План обследования:", data.examination_plan, PRIMARY_MARKERS)
        editor.replace_block(["План лечения"], "План лечения:", data.treatment_plan, PRIMARY_MARKERS)

        diagnosis_sentence = ""
        diagnosis = sanitize_diagnosis(data.diagnosis)
        if diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, профильного статуса, "
                f"данных клинических исследований установлен диагноз: {diagnosis}"
            )
        editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, PRIMARY_MARKERS)
        editor.replace_block(["Эпидемиологический анамнез"], "Эпидемиологический анамнез:", data.epidemiology, PRIMARY_MARKERS)
        put_expert_anamnesis(
            editor,
            data,
            PRIMARY_MARKERS,
            ["Врач", "Врач", "Зав. отделением", "Зав. отд."],
            include_sick_leave_number=False,
            include_sick_leave=False,
            include_return_to_work=False,
            replace_existing=False,
        )
        editor.replace_block(["Врач", "Врач"], "Врач", data.doctor, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Зав. отделением", "Зав. отд."], "Зав. отделением", data.head, PRIMARY_MARKERS, allow_empty=True)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_discharge(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        header_date = data.discharge_date or data.admission_date or "Дата, время"
        header = f"{header_date}      Выписной эпикриз № {data.case_number}".rstrip()
        editor.replace_first_matching_paragraph(["Дата, время"], header)
        birth_text = format_birth_for_person_line(data.birth)
        person_line = f"{data.fio}, {birth_text}, зарегистрирован по адресу: {data.registered}".strip(" ,")
        editor.replace_first_matching_paragraph(["г.р.,", "зарегистрирован по адресу"], person_line)
        period = f"Находился на лечении в медицинской организации по профилю с {data.admission_date} по {data.discharge_date}".strip()
        editor.replace_first_matching_paragraph(["Находился на лечении"], period)

        put_expert_anamnesis(editor, data, DISCHARGE_MARKERS, ["В 3 отделение КДП поступает"])

        editor.replace_block(["В 3 отделение КДП поступает", "Поступает"], "Поступает:", data.admission, DISCHARGE_MARKERS)
        editor.replace_block(["Жалобы при поступлении", "Жалобы"], "Жалобы при поступлении:", data.complaints, DISCHARGE_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, DISCHARGE_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, DISCHARGE_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Профильный статус при поступлении:", data.mental_status, DISCHARGE_MARKERS)
        diagnosis = sanitize_diagnosis(data.diagnosis)
        if diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, профильного статуса, "
                f"данных клинических исследований установлен диагноз: {diagnosis}"
            )
            if not editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, DISCHARGE_MARKERS):
                editor.insert_before_first_matching_paragraph(["Сомато-неврологический статус", "Соматический статус"], diagnosis_sentence)
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, DISCHARGE_MARKERS)
        self._replace_lab_lines(editor, dates, data=data, all_markers=DISCHARGE_MARKERS)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ –", data.epi_text, DISCHARGE_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        if data.treatment_plan:
            editor.replace_block(["Лечение"], "Лечение:", data.treatment_plan, DISCHARGE_MARKERS)
        signature = f"  Зав. отд. {data.head}                                                                                                 Врач\t{data.doctor}"
        editor.replace_first_matching_paragraph(["Зав. отд.", "Врач"], signature)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))
