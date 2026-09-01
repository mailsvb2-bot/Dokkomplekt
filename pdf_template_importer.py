from __future__ import annotations

from contextlib import suppress
import importlib
import os
from pathlib import Path
from typing import Iterable

from docx import Document

from diagnostic_logging import record_soft_exception
from document_intelligence.analyzer import DocumentIntelligenceCore
from document_intelligence.models import DocumentSource
from document_intelligence.pdf_reader import read_pdf_text
from personal_document_buttons import stable_document_id, unique_button_label
from universal_profiles import DocumentPack, DocumentTemplateSpec

PDF_TEMPLATE_IMPORT_LOCK_VERSION = "v1.0"


def import_pdf_templates_to_pack(
    pack: DocumentPack,
    pdf_paths: Iterable[str | Path],
    profile_base_dir: str | Path,
) -> tuple[str, ...]:
    base = Path(profile_base_dir).expanduser()
    templates_dir = base / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    existing_labels = {str(document.button_label or "").casefold() for document in pack.documents}
    imported: list[str] = []
    for raw_path in pdf_paths:
        source = Path(raw_path).expanduser()
        if source.suffix.lower() != ".pdf":
            raise ValueError(f"Unsupported source format: {source.name}")
        pdf_text = read_pdf_text(source)
        if not pdf_text.has_text:
            details = "; ".join(pdf_text.warnings) or "no extractable text"
            raise ValueError(f"PDF has no extractable text: {source.name}; {details}")
        blueprint = DocumentIntelligenceCore().analyze_source(DocumentSource(str(source), user_label=source.stem))
        label = unique_button_label(blueprint.title or source.stem, existing_labels)
        target = _available_path(templates_dir / (_safe_stem(label or source.stem) + ".docx"))
        conversion_mode = _write_pdf_template(source, target, pdf_text.text)
        spec = DocumentTemplateSpec(
            id=stable_document_id("pdf", label, source),
            button_label=label,
            template=target.name,
            output_name="{{patient.fio}} " + label + ".docx",
            required_fields=blueprint.required_field_ids,
            optional_fields=(),
            category=blueprint.domain or "custom",
            description=(
                "Converted from a user-owned PDF through Microsoft Word with layout preservation."
                if conversion_mode == "word"
                else "Reconstructed from extractable PDF text; complex PDF layout may require manual adjustment."
            ),
            role_id="pdf_source",
            button_language="auto",
            source_language="auto",
            button_label_source="pdf_document_intelligence",
        )
        pack.add_document(spec)
        existing_labels.add(label.casefold())
        imported.append(label)
    return tuple(imported)


def _write_pdf_template(source: Path, path: Path, text: str) -> str:
    """Prefer Word's PDF reflow on Windows; fall back to explicit text rebuild."""

    if os.name == "nt":
        word = None
        document = None
        try:
            win32com_client = importlib.import_module("win32com.client")
            word = win32com_client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(
                str(source.resolve()),
                ReadOnly=True,
                AddToRecentFiles=False,
                ConfirmConversions=False,
            )
            path.parent.mkdir(parents=True, exist_ok=True)
            document.SaveAs2(str(path.resolve()), FileFormat=16)
            if path.exists() and path.stat().st_size > 0:
                return "word"
        except Exception as exc:
            record_soft_exception("pdf_template_importer.word_pdf_reflow", exc, detail=str(source))
        finally:
            with suppress(Exception):
                if document is not None:
                    document.Close(False)
            with suppress(Exception):
                if word is not None:
                    word.Quit()

    _write_pdf_text_template(path, text)
    return "text"


def _write_pdf_text_template(path: Path, text: str) -> None:
    doc = Document()
    for raw_line in str(text or "").splitlines():
        line = " ".join(raw_line.split())
        if line:
            doc.add_paragraph(line)
    if not doc.paragraphs:
        doc.add_paragraph(path.stem)
    doc.save(path)


def _safe_stem(value: str) -> str:
    import re

    text = re.sub(r"[^a-zA-Z0-9._ -]+", "_", str(value or "pdf_template")).strip(" ._")
    return text[:80] or "pdf_template"


def _available_path(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(2, 1000):
        candidate = path.with_name(f"{path.stem} ({index}){path.suffix}")
        if not candidate.exists():
            return candidate
    raise FileExistsError(f"Could not allocate a free template path: {path}")


def assert_pdf_template_import_lock() -> None:
    if PDF_TEMPLATE_IMPORT_LOCK_VERSION != "v1.0":
        raise AssertionError("PDF template import lock changed unexpectedly")
