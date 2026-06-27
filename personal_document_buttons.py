"""Persistent doctor-owned document buttons for block 03.

The old production UI has built-in Russian tiles.  The universal product
needs a different contract: when a doctor adds a regular document template,
the button text becomes profile-owned data, preferably in the doctor's
working language.  This module only decides names/ids/metadata; it does not
render DOCX and does not mutate the legacy DOCUMENT_ORDER.
"""

from __future__ import annotations

from diagnostic_logging import record_soft_exception
from dataclasses import asdict, dataclass
from pathlib import Path
import hashlib
import re
from typing import Iterable, Mapping

from medical_language_catalog import normalize_language_id
from regulatory_document_roles import DocumentRole, default_document_role_registry

PERSONAL_DOCUMENT_BUTTON_LOCK_VERSION = "v1.4"
BUTTON_LABEL_IS_PROFILE_DATA = True
BUTTONS_DO_NOT_REQUIRE_HARDCODED_UI_TRANSLATION = True
BUTTON_ROLE_CONFIDENCE_THRESHOLD = 0.25
PROFILE_BUTTON_LABELS_ARE_PERSISTED = True
TOP_OF_DOCX_TITLE_RECOGNITION_ENABLED = True
SPECIALTY_NEUTRAL_TITLE_RECOGNITION_ENABLED = True
PROFILE_BUTTONS_ARE_NOT_PSYCHIATRY_BOUND = True


def safe_profile_filename(label: str) -> str:
    text = re.sub(r"[^0-9A-Za-zА-Яа-яЁё._-]+", "_", str(label or "doctor_profile")).strip("._-")
    return (text or "doctor_profile")[:80] + ".medpack.json"


def safe_profile_pack_id(label: str, target: str | Path | None = None) -> str:
    raw = str(label or "doctor_profile").strip() or "doctor_profile"
    base = re.sub(r"[^a-z0-9_]+", "_", raw.lower()).strip("_") or "doctor_profile"
    material = f"{raw}|{Path(target).as_posix() if target else ''}"
    suffix = hashlib.sha1(material.encode("utf-8", errors="replace")).hexdigest()[:8]
    return f"doctor.{base[:48]}_{suffix}"


def available_profile_path(path: Path) -> Path:
    candidate = Path(path).expanduser()
    if not candidate.exists():
        return candidate
    stem = candidate.name[:-len(".medpack.json")] if candidate.name.endswith(".medpack.json") else candidate.stem
    suffix = ".medpack.json" if candidate.name.endswith(".medpack.json") else candidate.suffix
    for index in range(2, 1000):
        next_candidate = candidate.with_name(f"{stem} ({index}){suffix}")
        if not next_candidate.exists():
            return next_candidate
    raise FileExistsError(f"Не удалось создать свободное имя профиля: {candidate}")


def unique_button_label(label: str, existing_labels: set[str]) -> str:
    base = re.sub(r"\s+", " ", str(label or "Документ")).strip() or "Документ"
    if base.casefold() not in existing_labels:
        return base
    for index in range(2, 1000):
        candidate = f"{base} ({index})"
        if candidate.casefold() not in existing_labels:
            return candidate
    raise ValueError(f"Слишком много кнопок с одинаковым названием: {base}")


# Human-facing labels for recurring medical document roles.  They are not a
# legal source of truth and they are intentionally editable by the doctor in
# the profile.  If a language/role is missing, the program falls back to the
# template title/file name rather than inventing a hidden hard binding.
LOCALIZED_ROLE_LABELS: dict[str, dict[str, str]] = {
    "hospitalization_referral": {
        "ru": "Направление на госпитализацию",
        "uk": "Направлення на госпіталізацію",
        "be": "Накіраванне на шпіталізацыю",
        "kk": "Стационарға жатқызуға жолдама",
        "uz": "Shifoxonaga yotqizish uchun yo‘llanma",
        "az": "Hospitalizasiya üçün göndəriş",
        "hy": "Հոսպիտալացման ուղեգիր",
        "ka": "ჰოსპიტალიზაციის მიმართვა",
        "tg": "Роҳхат барои бистарӣ кардан",
        "tk": "Hassahana ýerleşdirmek üçin ugrukdyrma",
        "en": "Hospitalization referral",
    },
    "admission_doctor_exam": {
        "ru": "Осмотр врача приёмного покоя",
        "uk": "Огляд лікаря приймального відділення",
        "be": "Агляд лекара прыёмнага аддзялення",
        "kk": "Қабылдау бөлімінің дәрігерлік қарауы",
        "uz": "Qabul bo‘limi shifokori ko‘rigi",
        "az": "Qəbul şöbəsi həkiminin müayinəsi",
        "hy": "Ընդունարանի բժշկի զննում",
        "ka": "მიმღები განყოფილების ექიმის გასინჯვა",
        "tg": "Муоинаи духтури шуъбаи қабул",
        "tk": "Kabul ediş bölüminiň lukmanynyň gözden geçirmesi",
        "en": "Admission doctor examination",
    },
    "primary_exam": {
        "ru": "Первичный осмотр",
        "uk": "Первинний огляд",
        "be": "Першасны агляд",
        "kk": "Алғашқы қарау",
        "uz": "Birlamchi ko‘rik",
        "az": "İlkin müayinə",
        "hy": "Առաջնային զննում",
        "ka": "პირველადი გასინჯვა",
        "tg": "Муоинаи аввалия",
        "tk": "Ilkinji gözden geçiriş",
        "en": "Initial examination",
    },
    "inpatient_record": {
        "ru": "История болезни",
        "uk": "Історія хвороби",
        "be": "Гісторыя хваробы",
        "kk": "Ауру тарихы",
        "uz": "Kasallik tarixi",
        "az": "Xəstəlik tarixi",
        "hy": "Հիվանդության պատմություն",
        "ka": "ავადმყოფობის ისტორია",
        "tg": "Таърихи беморӣ",
        "tk": "Kesel taryhy",
        "en": "Medical record",
    },
    "daily_diary": {
        "ru": "Дневник наблюдения",
        "uk": "Щоденник спостереження",
        "be": "Дзённік назірання",
        "kk": "Бақылау күнделігі",
        "uz": "Kuzatuv kundaligi",
        "az": "Müşahidə gündəliyi",
        "hy": "Դիտարկման օրագիր",
        "ka": "დაკვირვების დღიური",
        "tg": "Рӯзномаи мушоҳида",
        "tk": "Gözegçilik gündeligi",
        "en": "Observation diary",
    },
    "discharge_epicrisis": {
        "ru": "Выписной эпикриз",
        "uk": "Виписний епікриз",
        "be": "Выпісны эпікрыз",
        "kk": "Шығару эпикризі",
        "uz": "Chiqish epikrizi",
        "az": "Çıxarış epikrizi",
        "hy": "Դուրսգրման էպիկրիզ",
        "ka": "გაწერის ეპიკრიზი",
        "tg": "Эпикризи ҷавобшавӣ",
        "tk": "Çykaryş epikrizi",
        "en": "Discharge summary",
    },
    "transfer_epicrisis": {
        "ru": "Переводной эпикриз",
        "uk": "Перевідний епікриз",
        "be": "Пераводны эпікрыз",
        "kk": "Ауыстыру эпикризі",
        "uz": "O‘tkazish epikrizi",
        "az": "Köçürülmə epikrizi",
        "hy": "Փոխադրման էպիկրիզ",
        "ka": "გადაყვანის ეპიკრიზი",
        "tg": "Эпикризи гузаронидан",
        "tk": "Geçiriş epikrizi",
        "en": "Transfer summary",
    },
    "specialist_consultation": {
        "ru": "Консультационное заключение",
        "uk": "Консультаційний висновок",
        "be": "Кансультацыйнае заключэнне",
        "kk": "Консультациялық қорытынды",
        "uz": "Konsultativ xulosa",
        "az": "Konsultativ rəy",
        "hy": "Խորհրդատվական եզրակացություն",
        "ka": "კონსულტაციის დასკვნა",
        "tg": "Хулосаи машваратӣ",
        "tk": "Maslahat netijesi",
        "en": "Consultation report",
    },
    "operation_protocol": {
        "ru": "Протокол операции",
        "uk": "Протокол операції",
        "be": "Пратакол аперацыі",
        "kk": "Операция хаттамасы",
        "uz": "Operatsiya bayonnomasi",
        "az": "Əməliyyat protokolu",
        "hy": "Վիրահատության արձանագրություն",
        "ka": "ოპერაციის პროტოკოლი",
        "tg": "Протоколи амалиёт",
        "tk": "Operasiýa teswirnamasy",
        "en": "Operation protocol",
    },
    "anesthesia_preop": {
        "ru": "Осмотр анестезиолога",
        "uk": "Огляд анестезіолога",
        "be": "Агляд анестэзіёлага",
        "kk": "Анестезиолог қарауы",
        "uz": "Anesteziolog ko‘rigi",
        "az": "Anestezioloqun müayinəsi",
        "hy": "Անեսթեզիոլոգի զննում",
        "ka": "ანესთეზიოლოგის გასინჯვა",
        "tg": "Муоинаи анестезиолог",
        "tk": "Anesteziologyň gözden geçirmesi",
        "en": "Anesthesiologist examination",
    },
    "informed_consent": {
        "ru": "Информированное согласие",
        "uk": "Інформована згода",
        "be": "Інфармаваная згода",
        "kk": "Ақпараттандырылған келісім",
        "uz": "Axborotli rozilik",
        "az": "Məlumatlandırılmış razılıq",
        "hy": "Տեղեկացված համաձայնություն",
        "ka": "ინფორმირებული თანხმობა",
        "tg": "Розигии огоҳона",
        "tk": "Habarly razylyk",
        "en": "Informed consent",
    },
    "medical_commission": {
        "ru": "Врачебная комиссия",
        "uk": "Лікарська комісія",
        "be": "Урачэбная камісія",
        "kk": "Дәрігерлік комиссия",
        "uz": "Shifokorlar komissiyasi",
        "az": "Həkim komissiyası",
        "hy": "Բժշկական հանձնաժողով",
        "ka": "სამედიცინო კომისია",
        "tg": "Комиссияи тиббӣ",
        "tk": "Lukmançylyk topary",
        "en": "Medical commission",
    },
    "mse_referral": {
        "ru": "Направление на МСЭ",
        "uk": "Направлення на МСЕК",
        "be": "Накіраванне на МСЭ",
        "kk": "МӘС-ке жолдама",
        "uz": "TMEK uchun yo‘llanma",
        "az": "Tibbi-sosial ekspertizaya göndəriş",
        "hy": "Ուղեգիր բժշկասոցիալական փորձաքննության",
        "ka": "სამედიცინო-სოციალურ ექსპერტიზაზე მიმართვა",
        "tg": "Роҳхат ба экспертизаи тиббӣ-иҷтимоӣ",
        "tk": "Lukmançylyk-durmuş ekspertizasyna ugrukdyrma",
        "en": "Disability assessment referral",
    },
    "joint_medical_exam": {
        "ru": "Совместный осмотр",
        "en": "Joint medical examination",
    },
    "vk_mse": {
        "ru": "ВК на МСЭ",
        "en": "Medical commission for disability assessment",
    },
    "sick_leave_vk": {
        "ru": "ВК больничный",
        "en": "Sick-leave commission",
    },
    "military_commissariat_act": {
        "ru": "Акт для РВК",
        "en": "Military commissariat medical act",
    },
    "lab_results": {
        "ru": "Лабораторные результаты",
        "uk": "Лабораторні результати",
        "be": "Лабараторныя вынікі",
        "kk": "Зертханалық нәтижелер",
        "uz": "Laboratoriya natijalari",
        "az": "Laborator nəticələr",
        "hy": "Լաբորատոր արդյունքներ",
        "ka": "ლაბორატორიული შედეგები",
        "tg": "Натиҷаҳои лабораторӣ",
        "tk": "Laboratoriýa netijeleri",
        "en": "Laboratory results",
    },
    "instrumental_study": {
        "ru": "Инструментальное исследование",
        "uk": "Інструментальне дослідження",
        "be": "Інструментальнае даследаванне",
        "kk": "Аспаптық зерттеу",
        "uz": "Instrumental tekshiruv",
        "az": "İnstrumental müayinə",
        "hy": "Գործիքային հետազոտություն",
        "ka": "ინსტრუმენტული კვლევა",
        "tg": "Тадқиқоти инструменталӣ",
        "tk": "Instrumental barlag",
        "en": "Instrumental study",
    },
}


@dataclass(frozen=True)
class ButtonLabelSuggestion:
    label: str
    role_id: str = "unknown"
    document_id: str = ""
    language_id: str = "auto"
    source_language: str = "auto"
    source: str = "fallback"
    confidence: float = 0.0
    matched_markers: tuple[str, ...] = ()

    def to_dict(self) -> dict:
        return asdict(self)


def localized_role_label(role_id: str, language_id: str | None, *, fallback: str = "") -> str:
    role = str(role_id or "").strip().lower()
    lang = normalize_language_id(language_id, default="ru")
    by_lang = LOCALIZED_ROLE_LABELS.get(role, {})
    if lang in by_lang:
        return by_lang[lang]
    if "ru" in by_lang:
        return by_lang["ru"]
    registry_role = default_document_role_registry().get(role)
    if registry_role:
        return registry_role.label
    return normalize_button_label(fallback) or "Документ"


def regular_document_role_choices(language_id: str | None = "ru") -> tuple[str, ...]:
    lang = normalize_language_id(language_id, default="ru")
    choices: list[str] = []
    for role in default_document_role_registry().roles():
        label = localized_role_label(role.id, lang, fallback=role.label)
        choices.append(f"{label} [{role.id}]")
    return tuple(choices)


def role_id_from_choice(choice: str) -> str:
    marker = str(choice or "").rsplit("[", 1)
    if len(marker) == 2 and marker[1].endswith("]"):
        return marker[1][:-1].strip().lower()
    text = choice.strip().lower()
    for role in default_document_role_registry().roles():
        if text == role.id or text == role.label.lower():
            return role.id
    return "unknown"


def choose_button_language(
    *,
    preferred_language: str | None = None,
    source_language: str | None = None,
    ui_language: str | None = None,
) -> str:
    preferred = str(preferred_language or "").strip().lower()
    if preferred and preferred not in {"auto", "same_as_source"}:
        return normalize_language_id(preferred, default="ru")
    source = normalize_language_id(source_language, default="auto")
    if source != "auto":
        return source
    ui = normalize_language_id(ui_language, default="ru")
    return "ru" if ui == "auto" else ui



def neutralize_weak_role_for_generic_profile(
    role_id: str,
    confidence: float,
    *,
    top_title: str,
    explicit_specialty: str = "",
) -> str:
    """Keep role metadata from hijacking a neutral doctor's custom button.

    In a specialty-neutral profile, a title like "Справка для бассейна" or
    "Протокол эндоскопического исследования" must remain the doctor's document,
    not be silently attached to a legacy/commission workflow because one
    generic word matched a legacy role.  Strong universal roles are preserved
    only when the title itself says so.
    """

    role = str(role_id or "").strip().lower() or "unknown"
    specialty = str(explicit_specialty or "").strip().lower()
    title = str(top_title or "").strip().lower().replace("ё", "е")
    if specialty not in {"", "generic", "custom", "any"} or not top_title:
        return role
    if any(word in title for word in ("дневник", "наблюдени")):
        return "daily_diary"
    if "перевод" in title and "эпикриз" in title:
        return "transfer_epicrisis"
    if any(word in title for word in ("выпис", "эпикриз")):
        return "discharge_epicrisis"
    if "операц" in title and any(word in title for word in ("протокол", "ход", "операц")):
        return "operation_protocol"
    if "анестезиолог" in title:
        return "anesthesia_preop"
    if "соглас" in title:
        return "informed_consent"
    if "рвк" in title or "военком" in title or "военный комиссариат" in title:
        return "military_commissariat_act"
    if ("больнич" in title or "нетрудоспособ" in title) and ("вк" in title or "комисс" in title or "протокол" in title):
        return "sick_leave_vk"
    if any(word in title for word in ("вк на мсэ", "мсэ", "мсек", "медико-социаль")):
        return "vk_mse"
    if "совместный осмотр" in title or "комиссионный осмотр" in title:
        return "joint_medical_exam"
    if any(word in title for word in ("консультац", "консультатив", "заключение специалист")):
        return "specialist_consultation"
    if confidence < 0.45:
        return "unknown"
    return role


def suggest_button_label_for_template(
    template_path: str | Path,
    *,
    preferred_language: str | None = None,
    ui_language: str | None = "ru",
    explicit_specialty: str = "",
    explicit_role_id: str | None = None,
    fallback_label: str | None = None,
) -> ButtonLabelSuggestion:
    """Implement the suggest_button_label_for_template workflow with validation, UI state updates and diagnostics."""
    path = Path(template_path).expanduser()
    source_language = _safe_detect_language(path)
    language_id = choose_button_language(preferred_language=preferred_language, source_language=source_language, ui_language=ui_language)
    top_title = recognize_document_title_from_template(path)
    fallback = normalize_button_label(fallback_label or top_title or _template_title(path) or path.stem)
    role_id = str(explicit_role_id or "").strip().lower()
    confidence = 1.0 if role_id and role_id != "unknown" else 0.0
    matched: tuple[str, ...] = ()
    if not role_id or role_id == "unknown":
        try:
            from regulatory_document_classifier import classify_docx
            classification = classify_docx(path, explicit_specialty=explicit_specialty)
            role_id = classification.role_id
            confidence = classification.confidence
            matched = classification.best.matched_markers if classification.best else ()
        except Exception as exc:
            record_soft_exception("personal_document_buttons.classify_docx", exc, detail=str(path))
            role_id = "unknown"
            confidence = 0.0
    role_id = neutralize_weak_role_for_generic_profile(
        role_id,
        confidence,
        top_title=top_title,
        explicit_specialty=explicit_specialty,
    )
    # The 10-second doctor setup flow must preserve the title the doctor sees at
    # the top of the Word sheet: "Первичный осмотр", "Осмотр хирурга",
    # "Протокол операции" and so on.  Role i18n is still useful metadata, but it
    # must not rename a doctor's own button into a generic label.
    if top_title and not fallback_label:
        label = top_title
        source = "template_top_title"
    elif role_id != "unknown" and confidence >= BUTTON_ROLE_CONFIDENCE_THRESHOLD:
        label = localized_role_label(role_id, language_id, fallback=fallback)
        source = "role_i18n"
    else:
        label = fallback
        role_id = role_id if role_id else "unknown"
        source = "template_title"
    label = normalize_button_label(label)
    return ButtonLabelSuggestion(
        label=label,
        role_id=role_id or "unknown",
        document_id=stable_document_id(role_id or label, label, path),
        language_id=language_id,
        source_language=source_language,
        source=source,
        confidence=max(confidence, 0.82 if source == "template_top_title" else 0.0),
        matched_markers=tuple(matched),
    )




_DOCUMENT_TITLE_KEYWORDS = (
    "осмотр", "первичный", "повторный", "хирург", "терапевт", "невролог", "кардиолог",
    "уролог", "гинеколог", "травматолог", "эндокринолог", "гастроэнтеролог", "педиатр", "офтальмолог", "лор",
    "протокол", "операци", "манипуляц", "процедур", "эпикриз", "выпис", "переводн", "консультац", "заключение",
    "дневник", "наблюдени", "история болезни", "медицинская карта", "акт", "справка", "лист", "назначени",
    "исследован", "обследован", "описание", "консилиум", "освидетельств", "направление", "согласие", "рекомендац",
    "рвк", "вк", "мсэ", "комисси", "operation", "protocol", "discharge", "summary", "consultation", "examination", "exam",
    "certificate", "report", "referral", "consent", "procedure", "surgery",
)

_DOCUMENT_TITLE_STOPWORDS = (
    "министерство", "департамент", "комитет", "больница", "клиника", "диспансер", "поликлиника",
    "стационар", "отделение", "адрес", "телефон", "e-mail", "email", "инн", "огрн", "окпо",
    "утверждаю", "главный врач", "заведующий", "м.п.", "печать", "форма", "код формы",
    "ф.и.о", "фио", "пациент", "дата рождения", "история болезни", "палата", "пол", "возраст",
    "дата поступления", "дата выписки", "лечащий врач", "подпись", "страница",
)


def recognize_document_title_from_template(template_path: str | Path, *, max_blocks: int = 14) -> str:
    """Return a human document title from the visible upper DOCX area.

    Critical production rule: block-03 buttons are created from the document
    title the doctor sees at the top of the Word sheet.  The old implementation
    reused ``extract_docx_blocks()``, whose order is useful for scanning but not
    visual: all body paragraphs were returned before tables.  Many hospital
    forms keep the title in the first table row, so the button could be created
    from a random body phrase below.  This reader walks the DOCX body XML in
    real paragraph/table order and only then falls back to conservative scanner
    blocks.
    """

    path = Path(template_path).expanduser()
    try:
        blocks = _extract_visual_top_blocks(path, max_blocks=max_blocks)
    except Exception as exc:
        record_soft_exception("personal_document_buttons:recognize_title_visual", exc)
        blocks = ()
    title = _best_title_from_top_blocks(blocks)
    if title:
        return title
    try:
        from universal_scanner import extract_docx_blocks
        fallback_blocks = tuple((block.path_hint, block.text) for block in extract_docx_blocks(path)[:max_blocks])
    except Exception as exc:
        record_soft_exception("personal_document_buttons:recognize_title", exc)
        return ""
    return _best_title_from_top_blocks(fallback_blocks)


def _extract_visual_top_blocks(path: Path, *, max_blocks: int = 14) -> tuple[tuple[str, str], ...]:
    """Read non-empty top paragraphs/table rows in the order they appear on the page."""

    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(path))
    blocks: list[tuple[str, str]] = []

    def add_text(hint: str, raw: str) -> None:
        text = normalize_button_label(str(raw or ""))
        if text:
            blocks.append((hint, text))

    for child_index, child in enumerate(doc.element.body.iterchildren()):
        if len(blocks) >= max_blocks:
            break
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            add_text(f"body.paragraph[{child_index}]", paragraph.text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row_index, row in enumerate(table.rows):
                if len(blocks) >= max_blocks:
                    break
                cells = [normalize_button_label(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                add_text(f"body.table[{child_index}].row[{row_index}]", row_text)

    # Fallback only: some clinics put the title in a Word header.  It is scanned
    # after body top rows so hospital letterheads do not overtake a visible body
    # title.
    if len(blocks) < max_blocks:
        for section_index, section in enumerate(doc.sections):
            for area_name, area in (("header", section.header), ("footer", section.footer)):
                for paragraph_index, paragraph in enumerate(area.paragraphs):
                    if len(blocks) >= max_blocks:
                        break
                    add_text(f"section[{section_index}].{area_name}.paragraph[{paragraph_index}]", paragraph.text)
                for table_index, table in enumerate(area.tables):
                    for row_index, row in enumerate(table.rows):
                        if len(blocks) >= max_blocks:
                            break
                        cells = [normalize_button_label(cell.text) for cell in row.cells]
                        row_text = " | ".join(cell for cell in cells if cell)
                        add_text(f"section[{section_index}].{area_name}.table[{table_index}].row[{row_index}]", row_text)
    return tuple(blocks[:max_blocks])


def _best_title_from_top_blocks(blocks: Iterable[tuple[str, str]]) -> str:
    best: tuple[float, str] | None = None
    for block_index, (_hint, text) in enumerate(blocks):
        for part in _split_title_candidates(text):
            title = _normalize_detected_document_title(part)
            if not title:
                continue
            score = _document_title_score(title, block_index)
            if score <= 0:
                continue
            if best is None or score > best[0]:
                best = (score, title)
    return best[1] if best else ""


def _split_title_candidates(text: str) -> tuple[str, ...]:
    raw = re.sub(r"\{\{[^}]+\}\}", " ", str(text or ""))
    parts: list[str] = []
    for line in re.split(r"[\n\r]+", raw):
        line = line.strip()
        if not line:
            continue
        if "|" in line:
            parts.extend(piece.strip() for piece in line.split("|") if piece.strip())
        else:
            parts.append(line)
    return tuple(parts)


def _normalize_detected_document_title(value: str) -> str:
    text = normalize_button_label(value)
    text = re.sub(r"^(?:документ|форма|приложение)\s*(?:№|n|no)?\s*[0-9A-Za-zА-Яа-я._/-]*\s*[:.\-–—]+\s*", "", text, flags=re.IGNORECASE)
    text = normalize_button_label(text.strip("«»\"'()[]{}"))
    if len(text) < 3 or len(text) > 90:
        return ""
    lower = text.lower()
    if _looks_like_field_label_not_document_title(lower):
        return ""
    if _looks_like_vague_option_not_title(lower):
        return ""
    if _looks_like_timed_body_note_not_title(lower):
        return ""
    if _looks_like_service_header_line(lower):
        return ""
    if re.search(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", lower):
        return ""
    if sum(ch.isdigit() for ch in text) > max(2, len(text) // 5):
        return ""
    if not re.search(r"[A-Za-zА-Яа-яЁёІіЇїЄєӘәҒғҚқҢңӨөҰұҮүҺһЎўІі]", text):
        return ""
    # Specialty-neutral constructor rule: a doctor's own button name is taken
    # from the visible upper title of the DOCX.  Do not require a narrow specialty or
    # role keyword here; reject obvious field/header rows above, then let
    # the top-of-page score prefer real document titles for surgeons, therapists
    # and any other medical specialty.
    if text.isupper() and len(text.split()) <= 8:
        text = text.lower().capitalize()
        replacements = {"мсэ": "МСЭ", "рвк": "РВК", "вк": "ВК", "ф.и.о.": "Ф.И.О."}
        for src, dst in replacements.items():
            text = re.sub(rf"\b{re.escape(src)}\b", dst, text, flags=re.IGNORECASE)
    return normalize_button_label(text)



def _looks_like_timed_body_note_not_title(lower_text: str) -> bool:
    """Reject dated diary/body notes that look like titles only by containing 'осмотр'."""

    lower = str(lower_text or "").strip().lower().replace("ё", "е")
    if not lower:
        return True
    # Examples from real regression: "2023г. 10:00 Совместный осмотр ...".
    if re.match(r"^(?:19|20)\d{2}\s*г?\.?\s+(?:[01]?\d|2[0-3])[:.][0-5]\d\b", lower):
        return True
    if re.match(r"^\d{1,2}\s*[./-]\s*\d{1,2}\s*[./-]\s*\d{2,4}\s+(?:[01]?\d|2[0-3])[:.][0-5]\d\b", lower):
        return True
    if re.match(r"^(?:[01]?\d|2[0-3])[:.][0-5]\d\b", lower) and any(word in lower for word in ("осмотр", "консультац", "назнач", "обход")):
        return True
    body_markers = (
        "совместный осмотр", "осмотр зав", "зав. отдел", "заместител", "глав врача", "главного врача",
        "дежурный врач", "лечащий врач", "обход", "динамик", "состояние пациента", "жалобы на момент",
    )
    if any(marker in lower for marker in body_markers) and re.search(r"\b(?:19|20)\d{2}\s*г|\b(?:[01]?\d|2[0-3])[:.][0-5]\d\b", lower):
        return True
    return False


def _looks_like_vague_option_not_title(lower_text: str) -> bool:
    lower = str(lower_text or "").strip().lower().replace("ё", "е")
    lower = re.sub(r"\s+", " ", lower).strip(" .;:")
    if not lower:
        return True
    # Regression from template forms: a small selector row was used as a button:
    # "первичный, повторный".  This is not a document title; it is a visit type.
    if re.fullmatch(r"первичн\w*\s*[,/;]\s*повторн\w*", lower):
        return True
    if re.fullmatch(r"повторн\w*\s*[,/;]\s*первичн\w*", lower):
        return True
    if lower in {"первичный", "повторный", "плановый", "экстренный", "амбулаторный", "стационарный"}:
        return True
    return False

def _looks_like_field_label_not_document_title(lower_text: str) -> bool:
    """Reject patient-field rows that contain document-ish words by accident."""

    lower = str(lower_text or "").strip().lower().strip(" .;:")
    if not lower:
        return True
    if re.match(r"^история\s+болезни\s*(?:№|номер|n|no|:|[-–—]|\d)", lower):
        return True
    field_prefixes = (
        "дата поступления", "дата выписки", "дата перевода", "диагноз",
        "диагноз при поступлении", "диагноз при выписке", "лечение",
        "назначенное лечение", "проведенное лечение", "жалобы", "анамнез",
        "объективный статус", "соматический статус", "психический статус", "неврологический статус", "локальный статус",
        "состояние при выписке", "рекомендации", "пациент", "ф.и.о", "фио", "адрес", "паспорт", "снилс",
        "лечащий врач", "заведующий отделением", "глав врач", "главный врач", "заместитель главного врача", "температура", "пульс", "ад",
    )
    for prefix in field_prefixes:
        if lower == prefix or lower.startswith(prefix + ":") or lower.startswith(prefix + " —") or lower.startswith(prefix + " -"):
            return True
    return False


def _looks_like_service_header_line(lower_text: str) -> bool:
    """Reject hospital header rows without rejecting real document titles.

    Earlier title recognition dropped any line that merely contained words like
    ``отделение`` or ``пациент``.  Real titles can include them, for example
    ``Осмотр врача приёмного отделения``.  Here a stopword blocks only a line
    that has no title keyword, or a very bureaucratic header line.
    """

    lower = str(lower_text or "").strip().lower()
    if not lower:
        return True
    has_title_keyword = any(keyword in lower for keyword in _DOCUMENT_TITLE_KEYWORDS)
    has_stopword = any(stop in lower for stop in _DOCUMENT_TITLE_STOPWORDS)
    if not has_stopword:
        return False
    if not has_title_keyword:
        return True
    strong_header_markers = ("министерство", "департамент", "комитет", "больница", "клиника", "адрес", "телефон", "инн", "огрн", "окпо")
    return any(marker in lower for marker in strong_header_markers)


def _document_title_score(title: str, block_index: int) -> float:
    lower = title.lower()
    if _looks_like_vague_option_not_title(lower) or _looks_like_timed_body_note_not_title(lower):
        return 0.0
    score = 80.0 - (block_index * 4.0)
    for keyword in _DOCUMENT_TITLE_KEYWORDS:
        if keyword in lower:
            score += 18.0
    if 8 <= len(title) <= 45:
        score += 12.0
    if title[:1].isupper() or title.isupper():
        score += 4.0
    if lower.startswith(("первичный", "повторный", "выписной", "переводной", "протокол", "осмотр", "акт", "справка", "заключение", "направление", "лист")):
        score += 16.0
    if any(keyword in lower for keyword in _DOCUMENT_TITLE_KEYWORDS):
        score += 10.0
    if _looks_like_service_header_line(lower):
        return 0.0
    return score


def stable_document_id(role_id: str, label: str, template_path: str | Path) -> str:
    """Return a stable id seed that survives non-Latin filenames.

    Role semantics stay visible, but a short deterministic hash is always added
    from role/label/template material.  This prevents Georgian/Armenian or other
    non-Latin stems from collapsing to the same ASCII id after normalization.
    """

    role = str(role_id or "").strip().lower()
    stem = Path(template_path).stem if template_path else label
    material = f"{role}|{label}|{Path(template_path).as_posix() if template_path else ''}"
    suffix = hashlib.sha1(material.encode("utf-8", errors="replace")).hexdigest()[:8]
    if role and role != "unknown":
        base = _safe_id(f"{role}_{stem or label or 'document'}")
    else:
        base = _safe_id(stem or label or "document")
    return _safe_id(f"{base}_{suffix}")


def normalize_button_label(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" ._—-\t\n\r")
    return text[:80].strip() or "Документ"


def _safe_detect_language(path: Path) -> str:
    try:
        from medical_language_detector import detect_docx_language
        detected = detect_docx_language(path).language_id
        return normalize_language_id(detected, default="auto")
    except Exception as exc:
        record_soft_exception("personal_document_buttons.detect_language", exc, detail=str(path))
        return "auto"


def _template_title(path: Path) -> str:
    try:
        from universal_scanner import extract_docx_blocks
        for block in extract_docx_blocks(path):
            text = re.sub(r"\{\{[^}]+\}\}", " ", block.text or "")
            text = normalize_button_label(text)
            if len(text) >= 3 and not _looks_like_only_placeholders(text):
                return text
    except Exception as exc:
        record_soft_exception("personal_document_buttons:377", exc)
    return ""


def _looks_like_only_placeholders(text: str) -> bool:
    return not re.search(r"[A-Za-zА-Яа-яЁёІіЇїЄєӘәҒғҚқҢңӨөҰұҮүҺһЎўІіԱ-ֆა-ჰ]", text or "")


def _safe_id(value: str) -> str:
    text = str(value or "").strip().lower()
    translit = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e", "ж": "zh", "з": "z", "и": "i", "й": "y",
        "к": "k", "л": "l", "м": "m", "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u", "ф": "f",
        "х": "h", "ц": "c", "ч": "ch", "ш": "sh", "щ": "sch", "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    }
    text = "".join(translit.get(ch, ch) for ch in text)
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return text[:64] or "document"


def assert_personal_document_button_lock() -> None:
    if PERSONAL_DOCUMENT_BUTTON_LOCK_VERSION != "v1.4":
        raise AssertionError("Personal document button lock changed unexpectedly")
    if not BUTTON_LABEL_IS_PROFILE_DATA:
        raise AssertionError("Dynamic document button labels must remain profile-owned data")
    if not BUTTONS_DO_NOT_REQUIRE_HARDCODED_UI_TRANSLATION:
        raise AssertionError("Document buttons must not require hardcoded full-UI translation")
    if not PROFILE_BUTTON_LABELS_ARE_PERSISTED:
        raise AssertionError("Profile button labels must stay persisted in medpack data")
    if not TOP_OF_DOCX_TITLE_RECOGNITION_ENABLED:
        raise AssertionError("Doctor template titles must be recognized from the top of DOCX sheets")
    if not SPECIALTY_NEUTRAL_TITLE_RECOGNITION_ENABLED or not PROFILE_BUTTONS_ARE_NOT_PSYCHIATRY_BOUND:
        raise AssertionError("Doctor-owned buttons must stay specialty-neutral and must not be narrow-specialty-bound")
    if neutralize_weak_role_for_generic_profile("medical_commission", 0.32, top_title="Справка для бассейна", explicit_specialty="generic") != "unknown":
        raise AssertionError("Generic custom titles must not inherit weak legacy roles")
    if _normalize_detected_document_title("первичный, повторный"):
        raise AssertionError("Visit-type selectors must not become block-03 button titles")
    if _normalize_detected_document_title("2023г. 10:00 Совместный осмотр зам. глав врача Зубковой А.А. №"):
        raise AssertionError("Dated body notes must not become block-03 button titles")
    if stable_document_id("operation_protocol", "Протокол операции", "first.docx") == stable_document_id("operation_protocol", "Протокол операции", "second.docx"):
        raise AssertionError("Several documents of the same role must not overwrite each other")
    if stable_document_id("operation_protocol", "ოპერაციის პროტოკოლი", "ოპერაცია.docx") == stable_document_id("operation_protocol", "ოპერაციის პროტოკოლი", "მეორე.docx"):
        raise AssertionError("Non-Latin document names must not collapse to one id")
    for legacy_title, expected_role in (
        ("ВК больничный", "sick_leave_vk"),
        ("Акт для РВК", "military_commissariat_act"),
        ("Совместный осмотр", "joint_medical_exam"),
        ("ВК на МСЭ", "vk_mse"),
    ):
        if neutralize_weak_role_for_generic_profile("medical_commission", 0.32, top_title=legacy_title, explicit_specialty="generic") != expected_role:
            raise AssertionError(f"Early-production role title was not preserved: {legacy_title}")
    for role_id in ("operation_protocol", "discharge_epicrisis", "primary_exam", "informed_consent"):
        labels = LOCALIZED_ROLE_LABELS.get(role_id, {})
        for lang in ("ru", "uk", "az", "hy", "ka", "en"):
            if not labels.get(lang):
                raise AssertionError(f"Missing localized button label for {role_id}/{lang}")
    for role_id in ("joint_medical_exam", "vk_mse", "sick_leave_vk", "military_commissariat_act"):
        if not LOCALIZED_ROLE_LABELS.get(role_id, {}).get("ru"):
            raise AssertionError(f"Missing legacy-compatible role label for {role_id}/ru")

