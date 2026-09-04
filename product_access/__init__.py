from __future__ import annotations

"""Local product access, tariff, trial and license contract for Dokkomplekt.

Stores only product metadata. Never store/read/send patient documents, names,
diagnoses, template contents or patient file names here. Runtime UI imports are
lazy so headless CI can test the product contract without a Tk display stack.
"""

from dataclasses import dataclass
import base64
from datetime import datetime, timedelta, timezone
import hashlib
import hmac
import json
import os
import platform
import secrets
from pathlib import Path
import sys
from typing import Any, Iterable, Mapping
import uuid

from diagnostic_logging import record_soft_exception
from medical_paths import atomic_write_json as durable_atomic_write_json
from medical_paths import interprocess_file_lock

PRODUCT_ACCESS_CONTRACT_VERSION = "v1.0"
WATERMARK_CONTRACT_VERSION = "v1.0"
NO_PATIENT_DATA_IN_LICENSE_STATE = True
LOCAL_ONLY_PRODUCT_ACCESS = True
FOOTER_WATERMARK_ENABLED = True
NO_WATERMARK_FOR_PAID_LICENSES = True
TEST_PRODUCT_ACCESS_DISABLED_ENV = "DOKKOMPLEKT_TEST_DISABLE_PRODUCT_ACCESS"
PRODUCT_ACCESS_STATE_VERSION = 3
PRODUCTION_TRIAL_EPOCH = "v1.4.91-public"
PRODUCTION_TRIAL_CUTOFF = datetime(2026, 9, 3, 20, 2, 48, tzinfo=timezone.utc)
TRIAL_WATERMARK_TEXT = "ПРОБНАЯ ВЕРСИЯ. НЕ ИСПОЛЬЗОВАТЬ КАК МЕДИЦИНСКИЙ ДОКУМЕНТ."
EXPIRED_DEMO_WATERMARK_TEXT = "ДЕМО-ДОКУМЕНТ. ЛИЦЕНЗИЯ НЕ АКТИВНА."


def _env_flag(name: str) -> bool:
    return os.getenv(name, "").strip().lower() in {"1", "true", "yes", "on"}


def product_access_enforcement_enabled() -> bool:
    """Return whether runtime document creation should enforce product access."""
    if getattr(sys, "frozen", False):
        return True
    return not _env_flag(TEST_PRODUCT_ACCESS_DISABLED_ENV)


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def parse_dt(value: str | None) -> datetime | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        raw = raw[:-1] + "+00:00" if raw.endswith("Z") else raw
        dt = datetime.fromisoformat(raw)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).replace(microsecond=0)
    except ValueError:
        return None


def iso(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat()


def month_key(now: datetime | None = None) -> str:
    return (now or utc_now()).strftime("%Y-%m")


def stable_json(payload: Mapping[str, Any]) -> str:
    return json.dumps(dict(payload), ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def machine_fingerprint() -> str:
    raw = "|".join(str(value or "").lower() for value in (platform.system(), platform.machine(), platform.node(), os.getenv("COMPUTERNAME"), os.getenv("PROCESSOR_IDENTIFIER"), uuid.getnode()))
    return hashlib.sha256(raw.encode("utf-8", errors="replace")).hexdigest()[:32]


@dataclass(frozen=True)
class PlanLimits:
    plan_id: str
    title: str
    monthly_price_rub: int
    yearly_price_rub: int
    included_machines: int
    included_users: int
    profile_limit: int
    template_limit: int
    document_limit_month: int
    max_documents_per_run: int
    watermark_mode: str = "none"
    batch_generation: bool = False
    batch_print: bool = False
    shared_department_profile: bool = False
    role_management: bool = False
    offline_activation: bool = True
    local_license_server: bool = False
    overage_percent: int = 20
    grace_days: int = 7
    support_level: str = "base"


PLAN_LIMITS: dict[str, PlanLimits] = {
    "trial": PlanLimits("trial", "Trial", 0, 0, 1, 1, 1, 5, 30, 30, "trial", offline_activation=False, overage_percent=0, grace_days=0, support_level="knowledge_base"),
    "doctor_start": PlanLimits("doctor_start", "Doctor Start", 1490, 14900, 1, 1, 1, 30, 600, 10),
    "doctor_pro": PlanLimits("doctor_pro", "Doctor Pro", 3900, 29900, 2, 1, 3, 150, 3000, 50, batch_generation=True, batch_print=True, support_level="priority"),
    "department": PlanLimits("department", "Department", 14900, 149000, 5, 10, 10, 500, 20000, 100, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, grace_days=14, support_level="department"),
    "clinic": PlanLimits("clinic", "Clinic", 49000, 490000, 20, 50, 50, 2000, 100000, 250, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, local_license_server=True, grace_days=30, support_level="sla"),
    "enterprise": PlanLimits("enterprise", "Enterprise", 0, 900000, 9999, 9999, 9999, 999999, 9999999, 1000, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, local_license_server=True, grace_days=45, support_level="enterprise_sla"),
}


@dataclass(frozen=True)
class LicenseEntitlement:
    license_id: str
    plan: str
    owner_name: str = ""
    organization_name: str = ""
    seats: int = 1
    allowed_machines: tuple[str, ...] = ()
    valid_until: str = ""
    issued_at: str = ""
    generation_limit_month: int | None = None
    template_limit: int | None = None
    profile_limit: int | None = None
    watermark_mode: str | None = None
    offline_grace_days: int | None = None
    features: tuple[str, ...] = ()
    signature: str = ""

    @classmethod
    def from_mapping(cls, payload: Mapping[str, Any]) -> "LicenseEntitlement":
        return cls(
            str(payload.get("license_id") or "").strip(),
            str(payload.get("plan") or "").lower().strip(),
            str(payload.get("owner_name") or "").strip(),
            str(payload.get("organization_name") or "").strip(),
            max(1, int(payload.get("seats") or 1)),
            tuple(str(item).lower().strip() for item in payload.get("allowed_machines", ()) if str(item).strip()),
            str(payload.get("valid_until") or "").strip(),
            str(payload.get("issued_at") or "").strip(),
            int(payload["generation_limit_month"]) if payload.get("generation_limit_month") is not None else None,
            int(payload["template_limit"]) if payload.get("template_limit") is not None else None,
            int(payload["profile_limit"]) if payload.get("profile_limit") is not None else None,
            str(payload.get("watermark_mode")).lower().strip() if payload.get("watermark_mode") is not None else None,
            int(payload["offline_grace_days"]) if payload.get("offline_grace_days") is not None else None,
            tuple(str(item).strip() for item in payload.get("features", ()) if str(item).strip()),
            str(payload.get("signature") or "").strip(),
        )

    def unsigned_payload(self) -> dict[str, Any]:
        return {
            "license_id": self.license_id,
            "plan": self.plan,
            "owner_name": self.owner_name,
            "organization_name": self.organization_name,
            "seats": self.seats,
            "allowed_machines": list(self.allowed_machines),
            "valid_until": self.valid_until,
            "issued_at": self.issued_at,
            "generation_limit_month": self.generation_limit_month,
            "template_limit": self.template_limit,
            "profile_limit": self.profile_limit,
            "watermark_mode": self.watermark_mode,
            "offline_grace_days": self.offline_grace_days,
            "features": list(self.features),
        }

    def plan_limits(self) -> PlanLimits:
        if self.plan not in PLAN_LIMITS or self.plan == "trial":
            raise ValueError(f"Unknown paid license plan: {self.plan}")
        return PLAN_LIMITS[self.plan]

    def valid_until_dt(self) -> datetime | None:
        return parse_dt(self.valid_until)

    def is_expired(self, now: datetime | None = None) -> bool:
        valid_until = self.valid_until_dt()
        return valid_until is None or (now or utc_now()) > valid_until

    def signature_expected(self, secret: str) -> str:
        return hmac.new(secret.encode(), stable_json(self.unsigned_payload()).encode(), hashlib.sha256).hexdigest()

    def signature_valid(self, secret: str) -> bool:
        return bool(self.signature and secret and hmac.compare_digest(self.signature, self.signature_expected(secret)))


@dataclass(frozen=True)
class LicenseState:
    plan: str
    title: str
    active: bool
    reason: str
    license_id: str = ""
    owner_label: str = ""
    valid_until: str = ""
    trial_started_at: str = ""
    trial_ends_at: str = ""
    days_left: int = 0
    documents_used_month: int = 0
    documents_limit_month: int = 0
    documents_used_total_trial: int = 0
    remaining_documents_month: int = 0
    template_limit: int = 0
    profile_limit: int = 0
    included_machines: int = 1
    watermark_mode: str = "none"
    warning: str = ""

    @property
    def watermark_required(self) -> bool:
        return self.watermark_mode in {"trial", "expired_demo"}

    def watermark_text(self) -> str:
        if self.watermark_mode == "trial":
            return TRIAL_WATERMARK_TEXT
        if self.watermark_mode == "expired_demo":
            return EXPIRED_DEMO_WATERMARK_TEXT
        return ""


@dataclass(frozen=True)
class AccessDecision:
    allowed: bool
    code: str
    title: str
    message: str
    state: LicenseState
    warning: str = ""


@dataclass(frozen=True)
class WatermarkResult:
    path: str
    changed: bool
    error: str = ""


@dataclass(frozen=True)
class WatermarkBatchResult:
    results: tuple[WatermarkResult, ...]

    @property
    def errors(self) -> tuple[str, ...]:
        return tuple(f"{item.path}: {item.error}" for item in self.results if item.error)

    @property
    def changed_count(self) -> int:
        return sum(1 for item in self.results if item.changed)


@dataclass(frozen=True)
class UsageReservation:
    """One fail-closed usage reservation held while output is still staged."""

    token: str
    count: int
    month: str
    trial: bool


class ProductAccessManager:
    def __init__(self, storage_dir: str | Path | None = None, now: datetime | None = None):
        explicit_storage = storage_dir is not None or bool(os.getenv("DOKKOMPLEKT_LICENSE_DIR"))
        self.storage_dir = Path(storage_dir) if storage_dir else self.default_storage_dir()
        # The machine-wide HKCU guard is an anti-tamper owner for the normal
        # production location.  Explicit/portable storage is an intentionally
        # isolated namespace (tests, recovery, portable diagnostics) and must not
        # inherit or overwrite another storage root's machine-wide trial state.
        self._registry_guard_enabled = os.name == "nt" and not explicit_storage
        self.now = now
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.state_path = self.storage_dir / "product_access_state.json"
        self.state_guard_path = self.storage_dir / "product_access_guard.json"
        self.license_path = Path(os.getenv("DOKKOMPLEKT_LICENSE_FILE") or self.storage_dir / "license.json")
        self.state_lock_path = self.storage_dir / "product_access_state.lock"
        self.integrity_key_path = self.storage_dir / "product_access_integrity.key"
        self.integrity_key_guard_path = self.storage_dir / "product_access_integrity.guard"
        self._live_usage_reservations: set[str] = set()

    def _state_mutation_lock(self):
        return interprocess_file_lock(self.state_lock_path, timeout_seconds=10.0, stale_seconds=120.0)

    @staticmethod
    def default_storage_dir() -> Path:
        if os.getenv("DOKKOMPLEKT_LICENSE_DIR"):
            return Path(os.environ["DOKKOMPLEKT_LICENSE_DIR"]).expanduser()
        if os.getenv("LOCALAPPDATA"):
            return Path(os.environ["LOCALAPPDATA"]) / "Dokkomplekt"
        return Path.home() / ".dokkomplekt"

    def _now(self) -> datetime:
        return self.now or utc_now()

    @staticmethod
    def _legacy_state_integrity_key() -> bytes:
        seed = f"dokkomplekt-product-access-v2|{machine_fingerprint()}"
        return hashlib.sha256(seed.encode("utf-8", errors="replace")).digest()

    @staticmethod
    def _dpapi_protect(value: bytes) -> bytes:
        if os.name != "nt":
            return value
        import ctypes
        from ctypes import wintypes

        class DataBlob(ctypes.Structure):
            _fields_ = [("cbData", wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_ubyte))]

        source_buffer = (ctypes.c_ubyte * len(value)).from_buffer_copy(value)
        source = DataBlob(len(value), source_buffer)
        protected = DataBlob()
        if not ctypes.windll.crypt32.CryptProtectData(
            ctypes.byref(source), None, None, None, None, 0, ctypes.byref(protected)
        ):
            raise OSError("Windows DPAPI could not protect the product-access integrity key.")
        try:
            return ctypes.string_at(protected.pbData, protected.cbData)
        finally:
            ctypes.windll.kernel32.LocalFree(protected.pbData)

    @staticmethod
    def _dpapi_unprotect(value: bytes) -> bytes:
        if os.name != "nt":
            return value
        import ctypes
        from ctypes import wintypes

        class DataBlob(ctypes.Structure):
            _fields_ = [("cbData", wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_ubyte))]

        source_buffer = (ctypes.c_ubyte * len(value)).from_buffer_copy(value)
        source = DataBlob(len(value), source_buffer)
        clear = DataBlob()
        if not ctypes.windll.crypt32.CryptUnprotectData(
            ctypes.byref(source), None, None, None, None, 0, ctypes.byref(clear)
        ):
            raise OSError("Windows DPAPI could not unprotect the product-access integrity key.")
        try:
            return ctypes.string_at(clear.pbData, clear.cbData)
        finally:
            ctypes.windll.kernel32.LocalFree(clear.pbData)

    def _encode_integrity_key(self, key: bytes) -> str:
        protected = self._dpapi_protect(key)
        prefix = "dpapi:" if os.name == "nt" else "raw:"
        return prefix + base64.b64encode(protected).decode("ascii")

    def _decode_integrity_key(self, text: str) -> bytes:
        raw = str(text or "").strip()
        expected_prefix = "dpapi:" if os.name == "nt" else "raw:"
        if not raw.startswith(expected_prefix):
            raise ValueError("Unexpected product-access integrity-key format.")
        protected = base64.b64decode(raw[len(expected_prefix) :], validate=True)
        key = self._dpapi_unprotect(protected)
        if len(key) != 32:
            raise ValueError("Invalid product-access integrity-key length.")
        return key

    def _read_integrity_key_copy(self, path: Path) -> tuple[bytes | None, bool]:
        if not path.exists():
            return None, False
        try:
            return self._decode_integrity_key(path.read_text("ascii")), False
        except Exception as exc:
            record_soft_exception("product_access.read_integrity_key", exc, detail=str(path))
            return None, True

    def _write_integrity_key_copies(self, key: bytes) -> None:
        from medical_paths import atomic_write_text

        encoded = self._encode_integrity_key(key)
        for path in (self.integrity_key_path, self.integrity_key_guard_path):
            atomic_write_text(path, encoded)
            try:
                path.chmod(0o600)
            except OSError:
                pass

    def _state_integrity_key(self) -> bytes:
        primary, primary_bad = self._read_integrity_key_copy(self.integrity_key_path)
        guard, guard_bad = self._read_integrity_key_copy(self.integrity_key_guard_path)
        valid = [key for key in (primary, guard) if key is not None]
        if primary is not None and guard is not None and not hmac.compare_digest(primary, guard):
            raise ValueError("Product-access integrity key copies disagree.")
        if valid:
            key = valid[0]
            if primary_bad or guard_bad or primary is None or guard is None:
                self._write_integrity_key_copies(key)
            return key
        if primary_bad or guard_bad:
            raise ValueError("Product-access integrity key is damaged.")
        key = secrets.token_bytes(32)
        self._write_integrity_key_copies(key)
        return key

    def _with_state_integrity(self, payload: Mapping[str, Any]) -> dict[str, Any]:
        result = dict(payload)
        result["state_version"] = PRODUCT_ACCESS_STATE_VERSION
        result["usage_sequence"] = max(0, int(result.get("usage_sequence", 0) or 0)) + 1
        result.pop("_state_mac", None)
        message = stable_json(result).encode("utf-8")
        result["_state_mac"] = hmac.new(self._state_integrity_key(), message, hashlib.sha256).hexdigest()
        return result

    def _verify_state_integrity(self, payload: Mapping[str, Any]) -> bool:
        version = int(payload.get("state_version") or 1)
        actual = str(payload.get("_state_mac") or "").strip()
        if version < 2:
            return not actual
        if not actual:
            return False
        unsigned = dict(payload)
        unsigned.pop("_state_mac", None)
        key = self._legacy_state_integrity_key() if version == 2 else self._state_integrity_key()
        expected = hmac.new(key, stable_json(unsigned).encode("utf-8"), hashlib.sha256).hexdigest()
        return hmac.compare_digest(actual, expected)

    def _read_state_copy(self, path: Path) -> tuple[dict[str, Any] | None, bool]:
        if not path.exists():
            return None, False
        try:
            payload = json.loads(path.read_text("utf-8"))
            if not isinstance(payload, dict) or not self._verify_state_integrity(payload):
                return None, True
            return payload, False
        except Exception as exc:
            record_soft_exception("product_access.read_state_copy", exc, detail=str(path))
            return None, True

    def _read_registry_guard(self) -> tuple[dict[str, Any] | None, bool]:
        if os.name != "nt":
            return None, False
        try:
            import winreg  # type: ignore

            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Dokkomplekt\ProductAccess") as key:
                raw, _kind = winreg.QueryValueEx(key, "TrialGuard")
            payload = json.loads(str(raw or "{}"))
            if not isinstance(payload, dict) or not self._verify_state_integrity(payload):
                return None, True
            return payload, False
        except FileNotFoundError:
            return None, False
        except Exception as exc:
            record_soft_exception("product_access.read_registry_guard", exc)
            return None, True

    @staticmethod
    def _state_semantic_payload(payload: Mapping[str, Any]) -> dict[str, Any]:
        result = dict(payload)
        result.pop("_state_mac", None)
        return result

    @staticmethod
    def _merge_state_copies(copies: Iterable[Mapping[str, Any]]) -> dict[str, Any]:
        items = [dict(item) for item in copies if isinstance(item, Mapping)]
        if not items:
            return {}

        # A successful first-public-release reset creates an integrity-protected
        # epoch marker.  A stale legacy Registry/file copy from before that reset
        # must never drag the canonical trial start/usage backwards after one of
        # the redundant writes failed.  Once at least one valid current-epoch
        # copy exists, merge only within that epoch.  The MAC makes it impossible
        # to fabricate this marker by editing JSON/Registry manually.
        current_epoch_items = [
            item for item in items
            if str(item.get("trial_epoch") or "") == PRODUCTION_TRIAL_EPOCH
        ]
        merge_items = current_epoch_items or items
        result = dict(max(merge_items, key=lambda item: int(item.get("usage_sequence", 0) or 0)))

        starts = [parse_dt(str(item.get("trial_started_at") or "")) for item in merge_items]
        starts = [item for item in starts if item is not None]
        if starts:
            result["trial_started_at"] = iso(min(starts))
        result["trial_created_total"] = max(int(item.get("trial_created_total", 0) or 0) for item in merge_items)
        result["usage_sequence"] = max(int(item.get("usage_sequence", 0) or 0) for item in merge_items)
        merged_usage: dict[str, int] = {}
        for item in merge_items:
            usage = item.get("usage_by_month") if isinstance(item.get("usage_by_month"), dict) else {}
            for key, value in usage.items():
                merged_usage[str(key)] = max(merged_usage.get(str(key), 0), int(value or 0))
        result["usage_by_month"] = merged_usage
        result.pop("_state_mac", None)
        result["state_version"] = PRODUCT_ACCESS_STATE_VERSION
        return result

    def _load_state_payload(self) -> dict[str, Any]:
        primary, primary_corrupt = self._read_state_copy(self.state_path)
        guard, guard_corrupt = self._read_state_copy(self.state_guard_path)
        if self._registry_guard_enabled:
            registry, registry_corrupt = self._read_registry_guard()
        else:
            registry, registry_corrupt = None, False
        valid = [item for item in (primary, guard, registry) if item is not None]
        if valid:
            merged = self._merge_state_copies(valid)
            canonical = self._state_semantic_payload(merged)
            divergent = any(self._state_semantic_payload(item) != canonical for item in valid)
            if (
                primary_corrupt
                or guard_corrupt
                or registry_corrupt
                or primary is None
                or guard is None
                or divergent
            ):
                # Heal every redundant owner from the canonical merged state.  In
                # particular this repairs a stale Windows Registry guard left by
                # a partially-successful public-trial reset.
                self._save_state_payload(merged)
            return merged
        if primary_corrupt or guard_corrupt or registry_corrupt:
            return {"_state_corrupt": True}
        return {}

    @staticmethod
    def _atomic_write_json(path: Path, payload: Mapping[str, Any]) -> None:
        durable_atomic_write_json(path, payload, sort_keys=True)

    def _write_registry_guard(self, payload: Mapping[str, Any]) -> None:
        if os.name != "nt":
            return
        try:
            import winreg  # type: ignore

            with winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\Dokkomplekt\ProductAccess") as key:
                winreg.SetValueEx(key, "TrialGuard", 0, winreg.REG_SZ, json.dumps(dict(payload), ensure_ascii=False, sort_keys=True))
        except Exception as exc:
            record_soft_exception("product_access.write_registry_guard", exc)

    def _save_state_payload(self, payload: Mapping[str, Any]) -> None:
        protected = self._with_state_integrity(payload)
        self._atomic_write_json(self.state_path, protected)
        self._atomic_write_json(self.state_guard_path, protected)
        if self._registry_guard_enabled:
            self._write_registry_guard(protected)

    def _ensure_trial_started(self, payload: dict[str, Any]) -> dict[str, Any]:
        if payload.get("_state_corrupt"):
            return payload
        payload = dict(payload)
        changed = False
        started_at = parse_dt(str(payload.get("trial_started_at") or ""))
        # v1.4.91 is the first public production release.  Internal/pre-release
        # builds used the same machine-wide trial storage, so months of developer
        # testing could otherwise make a freshly downloaded production EXE say
        # "Пробный период завершён" on its very first real launch.  Grant the
        # public trial exactly once for state that demonstrably predates that
        # release.  The epoch marker is integrity-protected and persisted to all
        # redundant copies, so reinstalling/updating cannot reset the trial again.
        if (
            not self.license_path.exists()
            and payload.get("trial_epoch") != PRODUCTION_TRIAL_EPOCH
            and started_at is not None
            and started_at < PRODUCTION_TRIAL_CUTOFF
        ):
            payload["trial_started_at"] = iso(self._now())
            payload["trial_created_total"] = 0
            payload["usage_by_month"] = {}
            payload["trial_epoch"] = PRODUCTION_TRIAL_EPOCH
            payload["trial_public_reset_at"] = iso(self._now())
            started_at = self._now()
            changed = True
        if not payload.get("trial_started_at"):
            payload["trial_started_at"] = iso(self._now())
            payload.setdefault("usage_by_month", {})
            payload.setdefault("trial_created_total", 0)
            changed = True
        if payload.get("trial_epoch") != PRODUCTION_TRIAL_EPOCH:
            payload["trial_epoch"] = PRODUCTION_TRIAL_EPOCH
            changed = True
        if changed:
            self._save_state_payload(payload)
        return payload

    @staticmethod
    def _license_secret() -> str:
        if os.getenv("DOKKOMPLEKT_LICENSE_VERIFY_SECRET"):
            return os.environ["DOKKOMPLEKT_LICENSE_VERIFY_SECRET"].strip()
        try:
            from product_license_secret import LICENSE_VERIFY_SECRET  # type: ignore
            return str(LICENSE_VERIFY_SECRET or "").strip()
        except Exception:
            return ""

    def load_license(self) -> LicenseEntitlement | None:
        if not self.license_path.exists():
            return None
        try:
            payload = json.loads(self.license_path.read_text("utf-8"))
        except Exception as exc:
            raise ValueError(f"Файл лицензии повреждён: {exc}") from exc
        if not isinstance(payload, dict):
            raise ValueError("Файл лицензии должен содержать JSON-объект.")
        return LicenseEntitlement.from_mapping(payload)

    def install_license_text(self, text: str) -> LicenseState:
        payload = json.loads(text or "{}")
        if not isinstance(payload, dict):
            raise ValueError("Файл лицензии должен быть JSON-объектом.")
        entitlement = LicenseEntitlement.from_mapping(payload)
        self._validate_license(entitlement, require_not_expired=False)
        stored = entitlement.unsigned_payload()
        stored["signature"] = entitlement.signature
        durable_atomic_write_json(self.license_path, stored, sort_keys=True)
        return self.current_state()

    def _validate_license(self, entitlement: LicenseEntitlement, *, require_not_expired: bool = True) -> None:
        if not entitlement.license_id:
            raise ValueError("В лицензии нет license_id.")
        if entitlement.plan not in PLAN_LIMITS or entitlement.plan == "trial":
            raise ValueError(f"Неизвестный тариф лицензии: {entitlement.plan!r}.")
        secret = self._license_secret()
        unsigned_ok = _env_flag("DOKKOMPLEKT_ALLOW_UNSIGNED_LICENSES") and not getattr(sys, "frozen", False)
        if secret and not entitlement.signature_valid(secret):
            raise ValueError("Подпись лицензии не прошла проверку.")
        if not secret and not unsigned_ok:
            raise ValueError("Подпись лицензии не может быть проверена в этой сборке.")
        if require_not_expired and entitlement.is_expired(self._now()):
            raise ValueError("Срок действия лицензии истёк.")
        if entitlement.allowed_machines and machine_fingerprint() not in entitlement.allowed_machines:
            raise ValueError("Лицензия не привязана к этому компьютеру.")

    def _state_from_payload(self, payload: Mapping[str, Any]) -> LicenseState:
        usage = payload.get("usage_by_month") if isinstance(payload.get("usage_by_month"), dict) else {}
        used = int(usage.get(month_key(self._now()), 0) or 0)
        trial_total = int(payload.get("trial_created_total", 0) or 0)
        try:
            entitlement = self.load_license()
        except Exception as exc:
            return self._blocked_state(f"Файл лицензии повреждён или не может быть проверен: {exc}", used, trial_total)
        if entitlement:
            try:
                self._validate_license(entitlement, require_not_expired=False)
                return self._paid_state(entitlement, used)
            except ValueError as exc:
                return self._blocked_state(str(exc), used, trial_total)
        if payload.get("_state_corrupt"):
            return self._blocked_state(
                "Локальное состояние пробного периода повреждено. Восстановите файл состояния/резервную копию или установите лицензию.",
                used,
                trial_total,
            )
        return self._trial_state(payload, used, trial_total)

    def current_state(self) -> LicenseState:
        with self._state_mutation_lock():
            payload = self._ensure_trial_started(self._load_state_payload())
            return self._state_from_payload(payload)

    def _paid_state(self, entitlement: LicenseEntitlement, used: int) -> LicenseState:
        limits = entitlement.plan_limits()
        valid_until = entitlement.valid_until_dt()
        now = self._now()
        grace = int(entitlement.offline_grace_days if entitlement.offline_grace_days is not None else limits.grace_days)
        expired = valid_until is None or now > valid_until
        in_grace = bool(expired and valid_until and now <= valid_until + timedelta(days=grace))
        if expired and not in_grace:
            return self._blocked_state("Срок действия лицензии истёк.", used, 0)
        monthly_limit = int(entitlement.generation_limit_month or limits.document_limit_month)
        return LicenseState(
            plan=entitlement.plan,
            title=limits.title,
            active=True,
            reason="active_grace" if in_grace else "active",
            license_id=entitlement.license_id,
            owner_label=entitlement.organization_name or entitlement.owner_name,
            valid_until=entitlement.valid_until,
            documents_used_month=used,
            documents_limit_month=monthly_limit,
            remaining_documents_month=max(0, monthly_limit - used),
            template_limit=int(entitlement.template_limit or limits.template_limit),
            profile_limit=int(entitlement.profile_limit or limits.profile_limit),
            included_machines=int(entitlement.seats or limits.included_machines),
            watermark_mode=str(entitlement.watermark_mode or limits.watermark_mode),
            warning=f"Лицензия истекла, действует льготный период {grace} дн." if in_grace else "",
        )

    def _trial_state(self, payload: Mapping[str, Any], used: int, trial_total: int) -> LicenseState:
        limits = PLAN_LIMITS["trial"]
        started_at = parse_dt(str(payload.get("trial_started_at") or "")) or self._now()
        ends_at = started_at + timedelta(days=14)
        active = self._now() <= ends_at and trial_total < limits.document_limit_month
        reason = "trial_active" if active else "trial_document_limit" if trial_total >= limits.document_limit_month else "trial_expired"
        remaining_seconds = max(0.0, (ends_at - self._now()).total_seconds())
        days_left = int((remaining_seconds + 86399) // 86400) if active else 0
        return LicenseState(
            plan="trial",
            title=limits.title,
            active=active,
            reason=reason,
            trial_started_at=iso(started_at),
            trial_ends_at=iso(ends_at),
            days_left=days_left,
            documents_used_month=used,
            documents_limit_month=limits.document_limit_month,
            documents_used_total_trial=trial_total,
            remaining_documents_month=max(0, limits.document_limit_month - trial_total),
            template_limit=limits.template_limit,
            profile_limit=limits.profile_limit,
            included_machines=limits.included_machines,
            watermark_mode=limits.watermark_mode if active else "expired_demo",
            warning="Пробная версия создаёт документы только с водяным знаком." if active else "Пробный период завершён.",
        )

    def _blocked_state(self, reason: str, used: int, trial_total: int) -> LicenseState:
        return LicenseState("blocked", "Лицензия не активна", False, "blocked", documents_used_month=used, documents_used_total_trial=trial_total, watermark_mode="expired_demo", warning=reason)

    def _check_document_creation_for_state(
        self,
        state: LicenseState,
        requested_count: int,
        *,
        template_count: int | None = None,
        profile_count: int | None = None,
        enforce_per_run: bool = True,
    ) -> AccessDecision:
        count = max(1, int(requested_count or 1))
        if not state.active:
            return AccessDecision(False, "license_inactive", "Лицензия не активна", state.warning or "Создание рабочих документов заблокировано.", state)
        limits = PLAN_LIMITS.get(state.plan, PLAN_LIMITS["trial"])
        if enforce_per_run and count > limits.max_documents_per_run:
            return AccessDecision(False, "per_run_limit", "Слишком много документов за один запуск", f"Тариф разрешает до {limits.max_documents_per_run} документов за один запуск. Выбрано: {count}.", state)
        if template_count is not None and int(template_count) > state.template_limit:
            return AccessDecision(False, "template_limit", "Превышен лимит шаблонов", f"Лимит тарифа: {state.template_limit} шаблонов.", state)
        if profile_count is not None and int(profile_count) > state.profile_limit:
            return AccessDecision(False, "profile_limit", "Превышен лимит профилей", f"Лимит тарифа: {state.profile_limit} профилей.", state)
        if state.plan == "trial":
            if count > state.remaining_documents_month:
                return AccessDecision(False, "trial_limit", "Пробный лимит исчерпан", "Пробная версия разрешает 30 созданных документов всего.", state)
            return AccessDecision(True, "ok_trial", "Пробная версия", "Документы будут созданы с пробным водяным знаком.", state, state.warning)
        hard_limit = state.documents_limit_month + int(state.documents_limit_month * max(0, limits.overage_percent) / 100)
        projected = state.documents_used_month + count
        if state.documents_limit_month and projected > hard_limit:
            return AccessDecision(False, "monthly_limit", "Месячный лимит документов исчерпан", f"Использовано {state.documents_used_month}/{state.documents_limit_month}; льготный перерасход исчерпан.", state)
        warning = state.warning
        if not warning and state.documents_limit_month and projected > state.documents_limit_month:
            warning = f"Будет превышен месячный лимит {state.documents_limit_month}; действует льготный перерасход до {hard_limit}."
        if not warning and state.documents_limit_month and projected >= int(state.documents_limit_month * 0.8):
            warning = f"Использовано более 80% месячного лимита: после создания будет {projected}/{state.documents_limit_month}."
        return AccessDecision(True, "ok", "Доступ разрешён", "Создание документов разрешено.", state, warning)

    def check_document_creation(self, requested_count: int, *, template_count: int | None = None, profile_count: int | None = None) -> AccessDecision:
        state = self.current_state()
        return self._check_document_creation_for_state(
            state, requested_count, template_count=template_count, profile_count=profile_count
        )

    def check_configuration_limits(self, *, template_count: int | None = None, profile_count: int | None = None) -> AccessDecision:
        state = self.current_state()
        if not state.active:
            return AccessDecision(False, "license_inactive", "Лицензия не активна", state.warning or "Изменение рабочего профиля заблокировано.", state)
        if template_count is not None and int(template_count) > state.template_limit:
            return AccessDecision(False, "template_limit", "Превышен лимит шаблонов", f"Лимит тарифа: {state.template_limit} шаблонов.", state)
        if profile_count is not None and int(profile_count) > state.profile_limit:
            return AccessDecision(False, "profile_limit", "Превышен лимит профилей", f"Лимит тарифа: {state.profile_limit} профилей.", state)
        return AccessDecision(True, "ok", "Доступ разрешён", "Изменение профиля разрешено.", state, state.warning)

    @staticmethod
    def _increment_usage_payload(payload: dict[str, Any], *, month: str, count: int, trial: bool, now_text: str) -> None:
        usage = payload.get("usage_by_month") if isinstance(payload.get("usage_by_month"), dict) else {}
        usage[month] = int(usage.get(month, 0) or 0) + count
        payload["usage_by_month"] = usage
        if trial:
            payload["trial_created_total"] = int(payload.get("trial_created_total", 0) or 0) + count
        payload["updated_at"] = now_text

    def record_created_documents(self, count: int) -> None:
        delta = max(0, int(count or 0))
        if not delta:
            return
        with self._state_mutation_lock():
            payload = self._ensure_trial_started(self._load_state_payload())
            state = self._state_from_payload(payload)
            # record_created_documents is the accounting primitive for already-created
            # legacy outputs and migration/tests. Per-run size is a pre-creation
            # policy; at this point we must only prevent exceeding the global
            # trial/monthly allowance atomically.
            decision = self._check_document_creation_for_state(state, delta, enforce_per_run=False)
            if not decision.allowed:
                raise PermissionError(decision.message)
            self._increment_usage_payload(
                payload, month=month_key(self._now()), count=delta, trial=state.plan == "trial", now_text=iso(self._now())
            )
            self._save_state_payload(payload)

    def reserve_created_documents(self, count: int) -> UsageReservation | None:
        """Atomically validate and charge the actual staged output count."""
        delta = max(0, int(count or 0))
        if not delta:
            return None
        with self._state_mutation_lock():
            payload = self._ensure_trial_started(self._load_state_payload())
            state = self._state_from_payload(payload)
            decision = self._check_document_creation_for_state(state, delta)
            if not decision.allowed:
                raise PermissionError(decision.message)
            reservation = UsageReservation(uuid.uuid4().hex, delta, month_key(self._now()), state.plan == "trial")
            self._increment_usage_payload(
                payload, month=reservation.month, count=reservation.count, trial=reservation.trial, now_text=iso(self._now())
            )
            self._save_state_payload(payload)
        self._live_usage_reservations.add(reservation.token)
        return reservation

    def finalize_created_documents(self, reservation: UsageReservation | None) -> None:
        if reservation is not None:
            self._live_usage_reservations.discard(reservation.token)

    def release_created_documents(self, reservation: UsageReservation | None) -> None:
        if reservation is None or reservation.token not in self._live_usage_reservations:
            return
        with self._state_mutation_lock():
            payload = self._ensure_trial_started(self._load_state_payload())
            usage = payload.get("usage_by_month") if isinstance(payload.get("usage_by_month"), dict) else {}
            current = int(usage.get(reservation.month, 0) or 0)
            usage[reservation.month] = max(0, current - reservation.count)
            payload["usage_by_month"] = usage
            if reservation.trial:
                payload["trial_created_total"] = max(0, int(payload.get("trial_created_total", 0) or 0) - reservation.count)
            payload["updated_at"] = iso(self._now())
            self._save_state_payload(payload)
        self._live_usage_reservations.discard(reservation.token)

    def current_watermark_text(self) -> str:
        return self.current_state().watermark_text()

    def summary_text(self) -> str:
        state = self.current_state()
        used = state.documents_used_total_trial if state.plan == "trial" else state.documents_used_month
        lines = [f"Тариф: {state.title}", f"Статус: {'активен' if state.active else 'не активен'}"]
        if state.owner_label:
            lines.append(f"Владелец: {state.owner_label}")
        if state.license_id:
            lines.append(f"Лицензия: {state.license_id}")
        if state.valid_until:
            lines.append(f"Действует до: {state.valid_until}")
        if state.trial_ends_at:
            lines.append(f"Пробный период до: {state.trial_ends_at}")
        if state.documents_limit_month:
            lines.append(f"Документы: {used} / {state.documents_limit_month}")
        lines.extend([f"Шаблоны: до {state.template_limit}", f"Профили: до {state.profile_limit}", f"Компьютеры: до {state.included_machines}"])
        if state.watermark_required:
            lines.append("Водяной знак: включён")
        if state.warning:
            lines.append(f"Предупреждение: {state.warning}")
        return "\n".join(lines)


def sign_license_payload(payload: Mapping[str, Any], secret: str) -> dict[str, Any]:
    entitlement = LicenseEntitlement.from_mapping(payload)
    unsigned = entitlement.unsigned_payload()
    unsigned["signature"] = hmac.new(str(secret).encode(), stable_json(unsigned).encode(), hashlib.sha256).hexdigest()
    return unsigned


def apply_docx_footer_watermark(path: str | Path, text: str) -> WatermarkResult:
    target = Path(path)
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkResult(str(target), changed=False)
    if target.suffix.lower() != ".docx":
        return WatermarkResult(str(target), changed=False, error="watermark supports generated .docx files only")
    if not target.exists() or not target.is_file():
        return WatermarkResult(str(target), changed=False, error="file not found")
    try:
        from docx import Document
        document = Document(str(target))
        changed = False
        for section in document.sections:
            footer = section.footer
            existing = "\n".join(paragraph.text for paragraph in footer.paragraphs)
            if watermark in existing:
                continue
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if paragraph.text.strip():
                paragraph = footer.add_paragraph()
            paragraph.text = f" {watermark} "
            changed = True
        if changed:
            document.save(str(target))
        return WatermarkResult(str(target), changed=changed)
    except Exception as exc:
        return WatermarkResult(str(target), changed=False, error=str(exc))


def apply_watermark_to_files(paths: Iterable[str | Path], text: str) -> WatermarkBatchResult:
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkBatchResult(tuple())
    return WatermarkBatchResult(tuple(apply_docx_footer_watermark(path, watermark) for path in paths))


class ProductAccessMixin:
    """Wrap document creation with local licensing, limits and watermark policy."""

    def _estimate_selected_document_count(self, selected_medical: list[str], selected_diaries: bool, selected_custom: list[str]) -> int:
        return max(1, len(selected_medical or []) + len(selected_custom or []) + (1 if selected_diaries else 0))

    def _product_access_manager(self) -> ProductAccessManager:
        return ProductAccessManager()

    def _enforce_product_configuration_limits(self, *, template_count: int | None = None, profile_count: int | None = None) -> None:
        if not product_access_enforcement_enabled():
            return
        decision = self._product_access_manager().check_configuration_limits(
            template_count=template_count, profile_count=profile_count
        )
        if not decision.allowed:
            raise PermissionError(f"{decision.title}: {decision.message}")

    def _product_configuration_allowed(self, *, template_count: int | None = None, profile_count: int | None = None, show_warning: bool = True) -> bool:
        if not product_access_enforcement_enabled():
            return True
        decision = self._product_access_manager().check_configuration_limits(
            template_count=template_count, profile_count=profile_count
        )
        if decision.allowed:
            return True
        if show_warning:
            try:
                from tkinter import messagebox
                messagebox.showwarning(decision.title, decision.message, parent=getattr(self, "root", None))
            except Exception as exc:
                record_soft_exception("product_access.configuration_warning", exc)
        return False

    def _apply_product_watermark_to_docx(self, created_files: Iterable[str | Path]) -> list[Path]:
        paths = [Path(item) for item in created_files]
        if not paths or not product_access_enforcement_enabled():
            return paths
        manager = self._product_access_manager()
        watermark = manager.current_watermark_text()
        docx_paths = [path for path in paths if path.suffix.lower() == ".docx"]
        if watermark and docx_paths:
            result = apply_watermark_to_files(docx_paths, watermark)
            if result.errors:
                self._discard_unlicensed_outputs(paths)
                raise RuntimeError(
                    "Документы не выданы: не удалось гарантированно применить водяной знак trial/demo:\n"
                    + "\n".join(result.errors[:10])
                )
        return paths

    def _apply_product_watermark_before_pdf_export(self, docx_files: Iterable[str | Path]) -> list[Path]:
        """Watermark DOCX before PDF conversion so trial PDF is never unmarked."""

        return self._apply_product_watermark_to_docx(docx_files)

    def _reserve_product_access_for_staged_files(self, created_files: Iterable[str | Path]):
        paths = [Path(item) for item in created_files]
        if not paths or not product_access_enforcement_enabled():
            return None
        self._apply_product_watermark_to_docx(paths)
        manager = self._product_access_manager()
        try:
            reservation = manager.reserve_created_documents(len(paths))
        except Exception as exc:
            self._discard_unlicensed_outputs(paths)
            raise RuntimeError("Документы не выданы: не удалось надёжно зарезервировать счётчик лицензии.") from exc
        return manager, reservation

    @staticmethod
    def _release_product_access_reservation(handle) -> None:
        if not handle:
            return
        manager, reservation = handle
        manager.release_created_documents(reservation)

    @staticmethod
    def _finalize_product_access_reservation(handle) -> None:
        if not handle:
            return
        manager, reservation = handle
        manager.finalize_created_documents(reservation)

    def _enforce_product_access_on_created_files(self, created_files: Iterable[str | Path]) -> list[Path]:
        paths = [Path(item) for item in created_files]
        if not paths or not product_access_enforcement_enabled():
            return paths
        manager = self._product_access_manager()
        watermark = manager.current_watermark_text()
        if watermark:
            docx_paths = [path for path in paths if path.suffix.lower() == ".docx"]
            result = apply_watermark_to_files(docx_paths, watermark)
            if result.errors:
                self._discard_unlicensed_outputs(paths)
                raise RuntimeError(
                    "Документы не выданы: не удалось гарантированно применить водяной знак trial/demo:\n"
                    + "\n".join(result.errors[:10])
                )
        try:
            manager.record_created_documents(len(paths))
        except Exception as exc:
            self._discard_unlicensed_outputs(paths)
            raise RuntimeError("Документы не выданы: не удалось надёжно записать счётчик лицензии.") from exc
        return paths

    def create_selected_outputs(self, *, print_after: bool = False) -> None:
        if not product_access_enforcement_enabled():
            return super().create_selected_outputs(print_after=print_after)
        selected = self._selected_outputs_or_warn()
        if selected is None:
            return
        selected_medical, selected_diaries, selected_custom = selected
        manager = self._product_access_manager()
        decision = manager.check_document_creation(self._estimate_selected_document_count(selected_medical, selected_diaries, selected_custom))
        if not decision.allowed:
            from tkinter import messagebox
            messagebox.showwarning(decision.title, decision.message)
            try:
                self._log(f"\n⚠ {decision.title}: {decision.message}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_denied", exc)
            return
        if decision.warning:
            try:
                self._log(f"\n⚠ Лицензия: {decision.warning}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_warning", exc)
        return super().create_selected_outputs(print_after=print_after)

    def _created_files_from_results(self, created_medical: list[Path], created_custom: list[Path], diary_result):
        """Preserve fail-closed enforcement for legacy/non-transactional callers."""

        created_files = super()._created_files_from_results(created_medical, created_custom, diary_result)
        return self._enforce_product_access_on_created_files(created_files)

    @staticmethod
    def _discard_unlicensed_outputs(paths: Iterable[str | Path]) -> None:
        for raw_path in paths:
            path = Path(raw_path)
            try:
                if path.is_file():
                    path.unlink()
            except Exception as exc:
                record_soft_exception("product_access.discard_unlicensed_output", exc, detail=str(path))


class ProductLicenseMixin:
    def _initialize_app(self, root) -> None:
        super()._initialize_app(root)
        self._install_product_license_entrypoints()

    def _install_product_license_entrypoints(self) -> None:
        try:
            self.root.bind_all("<Control-l>", lambda _event: self.show_product_license_dialog())
            self.root.bind_all("<Control-L>", lambda _event: self.show_product_license_dialog())
        except Exception as exc:
            record_soft_exception("product_license.install_entrypoints", exc)

    def show_product_license_dialog(self) -> None:
        import tkinter as tk
        from tkinter import filedialog, messagebox
        manager = self._product_access_manager()
        window = tk.Toplevel(self.root)
        window.title("Лицензия Dokkomplekt")
        window.transient(self.root)
        window.grab_set()
        window.geometry("620x520")
        window.minsize(560, 460)
        outer = tk.Frame(window, padx=16, pady=14)
        outer.pack(fill="both", expand=True)
        outer.grid_columnconfigure(0, weight=1)
        outer.grid_rowconfigure(1, weight=1)
        tk.Label(outer, text="Лицензия и лимиты продукта", font=("Segoe UI", 13, "bold"), anchor="w").grid(row=0, column=0, sticky="ew")
        summary = tk.Text(outer, height=11, wrap="word")
        summary.grid(row=1, column=0, sticky="nsew", pady=(10, 10))
        summary.configure(state="normal")
        summary.insert("1.0", manager.summary_text())
        summary.configure(state="disabled")
        tk.Label(outer, text="Для offline-активации вставьте JSON лицензии или загрузите .json файл. Программа проверяет доступ локально и не отправляет документы пациента наружу.", justify="left", wraplength=560, anchor="w").grid(row=2, column=0, sticky="ew", pady=(0, 8))
        license_text = tk.Text(outer, height=7, wrap="word")
        license_text.grid(row=3, column=0, sticky="ew")
        buttons = tk.Frame(outer)
        buttons.grid(row=4, column=0, sticky="ew", pady=(12, 0))
        for column in range(4):
            buttons.grid_columnconfigure(column, weight=1)

        def refresh() -> None:
            fresh_manager = self._product_access_manager()
            summary.configure(state="normal")
            summary.delete("1.0", "end")
            summary.insert("1.0", fresh_manager.summary_text())
            summary.configure(state="disabled")

        def install_from_text() -> None:
            raw = license_text.get("1.0", "end").strip()
            if not raw:
                messagebox.showwarning("Лицензия", "Вставьте JSON лицензии или загрузите файл лицензии.")
                return
            try:
                self._product_access_manager().install_license_text(raw)
                refresh()
                messagebox.showinfo("Лицензия", "Лицензия установлена.")
            except Exception as exc:
                messagebox.showerror("Лицензия не установлена", str(exc))

        def load_file() -> None:
            path = filedialog.askopenfilename(title="Выберите файл лицензии", filetypes=(("Файл лицензии JSON", "*.json"), ("Все файлы", "*.*")))
            if not path:
                return
            try:
                license_text.delete("1.0", "end")
                license_text.insert("1.0", Path(path).read_text(encoding="utf-8"))
            except OSError as exc:
                messagebox.showerror("Лицензия", f"Не удалось прочитать файл лицензии:\n{exc}")
        tk.Button(buttons, text="Загрузить файл", command=load_file).grid(row=0, column=0, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Установить", command=install_from_text).grid(row=0, column=1, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Обновить", command=refresh).grid(row=0, column=2, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Закрыть", command=window.destroy).grid(row=0, column=3, sticky="ew")
