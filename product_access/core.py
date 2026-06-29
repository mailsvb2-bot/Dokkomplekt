from __future__ import annotations

"""Hardened local product-access core for Dokkomplekt.

Only product metadata is stored here. Patient names, diagnoses, source file names,
DOCX contents and template contents are intentionally outside this boundary.
"""

from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
import hashlib
import hmac
import json
import os
import platform
from pathlib import Path
from typing import Any, Iterable, Mapping
import uuid

from diagnostic_logging import record_soft_exception

PRODUCT_ACCESS_CONTRACT_VERSION = "v1.1-postmerge-hardening"
WATERMARK_CONTRACT_VERSION = "v1.1-postmerge-hardening"
NO_PATIENT_DATA_IN_LICENSE_STATE = True
LOCAL_ONLY_PRODUCT_ACCESS = True
FOOTER_WATERMARK_ENABLED = True
NO_WATERMARK_FOR_PAID_LICENSES = True
PRODUCT_ACCESS_DISABLED_ENV = "MEDICAL_AUTOFILL_DISABLE_PRODUCT_ACCESS"
TRIAL_WATERMARK_TEXT = "ПРОБНАЯ ВЕРСИЯ. НЕ ИСПОЛЬЗОВАТЬ КАК МЕДИЦИНСКИЙ ДОКУМЕНТ."
EXPIRED_DEMO_WATERMARK_TEXT = "ДЕМО-ДОКУМЕНТ. ЛИЦЕНЗИЯ НЕ АКТИВНА."
STATE_SCHEMA_VERSION = 2


class LicenseContractError(ValueError):
    pass


def _env_flag(name: str) -> bool:
    return os.getenv(name, "").strip().lower() in {"1", "true", "yes", "on"}


def product_access_enforcement_enabled() -> bool:
    return not (_env_flag(PRODUCT_ACCESS_DISABLED_ENV) or _env_flag("CI"))


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def parse_dt(value: str | None) -> datetime | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        raw = raw[:-1] + "+00:00" if raw.endswith("Z") else raw
        dt = datetime.fromisoformat(raw)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).replace(microsecond=0)
    except (TypeError, ValueError):
        return None


def iso(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat()


def month_key(now: datetime | None = None) -> str:
    return (now or utc_now()).astimezone(timezone.utc).strftime("%Y-%m")


def _stable(value: Any) -> Any:
    if isinstance(value, Mapping):
        return {str(k): _stable(value[k]) for k in sorted(value, key=str)}
    if isinstance(value, (list, tuple)):
        return [_stable(v) for v in value]
    return value


def stable_json(payload: Mapping[str, Any]) -> str:
    return json.dumps(_stable(dict(payload)), ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def _as_int(value: Any, default: int = 0, *, minimum: int | None = None, maximum: int | None = None) -> int:
    try:
        out = int(value)
    except (TypeError, ValueError, OverflowError):
        out = default
    if minimum is not None:
        out = max(minimum, out)
    if maximum is not None:
        out = min(maximum, out)
    return out


def _tuple(value: Any, *, lower: bool = False) -> tuple[str, ...]:
    seq: Iterable[Any] = (value,) if isinstance(value, str) else value if isinstance(value, Iterable) else ()
    items = (str(x).strip() for x in seq)
    return tuple((x.lower() if lower else x) for x in items if x)


def machine_fingerprint() -> str:
    raw = "|".join(str(v or "").lower() for v in (platform.system(), platform.machine(), platform.node(), os.getenv("COMPUTERNAME"), os.getenv("PROCESSOR_IDENTIFIER"), uuid.getnode()))
    return hashlib.sha256(raw.encode("utf-8", errors="replace")).hexdigest()[:32]


@dataclass(frozen=True)
class PlanLimits:
    plan_id: str
    title: str
    monthly_price_rub: int
    yearly_price_rub: int
    included_machines: int
    included_users: int
    profile_limit: int
    template_limit: int
    document_limit_month: int
    max_documents_per_run: int
    watermark_mode: str = "none"
    batch_generation: bool = False
    batch_print: bool = False
    shared_department_profile: bool = False
    role_management: bool = False
    offline_activation: bool = True
    local_license_server: bool = False
    overage_percent: int = 20
    grace_days: int = 7
    support_level: str = "base"


PLAN_LIMITS: dict[str, PlanLimits] = {
    "trial": PlanLimits("trial", "Trial", 0, 0, 1, 1, 1, 5, 30, 3, "trial", offline_activation=False, overage_percent=0, grace_days=0, support_level="knowledge_base"),
    "doctor_start": PlanLimits("doctor_start", "Doctor Start", 1490, 14900, 1, 1, 1, 30, 600, 10),
    "doctor_pro": PlanLimits("doctor_pro", "Doctor Pro", 3900, 29900, 2, 1, 3, 150, 3000, 50, batch_generation=True, batch_print=True, support_level="priority"),
    "department": PlanLimits("department", "Department", 14900, 149000, 5, 10, 10, 500, 20000, 100, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, grace_days=14, support_level="department"),
    "clinic": PlanLimits("clinic", "Clinic", 49000, 490000, 20, 50, 50, 2000, 100000, 250, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, local_license_server=True, grace_days=30, support_level="sla"),
    "enterprise": PlanLimits("enterprise", "Enterprise", 0, 900000, 9999, 9999, 9999, 999999, 9999999, 1000, batch_generation=True, batch_print=True, shared_department_profile=True, role_management=True, local_license_server=True, grace_days=45, support_level="enterprise_sla"),
}


@dataclass(frozen=True)
class LicenseEntitlement:
    license_id: str
    plan: str
    owner_name: str = ""
    organization_name: str = ""
    seats: int = 1
    allowed_machines: tuple[str, ...] = ()
    valid_until: str = ""
    issued_at: str = ""
    generation_limit_month: int | None = None
    template_limit: int | None = None
    profile_limit: int | None = None
    watermark_mode: str | None = None
    offline_grace_days: int | None = None
    features: tuple[str, ...] = ()
    signature: str = ""

    @classmethod
    def from_mapping(cls, payload: Mapping[str, Any]) -> "LicenseEntitlement":
        if not isinstance(payload, Mapping):
            raise LicenseContractError("Файл лицензии должен быть JSON-объектом.")
        return cls(
            str(payload.get("license_id") or "").strip(), str(payload.get("plan") or "").lower().strip(),
            str(payload.get("owner_name") or "").strip(), str(payload.get("organization_name") or "").strip(),
            _as_int(payload.get("seats"), 1, minimum=1, maximum=9999), _tuple(payload.get("allowed_machines"), lower=True),
            str(payload.get("valid_until") or "").strip(), str(payload.get("issued_at") or "").strip(),
            _as_int(payload.get("generation_limit_month"), 0, minimum=0) if payload.get("generation_limit_month") is not None else None,
            _as_int(payload.get("template_limit"), 0, minimum=0) if payload.get("template_limit") is not None else None,
            _as_int(payload.get("profile_limit"), 0, minimum=0) if payload.get("profile_limit") is not None else None,
            str(payload.get("watermark_mode")).lower().strip() if payload.get("watermark_mode") is not None else None,
            _as_int(payload.get("offline_grace_days"), 0, minimum=0, maximum=90) if payload.get("offline_grace_days") is not None else None,
            _tuple(payload.get("features")), str(payload.get("signature") or "").strip(),
        )

    def unsigned_payload(self) -> dict[str, Any]:
        return {"license_id": self.license_id, "plan": self.plan, "owner_name": self.owner_name, "organization_name": self.organization_name, "seats": self.seats, "allowed_machines": list(self.allowed_machines), "valid_until": self.valid_until, "issued_at": self.issued_at, "generation_limit_month": self.generation_limit_month, "template_limit": self.template_limit, "profile_limit": self.profile_limit, "watermark_mode": self.watermark_mode, "offline_grace_days": self.offline_grace_days, "features": list(self.features)}

    def plan_limits(self) -> PlanLimits:
        if self.plan not in PLAN_LIMITS or self.plan == "trial":
            raise LicenseContractError(f"Неизвестный тариф лицензии: {self.plan!r}.")
        return PLAN_LIMITS[self.plan]

    def valid_until_dt(self) -> datetime | None:
        return parse_dt(self.valid_until)

    def issued_at_dt(self) -> datetime | None:
        return parse_dt(self.issued_at)

    def is_expired(self, now: datetime | None = None) -> bool:
        valid_until = self.valid_until_dt()
        return valid_until is None or (now or utc_now()) > valid_until

    def signature_expected(self, secret: str) -> str:
        return hmac.new(secret.encode(), stable_json(self.unsigned_payload()).encode(), hashlib.sha256).hexdigest()

    def signature_valid(self, secret: str) -> bool:
        return bool(self.signature and secret and hmac.compare_digest(self.signature, self.signature_expected(secret)))


@dataclass(frozen=True)
class LicenseState:
    plan: str
    title: str
    active: bool
    reason: str
    license_id: str = ""
    owner_label: str = ""
    valid_until: str = ""
    trial_started_at: str = ""
    trial_ends_at: str = ""
    days_left: int = 0
    documents_used_month: int = 0
    documents_limit_month: int = 0
    documents_used_total_trial: int = 0
    remaining_documents_month: int = 0
    template_limit: int = 0
    profile_limit: int = 0
    included_machines: int = 1
    watermark_mode: str = "none"
    warning: str = ""

    @property
    def watermark_required(self) -> bool:
        return self.watermark_mode in {"trial", "expired_demo"}

    def watermark_text(self) -> str:
        return TRIAL_WATERMARK_TEXT if self.watermark_mode == "trial" else EXPIRED_DEMO_WATERMARK_TEXT if self.watermark_mode == "expired_demo" else ""


@dataclass(frozen=True)
class AccessDecision:
    allowed: bool
    code: str
    title: str
    message: str
    state: LicenseState
    warning: str = ""


@dataclass(frozen=True)
class WatermarkResult:
    path: str
    changed: bool
    error: str = ""


@dataclass(frozen=True)
class WatermarkBatchResult:
    results: tuple[WatermarkResult, ...]

    @property
    def errors(self) -> tuple[str, ...]:
        return tuple(f"{Path(r.path).name}: {r.error}" for r in self.results if r.error)

    @property
    def changed_count(self) -> int:
        return sum(1 for r in self.results if r.changed)


class ProductAccessManager:
    def __init__(self, storage_dir: str | Path | None = None, now: datetime | None = None):
        self.storage_dir = Path(storage_dir) if storage_dir else self.default_storage_dir()
        self.now = now
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.state_path = self.storage_dir / "product_access_state.json"
        self.license_path = Path(os.getenv("DOKKOMPLEKT_LICENSE_FILE") or self.storage_dir / "license.json")

    @staticmethod
    def default_storage_dir() -> Path:
        if os.getenv("DOKKOMPLEKT_LICENSE_DIR"):
            return Path(os.environ["DOKKOMPLEKT_LICENSE_DIR"]).expanduser()
        return Path(os.environ["LOCALAPPDATA"]) / "Dokkomplekt" if os.getenv("LOCALAPPDATA") else Path.home() / ".dokkomplekt"

    def _now(self) -> datetime:
        return self.now or utc_now()

    def _load_json_file(self, path: Path) -> dict[str, Any]:
        try:
            if not path.exists():
                return {}
            data = json.loads(path.read_text("utf-8"))
            return data if isinstance(data, dict) else {}
        except (OSError, json.JSONDecodeError):
            return {}

    def _migrated_state_payload(self, payload: Mapping[str, Any]) -> dict[str, Any]:
        out = dict(payload)
        out["schema_version"] = STATE_SCHEMA_VERSION
        out["contract_version"] = PRODUCT_ACCESS_CONTRACT_VERSION
        if not isinstance(out.get("usage_by_month"), dict):
            out["usage_by_month"] = {}
        out["trial_created_total"] = _as_int(out.get("trial_created_total"), 0, minimum=0)
        return out

    def _load_state_payload(self) -> dict[str, Any]:
        return self._migrated_state_payload(self._load_json_file(self.state_path))

    def _save_json_file(self, path: Path, payload: Mapping[str, Any]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_name(f"{path.name}.tmp")
        tmp.write_text(json.dumps(dict(payload), ensure_ascii=False, indent=2, sort_keys=True), "utf-8")
        os.replace(tmp, path)

    def _save_state_payload(self, payload: Mapping[str, Any]) -> None:
        data = self._migrated_state_payload(payload)
        data["updated_at"] = iso(self._now())
        self._save_json_file(self.state_path, data)

    def _ensure_trial_started(self, payload: dict[str, Any]) -> dict[str, Any]:
        data = self._migrated_state_payload(payload)
        if not data.get("trial_started_at"):
            data["trial_started_at"] = iso(self._now())
            self._save_state_payload(data)
        return data

    @staticmethod
    def _license_secret() -> str:
        if os.getenv("DOKKOMPLEKT_LICENSE_VERIFY_SECRET"):
            return os.environ["DOKKOMPLEKT_LICENSE_VERIFY_SECRET"].strip()
        try:
            from product_license_secret import LICENSE_VERIFY_SECRET  # type: ignore
            return str(LICENSE_VERIFY_SECRET or "").strip()
        except Exception:
            return ""

    def load_license(self) -> LicenseEntitlement | None:
        data = self._load_json_file(self.license_path)
        return LicenseEntitlement.from_mapping(data) if data else None

    def install_license_text(self, text: str) -> LicenseState:
        try:
            data = json.loads(text or "{}")
        except json.JSONDecodeError as exc:
            raise LicenseContractError("Файл лицензии должен быть корректным JSON.") from exc
        ent = LicenseEntitlement.from_mapping(data)
        self._validate_license(ent, require_not_expired=False)
        stored = ent.unsigned_payload()
        stored["signature"] = ent.signature
        self._save_json_file(self.license_path, stored)
        return self.current_state()

    def _validate_license(self, ent: LicenseEntitlement, *, require_not_expired: bool = True) -> None:
        if not ent.license_id:
            raise LicenseContractError("В лицензии нет license_id.")
        limits = ent.plan_limits()
        if ent.allowed_machines and len(set(ent.allowed_machines)) > ent.seats:
            raise LicenseContractError("Количество привязанных компьютеров превышает seats лицензии.")
        secret = self._license_secret()
        if secret and not ent.signature_valid(secret):
            raise LicenseContractError("Подпись лицензии не прошла проверку.")
        if not secret and not _env_flag("DOKKOMPLEKT_ALLOW_UNSIGNED_LICENSES"):
            raise LicenseContractError("Подпись лицензии не может быть проверена в этой сборке.")
        if require_not_expired and ent.is_expired(self._now()):
            raise LicenseContractError("Срок действия лицензии истёк.")
        if ent.issued_at_dt() and ent.issued_at_dt() > self._now() + timedelta(minutes=10):
            raise LicenseContractError("Дата выдачи лицензии находится в будущем.")
        if ent.allowed_machines and machine_fingerprint() not in ent.allowed_machines:
            raise LicenseContractError("Лицензия не привязана к этому компьютеру.")
        if ent.watermark_mode and ent.watermark_mode not in {"none", "trial", "expired_demo"}:
            raise LicenseContractError("В лицензии указан неизвестный режим водяного знака.")
        for value, msg in ((ent.generation_limit_month, "Месячный лимит документов"), (ent.template_limit, "Лимит шаблонов"), (ent.profile_limit, "Лимит профилей")):
            if value is not None and value < 1:
                raise LicenseContractError(f"{msg} должен быть положительным.")
        if ent.seats > max(limits.included_machines, 1) and ent.plan != "enterprise":
            raise LicenseContractError("Количество seats превышает тарифный предел.")

    def current_state(self) -> LicenseState:
        data = self._ensure_trial_started(self._load_state_payload())
        usage = data.get("usage_by_month") if isinstance(data.get("usage_by_month"), dict) else {}
        used = _as_int(usage.get(month_key(self._now())), 0, minimum=0)
        trial_total = _as_int(data.get("trial_created_total"), 0, minimum=0)
        ent = self.load_license()
        if ent:
            try:
                self._validate_license(ent, require_not_expired=False)
                return self._paid_state(ent, used)
            except ValueError as exc:
                return self._blocked_state(str(exc), used, trial_total)
        return self._trial_state(data, used, trial_total)

    def _paid_state(self, ent: LicenseEntitlement, used: int) -> LicenseState:
        limits = ent.plan_limits(); valid_until = ent.valid_until_dt(); now = self._now()
        grace = int(ent.offline_grace_days if ent.offline_grace_days is not None else limits.grace_days)
        expired = valid_until is None or now > valid_until
        in_grace = bool(expired and valid_until and now <= valid_until + timedelta(days=grace))
        if expired and not in_grace:
            return self._blocked_state("Срок действия лицензии истёк.", used, 0)
        monthly = int(ent.generation_limit_month or limits.document_limit_month)
        return LicenseState(ent.plan, limits.title, True, "active_grace" if in_grace else "active", ent.license_id, ent.organization_name or ent.owner_name, ent.valid_until, days_left=max(0, (valid_until.date() - now.date()).days) if valid_until else 0, documents_used_month=used, documents_limit_month=monthly, remaining_documents_month=max(0, monthly - used), template_limit=int(ent.template_limit or limits.template_limit), profile_limit=int(ent.profile_limit or limits.profile_limit), included_machines=int(ent.seats or limits.included_machines), watermark_mode=str(ent.watermark_mode or limits.watermark_mode), warning=f"Лицензия истекла, действует льготный период {grace} дн." if in_grace else "")

    def _trial_state(self, data: Mapping[str, Any], used: int, trial_total: int) -> LicenseState:
        limits = PLAN_LIMITS["trial"]; started = parse_dt(str(data.get("trial_started_at") or "")) or self._now(); ends = started + timedelta(days=14)
        active = self._now() <= ends and trial_total < limits.document_limit_month
        reason = "trial_active" if active else "trial_document_limit" if trial_total >= limits.document_limit_month else "trial_expired"
        return LicenseState("trial", limits.title, active, reason, trial_started_at=iso(started), trial_ends_at=iso(ends), days_left=max(0, (ends.date() - self._now().date()).days), documents_used_month=used, documents_limit_month=limits.document_limit_month, documents_used_total_trial=trial_total, remaining_documents_month=max(0, limits.document_limit_month - trial_total), template_limit=limits.template_limit, profile_limit=limits.profile_limit, included_machines=limits.included_machines, watermark_mode=limits.watermark_mode if active else "expired_demo", warning="Пробная версия создаёт документы только с водяным знаком." if active else "Пробный период завершён.")

    def _blocked_state(self, reason: str, used: int, trial_total: int) -> LicenseState:
        return LicenseState("blocked", "Лицензия не активна", False, "blocked", documents_used_month=used, documents_used_total_trial=trial_total, watermark_mode="expired_demo", warning=reason)

    def check_document_creation(self, requested_count: int, *, template_count: int | None = None, profile_count: int | None = None) -> AccessDecision:
        count = max(1, _as_int(requested_count, 1, minimum=1)); state = self.current_state(); limits = PLAN_LIMITS.get(state.plan, PLAN_LIMITS["trial"])
        if not state.active:
            return AccessDecision(False, "license_inactive", "Лицензия не активна", state.warning or "Создание рабочих документов заблокировано.", state)
        if count > limits.max_documents_per_run:
            return AccessDecision(False, "per_run_limit", "Слишком много документов за один запуск", f"Тариф разрешает до {limits.max_documents_per_run} документов за один запуск. Выбрано: {count}.", state)
        if template_count is not None and _as_int(template_count, 0, minimum=0) > state.template_limit:
            return AccessDecision(False, "template_limit", "Превышен лимит шаблонов", f"Лимит тарифа: {state.template_limit} шаблонов.", state)
        if profile_count is not None and _as_int(profile_count, 0, minimum=0) > state.profile_limit:
            return AccessDecision(False, "profile_limit", "Превышен лимит профилей", f"Лимит тарифа: {state.profile_limit} профилей.", state)
        if state.plan == "trial":
            if count > state.remaining_documents_month:
                return AccessDecision(False, "trial_limit", "Пробный лимит исчерпан", "Пробная версия разрешает 30 созданных документов всего.", state)
            return AccessDecision(True, "ok_trial", "Пробная версия", "Документы будут созданы с пробным водяным знаком.", state, state.warning)
        hard = state.documents_limit_month + int(state.documents_limit_month * max(0, limits.overage_percent) / 100); projected = state.documents_used_month + count
        if state.documents_limit_month and projected > hard:
            return AccessDecision(False, "monthly_limit", "Месячный лимит документов исчерпан", f"Использовано {state.documents_used_month}/{state.documents_limit_month}; льготный перерасход исчерпан.", state)
        warning = state.warning or (f"Будет превышен месячный лимит {state.documents_limit_month}; действует льготный перерасход до {hard}." if state.documents_limit_month and projected > state.documents_limit_month else "") or (f"Использовано более 80% месячного лимита: после создания будет {projected}/{state.documents_limit_month}." if state.documents_limit_month and projected >= int(state.documents_limit_month * 0.8) else "")
        return AccessDecision(True, "ok", "Доступ разрешён", "Создание документов разрешено.", state, warning)

    def record_created_documents(self, count: int) -> None:
        delta = _as_int(count, 0, minimum=0)
        if not delta:
            return
        data = self._ensure_trial_started(self._load_state_payload()); usage = data.get("usage_by_month") if isinstance(data.get("usage_by_month"), dict) else {}; key = month_key(self._now())
        usage[key] = _as_int(usage.get(key), 0, minimum=0) + delta; data["usage_by_month"] = usage
        if self.current_state().plan == "trial":
            data["trial_created_total"] = _as_int(data.get("trial_created_total"), 0, minimum=0) + delta
        self._save_state_payload(data)

    def current_watermark_text(self) -> str:
        return self.current_state().watermark_text()

    def summary_text(self) -> str:
        state = self.current_state(); used = state.documents_used_total_trial if state.plan == "trial" else state.documents_used_month
        lines = [f"Тариф: {state.title}", f"Статус: {'активен' if state.active else 'не активен'}"]
        for line in (f"Владелец: {state.owner_label}" if state.owner_label else "", f"Лицензия: {state.license_id}" if state.license_id else "", f"Действует до: {state.valid_until}" if state.valid_until else "", f"Пробный период до: {state.trial_ends_at}" if state.trial_ends_at else "", f"Документы: {used} / {state.documents_limit_month}" if state.documents_limit_month else ""):
            if line:
                lines.append(line)
        lines += [f"Шаблоны: до {state.template_limit}", f"Профили: до {state.profile_limit}", f"Компьютеры: до {state.included_machines}"]
        if state.watermark_required:
            lines.append("Водяной знак: включён")
        if state.warning:
            lines.append(f"Предупреждение: {state.warning}")
        return "\n".join(lines)


def sign_license_payload(payload: Mapping[str, Any], secret: str) -> dict[str, Any]:
    if not str(secret or "").strip():
        raise LicenseContractError("License signing secret is empty.")
    ent = LicenseEntitlement.from_mapping(payload); unsigned = ent.unsigned_payload()
    unsigned["signature"] = hmac.new(str(secret).encode(), stable_json(unsigned).encode(), hashlib.sha256).hexdigest()
    return unsigned


def apply_docx_footer_watermark(path: str | Path, text: str) -> WatermarkResult:
    target = Path(path); watermark = str(text or "").strip()
    if not watermark:
        return WatermarkResult(str(target), False)
    if target.suffix.lower() != ".docx":
        return WatermarkResult(str(target), False, "watermark supports generated .docx files only")
    if not target.exists() or not target.is_file():
        return WatermarkResult(str(target), False, "file not found")
    try:
        from docx import Document
        document = Document(str(target)); changed = False
        for section in document.sections:
            footer = section.footer; existing = "\n".join(p.text for p in footer.paragraphs)
            if watermark in existing:
                continue
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if paragraph.text.strip():
                paragraph = footer.add_paragraph()
            paragraph.text = f" {watermark} "; changed = True
        if changed:
            document.save(str(target))
        return WatermarkResult(str(target), changed)
    except Exception as exc:
        return WatermarkResult(str(target), False, str(exc))


def apply_watermark_to_files(paths: Iterable[str | Path], text: str) -> WatermarkBatchResult:
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkBatchResult(tuple())
    seen: set[str] = set(); unique: list[str | Path] = []
    for path in paths:
        key = str(Path(path))
        if key not in seen:
            seen.add(key); unique.append(path)
    return WatermarkBatchResult(tuple(apply_docx_footer_watermark(path, watermark) for path in unique))


class ProductAccessMixin:
    def _estimate_selected_document_count(self, selected_medical: list[str], selected_diaries: bool, selected_custom: list[str]) -> int:
        return max(1, len(selected_medical or []) + len(selected_custom or []) + (1 if selected_diaries else 0))

    def _estimate_product_template_count(self, selected_medical: list[str], selected_custom: list[str]) -> int:
        counts = [len(set(selected_medical or ())) + len(set(selected_custom or ())) ]
        for attr in ("custom_document_specs", "custom_documents", "document_specs"):
            value = getattr(self, attr, None)
            if isinstance(value, Mapping):
                counts.append(len(value))
            elif isinstance(value, (list, tuple, set)):
                counts.append(len(value))
        try:
            docs = self.custom_documents_for_main_ui()
            if isinstance(docs, (list, tuple, set)):
                counts.append(len(docs))
        except Exception as exc:
            record_soft_exception("product_access.template_count", exc)
        return max(counts or [1])

    def _estimate_product_profile_count(self) -> int:
        counts = [1]
        for attr in ("profiles", "document_profiles", "universal_profiles", "profile_registry"):
            value = getattr(self, attr, None)
            if isinstance(value, Mapping):
                counts.append(len(value))
            elif isinstance(value, (list, tuple, set)):
                counts.append(len(value))
        return max(counts)

    def _product_access_manager(self) -> ProductAccessManager:
        return ProductAccessManager()

    def create_selected_outputs(self, *, print_after: bool = False) -> None:
        if not product_access_enforcement_enabled():
            return super().create_selected_outputs(print_after=print_after)
        selected = self._selected_outputs_or_warn()
        if selected is None:
            return None
        selected_medical, selected_diaries, selected_custom = selected; manager = self._product_access_manager()
        decision = manager.check_document_creation(self._estimate_selected_document_count(selected_medical, selected_diaries, selected_custom), template_count=self._estimate_product_template_count(selected_medical, selected_custom), profile_count=self._estimate_product_profile_count())
        if not decision.allowed:
            from tkinter import messagebox
            messagebox.showwarning(decision.title, decision.message)
            try:
                self._log(f"\n⚠ {decision.title}: {decision.message}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_denied", exc)
            return None
        if decision.warning:
            try:
                self._log(f"\n⚠ Лицензия: {decision.warning}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_warning", exc)
        return super().create_selected_outputs(print_after=print_after)

    def _created_files_from_results(self, created_medical: list[Path], created_custom: list[Path], diary_result):
        created = super()._created_files_from_results(created_medical, created_custom, diary_result)
        if not created or not product_access_enforcement_enabled():
            return created
        manager = self._product_access_manager(); watermark = manager.current_watermark_text()
        if watermark:
            result = apply_watermark_to_files(created, watermark)
            if result.errors:
                try:
                    self._log("\n⚠ Водяной знак trial/demo применён не ко всем документам:\n" + "\n".join(result.errors[:10]) + "\n")
                except Exception as exc:
                    record_soft_exception("product_access.watermark_log", exc)
        try:
            manager.record_created_documents(len(created))
        except Exception as exc:
            record_soft_exception("product_access.record_created_documents", exc)
        return created


class ProductLicenseMixin:
    def _initialize_app(self, root) -> None:
        super()._initialize_app(root); self._install_product_license_entrypoints()

    def _install_product_license_entrypoints(self) -> None:
        try:
            self.root.bind_all("<Control-l>", lambda _event: self.show_product_license_dialog())
            self.root.bind_all("<Control-L>", lambda _event: self.show_product_license_dialog())
        except Exception as exc:
            record_soft_exception("product_license.install_entrypoints", exc)

    def show_product_license_dialog(self) -> None:
        import tkinter as tk
        from tkinter import filedialog, messagebox
        manager = self._product_access_manager(); window = tk.Toplevel(self.root); window.title("Лицензия Dokkomplekt"); window.transient(self.root); window.grab_set(); window.geometry("620x520"); window.minsize(560, 460)
        outer = tk.Frame(window, padx=16, pady=14); outer.pack(fill="both", expand=True); outer.grid_columnconfigure(0, weight=1); outer.grid_rowconfigure(1, weight=1)
        tk.Label(outer, text="Лицензия и лимиты продукта", font=("Segoe UI", 13, "bold"), anchor="w").grid(row=0, column=0, sticky="ew")
        summary = tk.Text(outer, height=11, wrap="word"); summary.grid(row=1, column=0, sticky="nsew", pady=(10, 10)); summary.configure(state="normal"); summary.insert("1.0", manager.summary_text()); summary.configure(state="disabled")
        tk.Label(outer, text="Для offline-активации вставьте JSON лицензии или загрузите .json файл. Программа проверяет доступ локально и не отправляет документы пациента наружу.", justify="left", wraplength=560, anchor="w").grid(row=2, column=0, sticky="ew", pady=(0, 8))
        license_text = tk.Text(outer, height=7, wrap="word"); license_text.grid(row=3, column=0, sticky="ew"); buttons = tk.Frame(outer); buttons.grid(row=4, column=0, sticky="ew", pady=(12, 0))
        for column in range(4):
            buttons.grid_columnconfigure(column, weight=1)
        def refresh() -> None:
            fresh = self._product_access_manager(); summary.configure(state="normal"); summary.delete("1.0", "end"); summary.insert("1.0", fresh.summary_text()); summary.configure(state="disabled")
        def install_from_text() -> None:
            raw = license_text.get("1.0", "end").strip()
            if not raw:
                messagebox.showwarning("Лицензия", "Вставьте JSON лицензии или загрузите файл лицензии."); return
            try:
                self._product_access_manager().install_license_text(raw); refresh(); messagebox.showinfo("Лицензия", "Лицензия установлена.")
            except Exception as exc:
                messagebox.showerror("Лицензия не установлена", str(exc))
        def load_file() -> None:
            path = filedialog.askopenfilename(title="Выберите файл лицензии", filetypes=(("License JSON", "*.json"), ("All files", "*.*")))
            if not path:
                return
            try:
                license_text.delete("1.0", "end"); license_text.insert("1.0", Path(path).read_text(encoding="utf-8"))
            except OSError as exc:
                messagebox.showerror("Лицензия", f"Не удалось прочитать файл лицензии:\n{exc}")
        tk.Button(buttons, text="Загрузить файл", command=load_file).grid(row=0, column=0, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Установить", command=install_from_text).grid(row=0, column=1, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Обновить", command=refresh).grid(row=0, column=2, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Закрыть", command=window.destroy).grid(row=0, column=3, sticky="ew")
