from __future__ import annotations

"""Local product access, tariff, trial and license contract for Dokkomplekt.

Stores only product metadata. Never store/read/send patient documents, names,
diagnoses, template contents or patient file names here.  The runtime mixins and
watermark helpers intentionally live in this one product module so the project
keeps the existing anti-microfile production contour.
"""

from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
import hashlib
import hmac
import json
import os
import platform
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
from typing import Any, Iterable, Mapping
import uuid

from diagnostic_logging import record_soft_exception

PRODUCT_ACCESS_CONTRACT_VERSION = "v1.0"
WATERMARK_CONTRACT_VERSION = "v1.0"
NO_PATIENT_DATA_IN_LICENSE_STATE = True
LOCAL_ONLY_PRODUCT_ACCESS = True
FOOTER_WATERMARK_ENABLED = True
NO_WATERMARK_FOR_PAID_LICENSES = True
TRIAL_WATERMARK_TEXT = "ПРОБНАЯ ВЕРСИЯ. НЕ ИСПОЛЬЗОВАТЬ КАК МЕДИЦИНСКИЙ ДОКУМЕНТ."
EXPIRED_DEMO_WATERMARK_TEXT = "ДЕМО-ДОКУМЕНТ. ЛИЦЕНЗИЯ НЕ АКТИВНА."


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def parse_dt(value: str | None) -> datetime | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        raw = raw[:-1] + "+00:00" if raw.endswith("Z") else raw
        dt = datetime.fromisoformat(raw)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).replace(microsecond=0)
    except ValueError:
        return None


def iso(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat()


def month_key(now: datetime | None = None) -> str:
    return (now or utc_now()).strftime("%Y-%m")


def stable_json(payload: Mapping[str, Any]) -> str:
    return json.dumps(dict(payload), ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def machine_fingerprint() -> str:
    raw = "|".join(
        str(value or "").lower()
        for value in (
            platform.system(),
            platform.machine(),
            platform.node(),
            os.getenv("COMPUTERNAME"),
            os.getenv("PROCESSOR_IDENTIFIER"),
            uuid.getnode(),
        )
    )
    return hashlib.sha256(raw.encode("utf-8", errors="replace")).hexdigest()[:32]


@dataclass(frozen=True)
class PlanLimits:
    plan_id: str
    title: str
    monthly_price_rub: int
    yearly_price_rub: int
    included_machines: int
    included_users: int
    profile_limit: int
    template_limit: int
    document_limit_month: int
    max_documents_per_run: int
    watermark_mode: str = "none"
    batch_generation: bool = False
    batch_print: bool = False
    shared_department_profile: bool = False
    role_management: bool = False
    offline_activation: bool = True
    local_license_server: bool = False
    overage_percent: int = 20
    grace_days: int = 7
    support_level: str = "base"


PLAN_LIMITS: dict[str, PlanLimits] = {
    "trial": PlanLimits(
        plan_id="trial",
        title="Trial",
        monthly_price_rub=0,
        yearly_price_rub=0,
        included_machines=1,
        included_users=1,
        profile_limit=1,
        template_limit=5,
        document_limit_month=30,
        max_documents_per_run=3,
        watermark_mode="trial",
        offline_activation=False,
        overage_percent=0,
        grace_days=0,
        support_level="knowledge_base",
    ),
    "doctor_start": PlanLimits("doctor_start", "Doctor Start", 1490, 14900, 1, 1, 1, 30, 600, 10),
    "doctor_pro": PlanLimits(
        "doctor_pro",
        "Doctor Pro",
        3900,
        29900,
        2,
        1,
        3,
        150,
        3000,
        50,
        batch_generation=True,
        batch_print=True,
        support_level="priority",
    ),
    "department": PlanLimits(
        "department",
        "Department",
        14900,
        149000,
        5,
        10,
        10,
        500,
        20000,
        100,
        batch_generation=True,
        batch_print=True,
        shared_department_profile=True,
        role_management=True,
        grace_days=14,
        support_level="department",
    ),
    "clinic": PlanLimits(
        "clinic",
        "Clinic",
        49000,
        490000,
        20,
        50,
        50,
        2000,
        100000,
        250,
        batch_generation=True,
        batch_print=True,
        shared_department_profile=True,
        role_management=True,
        local_license_server=True,
        grace_days=30,
        support_level="sla",
    ),
    "enterprise": PlanLimits(
        "enterprise",
        "Enterprise",
        0,
        900000,
        9999,
        9999,
        9999,
        999999,
        9999999,
        1000,
        batch_generation=True,
        batch_print=True,
        shared_department_profile=True,
        role_management=True,
        local_license_server=True,
        grace_days=45,
        support_level="enterprise_sla",
    ),
}


@dataclass(frozen=True)
class LicenseEntitlement:
    license_id: str
    plan: str
    owner_name: str = ""
    organization_name: str = ""
    seats: int = 1
    allowed_machines: tuple[str, ...] = ()
    valid_until: str = ""
    issued_at: str = ""
    generation_limit_month: int | None = None
    template_limit: int | None = None
    profile_limit: int | None = None
    watermark_mode: str | None = None
    offline_grace_days: int | None = None
    features: tuple[str, ...] = ()
    signature: str = ""

    @classmethod
    def from_mapping(cls, payload: Mapping[str, Any]) -> "LicenseEntitlement":
        return cls(
            license_id=str(payload.get("license_id") or "").strip(),
            plan=str(payload.get("plan") or "").lower().strip(),
            owner_name=str(payload.get("owner_name") or "").strip(),
            organization_name=str(payload.get("organization_name") or "").strip(),
            seats=max(1, int(payload.get("seats") or 1)),
            allowed_machines=tuple(
                str(item).lower().strip() for item in payload.get("allowed_machines", ()) if str(item).strip()
            ),
            valid_until=str(payload.get("valid_until") or "").strip(),
            issued_at=str(payload.get("issued_at") or "").strip(),
            generation_limit_month=int(payload["generation_limit_month"])
            if payload.get("generation_limit_month") is not None
            else None,
            template_limit=int(payload["template_limit"]) if payload.get("template_limit") is not None else None,
            profile_limit=int(payload["profile_limit"]) if payload.get("profile_limit") is not None else None,
            watermark_mode=str(payload.get("watermark_mode")).lower().strip()
            if payload.get("watermark_mode") is not None
            else None,
            offline_grace_days=int(payload["offline_grace_days"])
            if payload.get("offline_grace_days") is not None
            else None,
            features=tuple(str(item).strip() for item in payload.get("features", ()) if str(item).strip()),
            signature=str(payload.get("signature") or "").strip(),
        )

    def unsigned_payload(self) -> dict[str, Any]:
        return {
            "license_id": self.license_id,
            "plan": self.plan,
            "owner_name": self.owner_name,
            "organization_name": self.organization_name,
            "seats": self.seats,
            "allowed_machines": list(self.allowed_machines),
            "valid_until": self.valid_until,
            "issued_at": self.issued_at,
            "generation_limit_month": self.generation_limit_month,
            "template_limit": self.template_limit,
            "profile_limit": self.profile_limit,
            "watermark_mode": self.watermark_mode,
            "offline_grace_days": self.offline_grace_days,
            "features": list(self.features),
        }

    def plan_limits(self) -> PlanLimits:
        if self.plan not in PLAN_LIMITS or self.plan == "trial":
            raise ValueError(f"Unknown paid license plan: {self.plan}")
        return PLAN_LIMITS[self.plan]

    def valid_until_dt(self) -> datetime | None:
        return parse_dt(self.valid_until)

    def is_expired(self, now: datetime | None = None) -> bool:
        valid_until = self.valid_until_dt()
        return valid_until is None or (now or utc_now()) > valid_until

    def signature_expected(self, secret: str) -> str:
        return hmac.new(secret.encode(), stable_json(self.unsigned_payload()).encode(), hashlib.sha256).hexdigest()

    def signature_valid(self, secret: str) -> bool:
        return bool(self.signature and secret and hmac.compare_digest(self.signature, self.signature_expected(secret)))


@dataclass(frozen=True)
class LicenseState:
    plan: str
    title: str
    active: bool
    reason: str
    license_id: str = ""
    owner_label: str = ""
    valid_until: str = ""
    trial_started_at: str = ""
    trial_ends_at: str = ""
    days_left: int = 0
    documents_used_month: int = 0
    documents_limit_month: int = 0
    documents_used_total_trial: int = 0
    remaining_documents_month: int = 0
    template_limit: int = 0
    profile_limit: int = 0
    included_machines: int = 1
    watermark_mode: str = "none"
    warning: str = ""

    @property
    def watermark_required(self) -> bool:
        return self.watermark_mode in {"trial", "expired_demo"}

    def watermark_text(self) -> str:
        if self.watermark_mode == "trial":
            return TRIAL_WATERMARK_TEXT
        if self.watermark_mode == "expired_demo":
            return EXPIRED_DEMO_WATERMARK_TEXT
        return ""


@dataclass(frozen=True)
class AccessDecision:
    allowed: bool
    code: str
    title: str
    message: str
    state: LicenseState
    warning: str = ""


@dataclass(frozen=True)
class WatermarkResult:
    path: str
    changed: bool
    error: str = ""


@dataclass(frozen=True)
class WatermarkBatchResult:
    results: tuple[WatermarkResult, ...]

    @property
    def errors(self) -> tuple[str, ...]:
        return tuple(f"{item.path}: {item.error}" for item in self.results if item.error)

    @property
    def changed_count(self) -> int:
        return sum(1 for item in self.results if item.changed)


class ProductAccessManager:
    def __init__(self, storage_dir: str | Path | None = None, now: datetime | None = None):
        self.storage_dir = Path(storage_dir) if storage_dir else self.default_storage_dir()
        self.now = now
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.state_path = self.storage_dir / "product_access_state.json"
        self.license_path = Path(os.getenv("DOKKOMPLEKT_LICENSE_FILE") or self.storage_dir / "license.json")

    @staticmethod
    def default_storage_dir() -> Path:
        if os.getenv("DOKKOMPLEKT_LICENSE_DIR"):
            return Path(os.environ["DOKKOMPLEKT_LICENSE_DIR"]).expanduser()
        if os.getenv("LOCALAPPDATA"):
            return Path(os.environ["LOCALAPPDATA"]) / "Dokkomplekt"
        return Path.home() / ".dokkomplekt"

    def _now(self) -> datetime:
        return self.now or utc_now()

    def _load_state_payload(self) -> dict[str, Any]:
        try:
            if not self.state_path.exists():
                return {}
            payload = json.loads(self.state_path.read_text("utf-8"))
            return payload if isinstance(payload, dict) else {}
        except Exception:
            return {}

    def _save_state_payload(self, payload: Mapping[str, Any]) -> None:
        tmp_path = self.state_path.with_suffix(".tmp")
        tmp_path.write_text(json.dumps(dict(payload), ensure_ascii=False, indent=2, sort_keys=True), "utf-8")
        os.replace(tmp_path, self.state_path)

    def _ensure_trial_started(self, payload: dict[str, Any]) -> dict[str, Any]:
        if not payload.get("trial_started_at"):
            payload = dict(payload)
            payload["trial_started_at"] = iso(self._now())
            payload.setdefault("usage_by_month", {})
            payload.setdefault("trial_created_total", 0)
            self._save_state_payload(payload)
        return payload

    @staticmethod
    def _license_secret() -> str:
        if os.getenv("DOKKOMPLEKT_LICENSE_VERIFY_SECRET"):
            return os.environ["DOKKOMPLEKT_LICENSE_VERIFY_SECRET"].strip()
        try:
            from product_license_secret import LICENSE_VERIFY_SECRET  # type: ignore

            return str(LICENSE_VERIFY_SECRET or "").strip()
        except Exception:
            return ""

    def load_license(self) -> LicenseEntitlement | None:
        try:
            if not self.license_path.exists():
                return None
            payload = json.loads(self.license_path.read_text("utf-8"))
            return LicenseEntitlement.from_mapping(payload) if isinstance(payload, dict) else None
        except Exception:
            return None

    def install_license_text(self, text: str) -> LicenseState:
        payload = json.loads(text or "{}")
        if not isinstance(payload, dict):
            raise ValueError("Файл лицензии должен быть JSON-объектом.")
        entitlement = LicenseEntitlement.from_mapping(payload)
        self._validate_license(entitlement, require_not_expired=False)
        tmp_path = self.license_path.with_suffix(".tmp")
        tmp_path.parent.mkdir(parents=True, exist_ok=True)
        stored = entitlement.unsigned_payload()
        stored["signature"] = entitlement.signature
        tmp_path.write_text(json.dumps(stored, ensure_ascii=False, indent=2, sort_keys=True), "utf-8")
        os.replace(tmp_path, self.license_path)
        return self.current_state()

    def _validate_license(self, entitlement: LicenseEntitlement, *, require_not_expired: bool = True) -> None:
        if not entitlement.license_id:
            raise ValueError("В лицензии нет license_id.")
        if entitlement.plan not in PLAN_LIMITS or entitlement.plan == "trial":
            raise ValueError(f"Неизвестный тариф лицензии: {entitlement.plan!r}.")
        secret = self._license_secret()
        unsigned_ok = os.getenv("DOKKOMPLEKT_ALLOW_UNSIGNED_LICENSES", "").lower() in {"1", "true", "yes", "on"}
        if secret and not entitlement.signature_valid(secret):
            raise ValueError("Подпись лицензии не прошла проверку.")
        if not secret and not unsigned_ok:
            raise ValueError("Подпись лицензии не может быть проверена в этой сборке.")
        if require_not_expired and entitlement.is_expired(self._now()):
            raise ValueError("Срок действия лицензии истёк.")
        if entitlement.allowed_machines and machine_fingerprint() not in entitlement.allowed_machines:
            raise ValueError("Лицензия не привязана к этому компьютеру.")

    def current_state(self) -> LicenseState:
        payload = self._ensure_trial_started(self._load_state_payload())
        usage = payload.get("usage_by_month") if isinstance(payload.get("usage_by_month"), dict) else {}
        used = int(usage.get(month_key(self._now()), 0) or 0)
        trial_total = int(payload.get("trial_created_total", 0) or 0)
        entitlement = self.load_license()
        if entitlement:
            try:
                self._validate_license(entitlement, require_not_expired=False)
                return self._paid_state(entitlement, used)
            except ValueError as exc:
                return self._blocked_state(str(exc), used, trial_total)
        return self._trial_state(payload, used, trial_total)

    def _paid_state(self, entitlement: LicenseEntitlement, used: int) -> LicenseState:
        limits = entitlement.plan_limits()
        valid_until = entitlement.valid_until_dt()
        now = self._now()
        grace = int(entitlement.offline_grace_days if entitlement.offline_grace_days is not None else limits.grace_days)
        expired = valid_until is None or now > valid_until
        in_grace = bool(expired and valid_until and now <= valid_until + timedelta(days=grace))
        if expired and not in_grace:
            return self._blocked_state("Срок действия лицензии истёк.", used, 0)
        monthly_limit = int(entitlement.generation_limit_month or limits.document_limit_month)
        days_left = max(0, int(((valid_until or now) - now).total_seconds() // 86400))
        return LicenseState(
            plan=entitlement.plan,
            title=limits.title,
            active=True,
            reason="active_grace" if in_grace else "active",
            license_id=entitlement.license_id,
            owner_label=entitlement.organization_name or entitlement.owner_name,
            valid_until=entitlement.valid_until,
            days_left=days_left,
            documents_used_month=used,
            documents_limit_month=monthly_limit,
            remaining_documents_month=max(0, monthly_limit - used),
            template_limit=int(entitlement.template_limit or limits.template_limit),
            profile_limit=int(entitlement.profile_limit or limits.profile_limit),
            included_machines=int(entitlement.seats or limits.included_machines),
            watermark_mode=str(entitlement.watermark_mode or limits.watermark_mode),
            warning=f"Лицензия истекла, действует льготный период {grace} дн." if in_grace else "",
        )

    def _trial_state(self, payload: Mapping[str, Any], used: int, trial_total: int) -> LicenseState:
        limits = PLAN_LIMITS["trial"]
        started_at = parse_dt(str(payload.get("trial_started_at") or "")) or self._now()
        ends_at = started_at + timedelta(days=14)
        active = self._now() <= ends_at and trial_total < limits.document_limit_month
        reason = "trial_active" if active else "trial_document_limit" if trial_total >= limits.document_limit_month else "trial_expired"
        return LicenseState(
            plan="trial",
            title=limits.title,
            active=active,
            reason=reason,
            trial_started_at=iso(started_at),
            trial_ends_at=iso(ends_at),
            days_left=max(0, int((ends_at - self._now()).total_seconds() // 86400) + (1 if self._now() <= ends_at else 0)),
            documents_used_month=used,
            documents_limit_month=limits.document_limit_month,
            documents_used_total_trial=trial_total,
            remaining_documents_month=max(0, limits.document_limit_month - trial_total),
            template_limit=limits.template_limit,
            profile_limit=limits.profile_limit,
            included_machines=limits.included_machines,
            watermark_mode=limits.watermark_mode if active else "expired_demo",
            warning="Пробная версия создаёт документы только с водяным знаком." if active else "Пробный период завершён.",
        )

    def _blocked_state(self, reason: str, used: int, trial_total: int) -> LicenseState:
        return LicenseState(
            plan="blocked",
            title="Лицензия не активна",
            active=False,
            reason="blocked",
            documents_used_month=used,
            documents_used_total_trial=trial_total,
            watermark_mode="expired_demo",
            warning=reason,
        )

    def check_document_creation(
        self,
        requested_count: int,
        *,
        template_count: int | None = None,
        profile_count: int | None = None,
    ) -> AccessDecision:
        count = max(1, int(requested_count or 1))
        state = self.current_state()
        if not state.active:
            return AccessDecision(False, "license_inactive", "Лицензия не активна", state.warning or "Создание рабочих документов заблокировано.", state)
        limits = PLAN_LIMITS.get(state.plan, PLAN_LIMITS["trial"])
        if count > limits.max_documents_per_run:
            return AccessDecision(
                False,
                "per_run_limit",
                "Слишком много документов за один запуск",
                f"Тариф разрешает до {limits.max_documents_per_run} документов за один запуск. Выбрано: {count}.",
                state,
            )
        if template_count is not None and int(template_count) > state.template_limit:
            return AccessDecision(False, "template_limit", "Превышен лимит шаблонов", f"Лимит тарифа: {state.template_limit} шаблонов.", state)
        if profile_count is not None and int(profile_count) > state.profile_limit:
            return AccessDecision(False, "profile_limit", "Превышен лимит профилей", f"Лимит тарифа: {state.profile_limit} профилей.", state)
        if state.plan == "trial":
            if count > state.remaining_documents_month:
                return AccessDecision(False, "trial_limit", "Пробный лимит исчерпан", "Пробная версия разрешает 30 созданных документов всего.", state)
            return AccessDecision(True, "ok_trial", "Пробная версия", "Документы будут созданы с пробным водяным знаком.", state, state.warning)
        hard_limit = state.documents_limit_month + int(state.documents_limit_month * max(0, limits.overage_percent) / 100)
        projected = state.documents_used_month + count
        if state.documents_limit_month and projected > hard_limit:
            return AccessDecision(
                False,
                "monthly_limit",
                "Месячный лимит документов исчерпан",
                f"Использовано {state.documents_used_month}/{state.documents_limit_month}; льготный перерасход исчерпан.",
                state,
            )
        warning = state.warning
        if not warning and state.documents_limit_month and projected > state.documents_limit_month:
            warning = f"Будет превышен месячный лимит {state.documents_limit_month}; действует льготный перерасход до {hard_limit}."
        if not warning and state.documents_limit_month and projected >= int(state.documents_limit_month * 0.8):
            warning = f"Использовано более 80% месячного лимита: после создания будет {projected}/{state.documents_limit_month}."
        return AccessDecision(True, "ok", "Доступ разрешён", "Создание документов разрешено.", state, warning)

    def record_created_documents(self, count: int) -> None:
        delta = max(0, int(count or 0))
        if not delta:
            return
        payload = self._ensure_trial_started(self._load_state_payload())
        usage = payload.get("usage_by_month") if isinstance(payload.get("usage_by_month"), dict) else {}
        key = month_key(self._now())
        usage[key] = int(usage.get(key, 0) or 0) + delta
        payload["usage_by_month"] = usage
        if self.current_state().plan == "trial":
            payload["trial_created_total"] = int(payload.get("trial_created_total", 0) or 0) + delta
        payload["updated_at"] = iso(self._now())
        self._save_state_payload(payload)

    def current_watermark_text(self) -> str:
        return self.current_state().watermark_text()

    def summary_text(self) -> str:
        state = self.current_state()
        used = state.documents_used_total_trial if state.plan == "trial" else state.documents_used_month
        lines = [f"Тариф: {state.title}", f"Статус: {'активен' if state.active else 'не активен'}"]
        if state.owner_label:
            lines.append(f"Владелец: {state.owner_label}")
        if state.license_id:
            lines.append(f"Лицензия: {state.license_id}")
        if state.valid_until:
            lines.append(f"Действует до: {state.valid_until}")
        if state.trial_ends_at:
            lines.append(f"Пробный период до: {state.trial_ends_at}")
        if state.documents_limit_month:
            lines.append(f"Документы: {used} / {state.documents_limit_month}")
        lines.extend(
            [
                f"Шаблоны: до {state.template_limit}",
                f"Профили: до {state.profile_limit}",
                f"Компьютеры: до {state.included_machines}",
            ]
        )
        if state.watermark_required:
            lines.append("Водяной знак: включён")
        if state.warning:
            lines.append(f"Предупреждение: {state.warning}")
        return "\n".join(lines)


def sign_license_payload(payload: Mapping[str, Any], secret: str) -> dict[str, Any]:
    entitlement = LicenseEntitlement.from_mapping(payload)
    unsigned = entitlement.unsigned_payload()
    unsigned["signature"] = hmac.new(str(secret).encode(), stable_json(unsigned).encode(), hashlib.sha256).hexdigest()
    return unsigned


def apply_docx_footer_watermark(path: str | Path, text: str) -> WatermarkResult:
    target = Path(path)
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkResult(str(target), changed=False)
    if target.suffix.lower() != ".docx":
        return WatermarkResult(str(target), changed=False, error="watermark supports generated .docx files only")
    if not target.exists() or not target.is_file():
        return WatermarkResult(str(target), changed=False, error="file not found")
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency is part of app requirements
        return WatermarkResult(str(target), changed=False, error=f"python-docx unavailable: {exc}")
    try:
        document = Document(str(target))
        changed = False
        for section in document.sections:
            footer = section.footer
            existing = "\n".join(paragraph.text for paragraph in footer.paragraphs)
            if watermark in existing:
                continue
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if paragraph.text.strip():
                paragraph = footer.add_paragraph()
            paragraph.text = f" {watermark} "
            changed = True
        if changed:
            document.save(str(target))
        return WatermarkResult(str(target), changed=changed)
    except Exception as exc:
        return WatermarkResult(str(target), changed=False, error=str(exc))


def apply_watermark_to_files(paths: Iterable[str | Path], text: str) -> WatermarkBatchResult:
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkBatchResult(tuple())
    return WatermarkBatchResult(tuple(apply_docx_footer_watermark(path, watermark) for path in paths))


class ProductAccessMixin:
    """Wrap document creation with local licensing, limits and watermark policy."""

    def _estimate_selected_document_count(
        self,
        selected_medical: list[str],
        selected_diaries: bool,
        selected_custom: list[str],
    ) -> int:
        return max(1, len(selected_medical or []) + len(selected_custom or []) + (1 if selected_diaries else 0))

    def _product_access_manager(self) -> ProductAccessManager:
        return ProductAccessManager()

    def create_selected_outputs(self, *, print_after: bool = False) -> None:
        selected = self._selected_outputs_or_warn()
        if selected is None:
            return
        selected_medical, selected_diaries, selected_custom = selected
        estimated_count = self._estimate_selected_document_count(selected_medical, selected_diaries, selected_custom)
        manager = self._product_access_manager()
        decision = manager.check_document_creation(estimated_count)
        if not decision.allowed:
            messagebox.showwarning(decision.title, decision.message)
            try:
                self._log(f"\n⚠ {decision.title}: {decision.message}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_denied", exc)
            return
        if decision.warning:
            try:
                self._log(f"\n⚠ Лицензия: {decision.warning}\n")
            except Exception as exc:
                record_soft_exception("product_access.log_warning", exc)
        return super().create_selected_outputs(print_after=print_after)

    def _created_files_from_results(self, created_medical: list[Path], created_custom: list[Path], diary_result):
        created_files = super()._created_files_from_results(created_medical, created_custom, diary_result)
        if not created_files:
            return created_files
        manager = self._product_access_manager()
        watermark = manager.current_watermark_text()
        if watermark:
            result = apply_watermark_to_files(created_files, watermark)
            if result.errors:
                try:
                    self._log(
                        "\n⚠ Водяной знак trial/demo применён не ко всем документам:\n"
                        + "\n".join(result.errors[:10])
                        + "\n"
                    )
                except Exception as exc:
                    record_soft_exception("product_access.watermark_log", exc)
        try:
            manager.record_created_documents(len(created_files))
        except Exception as exc:
            record_soft_exception("product_access.record_created_documents", exc)
        return created_files


class ProductLicenseMixin:
    def _initialize_app(self, root: tk.Tk) -> None:
        super()._initialize_app(root)
        self._install_product_license_entrypoints()

    def _install_product_license_entrypoints(self) -> None:
        try:
            self.root.bind_all("<Control-l>", lambda _event: self.show_product_license_dialog())
            self.root.bind_all("<Control-L>", lambda _event: self.show_product_license_dialog())
        except Exception as exc:
            record_soft_exception("product_license.install_entrypoints", exc)

    def show_product_license_dialog(self) -> None:
        manager = self._product_access_manager()
        window = tk.Toplevel(self.root)
        window.title("Лицензия Dokkomplekt")
        window.transient(self.root)
        window.grab_set()
        window.geometry("620x520")
        window.minsize(560, 460)

        outer = tk.Frame(window, padx=16, pady=14)
        outer.pack(fill="both", expand=True)
        outer.grid_columnconfigure(0, weight=1)
        outer.grid_rowconfigure(1, weight=1)

        tk.Label(outer, text="Лицензия и лимиты продукта", font=("Segoe UI", 13, "bold"), anchor="w").grid(
            row=0,
            column=0,
            sticky="ew",
        )

        summary = tk.Text(outer, height=11, wrap="word")
        summary.grid(row=1, column=0, sticky="nsew", pady=(10, 10))
        summary.configure(state="normal")
        summary.insert("1.0", manager.summary_text())
        summary.configure(state="disabled")

        tk.Label(
            outer,
            text=(
                "Для offline-активации вставьте JSON лицензии или загрузите .json файл. "
                "Программа проверяет доступ локально и не отправляет документы пациента наружу."
            ),
            justify="left",
            wraplength=560,
            anchor="w",
        ).grid(row=2, column=0, sticky="ew", pady=(0, 8))

        license_text = tk.Text(outer, height=7, wrap="word")
        license_text.grid(row=3, column=0, sticky="ew")

        buttons = tk.Frame(outer)
        buttons.grid(row=4, column=0, sticky="ew", pady=(12, 0))
        for column in range(4):
            buttons.grid_columnconfigure(column, weight=1)

        def refresh() -> None:
            fresh_manager = self._product_access_manager()
            summary.configure(state="normal")
            summary.delete("1.0", "end")
            summary.insert("1.0", fresh_manager.summary_text())
            summary.configure(state="disabled")

        def install_from_text() -> None:
            raw = license_text.get("1.0", "end").strip()
            if not raw:
                messagebox.showwarning("Лицензия", "Вставьте JSON лицензии или загрузите файл лицензии.")
                return
            try:
                self._product_access_manager().install_license_text(raw)
                refresh()
                messagebox.showinfo("Лицензия", "Лицензия установлена.")
            except Exception as exc:
                messagebox.showerror("Лицензия не установлена", str(exc))

        def load_file() -> None:
            path = filedialog.askopenfilename(
                title="Выберите файл лицензии",
                filetypes=(("License JSON", "*.json"), ("All files", "*.*")),
            )
            if not path:
                return
            try:
                data = Path(path).read_text(encoding="utf-8")
                license_text.delete("1.0", "end")
                license_text.insert("1.0", data)
            except OSError as exc:
                messagebox.showerror("Лицензия", f"Не удалось прочитать файл лицензии:\n{exc}")

        tk.Button(buttons, text="Загрузить файл", command=load_file).grid(row=0, column=0, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Установить", command=install_from_text).grid(row=0, column=1, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Обновить", command=refresh).grid(row=0, column=2, sticky="ew", padx=(0, 6))
        tk.Button(buttons, text="Закрыть", command=window.destroy).grid(row=0, column=3, sticky="ew")
