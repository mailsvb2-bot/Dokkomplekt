from __future__ import annotations

"""Trial/demo watermark helpers for generated DOCX files.

This module is intentionally post-processing only: it never decides whether a
user is allowed to generate documents. Product access decisions live in
``product_licensing``. Watermarking is best-effort and local-only; failures are
reported to the caller without sending document contents anywhere.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

WATERMARK_CONTRACT_VERSION = "v1.0"
FOOTER_WATERMARK_ENABLED = True
NO_WATERMARK_FOR_PAID_LICENSES = True


@dataclass(frozen=True)
class WatermarkResult:
    path: str
    changed: bool
    error: str = ""


@dataclass(frozen=True)
class WatermarkBatchResult:
    results: tuple[WatermarkResult, ...]

    @property
    def errors(self) -> tuple[str, ...]:
        return tuple(f"{item.path}: {item.error}" for item in self.results if item.error)

    @property
    def changed_count(self) -> int:
        return sum(1 for item in self.results if item.changed)


def apply_docx_footer_watermark(path: str | Path, text: str) -> WatermarkResult:
    target = Path(path)
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkResult(str(target), changed=False)
    if target.suffix.lower() != ".docx":
        return WatermarkResult(str(target), changed=False, error="watermark supports generated .docx files only")
    if not target.exists() or not target.is_file():
        return WatermarkResult(str(target), changed=False, error="file not found")
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency is normally present in app requirements
        return WatermarkResult(str(target), changed=False, error=f"python-docx unavailable: {exc}")

    try:
        document = Document(str(target))
        changed = False
        for section in document.sections:
            footer = section.footer
            existing = "\n".join(paragraph.text for paragraph in footer.paragraphs)
            if watermark in existing:
                continue
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if paragraph.text.strip():
                paragraph = footer.add_paragraph()
            paragraph.text = f" {watermark} "
            changed = True
        if changed:
            document.save(str(target))
        return WatermarkResult(str(target), changed=changed)
    except Exception as exc:
        return WatermarkResult(str(target), changed=False, error=str(exc))


def apply_watermark_to_files(paths: Iterable[str | Path], text: str) -> WatermarkBatchResult:
    watermark = str(text or "").strip()
    if not watermark:
        return WatermarkBatchResult(tuple())
    return WatermarkBatchResult(tuple(apply_docx_footer_watermark(path, watermark) for path in paths))
