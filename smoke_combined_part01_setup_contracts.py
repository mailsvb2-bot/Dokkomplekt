from pathlib import Path
import shutil
from itertools import combinations

from docx import Document

import medical_documents as _medical_documents_module
from medical_documents import MedicalDocumentService, MedicalTextParser, DOCUMENT_ORDER, PatientData, build_expert_anamnesis, extract_docx_text, sanitize_diagnosis, treatment_period_text, parse_date, format_military_commissariat_area, format_military_commissariat_referral, format_date_with_russian_year_suffix, format_birth_for_person_line
assert not hasattr(_medical_documents_module, "MedicalApp"), "medical_documents.py must not contain old UI class"
from diary_filler import fill_diary_batch, extract_statuses_from_docx, parse_full_date, parse_month_year, safe_filename_part
from icd10_f import search_icd10_f
from medical_docx_reader import _first_valid_full_date

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "test_run_combined"
if OUT.exists():
    shutil.rmtree(OUT)
OUT.mkdir(parents=True, exist_ok=True)

nav = OUT / "Направление_тест.docx"
nav_doc = Document()
nav_doc.add_paragraph("10.06.2026 Первичный осмотр")
nav_doc.add_paragraph("История болезни № 123")
nav_doc.add_paragraph("Ф.И.О.: Иванова Ирина Ивановна")
nav_doc.add_paragraph("Год рождения: 1980")
nav_doc.add_paragraph("Зарегистрирован: город N, тестовый район")
nav_doc.add_paragraph("Работает в организации: не работает")
nav_doc.add_paragraph("В 3 отделение КДП поступает добровольно")
nav_doc.add_paragraph("Жалобы на момент осмотра: тревога, нарушение сна")
nav_doc.add_paragraph("Анамнез жизни: Со слов пациентки, росла и развивалась без особенностей.")
nav_doc.add_paragraph("Анамнез заболевания: Ухудшение состояния в течение месяца.")
nav_doc.add_paragraph("Профильный статус: В сознании, ориентирована, контакт доступен.")
nav_doc.add_paragraph("Соматический статус: Нормального питания, без грубой соматической патологии.")
nav_doc.add_paragraph("План лечения: терапия по назначению врача")
nav_doc.add_paragraph("На основании данных анамнеза жизни и заболевания, профильного статуса, данных клинических исследований был выставлен диагноз: F06.7 тестовый диагноз")
nav_doc.save(nav)

epi = OUT / "ЭПИ.docx"
epi_doc = Document()
epi_doc.add_paragraph("ЭПИ: ЭПИ тестовая информация.")
epi_doc.save(epi)



# --- Universal profile foundation smoke ---
from universal_fields import PatientCase, default_field_registry
from universal_profiles import DocumentPack, DocumentTemplateSpec, default_document_pack, load_document_pack, save_document_pack
from universal_scanner import learn_rule_from_selection, merge_scan_results, scan_docx, scan_many_docx
from universal_template_engine import (
    attach_template_to_pack,
    export_document_pack_zip,
    import_document_pack_zip,
    infer_document_spec_from_template,
    render_template_to_docx,
    validate_document_pack,
    validate_template,
)
from universal_generation import analyze_pack_readiness, render_documents_from_pack
from universal_main_documents import (
    CUSTOM_DOCUMENT_KIND_PREFIX,
    assert_dynamic_medpack_button_lock,
    custom_document_id_from_kind,
    custom_documents_for_main_ui,
    custom_kind,
    is_custom_kind,
    selected_custom_document_ids,
)
from universal_case_adapter import merge_case_values, patient_data_to_case
from universal_profile_builder import (
    build_profile_from_sources_and_templates,
    create_pack_from_preset,
    get_specialty_preset,
    ingest_templates_into_pack,
    recognize_template_buttons,
    profile_setup_checklist,
    specialty_presets,
)
from regulatory_advisory_policy import (
    ADVISORY_IS_NEVER_BLOCKING,
    DECLINE_LABEL,
    REGULATORY_SOFT_ADVISORY_LOCK_VERSION,
    assert_soft_advisory_lock,
)
from regulatory_document_classifier import classify_docx, classify_document_text
from regulatory_document_roles import default_document_role_registry
from regulatory_section_registry import default_section_registry
from regulatory_specialty_overlays import default_specialty_overlay_registry
from regulatory_template_advisor import advise_document, advise_template, save_advice_report
from regulatory_caucasus_aliases import assert_caucasus_alias_lock, caucasus_context_report
from regulatory_completion_blocks import (
    COMPLETION_POPUP_LOCK_VERSION,
    COMPLETION_POPUP_TITLE,
    assert_completion_popup_lock,
    apply_completion_values,
    completion_inputs_from_advice,
    completion_values_from_raw,
)
from medical_language_catalog import assert_language_catalog_lock, language_id_from_choice, language_profile
from medical_language_detector import assert_language_detection_lock, detect_text_language
from i18n_strings import assert_i18n_strings_lock, tr
from medical_orthography import assert_orthography_medical_safe_lock, correct_medical_text
from personal_document_buttons import (
    assert_personal_document_button_lock,
    localized_role_label,
    regular_document_role_choices,
    recognize_document_title_from_template,
    role_id_from_choice,
    suggest_button_label_for_template,
)
from language_preferences import LanguagePreferences, assert_language_preferences_lock

registry = default_field_registry()
assert "patient.fio" in registry
assert "patient.snils" in registry
pack = default_document_pack()
assert pack.document_labels() == (), pack.document_labels()
assert pack.documents == (), "doctor-owned default pack must start empty"
scan = scan_docx(nav, registry=pack.registry(), rules=pack.extraction_rules)
best = scan.best_matches()
assert best["patient.fio"].value == "Иванова Ирина Ивановна", best.get("patient.fio")
assert best["case.number"].value == "123", best.get("case.number")
assert best["admission.date"].value == "10.06.2026", best.get("admission.date")
assert "diagnosis.main" in best, best.keys()
assert best["diagnosis.icd10"].value == "F06.7", best.get("diagnosis.icd10")
manual_rule = learn_rule_from_selection(scan.blocks, field_id="complaints", selected_text="тревога, нарушение сна", registry=pack.registry())
assert manual_rule.field_id == "complaints"
pack.add_rule(manual_rule)
profile_path = OUT / "universal_profile.medpack.json"
save_document_pack(pack, profile_path)
loaded_pack = load_document_pack(profile_path)
assert loaded_pack.extraction_rules and loaded_pack.extraction_rules[-1].field_id == "complaints"
case = PatientCase()
for field_id, match in best.items():
    case.set(field_id, match.value, confidence=match.confidence)
assert case.get("patient.fio") == "Иванова Ирина Ивановна"
assert scan.review_rows(), "scanner must expose doctor review rows"
merged_scan = merge_scan_results(scan, scan)
assert merged_scan.best_matches()["patient.fio"].value == "Иванова Ирина Ивановна"
assert "procedure.anesthesia" in registry

custom_template = OUT / "custom_template.docx"
custom_doc = Document()
custom_doc.add_paragraph("Пациент: {{patient.fio}}")
custom_doc.add_paragraph("Диагноз: {{diagnosis.main}}")
custom_doc.add_paragraph("Документ: {{document.label}}")
custom_doc.add_paragraph("Анестезия: {{procedure.anesthesia}}")
custom_doc.save(custom_template)
custom_spec = infer_document_spec_from_template(custom_template, button_label="Тестовый пользовательский документ", registry=registry)
assert "patient.fio" in custom_spec.required_fields
assert "diagnosis.main" in custom_spec.required_fields
custom_validation = validate_template(custom_template, required_fields=custom_spec.required_fields, registry=registry)
assert custom_validation.ok, custom_validation.to_dict()
case.set("procedure.anesthesia", "местная", confidence=1.0)
rendered_custom = OUT / "rendered_custom.docx"
render_result = render_template_to_docx(template_path=custom_template, output_path=rendered_custom, case=case, document=custom_spec)
assert render_result.ok and rendered_custom.exists()
rendered_text = "\n".join(p.text for p in Document(rendered_custom).paragraphs)
assert "Иванова Ирина Ивановна" in rendered_text
assert "Тестовый пользовательский документ" in rendered_text
assert "местная" in rendered_text
pack.add_document(custom_spec)
profile_owned_pack = default_document_pack()
owned_spec, owned_template = attach_template_to_pack(profile_owned_pack, custom_template, OUT / "profile_owned", registry=registry)
assert owned_template.exists() and owned_spec.template.startswith("templates/"), owned_spec
readiness = analyze_pack_readiness(profile_owned_pack, case, base_dir=OUT / "profile_owned")
assert readiness.ready_count >= 1, readiness.human_report()
render_pack_dir = OUT / "rendered_pack"
pack_render = render_documents_from_pack(pack=profile_owned_pack, case=case, document_ids=[owned_spec.id], output_dir=render_pack_dir, base_dir=OUT / "profile_owned")
assert pack_render.ok and pack_render.created_files, pack_render.human_report()

# Dynamic medpack buttons must be exposed in a separate namespace for block 03.
assert_dynamic_medpack_button_lock()
assert CUSTOM_DOCUMENT_KIND_PREFIX == "custom_profile:"
custom_kind_value = custom_kind(owned_spec.id)
assert is_custom_kind(custom_kind_value)
assert custom_document_id_from_kind(custom_kind_value) == owned_spec.id
main_custom_docs = custom_documents_for_main_ui(profile_owned_pack, base_dir=OUT / "profile_owned")
assert any(item.document_id == owned_spec.id and item.kind == custom_kind_value for item in main_custom_docs), [item.to_dict() for item in main_custom_docs]

class _FakeVar:
    def __init__(self, value): self._value = value
    def get(self): return self._value

assert selected_custom_document_ids({custom_kind_value: _FakeVar(True), "primary": _FakeVar(True)}) == (owned_spec.id,)
pack_validation = validate_document_pack(profile_owned_pack, base_dir=OUT / "profile_owned")
assert pack_validation.ok or pack_validation.warnings
exported_pack = OUT / "profile_export.medpack.zip"
export_document_pack_zip(profile_owned_pack, exported_pack, template_base_dir=OUT / "profile_owned")
assert exported_pack.exists() and exported_pack.stat().st_size > 0
imported_pack, imported_manifest = import_document_pack_zip(exported_pack, OUT / "imported_profile")
assert imported_manifest.exists(), imported_manifest
assert any(str(doc.template).replace("\\", "/").startswith("templates/") for doc in imported_pack.documents if doc.id == owned_spec.id)
assert (OUT / "imported_profile" / "templates" / owned_template.name).exists()

# A custom template without placeholders must be rejected, not silently accepted.
empty_template = OUT / "empty_custom_template.docx"
empty_doc = Document()
empty_doc.add_paragraph("Обычный текст без placeholders")
empty_doc.save(empty_template)
empty_validation = validate_template(empty_template, registry=registry)
assert not empty_validation.ok, empty_validation.to_dict()

# Header/footer placeholders are common in real hospital templates and must be rendered.
header_template = OUT / "header_template.docx"
header_doc = Document()
header_doc.sections[0].header.paragraphs[0].text = "Пациент в шапке: {{patient.fio}}"
header_doc.add_paragraph("Тело документа: {{diagnosis.main}}")
header_doc.sections[0].footer.paragraphs[0].text = "Подвал: {{document.label}}"
header_doc.save(header_template)
header_spec = infer_document_spec_from_template(header_template, button_label="Шапка и подвал", registry=registry)
header_placeholders = validate_template(header_template, required_fields=header_spec.required_fields, registry=registry)
assert "patient.fio" in header_placeholders.field_ids(), header_placeholders.to_dict()
header_rendered = OUT / "header_rendered.docx"
header_result = render_template_to_docx(template_path=header_template, output_path=header_rendered, case=case, document=header_spec)
assert header_result.ok and header_rendered.exists()
header_result_doc = Document(header_rendered)
assert "Иванова Ирина Ивановна" in header_result_doc.sections[0].header.paragraphs[0].text
assert "Шапка и подвал" in header_result_doc.sections[0].footer.paragraphs[0].text

# Unknown selected document ids and invalid output_dir must be explicit failures/skips.
unknown_render = render_documents_from_pack(pack=profile_owned_pack, case=case, document_ids=["no_such_doc"], output_dir=OUT / "unknown_render", base_dir=OUT / "profile_owned")
assert unknown_render.skipped_documents and "no_such_doc" in unknown_render.skipped_documents[0]
output_file = OUT / "not_a_directory.txt"
output_file.write_text("x", encoding="utf-8")
try:
    render_documents_from_pack(pack=profile_owned_pack, case=case, document_ids=[owned_spec.id], output_dir=output_file, base_dir=OUT / "profile_owned")
except ValueError as exc:
    assert "Папка результата" in str(exc)
else:
    raise AssertionError("render_documents_from_pack must reject output_dir pointing to a file")

# Medpack import must reject traversal/backslash paths before extraction.
malicious_pack = OUT / "malicious.medpack.zip"
import zipfile, json
with zipfile.ZipFile(malicious_pack, "w") as zf:
    zf.writestr("pack.json", json.dumps(profile_owned_pack.to_dict(), ensure_ascii=False))
    zf.writestr("templates\\evil.docx", "bad")
try:
    import_document_pack_zip(malicious_pack, OUT / "malicious_import")
except ValueError as exc:
    assert "Небезопасный путь" in str(exc)
else:
    raise AssertionError("medpack import must reject backslash paths")


# --- Universal profile builder expansion smoke ---
assert any(preset.id == "surgery_base" for preset in specialty_presets())
surgery_preset = get_specialty_preset("surgery")
assert surgery_preset.specialty == "surgery"
surgery_pack = create_pack_from_preset("surgery_base", pack_id="custom.surgery_test", name="Хирургия тест")
assert surgery_pack.document_by_id("operation_protocol") is not None
assert "procedure.anesthesia" in surgery_pack.required_field_ids()
assert "status.neurological" in default_field_registry().ids()

# Scanner must read header/footer in source documents too, not only template docs.
header_source = OUT / "source_with_header_patient.docx"
header_src_doc = Document()
header_src_doc.sections[0].header.paragraphs[0].text = "Ф.И.О.: Петров Пётр Петрович"
header_src_doc.add_paragraph("История болезни № 456")
header_src_doc.add_paragraph("Дата поступления: 11.06.2026")
header_src_doc.add_paragraph("Диагноз: K35 Острый аппендицит")
header_src_doc.save(header_source)
header_scan = scan_docx(header_source, registry=registry)
assert header_scan.best_matches()["patient.fio"].value == "Петров Пётр Петрович"
combined_sources = scan_many_docx([nav, header_source], registry=registry)
assert combined_sources.best_matches()["case.number"].value in {"123", "456"}

operation_template = OUT / "operation_protocol_template.docx"
op_doc = Document()
op_doc.add_paragraph("Пациент: {{patient.fio}}")
op_doc.add_paragraph("Операция: {{procedure.name}}")
op_doc.add_paragraph("Дата операции: {{procedure.date}}")
op_doc.add_paragraph("Анестезия: {{procedure.anesthesia}}")
op_doc.save(operation_template)
surgery_case = PatientCase()
surgery_case.set("patient.fio", "Петров Пётр Петрович")
surgery_case.set("case.number", "456")
surgery_case.set("admission.date", "11.06.2026")
surgery_case.set("diagnosis.main", "K35 Острый аппендицит")
surgery_case.set("procedure.name", "Аппендэктомия")
surgery_case.set("procedure.date", "12.06.2026")
surgery_case.set("procedure.anesthesia", "эндотрахеальный наркоз")
ingestion = ingest_templates_into_pack(surgery_pack, [operation_template, operation_template], OUT / "surgery_profile")
assert ingestion.added_document_ids and ingestion.warnings, ingestion.human_report()
surgery_ready = analyze_pack_readiness(surgery_pack, surgery_case, base_dir=OUT / "surgery_profile")
assert "operation_protocol_template" in " ".join(surgery_ready.ready_document_ids), surgery_ready.human_report()
checklist = profile_setup_checklist(surgery_pack, base_dir=OUT / "surgery_profile")
assert "3–5 примерах" in checklist and "placeholders" in checklist
full_build = build_profile_from_sources_and_templates(
    source_paths=[nav, header_source],
    template_paths=[operation_template],
    profile_dir=OUT / "builder_profile",
    preset_id="surgery_base",
    pack_id="custom.builder_surgery",
    name="Хирургия builder",
)
assert full_build.saved_pack_path and Path(full_build.saved_pack_path).exists()
assert full_build.report.detected_fields
assert full_build.ingestion.added_document_ids
assert "Мастер профиля" in full_build.human_report()

# --- Regulatory soft-advisory knowledge smoke ---
assert REGULATORY_SOFT_ADVISORY_LOCK_VERSION.startswith("v")
assert ADVISORY_IS_NEVER_BLOCKING is True
assert DECLINE_LABEL == "Нет, не буду, делай как есть"
assert_soft_advisory_lock()
roles = default_document_role_registry().roles()
assert any(role.id == "discharge_epicrisis" for role in roles)
assert any(role.id == "operation_protocol" for role in roles)
sections = default_section_registry().detect_sections("Жалобы. Анамнез заболевания. Диагноз. Рекомендации.")
assert "complaints" in sections and "diagnosis" in sections
assert default_specialty_overlay_registry().detect("Протокол операции. Анестезия. Хирург.").id == "surgery"
assert_caucasus_alias_lock()
caucasus_report = caucasus_context_report()
assert "Армения" in caucasus_report and "Грузия" in caucasus_report and "Азербайджан" in caucasus_report
assert "patient.fio" in registry and "Պացիենտ" in registry.aliases_for("patient.fio")
assert "case_admin" in default_section_registry().detect_sections("Հիվանդության պատմություն. სამედიცინო ბარათი. Xəstəlik tarixi.")
assert classify_document_text("Դուրսգրման էպիկրիզ. Ախտորոշում. Խորհուրդներ.").role_id == "discharge_epicrisis"
assert classify_document_text("გაწერის ეპიკრიზი. დიაგნოზი. რეკომენდაციები.").role_id == "discharge_epicrisis"
assert classify_document_text("Xəstəlik tarixi. Diaqnoz. Müalicə.").role_id in {"inpatient_record", "primary_exam"}
classification_text = classify_document_text("Выписной эпикриз. Дата поступления. Дата выписки. Проведённое лечение. Рекомендации.")
assert classification_text.role_id == "discharge_epicrisis", classification_text.human_report()

soft_discharge_template = OUT / "soft_discharge_template.docx"
soft_doc = Document()
soft_doc.add_paragraph("Выписной эпикриз")
soft_doc.add_paragraph("Пациент: {{patient.fio}}")
soft_doc.add_paragraph("История болезни: {{case.number}}")
soft_doc.add_paragraph("Дата поступления: {{admission.date}}")
soft_doc.add_paragraph("Дата выписки: {{discharge.date}}")
soft_doc.add_paragraph("Диагноз: {{diagnosis.main}}")
soft_doc.save(soft_discharge_template)
soft_classification = classify_docx(soft_discharge_template)
assert soft_classification.role_id == "discharge_epicrisis", soft_classification.human_report()
soft_advice = advise_template(soft_discharge_template, registry=registry)
assert soft_advice.has_suggestions, soft_advice.human_report()
assert "recommendations" in soft_advice.suggested_field_ids() or "treatment.summary" in soft_advice.suggested_field_ids()
assert soft_advice.should_block_generation is False
assert "Нет, не буду, делай как есть" in soft_advice.soft_prompt_text()
assert any("{{recommendations}}" in block or "{{treatment.summary}}" in block for block in soft_advice.completion_blocks())
assert COMPLETION_POPUP_LOCK_VERSION.startswith("v")
assert COMPLETION_POPUP_TITLE == "Дополнить документ"
assert_completion_popup_lock()
completion_inputs = completion_inputs_from_advice(soft_advice, existing_case=soft_case if "soft_case" in globals() else PatientCase())
assert completion_inputs and all(item.placeholder.startswith("{{") for item in completion_inputs)
completion_raw = {completion_inputs[0].field_id: "Тестовое мягкое дополнение", "unknown.field": "ignore"}
completion_values = completion_values_from_raw(completion_inputs, completion_raw)
assert list(completion_values.values()) == ["Тестовое мягкое дополнение"]
completed_case = apply_completion_values(PatientCase(), completion_values)
assert completed_case.get(completion_inputs[0].field_id) == "Тестовое мягкое дополнение"
soft_report_path = save_advice_report(soft_advice, OUT / "soft_advice_report.txt")
assert soft_report_path.exists() and "Возможно, здесь стоит указать" in soft_report_path.read_text(encoding="utf-8")

# Declining soft advice must not block generation. The template requires only its own placeholders.
soft_spec = infer_document_spec_from_template(soft_discharge_template, button_label="Мягкий выписной", registry=registry)
soft_case = PatientCase()
for field_id in ["patient.fio", "case.number", "admission.date", "discharge.date", "diagnosis.main"]:
    soft_case.set(field_id, case.get(field_id) or "тест")
soft_rendered = OUT / "soft_decline_rendered.docx"
soft_render_result = render_template_to_docx(template_path=soft_discharge_template, output_path=soft_rendered, case=soft_case, document=soft_spec)
assert soft_render_result.ok and soft_rendered.exists(), soft_advice.human_report()

operation_missing_template = OUT / "operation_missing_template.docx"
op_missing_doc = Document()
op_missing_doc.add_paragraph("Протокол операции")
op_missing_doc.add_paragraph("Пациент: {{patient.fio}}")
op_missing_doc.add_paragraph("Операция: {{procedure.name}}")
op_missing_doc.save(operation_missing_template)
op_advice = advise_template(operation_missing_template, registry=registry, explicit_specialty="surgery")
assert op_advice.role_id == "operation_protocol", op_advice.human_report()
assert "procedure.anesthesia" in op_advice.suggested_field_ids() or "procedure.description" in op_advice.suggested_field_ids()
loaded_advice = advise_document(nav, registry=registry)
assert loaded_advice.should_block_generation is False

# --- Medical documents smoke with EPI and manual UI-like реквизиты ---
service = MedicalDocumentService()
_adapter_data = PatientData(fio="Тестов Тест", case_number="999", admission_date="01.02.2026", diagnosis="A00 тест", treatment_plan="лечение")
_adapter_case = patient_data_to_case(_adapter_data, source_document="adapter")
assert _adapter_case.get("patient.fio") == "Тестов Тест" and _adapter_case.get("case.number") == "999"
assert merge_case_values(_adapter_case, {"recommendations": "наблюдение"}).get("recommendations") == "наблюдение"

# --- Two-digit title date must normalize to the 1900/2000 century once, not twice ---
assert _first_valid_full_date("12.01.26") == "12.01.2026"
assert _first_valid_full_date("12.01.76") == "12.01.1976"
assert _first_valid_full_date("10052026 Первичный осмотр") == "10.05.2026"
assert _first_valid_full_date("100526 Первичный осмотр") == "10.05.2026"
assert _first_valid_full_date("1126 Первичный осмотр") == "01.01.2026"

# --- Admission date regression: title date must win over birth date ---
title_date_doc = OUT / "Дата_заголовок_против_рождения.docx"
td = Document()
td.add_paragraph("12.01.2026 Первичный осмотр")
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович, Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: K35.8 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Профильный статус: тест")
td.save(title_date_doc)
parsed_title_date = service.parse_primary_document(title_date_doc)
assert parsed_title_date.admission_date == "12.01.2026", parsed_title_date.admission_date
assert parsed_title_date.birth in {"04.01.2000", "2000", ""}

table_title_date_doc = OUT / "Дата_заголовок_таблица.docx"
td = Document()
table = td.add_table(rows=1, cols=2)
table.cell(0, 0).text = "13.02.2026"
table.cell(0, 1).text = "Первичный осмотр"
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович")
td.add_paragraph("Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: K35.8 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Профильный статус: тест")
td.save(table_title_date_doc)
parsed_table_title_date = service.parse_primary_document(table_title_date_doc)
assert parsed_table_title_date.admission_date == "13.02.2026", parsed_table_title_date.admission_date

compact_title_date_doc = OUT / "Дата_заголовок_без_точек.docx"
td = Document()
td.add_paragraph("1126 Первичный осмотр")
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович")
td.add_paragraph("Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: K35.8 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Профильный статус: тест")
td.save(compact_title_date_doc)
parsed_compact_title_date = service.parse_primary_document(compact_title_date_doc)
assert parsed_compact_title_date.admission_date == "01.01.2026", parsed_compact_title_date.admission_date


# --- UI source regression: sick-leave selector defaults to "нет" and popup has no duplicate yes/no fields ---
import main as _main_module
assert hasattr(_main_module.CombinedMedicalDiaryApp, "_prompt_assigned_treatment_if_needed")

def _project_python_source() -> str:
    return "\n".join(
        path.read_text(encoding="utf-8")
        for path in sorted(ROOT.glob("*.py"))
        if not path.name.startswith(("smoke_test", "smoke_combined_"))
    )

main_source = _project_python_source()
assert 'expert_sick_leave_needed_var = tk.StringVar(value="нет")' in main_source
assert '("Место работы", self.vk_mse_work_org_var.get().strip() or shared_org)' in main_source
assert '("Должность", self.vk_mse_position_var.get().strip() or shared_position)' in main_source
assert '("Место работы", self.sick_leave_vk_work_org_var.get().strip() or shared_org)' in main_source
assert '("Должность", self.sick_leave_vk_position_var.get().strip() or shared_position)' in main_source
assert '("Место работы, должность", self.vk_mse_work_org_var.get().strip())' not in main_source
assert '_apply_primary_work_defaults(data)' in main_source
assert 'self._primary_work_org_default' in main_source
assert 'self._work_details_manually_edited' in main_source
assert '"городской военный комиссариат"' in main_source and '"областной военный комиссариат"' in main_source
assert 'kind == "discharge"' in main_source
assert 'def _ensure_discharge_date' in main_source
assert 'title="Дата выписки"' in main_source
assert 'self.output_vars[DIARY_KIND] = tk.BooleanVar(value=False)' in main_source
assert 'admission_doctor_referral' in main_source
assert 'Перетащите сюда первичный осмотр/направление на госпитализацию' in main_source
assert 'text="Нужен больничный лист?"' in main_source
assert 'command=self._on_expert_sick_leave_fill' in main_source
assert 'text="Да"' in main_source
assert 'text="Нет"' in main_source
assert '("Работает? да/нет"' not in main_source
assert '("Нужен больничный лист? да/нет"' not in main_source
assert '("С какого числа больничный"' in main_source
assert '("Где работает / организация"' in main_source
assert '("Должность"' in main_source
assert "Ничего не запоминаем между разными popup-окнами" in main_source
assert "self._last_committee_date,\n            self.vk_date_var.get().strip()" not in main_source
assert "self.vk_date_var.get().strip(),\n            self.sick_leave_vk_commission_date_var.get().strip()" not in main_source
assert 'single_line=self._compact_ui' in main_source
assert 'suffix not in {".docx", ".docm"}' in main_source
assert 'txt_low.strip().startswith("эпи")' in main_source

release_zip_source = (ROOT / "make_release_zip.py").read_text(encoding="utf-8")
for snippet in ['".spec"', '".DS_Store"', '"Thumbs.db"', '".vscode"', '".idea"']:
    assert snippet in release_zip_source, snippet



# --- Multilingual UI/document-language/orthography smoke ---
assert_language_catalog_lock()
assert_language_detection_lock()
assert_i18n_strings_lock()
assert_orthography_medical_safe_lock()
assert_language_preferences_lock()
assert language_id_from_choice(language_profile("az").choice_label()) == "az"
assert tr("button.language", "en") == "Language"
assert detect_text_language("Հիվանդության պատմություն. Ախտորոշում").language_id == "hy"
assert detect_text_language("სამედიცინო ბარათი. დიაგნოზი").language_id == "ka"
assert detect_text_language("Xəstəlik tarixi. Diaqnoz və müalicə").language_id == "az"
ru_fixed = correct_medical_text("Рекоммендации: наблюдение. Диагноз K35.8 от 10.06.2026 {{patient.fio}}", language_id="ru").corrected
assert "Рекомендации" in ru_fixed
assert "K35.8" in ru_fixed and "10.06.2026" in ru_fixed and "{{patient.fio}}" in ru_fixed
prefs = LanguagePreferences.from_settings({"ui_language": "ka", "document_language": "auto", "output_language": "same_as_source", "spellcheck_enabled": True})
assert prefs.ui_language == "ka" and prefs.document_language == "auto" and prefs.output_language == "same_as_source"


# --- Persistent national document buttons smoke ---
assert_personal_document_button_lock()
assert localized_role_label("operation_protocol", "ru") == "Протокол операции"
assert localized_role_label("operation_protocol", "az") == "Əməliyyat protokolu"
assert localized_role_label("operation_protocol", "hy") == "Վիրահատության արձանագրություն"
assert localized_role_label("operation_protocol", "ka") == "ოპერაციის პროტოკოლი"
choices_az = regular_document_role_choices("az")
assert any("Əməliyyat protokolu [operation_protocol]" == item for item in choices_az)
assert role_id_from_choice("Əməliyyat protokolu [operation_protocol]") == "operation_protocol"
button_template = OUT / "button_operation_template.docx"
button_doc = Document()
button_doc.add_paragraph("Протокол операции")
button_doc.add_paragraph("Пациент: {{patient.fio}}")
button_doc.add_paragraph("Операция: {{procedure.name}}")
button_doc.add_paragraph("Ход операции: {{procedure.description}}")
button_doc.save(button_template)
button_suggestion = suggest_button_label_for_template(button_template, preferred_language="ru", explicit_role_id="operation_protocol")
assert button_suggestion.label == "Протокол операции"
assert button_suggestion.document_id.startswith("operation_protocol_")
button_pack = DocumentPack(pack_id="button.pack", name="Button pack", documents=())
button_spec, button_copy = attach_template_to_pack(
    button_pack,
    button_template,
    OUT / "button_profile",
    button_label=button_suggestion.label,
    document_id=button_suggestion.document_id,
    registry=registry,
    role_id=button_suggestion.role_id,
    button_language=button_suggestion.language_id,
    source_language=button_suggestion.source_language,
    button_label_source=button_suggestion.source,
)
assert button_spec.button_label == "Протокол операции"
assert button_spec.role_id == "operation_protocol"
assert button_spec.button_language == "ru"
assert button_spec.button_label_source in {"role_i18n", "manual", "template_top_title"}
button_profile_path = OUT / "button_profile" / "button_pack.json"
save_document_pack(button_pack, button_profile_path)
button_loaded = load_document_pack(button_profile_path)
assert button_loaded.document_by_id(button_spec.id).button_label == "Протокол операции"
assert button_loaded.document_by_id(button_spec.id).role_id == "operation_protocol"
from universal_main_documents import custom_documents_for_main_ui, PROFILE_BUTTON_LABELS_ARE_PERSISTED
assert PROFILE_BUTTON_LABELS_ARE_PERSISTED is True
button_tiles = custom_documents_for_main_ui(button_loaded, base_dir=button_profile_path.parent)
assert button_tiles and button_tiles[0].label == "Протокол операции" and button_tiles[0].role_id == "operation_protocol"


# The quick doctor setup must read names from the upper part of each Word sheet
# and create button labels from what the doctor actually sees, not from filenames.
quick_primary = OUT / "quick_primary_template.docx"
quick_doc = Document()
quick_doc.add_paragraph("ГБУЗ Городская больница № 1")
quick_doc.add_paragraph("Первичный осмотр")
quick_doc.add_paragraph("Пациент: {{patient.fio}}")
quick_doc.save(quick_primary)
quick_surgery = OUT / "quick_surgery_template.docx"
quick_doc = Document()
quick_doc.add_paragraph("Отделение хирургии")
quick_doc.add_paragraph("ОСМОТР ХИРУРГА")
quick_doc.add_paragraph("История болезни № {{case.number}}")
quick_doc.save(quick_surgery)
quick_operation = OUT / "quick_operation_template.docx"
quick_doc = Document()
quick_doc.add_paragraph("Протокол операции")
quick_doc.add_paragraph("Операция: {{procedure.name}}")
quick_doc.save(quick_operation)
quick_table_title = OUT / "quick_table_title_template.docx"
quick_doc = Document()
table = quick_doc.add_table(rows=1, cols=1)
table.cell(0, 0).text = "Выписной эпикриз"
quick_doc.add_paragraph("Дата выписки: {{discharge.date}}")
quick_doc.add_paragraph("Проведенное лечение: {{treatment.summary}}")
quick_doc.save(quick_table_title)
assert recognize_document_title_from_template(quick_primary) == "Первичный осмотр"
assert recognize_document_title_from_template(quick_surgery) == "Осмотр хирурга"
assert recognize_document_title_from_template(quick_table_title) == "Выписной эпикриз"
recognized_buttons = recognize_template_buttons([quick_primary, quick_surgery, quick_operation, quick_table_title], preferred_language="ru", ui_language="ru", specialty="surgery")
assert [item.label for item in recognized_buttons] == ["Первичный осмотр", "Осмотр хирурга", "Протокол операции", "Выписной эпикриз"]
assert all(item.source == "template_top_title" for item in recognized_buttons), [item.to_dict() for item in recognized_buttons]
quick_pack = DocumentPack(pack_id="quick.pack", name="Quick pack", documents=())
for item in recognized_buttons:
    spec, _copy = attach_template_to_pack(
        quick_pack,
        item.path,
        OUT / "quick_profile",
        button_label=item.label,
        document_id=item.document_id,
        registry=registry,
        role_id="" if item.role_id == "unknown" else item.role_id,
        button_language="ru",
        source_language="auto",
        button_label_source=item.source,
    )
    assert spec.button_label == item.label
    assert spec.button_label_source == "template_top_title"
quick_tiles = custom_documents_for_main_ui(quick_pack, base_dir=OUT / "quick_profile")
assert [tile.label for tile in quick_tiles] == ["Первичный осмотр", "Осмотр хирурга", "Протокол операции", "Выписной эпикриз"]


# Custom profile buttons must trigger the same required popups as legacy block-03
# buttons.  A doctor-owned diary asks for discharge date; a doctor-owned
# discharge epicrisis opens the discharge merged popup.
from actions_creation_orchestrator import ActionsCreationOrchestratorMixin


class _FakePopupVar:
    def __init__(self, value=""):
        self.value = value

    def get(self):
        return self.value


class _FakePopupApp(ActionsCreationOrchestratorMixin):
    def __init__(self, pack):
        self.pack = pack
        self.calls = []
        self.commission_date_var = _FakePopupVar("")
        self.commission_number_var = _FakePopupVar("")
        self.rvk_act_number_var = _FakePopupVar("")
        self.rvk_military_commissariat_var = _FakePopupVar("")
        self.sick_leave_vk_date_var = _FakePopupVar("")
        self.sick_leave_vk_protocol_number_var = _FakePopupVar("")
        self.sick_leave_vk_protocol_date_var = _FakePopupVar("")
        self.sick_leave_vk_commission_date_var = _FakePopupVar("")
        self.sick_leave_vk_work_org_var = _FakePopupVar("")
        self.sick_leave_vk_position_var = _FakePopupVar("")
        self.vk_date_var = _FakePopupVar("")
        self.vk_protocol_number_var = _FakePopupVar("")
        self.vk_protocol_date_var = _FakePopupVar("")
        self.vk_mse_work_org_var = _FakePopupVar("")

    def _load_or_create_universal_pack(self):
        return self.pack

    def _prompt_common_output_requirements(self, **kwargs):
        self.calls.append(("common", kwargs))
        return True

    def _prompt_discharge_output_requirements(self, **kwargs):
        self.calls.append(("discharge", kwargs))
        return True

    def _prompt_assigned_treatment_if_needed(self, **kwargs):
        self.calls.append(("treatment", kwargs))
        return True

    def _prompt_rvk_details(self):
        self.calls.append(("rvk", {}))
        return True

    def _prompt_commission_details(self):
        self.calls.append(("commission", {}))
        return True

    def _prompt_vk_mse_details(self):
        self.calls.append(("vk_mse", {}))
        return True

    def _prompt_sick_leave_vk_details(self):
        self.calls.append(("sick_leave_vk", {}))
        return True

    def _expert_anamnesis_needed_for_selection(self, selected_medical):
        return False

    def _manual_treatment_missing(self):
        return True

    def _hospitalization_details_missing(self):
        return True

    def _case_number_missing(self):
        return True

    def _should_prompt_discharge_date(self):
        return True

    def _vk_mse_details_complete(self):
        return False

    def _sick_leave_vk_details_complete(self):
        return False


diary_popup_pack = DocumentPack(
    pack_id="popup.diary",
    name="Popup diary",
    documents=(DocumentTemplateSpec("custom_diary", "Дневники", "templates/d.docx", category="diaries", role_id="daily_diary"),),
)
diary_popup_app = _FakePopupApp(diary_popup_pack)
assert diary_popup_app._collect_creation_requirements([], False, ["custom_diary"]) is True
assert diary_popup_app.calls and diary_popup_app.calls[0][0] == "common"
assert diary_popup_app.calls[0][1]["include_discharge_date"] is True

discharge_popup_pack = DocumentPack(
    pack_id="popup.discharge",
    name="Popup discharge",
    documents=(DocumentTemplateSpec("custom_discharge", "Выписной эпикриз", "templates/e.docx", role_id="discharge_epicrisis"),),
)
discharge_popup_app = _FakePopupApp(discharge_popup_pack)
assert discharge_popup_app._collect_creation_requirements([], False, ["custom_discharge"]) is True
assert any(call[0] == "discharge" for call in discharge_popup_app.calls), discharge_popup_app.calls

# --- Required diagnosis popup must reject unresolved free text while accepting safe ICD-10 matches ---
from icd10_f_search import normalize_required_diagnosis_with_icd10, diagnosis_has_icd10_code
assert normalize_required_diagnosis_with_icd10("тестовый диагноз без шифра") == ""
assert normalize_required_diagnosis_with_icd10("Острый аппендицит").startswith("K35")
assert diagnosis_has_icd10_code("К35.8 тестовый диагноз") is True

# --- Desktop intake must not scan arbitrary saved folders and must isolate popup selections ---
from desktop_intake import DESKTOP_INTAKE_FOLDER_NAME, DESKTOP_INTAKE_SETUP_PROMPT_VERSION, should_prompt_intake_setup
wrong_watched_folder = OUT / "desktop_wrong_saved_folder"
wrong_watched_folder.mkdir(parents=True, exist_ok=True)
assert should_prompt_intake_setup({
    "asked": True,
    "enabled": True,
    "folder": str(wrong_watched_folder),
    "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
}) is True
right_watched_folder = OUT / DESKTOP_INTAKE_FOLDER_NAME
right_watched_folder.mkdir(parents=True, exist_ok=True)
assert should_prompt_intake_setup({
    "asked": True,
    "enabled": True,
    "folder": str(right_watched_folder),
    "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
}) is False

from desktop_intake_mixin import DesktopIntakeMixin
from app_config import DIARY_KIND
from universal_main_documents import custom_kind

class _Var:
    def __init__(self, value=False):
        self.value = value
    def get(self):
        return self.value
    def set(self, value):
        self.value = bool(value)

class _FakeDesktopSelectionApp(DesktopIntakeMixin):
    def __init__(self):
        self.output_vars = {"primary": _Var(True), DIARY_KIND: _Var(True), custom_kind("old"): _Var(True), custom_kind("new"): _Var(False)}
        self.custom_output_vars = {custom_kind("old"): self.output_vars[custom_kind("old")], custom_kind("new"): self.output_vars[custom_kind("new")]}
        self.updated = False
    def _update_selected_outputs_status(self):
        self.updated = True

selection_app = _FakeDesktopSelectionApp()
selection_app._apply_desktop_intake_selected_kinds({custom_kind("new")})
assert selection_app.output_vars["primary"].get() is False
assert selection_app.output_vars[DIARY_KIND].get() is False
assert selection_app.output_vars[custom_kind("old")].get() is False
assert selection_app.output_vars[custom_kind("new")].get() is True
assert selection_app.updated is True

from medical_models import PatientData, build_patient_case_review, augment_patient_case_review_with_custom_flags
review_no_code = build_patient_case_review(
    PatientData(fio="Иванов Иван", output_fio="Иванов Иван", admission_date="01.01.2026", diagnosis="тестовый диагноз"),
    selected_medical=("primary",),
)
assert "diagnosis" in [field.key for field in review_no_code.critical_missing()], review_no_code.as_text()
review_diary_text_only = build_patient_case_review(
    PatientData(fio="Иванов Иван", output_fio="Иванов Иван", admission_date="01.01.2026", diagnosis="тестовый диагноз"),
    selected_diaries=True,
)
assert "diagnosis" not in [field.key for field in review_diary_text_only.critical_missing()], review_diary_text_only.as_text()
review_custom_no_code = augment_patient_case_review_with_custom_flags(
    build_patient_case_review(PatientData(fio="Иванов Иван", output_fio="Иванов Иван", admission_date="01.01.2026", diagnosis="тестовый диагноз"), selected_custom=("custom",)),
    {"requires_diagnosis": True},
    diagnosis="тестовый диагноз",
)
assert "diagnosis" in [field.key for field in review_custom_no_code.critical_missing()], review_custom_no_code.as_text()

# Required custom-document completion values must block generation instead of creating a half-empty DOCX.
from actions_universal_flow import ActionsUniversalFlowMixin
from regulatory_completion_blocks import CompletionInput
from universal_fields import PatientCase

class _FakeRequiredCompletionApp(ActionsUniversalFlowMixin):
    root = None
    def __init__(self, raw_values):
        self.raw_values = raw_values
    def _missing_custom_completion_inputs(self, current_pack, case, selected_custom_ids):
        return (CompletionInput("diagnosis.main", "Диагноз", reason="Не заполнено обязательное поле: Диагноз"),)
    def _prompt_regulatory_completion_values(self, inputs, *, parent):
        return self.raw_values
    def _log(self, value):
        pass

try:
    _FakeRequiredCompletionApp({})._offer_custom_completion_values(object(), PatientCase(), ["doc"])
    raise AssertionError("required completion cancellation must stop custom generation")
except ValueError as exc:
    assert "Диагноз" in str(exc)
completed_case = _FakeRequiredCompletionApp({"diagnosis.main": "K35 Острый аппендицит"})._offer_custom_completion_values(object(), PatientCase(), ["doc"])
assert completed_case.get("diagnosis.main") == "K35 Острый аппендицит"
