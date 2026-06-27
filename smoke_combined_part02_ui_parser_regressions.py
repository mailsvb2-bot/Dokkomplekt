from pathlib import Path as _SmokePath
if "_SMOKE_COMBINED_RUNNER_ACTIVE" not in globals() and "OUT" not in globals():
    from smoke_combined_runner import build_namespace_before as _build_smoke_namespace_before
    globals().update(_build_smoke_namespace_before(_SmokePath(__file__).name))

from app_config import DIARY_KIND
# --- UI sick-leave popup regression ---
class _FakeVar:
    def __init__(self, value=""):
        self.value = value

    def get(self):
        return self.value

    def set(self, value):
        self.value = value


ui_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
ui_logic.output_vars = {"discharge": _FakeVar(True)}
ui_logic.expert_work_status_var = _FakeVar("")
ui_logic.expert_work_org_var = _FakeVar("ООО Тест")
ui_logic.expert_position_var = _FakeVar("врач")
ui_logic.expert_sick_leave_needed_var = _FakeVar("нет")
ui_logic.expert_sick_leave_from_var = _FakeVar("15.06.2026")
ui_logic.expert_sick_leave_number_var = _FakeVar("")
ui_logic.discharge_date_var = _FakeVar("")
ui_logic._popup_discharge_date_override = ""
ui_logic._manual_discharge_date = False
ui_logic._set_ui_var = lambda var, value: var.set(value)
ui_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(ui_logic, _main_module.CombinedMedicalDiaryApp)
ui_logic._prompt_expert_anamnesis_details = lambda force=False: True
ui_logic._update_expert_sick_leave_display = lambda: None
ui_logic._redraw_selection_controls = lambda: None
number_popup_calls = []
date_popup_calls = []
ui_logic._prompt_discharge_sick_leave_number = lambda: number_popup_calls.append("number") or True
def _fake_discharge_date_popup():
    date_popup_calls.append("date")
    ui_logic.discharge_date_var.set("11.06.2026")
    ui_logic._popup_discharge_date_override = "11.06.2026"
    ui_logic._manual_discharge_date = True
    return True
ui_logic._prompt_discharge_date = _fake_discharge_date_popup
def _fake_discharge_output_requirements():
    date_popup_calls.append("date")
    number_popup_calls.append("number")
    ui_logic.discharge_date_var.set("11.06.2026")
    ui_logic._popup_discharge_date_override = "11.06.2026"
    ui_logic.expert_sick_leave_number_var.set("123456")
    ui_logic._manual_discharge_date = True
    return True
ui_logic._prompt_discharge_output_requirements = _fake_discharge_output_requirements
ui_logic._on_expert_sick_leave_fill()
assert ui_logic.expert_sick_leave_needed_var.get() == "да"
assert number_popup_calls == [], "Number popup must not open from the sick-leave Yes button"
assert ui_logic._ensure_discharge_sick_leave_number(prompt_if_needed=True) is True
assert number_popup_calls == ["number"], "Number popup must open only from discharge flow"
number_popup_calls.clear()
ui_logic.output_vars["discharge"].set(False)
assert ui_logic._ensure_discharge_sick_leave_number(prompt_if_needed=True) is True
assert number_popup_calls == [], "Number popup must not open when discharge is not selected"

# If discharge is already selected (default UI state), clicking the discharge
# tile must complete missing discharge requirements instead of silently
# deselecting the tile: first date, then sick-leave number when needed.
ui_logic.output_vars["discharge"].set(True)
ui_logic.expert_sick_leave_needed_var.set("да")
ui_logic.expert_sick_leave_number_var.set("")
ui_logic.discharge_date_var.set("")
ui_logic._popup_discharge_date_override = ""
number_popup_calls.clear()
date_popup_calls.clear()
ui_logic._update_selected_outputs_status = lambda: None
ui_logic._activate_output_tile = _main_module.CombinedMedicalDiaryApp._activate_output_tile.__get__(ui_logic, _main_module.CombinedMedicalDiaryApp)
ui_logic._activate_output_tile("discharge")
assert date_popup_calls == ["date"], "Discharge tile click must request discharge date when it is missing"
assert number_popup_calls == ["number"], "Discharge tile click must request sick-leave number when both conditions are true"
assert ui_logic.output_vars["discharge"].get() is True, "Discharge must remain selected while completing its required fields"

# Date of discharge is shared: it is required not only for discharge summary,
# but also for diaries and RVK act.
ui_logic.discharge_date_var.set("")
ui_logic._popup_discharge_date_override = ""
ui_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(True), DIARY_KIND: _FakeVar(False)}
assert ui_logic._should_prompt_discharge_date() is True, "RVK tile must request discharge date"
ui_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(True)}
assert ui_logic._should_prompt_discharge_date() is True, "Diaries tile must request discharge date"
ui_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(False)}
assert ui_logic._should_prompt_discharge_date() is False, "Discharge date popup must not open without required outputs"

# Common popup merge regression: for restored/programmatic selections without
# discharge/RVK, treatment/referral details and diary discharge date must be one
# popup, not two consecutive popups.
common_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
common_logic.case_number_var = _FakeVar("")
common_logic.assigned_treatment_var = _FakeVar("")
common_logic.diagnosis_var = _FakeVar("")
common_logic.discharge_date_var = _FakeVar("")
common_logic._popup_discharge_date_override = ""
common_logic._manual_discharge_date = False
common_logic._manual_diagnosis = False
common_logic.data = PatientData()
common_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(True)}
common_logic._set_ui_var = lambda var, value: var.set(value)
common_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(common_logic, _main_module.CombinedMedicalDiaryApp)
common_logic._selected_outputs_require_discharge_date = _main_module.CombinedMedicalDiaryApp._selected_outputs_require_discharge_date.__get__(common_logic, _main_module.CombinedMedicalDiaryApp)
common_logic._discharge_date_missing_or_invalid = _main_module.CombinedMedicalDiaryApp._discharge_date_missing_or_invalid.__get__(common_logic, _main_module.CombinedMedicalDiaryApp)
common_logic._store_discharge_date_value = _main_module.CombinedMedicalDiaryApp._store_discharge_date_value.__get__(common_logic, _main_module.CombinedMedicalDiaryApp)
common_logic._normalize_popup_diagnosis_value = lambda value: value
common_logic._case_number_popup_default = lambda: ""
common_logic._treatment_popup_default = lambda: ""
common_logic._discharge_popup_default = lambda: ""
common_logic._hospitalization_details_missing = lambda: True
common_logic._manual_treatment_missing = lambda: False
common_popup_calls = []
common_logic._prompt_fields = lambda title, rows, width=72: common_popup_calls.append((title, rows)) or ["123", "терапия", "K35.8 тест", "11062026"]
common_logic._prompt_common_output_requirements = _main_module.CombinedMedicalDiaryApp._prompt_common_output_requirements.__get__(common_logic, _main_module.CombinedMedicalDiaryApp)
assert common_logic._prompt_common_output_requirements(include_discharge_date=True) is True
assert len(common_popup_calls) == 1
assert [label for label, _default in common_popup_calls[0][1]] == ["Номер истории болезни", "Лечение", "Диагноз", "Дата выписки"]
assert common_logic.case_number_var.get() == "123"
assert common_logic.assigned_treatment_var.get() == "терапия"
assert common_logic.diagnosis_var.get() == "K35.8 тест"
assert common_logic.discharge_date_var.get() == "11.06.2026"

# Hospitalization referral popup must not request discharge date unless the
# selected outputs actually need it.
referral_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
referral_logic.primary_document_type_var = _FakeVar("hospitalization_referral")
referral_logic.case_number_var = _FakeVar("")
referral_logic.assigned_treatment_var = _FakeVar("")
referral_logic.diagnosis_var = _FakeVar("")
referral_logic.discharge_date_var = _FakeVar("")
referral_logic._popup_discharge_date_override = ""
referral_logic._manual_discharge_date = False
referral_logic._manual_diagnosis = False
referral_logic.data = PatientData()
referral_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(False)}
referral_logic.status_label = type("Status", (), {"config": lambda self, **kwargs: None})()
referral_logic._set_ui_var = lambda var, value: var.set(value)
referral_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(referral_logic, _main_module.CombinedMedicalDiaryApp)
referral_logic._selected_outputs_require_discharge_date = _main_module.CombinedMedicalDiaryApp._selected_outputs_require_discharge_date.__get__(referral_logic, _main_module.CombinedMedicalDiaryApp)
referral_logic._discharge_date_missing_or_invalid = _main_module.CombinedMedicalDiaryApp._discharge_date_missing_or_invalid.__get__(referral_logic, _main_module.CombinedMedicalDiaryApp)
referral_logic._store_discharge_date_value = _main_module.CombinedMedicalDiaryApp._store_discharge_date_value.__get__(referral_logic, _main_module.CombinedMedicalDiaryApp)
referral_logic._prompt_primary_exam_details_if_needed = lambda force=False: True
referral_logic._normalize_popup_diagnosis_value = lambda value: value
referral_logic._case_number_popup_default = lambda: ""
referral_logic._treatment_popup_default = lambda: ""
referral_logic._discharge_popup_default = lambda: ""
referral_popup_calls = []
referral_logic._prompt_fields = lambda title, rows, width=72: referral_popup_calls.append((title, rows)) or ["321", "лечение", "I10 тест"]
referral_logic._prompt_assigned_treatment_if_needed = _main_module.CombinedMedicalDiaryApp._prompt_assigned_treatment_if_needed.__get__(referral_logic, _main_module.CombinedMedicalDiaryApp)
assert referral_logic._prompt_assigned_treatment_if_needed(force=True) is True
assert [label for label, _default in referral_popup_calls[0][1]] == ["Номер истории болезни", "Лечение", "Диагноз"]


# Shared case-number regression: every block-03 medical popup includes the
# same «Номер истории болезни» field; diaries stay excluded.
case_dialog_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
case_dialog_logic.data = PatientData(case_number="77")
case_dialog_logic.case_number_var = _FakeVar("")
case_dialog_logic.commission_date_var = _FakeVar("")
case_dialog_logic.commission_number_var = _FakeVar("")
case_dialog_logic._today_str = lambda: "21.06.2026"
case_dialog_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(case_dialog_logic, _main_module.CombinedMedicalDiaryApp)
case_dialog_logic._case_number_popup_default = _main_module.CombinedMedicalDiaryApp._case_number_popup_default.__get__(case_dialog_logic, _main_module.CombinedMedicalDiaryApp)
case_dialog_logic._store_case_number_value = _main_module.CombinedMedicalDiaryApp._store_case_number_value.__get__(case_dialog_logic, _main_module.CombinedMedicalDiaryApp)
case_dialog_logic._remember_committee_dates = lambda **kwargs: None
commission_rows = []
case_dialog_logic._prompt_fields = lambda title, rows, linked_groups=None, width=28: commission_rows.append((title, rows)) or ["88", "21062026", "5"]
case_dialog_logic._prompt_commission_details = _main_module.CombinedMedicalDiaryApp._prompt_commission_details.__get__(case_dialog_logic, _main_module.CombinedMedicalDiaryApp)
assert case_dialog_logic._prompt_commission_details() is True
assert [label for label, _default in commission_rows[0][1]][0] == "Номер истории болезни"
assert commission_rows[0][1][0][1] == "77"
assert case_dialog_logic.case_number_var.get() == "88"
assert case_dialog_logic.data.case_number == "88"

vk_rows = []
case_dialog_logic.vk_date_var = _FakeVar("")
case_dialog_logic.vk_protocol_number_var = _FakeVar("")
case_dialog_logic.vk_protocol_date_var = _FakeVar("")
case_dialog_logic.vk_mse_work_org_var = _FakeVar("")
case_dialog_logic.vk_mse_position_var = _FakeVar("")
case_dialog_logic._shared_work_defaults = lambda: ("ООО Тест", "инженер")
case_dialog_logic._sync_shared_work_details = lambda org, position: None
case_dialog_logic._prompt_fields = lambda title, rows, width=64, linked_groups=None: vk_rows.append((title, rows, linked_groups)) or ["99", "22062026", "12", "22062026", "ООО Тест", "инженер"]
case_dialog_logic._prompt_vk_mse_details = _main_module.CombinedMedicalDiaryApp._prompt_vk_mse_details.__get__(case_dialog_logic, _main_module.CombinedMedicalDiaryApp)
assert case_dialog_logic._prompt_vk_mse_details() is True
assert [label for label, _default in vk_rows[0][1]][0] == "Номер истории болезни"
assert vk_rows[0][1][0][1] == "88"
assert vk_rows[0][2] == [(1, [3])]
assert case_dialog_logic.case_number_var.get() == "99"
assert case_dialog_logic.data.case_number == "99"

diary_only_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
diary_only_logic.case_number_var = _FakeVar("")
diary_only_logic.assigned_treatment_var = _FakeVar("")
diary_only_logic.diagnosis_var = _FakeVar("")
diary_only_logic.discharge_date_var = _FakeVar("")
diary_only_logic._popup_discharge_date_override = ""
diary_only_logic._manual_discharge_date = False
diary_only_logic.data = PatientData()
diary_only_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(True)}
diary_only_logic._set_ui_var = lambda var, value: var.set(value)
diary_only_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(diary_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_only_logic._selected_outputs_require_discharge_date = _main_module.CombinedMedicalDiaryApp._selected_outputs_require_discharge_date.__get__(diary_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_only_logic._discharge_date_missing_or_invalid = _main_module.CombinedMedicalDiaryApp._discharge_date_missing_or_invalid.__get__(diary_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_only_logic._store_discharge_date_value = _main_module.CombinedMedicalDiaryApp._store_discharge_date_value.__get__(diary_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_only_logic._case_number_missing = lambda: True
diary_only_logic._case_number_popup_default = lambda: ""
diary_only_logic._hospitalization_details_missing = lambda: False
diary_only_logic._manual_treatment_missing = lambda: False
diary_only_logic._discharge_popup_default = lambda: ""
diary_popup_calls = []
diary_only_logic._prompt_fields = lambda title, rows, width=72: diary_popup_calls.append((title, rows)) or ["11062026"]
diary_only_logic._prompt_common_output_requirements = _main_module.CombinedMedicalDiaryApp._prompt_common_output_requirements.__get__(diary_only_logic, _main_module.CombinedMedicalDiaryApp)
assert diary_only_logic._prompt_common_output_requirements(include_discharge_date=True, include_case_number=False, include_medical_details=False) is True
assert [label for label, _default in diary_popup_calls[0][1]] == ["Дата выписки"]
assert diary_only_logic.case_number_var.get() == ""

# Even if the primary document is a hospitalization referral with missing
# treatment/diagnosis, selecting only «Дневники» must ask only discharge date.
diary_referral_only_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
diary_referral_only_logic.case_number_var = _FakeVar("")
diary_referral_only_logic.assigned_treatment_var = _FakeVar("")
diary_referral_only_logic.diagnosis_var = _FakeVar("")
diary_referral_only_logic.discharge_date_var = _FakeVar("")
diary_referral_only_logic._popup_discharge_date_override = ""
diary_referral_only_logic._manual_discharge_date = False
diary_referral_only_logic.data = PatientData()
diary_referral_only_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(True)}
diary_referral_only_logic._set_ui_var = lambda var, value: var.set(value)
diary_referral_only_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(diary_referral_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_referral_only_logic._selected_outputs_require_discharge_date = _main_module.CombinedMedicalDiaryApp._selected_outputs_require_discharge_date.__get__(diary_referral_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_referral_only_logic._discharge_date_missing_or_invalid = _main_module.CombinedMedicalDiaryApp._discharge_date_missing_or_invalid.__get__(diary_referral_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_referral_only_logic._store_discharge_date_value = _main_module.CombinedMedicalDiaryApp._store_discharge_date_value.__get__(diary_referral_only_logic, _main_module.CombinedMedicalDiaryApp)
diary_referral_only_logic._case_number_missing = lambda: True
diary_referral_only_logic._case_number_popup_default = lambda: ""
diary_referral_only_logic._hospitalization_details_missing = lambda: True
diary_referral_only_logic._manual_treatment_missing = lambda: True
diary_referral_only_logic._discharge_popup_default = lambda: ""
diary_referral_only_calls = []
diary_referral_only_logic._prompt_fields = lambda title, rows, width=72: diary_referral_only_calls.append((title, rows)) or ["12062026"]
diary_referral_only_logic._prompt_common_output_requirements = _main_module.CombinedMedicalDiaryApp._prompt_common_output_requirements.__get__(diary_referral_only_logic, _main_module.CombinedMedicalDiaryApp)
assert diary_referral_only_logic._prompt_common_output_requirements(include_discharge_date=True, include_case_number=False, include_medical_details=False) is True
assert [label for label, _default in diary_referral_only_calls[0][1]] == ["Дата выписки"]
assert diary_referral_only_logic.assigned_treatment_var.get() == ""
assert diary_referral_only_logic.diagnosis_var.get() == ""

# --- Deep popup date contract: required dates must be normalized or rejected ---
from tkinter import messagebox as _date_contract_messagebox
_original_date_showwarning = _date_contract_messagebox.showwarning
_date_warnings: list[tuple[str, str]] = []
_date_contract_messagebox.showwarning = lambda title, message, **kwargs: _date_warnings.append((title, message))
try:
    assert case_dialog_logic.commission_date_var.get() == "21.06.2026", case_dialog_logic.commission_date_var.get()
    assert case_dialog_logic.vk_date_var.get() == "22.06.2026", case_dialog_logic.vk_date_var.get()
    assert case_dialog_logic.vk_protocol_date_var.get() == "22.06.2026", case_dialog_logic.vk_protocol_date_var.get()

    bad_date_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
    bad_date_logic.data = PatientData(case_number="77")
    bad_date_logic.case_number_var = _FakeVar("77")
    bad_date_logic.commission_date_var = _FakeVar("")
    bad_date_logic.commission_number_var = _FakeVar("")
    bad_date_logic._today_str = lambda: "21.06.2026"
    bad_date_logic._case_number_popup_default = _main_module.CombinedMedicalDiaryApp._case_number_popup_default.__get__(bad_date_logic, _main_module.CombinedMedicalDiaryApp)
    bad_date_logic._store_case_number_value = _main_module.CombinedMedicalDiaryApp._store_case_number_value.__get__(bad_date_logic, _main_module.CombinedMedicalDiaryApp)
    bad_date_logic._remember_committee_dates = lambda **kwargs: None
    bad_date_logic._prompt_fields = lambda title, rows, linked_groups=None, width=28: ["77", "99.99.2026", "5"]
    assert bad_date_logic._prompt_commission_details() is False
    assert bad_date_logic.commission_date_var.get() == ""
    assert any("Некорректная дата" in title for title, _message in _date_warnings)

    rvk_date_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
    rvk_date_logic.data = PatientData(admission_date="10.06.2026")
    rvk_date_logic.discharge_date_var = _FakeVar("")
    rvk_date_logic.admission_date_var = _FakeVar("10.06.2026")
    rvk_date_logic._popup_discharge_date_override = ""
    rvk_date_logic._manual_discharge_date = False
    rvk_date_logic._set_ui_var = lambda var, value: var.set(value)
    assert rvk_date_logic._store_discharge_date_value("09.06.2026") is False
    assert rvk_date_logic.discharge_date_var.get() == ""
    assert rvk_date_logic._store_discharge_date_value("11062026") is True
    assert rvk_date_logic.discharge_date_var.get() == "11.06.2026"
finally:
    _date_contract_messagebox.showwarning = _original_date_showwarning


# --- Primary selected status layout regression ---
layout_sources_text = Path("layout_sources.py").read_text(encoding="utf-8")
files_mixin_text = Path("files_mixin.py").read_text(encoding="utf-8")
assert 'primary_selected_status_var = tk.StringVar(value=" ")' in layout_sources_text
assert 'primary_drop_hint_label' in layout_sources_text
assert 'drop.grid_propagate(False)' in layout_sources_text
assert 'drop_height = self._px(96 if self._compact_ui else 106, 78)' in layout_sources_text
assert 'self.primary_drop_hint_label.config(text="", fg=FIELD)' in layout_sources_text
assert 'self.primary_drop_hint_label.grid_remove()' not in layout_sources_text
assert 'Path(path).name' in files_mixin_text
assert 'primary_selected_status_label.grid_remove()' not in files_mixin_text
assert 'primary_selected_status_label.grid()' not in files_mixin_text
assert 'def _truncate_label_text' in files_mixin_text
assert ('single_line=self._compact_ui' in Path("dnd_mixin.py").read_text(encoding="utf-8") or '_update_diary_text_label(success=True)' in Path("dnd_mixin.py").read_text(encoding="utf-8"))
from files_mixin import FilesMixin
_long_name = "Очень длинное название первичного документа пациента Иванова Ирина Ивановна 10052026.docx"
assert "…" in FilesMixin._truncate_label_text(_long_name, max_chars=40)

# --- Deep audit hardening regressions ---
from typing import get_type_hints
from diary_table_numbers import should_remove_holiday
from medical_docx_date_patterns import _first_valid_full_date as _title_date_first
assert get_type_hints(should_remove_holiday)["row_date"]
assert _title_date_first("10052026") == "10.05.2026"
assert _title_date_first("1126") == "01.01.2026"
assert "if not query:" in Path("diagnosis_widget.py").read_text(encoding="utf-8")
assert "if not query:" in Path("dialog_fields_popup.py").read_text(encoding="utf-8")
assert "_select_default_printer_sync" in Path("actions_creation_orchestrator.py").read_text(encoding="utf-8")

# --- Small parsing/formatting regressions fixed after audit ---
assert parse_date("12.01.26 г.").strftime("%d.%m.%Y") == "12.01.2026"
assert parse_date("12. 01.2026").strftime("%d.%m.%Y") == "12.01.2026"
assert parse_date("12 .01.26 г.").strftime("%d.%m.%Y") == "12.01.2026"
assert parse_date("10052026").strftime("%d.%m.%Y") == "10.05.2026"
assert parse_date("100526").strftime("%d.%m.%Y") == "10.05.2026"
assert parse_date("1126").strftime("%d.%m.%Y") == "01.01.2026"
assert parse_date("10526").strftime("%d.%m.%Y") == "01.05.2026"
assert parse_date("31126").strftime("%d.%m.%Y") == "31.01.2026"
assert format_date_with_russian_year_suffix("12.01.2026г.") == "12.01.2026 г."
assert format_birth_for_person_line("1980 г.р") == "1980 г.р"
assert format_birth_for_person_line("1980") == "1980 г.р."
assert parse_full_date("11.06.2026 г.").strftime("%d.%m.%Y") == "11.06.2026"
assert parse_full_date("11062026").strftime("%d.%m.%Y") == "11.06.2026"
assert parse_full_date("110626").strftime("%d.%m.%Y") == "11.06.2026"
assert parse_full_date("1126").strftime("%d.%m.%Y") == "01.01.2026"
assert parse_month_year("06.2026 г.") == (6, 2026)
assert format_military_commissariat_area("Канвинский") == "Канавинского района"
assert format_military_commissariat_area("Сормовский\\Московский") == "Сормовского и Московского района"
assert format_military_commissariat_referral("Канвинский") == "По направлению из Канавинского военкомата"
assert format_military_commissariat_referral("Сормовский/Московский") == "По направлению из Сормовского и Московского военкомата"

parser_after_audit = MedicalTextParser()
assert parser_after_audit.parse_text("Не работает").work_org == ""
assert parser_after_audit.parse_text("Работает в организации: не работает").work_org == ""
assert parser_after_audit.parse_text("Место работы: безработный").work_org == ""
parser_work_doctor = parser_after_audit.parse_text("Работает: ООО Ромашка, в должности врач")
assert parser_work_doctor.work_org == "ООО Ромашка", parser_work_doctor.work_org
assert parser_work_doctor.position == "врач", parser_work_doctor.position
parser_position_doctor = parser_after_audit.parse_text("Работает в организации: ООО Ромашка\nДолжность: врач")
assert parser_position_doctor.work_org == "ООО Ромашка", parser_position_doctor.work_org
assert parser_position_doctor.position == "врач", parser_position_doctor.position

# Case-number regression from real UI: a label-only "История болезни №" cell
# must not steal the neighbouring patient name and prefill popup field with ФИО.
from medical_text_utils import sanitize_case_number_candidate
case_spillover = parser_after_audit.parse_text("""
Первичный осмотр
История болезни №
Ф.И.О.
Михайлов Николай Иванович
Дата рождения: 1980
Диагноз: K35.8 тест
""")
assert case_spillover.fio == "Михайлов Николай Иванович", case_spillover.fio
assert case_spillover.case_number == "", case_spillover.case_number
assert sanitize_case_number_candidate("Михайлов Николай Иванович", patient_name="Михайлов Николай Иванович") == ""
assert sanitize_case_number_candidate("Михайлов Николай Иванович, 123", patient_name="Михайлов Николай Иванович") == "123"
assert sanitize_case_number_candidate("Михайлов Николай Иванович 123", patient_name="Михайлов Николай Иванович") == "123"
assert sanitize_case_number_candidate("Михайлов Н.И. №К-900", patient_name="Михайлов Николай Иванович") == "К-900"
assert sanitize_case_number_candidate("№ 123/45", patient_name="Михайлов Николай Иванович") == "123/45"
assert sanitize_case_number_candidate("ФИО: Михайлов Николай Иванович История болезни № 123", patient_name="Михайлов Николай Иванович") == "123"
case_spillover_with_tail = parser_after_audit.parse_text("""
Первичный осмотр
История болезни № Михайлов Николай Иванович 123
Ф.И.О.: Михайлов Николай Иванович
Диагноз: K35.8 тест
""")
assert case_spillover_with_tail.case_number == "123", case_spillover_with_tail.case_number
case_popup_guard = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
case_popup_guard.patient_name_var = _FakeVar("Михайлов Николай Иванович")
case_popup_guard.case_number_var = _FakeVar("Михайлов Николай Иванович")
case_popup_guard.navigation_path_var = _FakeVar("")
case_popup_guard.data = PatientData(case_number="Михайлов Николай Иванович", fio="Михайлов Николай Иванович")
case_popup_guard._case_number_popup_default = _main_module.CombinedMedicalDiaryApp._case_number_popup_default.__get__(case_popup_guard, _main_module.CombinedMedicalDiaryApp)
case_popup_guard._store_case_number_value = _main_module.CombinedMedicalDiaryApp._store_case_number_value.__get__(case_popup_guard, _main_module.CombinedMedicalDiaryApp)
assert case_popup_guard._case_number_popup_default() == ""
assert case_popup_guard.case_number_var.get() == ""
assert case_popup_guard.data.case_number == ""
assert case_popup_guard._store_case_number_value("Михайлов Николай Иванович") is False
assert case_popup_guard._store_case_number_value("К-900") is True
assert case_popup_guard.case_number_var.get() == "К-900"
assert case_popup_guard._store_case_number_value("№ 123/45") is True
assert case_popup_guard.case_number_var.get() == "123/45"

# Имена файлов должны сохраняться с пробелами, без подчеркиваний между словами.
assert _medical_documents_module.safe_filename("Сидоров Иван Михайлович") == "Сидоров Иван Михайлович"
assert _medical_documents_module.safe_filename("Сидоров/Иван:Михайлович") == "Сидоров Иван Михайлович"
assert _medical_documents_module.safe_filename("CON") == "CON_"
assert _medical_documents_module.safe_filename("CON.txt") == "CON.txt_"
assert safe_filename_part("LPT1.docx") == "LPT1.docx_"

from typing import get_type_hints
from medical_formatting import parse_date as _medical_parse_date
assert get_type_hints(_medical_parse_date)["return"]
assert safe_filename_part("NUL") == "NUL_"
assert "Первичный осмотр" in _medical_documents_module.OUTPUT_SUFFIXES["primary"]
assert "Осмотр врача приёмного покоя" in _medical_documents_module.OUTPUT_SUFFIXES["admission_doctor_referral"]
assert _medical_documents_module.TEMPLATE_FILES["admission_doctor_referral"] == "7 Направление врача приёмного покоя.docx"
assert "_" not in _medical_documents_module.OUTPUT_SUFFIXES["discharge"]

# --- Diagnosis parser regression: diagnosis line must not absorb neighboring sections ---
diag_cases = {
    "Диагноз: I10 Эссенциальная гипертензия Жалобы: нет": "I10 Эссенциальная гипертензия",
    "был выставлен диагноз: K35.8 Острый аппендицит План лечения: терапия": "K35.8 Острый аппендицит",
    "На основании данных анамнеза установлен диагноз: F06.6 Органическое эмоционально лабильное расстройство\nЭпидемиологический анамнез: без особенностей": "F06.6 Органическое эмоционально лабильное расстройство",
    "Диагноз: F": "",
}
for raw_diag, expected_diag in diag_cases.items():
    assert sanitize_diagnosis(raw_diag) == expected_diag, (raw_diag, sanitize_diagnosis(raw_diag))

diag_parse = MedicalTextParser().parse_text(
    "Диагноз: I10 Эссенциальная гипертензия Жалобы: нет Анамнез жизни: тест"
)
assert diag_parse.diagnosis == "I10 Эссенциальная гипертензия", diag_parse.diagnosis

# --- Parser styles regression: demographics in columns and in one compact line ---
parser_style_column = MedicalTextParser().parse_text("""
ФИО: Иванов Иван Иванович
возраст:34 года
Проживает : Г. Нижний Новгород, улица Ленина 34-15
Работает: ООО Завод
""")
assert parser_style_column.fio == "Иванов Иван Иванович", parser_style_column.fio
assert parser_style_column.birth == "34 года", parser_style_column.birth
assert parser_style_column.registered == "Г. Нижний Новгород, улица Ленина 34-15", parser_style_column.registered
assert parser_style_column.work_org == "ООО Завод", parser_style_column.work_org

parser_work_phrase = MedicalTextParser().parse_text("""
10.06.2026 Первичный осмотр
ФИО: Сидоров Иван Михайлович
Работает в Рассвет, в должности Уборщик.
Диагноз: K35.8 тест
""")
assert parser_work_phrase.work_org == "Рассвет", parser_work_phrase.work_org
assert parser_work_phrase.position == "Уборщик", parser_work_phrase.position

parser_work_label_combo = MedicalTextParser().parse_text("""
10.06.2026 Первичный осмотр
ФИО: Сидоров Иван Михайлович
Место работы: ООО «Привет», должность: начальник
Диагноз: K35.8 тест
""")
assert parser_work_label_combo.work_org == "ООО «Привет»", parser_work_label_combo.work_org
assert parser_work_label_combo.position == "начальник", parser_work_label_combo.position

parser_style_line = MedicalTextParser().parse_text(
    "Иванов Иван Иванович, 34 года, Г. Нижний Новгород, улица Ленина 34-15, ООО Завод"
)
assert parser_style_line.fio == "Иванов Иван Иванович", parser_style_line.fio
assert parser_style_line.birth == "34 года", parser_style_line.birth
assert parser_style_line.registered == "Г. Нижний Новгород, улица Ленина 34-15", parser_style_line.registered
assert parser_style_line.work_org == "ООО Завод", parser_style_line.work_org

parser_two_digit_birth = MedicalTextParser().parse_text(
    "Ф.И.О.: Иванов Иван Иванович, Дата рождения: 04.01.80, Место жительства: город N, ул. Тестовая, 1"
)
assert parser_two_digit_birth.birth == "04.01.80", parser_two_digit_birth.birth

# --- Compact demographics smoke: ФИО/возраст/адрес can be written in one line ---

# --- Deep audit regression: already-filled popup dates are revalidated, not trusted blindly ---
date_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
date_logic.data = PatientData(admission_date="10.06.2026")
date_logic.discharge_date_var = _FakeVar("09.06.2026")
date_logic._popup_discharge_date_override = ""
date_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(date_logic, _main_module.CombinedMedicalDiaryApp)
date_logic._admission_date_for_validation = _main_module.CombinedMedicalDiaryApp._admission_date_for_validation.__get__(date_logic, _main_module.CombinedMedicalDiaryApp)
date_logic._date_is_not_before_admission = _main_module.CombinedMedicalDiaryApp._date_is_not_before_admission.__get__(date_logic, _main_module.CombinedMedicalDiaryApp)
date_logic._current_discharge_date_value = _main_module.CombinedMedicalDiaryApp._current_discharge_date_value.__get__(date_logic, _main_module.CombinedMedicalDiaryApp)
date_logic._discharge_date_missing_or_invalid = _main_module.CombinedMedicalDiaryApp._discharge_date_missing_or_invalid.__get__(date_logic, _main_module.CombinedMedicalDiaryApp)
assert date_logic._discharge_date_missing_or_invalid() is True

date_logic.discharge_date_var.set("11.06.2026")
assert date_logic._discharge_date_missing_or_invalid() is False

expert_date_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
expert_date_logic.data = PatientData(admission_date="10.06.2026")
expert_date_logic.expert_sick_leave_needed_var = _FakeVar("да")
expert_date_logic.expert_sick_leave_from_var = _FakeVar("09.06.2026")
expert_date_logic.expert_work_org_var = _FakeVar("ООО Тест")
expert_date_logic.expert_position_var = _FakeVar("инженер")
expert_date_logic._normalize_yes_no = _main_module.CombinedMedicalDiaryApp._normalize_yes_no
expert_date_logic._normalize_date_for_ui = _main_module.CombinedMedicalDiaryApp._normalize_date_for_ui.__get__(expert_date_logic, _main_module.CombinedMedicalDiaryApp)
expert_date_logic._admission_date_for_validation = _main_module.CombinedMedicalDiaryApp._admission_date_for_validation.__get__(expert_date_logic, _main_module.CombinedMedicalDiaryApp)
expert_date_logic._date_is_not_before_admission = _main_module.CombinedMedicalDiaryApp._date_is_not_before_admission.__get__(expert_date_logic, _main_module.CombinedMedicalDiaryApp)
expert_date_logic._expert_details_are_complete = _main_module.CombinedMedicalDiaryApp._expert_details_are_complete.__get__(expert_date_logic, _main_module.CombinedMedicalDiaryApp)
assert expert_date_logic._expert_details_are_complete() is False
expert_date_logic.expert_sick_leave_from_var.set("10.06.2026")
assert expert_date_logic._expert_details_are_complete() is True

vk_preflight_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
vk_preflight_logic.data = PatientData(admission_date="10.06.2026")
vk_preflight_logic.vk_date_var = _FakeVar("09.06.2026")
vk_preflight_logic.vk_protocol_number_var = _FakeVar("42")
vk_preflight_logic.vk_protocol_date_var = _FakeVar("10.06.2026")
vk_preflight_logic.vk_mse_work_org_var = _FakeVar("ООО Тест")
vk_preflight_logic._admission_date_for_validation = _main_module.CombinedMedicalDiaryApp._admission_date_for_validation.__get__(vk_preflight_logic, _main_module.CombinedMedicalDiaryApp)
vk_preflight_logic._date_is_not_before_admission = _main_module.CombinedMedicalDiaryApp._date_is_not_before_admission.__get__(vk_preflight_logic, _main_module.CombinedMedicalDiaryApp)
vk_preflight_logic._popup_date_value_is_valid_and_in_episode = _main_module.CombinedMedicalDiaryApp._popup_date_value_is_valid_and_in_episode.__get__(vk_preflight_logic, _main_module.CombinedMedicalDiaryApp)
vk_preflight_logic._vk_mse_details_complete = _main_module.CombinedMedicalDiaryApp._vk_mse_details_complete.__get__(vk_preflight_logic, _main_module.CombinedMedicalDiaryApp)
assert vk_preflight_logic._vk_mse_details_complete() is False
vk_preflight_logic.vk_date_var.set("10.06.2026")
assert vk_preflight_logic._vk_mse_details_complete() is True

# --- v1.4.31 discharge date UI contract regressions ---
# Re-clicking an already selected discharge epicrisis must reopen the discharge
# details editor (with date forced into the popup) instead of silently deselecting
# the document. This is the direct-file workflow, not desktop patient-folder flow.
reclick_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
reclick_logic.output_vars = {"discharge": _FakeVar(True), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(False)}
reclick_logic.discharge_date_var = _FakeVar("11.06.2026")
reclick_logic._popup_discharge_date_override = "11.06.2026"
reclick_logic._update_selected_outputs_status = lambda: None
reclick_logic._redraw_selection_controls = lambda: None
reclick_calls = []
def _reclick_prompt(**kwargs):
    reclick_calls.append(kwargs)
    reclick_logic.discharge_date_var.set("12.06.2026")
    reclick_logic._popup_discharge_date_override = "12.06.2026"
    return True
reclick_logic._prompt_discharge_output_requirements = _reclick_prompt
reclick_logic._activate_output_tile = _main_module.CombinedMedicalDiaryApp._activate_output_tile.__get__(reclick_logic, _main_module.CombinedMedicalDiaryApp)
reclick_logic._activate_output_tile("discharge")
assert reclick_logic.output_vars["discharge"].get() is True
assert reclick_calls == [{"force_discharge_date": True}], reclick_calls
assert reclick_logic.discharge_date_var.get() == "12.06.2026"

# Turning on diaries must show the discharge-date popup even if a previous output
# already filled a date, because the same date controls the final diary table row.
diary_force_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
diary_force_logic.output_vars = {DIARY_KIND: _FakeVar(True)}
diary_force_logic._update_selected_outputs_status = lambda: None
diary_force_logic._redraw_selection_controls = lambda: None
diary_force_calls = []
def _force_ensure(**kwargs):
    diary_force_calls.append(kwargs)
    return True
diary_force_logic._ensure_discharge_date = _force_ensure
diary_force_logic._on_output_toggle = _main_module.CombinedMedicalDiaryApp._on_output_toggle.__get__(diary_force_logic, _main_module.CombinedMedicalDiaryApp)
diary_force_logic._on_output_toggle(DIARY_KIND)
assert diary_force_calls == [{"prompt_if_needed": True, "force_prompt": True}], diary_force_calls
assert diary_force_logic.output_vars[DIARY_KIND].get() is True

# Storing/typing a discharge date refreshes the explicit green UI status line.
status_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
status_logic._compact_ui = False
status_logic.navigation_path_var = _FakeVar(r"C:\\docs\\primary.docx")
status_logic.primary_document_type_var = _FakeVar("primary_exam")
status_logic.discharge_date_var = _FakeVar("11.06.2026")
status_logic.primary_selected_status_var = _FakeVar("")
status_logic._truncate_label_text = lambda text, max_chars: text if len(text) <= max_chars else text[: max_chars - 1] + "…"
status_logic._primary_selected_status_text = _main_module.CombinedMedicalDiaryApp._primary_selected_status_text.__get__(status_logic, _main_module.CombinedMedicalDiaryApp)
status_logic._refresh_primary_selected_status = _main_module.CombinedMedicalDiaryApp._refresh_primary_selected_status.__get__(status_logic, _main_module.CombinedMedicalDiaryApp)
status_logic._refresh_primary_selected_status()
assert "Дата выписки: 11.06.2026" in status_logic.primary_selected_status_var.get(), status_logic.primary_selected_status_var.get()

# --- v1.4.37 primary DOCX parser must read common neutral medical labels ---
# Real non-psychiatric primary exams often use table labels without the final
# dot in «Ф.И.О» and write «Клинический диагноз» instead of bare «Диагноз».
# These fields must populate the main UI after the primary document is loaded.
from docx import Document as _SmokeDocument
neutral_primary = OUT / "neutral_primary_demographics_diagnosis.docx"
doc = _SmokeDocument()
doc.add_paragraph("Первичный осмотр 24.06.2026")
table = doc.add_table(rows=4, cols=2)
table.cell(0, 0).text = "Ф.И.О"
table.cell(0, 1).text = "Иванов Иван Иванович"
table.cell(1, 0).text = "Дата рождения"
table.cell(1, 1).text = "01.01.1980"
table.cell(2, 0).text = "Клинический диагноз"
table.cell(2, 1).text = "I10 Эссенциальная гипертензия"
table.cell(3, 0).text = "Жалобы"
table.cell(3, 1).text = "головная боль"
doc.save(neutral_primary)
from medical_parser import MedicalTextParser as _SmokeMedicalTextParser
neutral_data = _SmokeMedicalTextParser().parse_docx(neutral_primary)
assert neutral_data.fio == "Иванов Иван Иванович", neutral_data.fio
assert neutral_data.diagnosis == "I10 Эссенциальная гипертензия", neutral_data.diagnosis
assert "Не найдено критическое поле: Ф.И.О." not in neutral_data.warnings, neutral_data.warnings

# --- Custom block-03 popup regression ---
# Doctor-created buttons must keep the same popup chain as old fixed buttons:
# custom diaries/discharge/RVK need discharge date, and placeholder-required
# custom documents must enter the required-field preflight.
from universal_profiles import DocumentTemplateSpec as _SmokeDocumentTemplateSpec
custom_popup_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
custom_popup_logic.case_number_var = _FakeVar("")
custom_popup_logic.assigned_treatment_var = _FakeVar("")
custom_popup_logic.diagnosis_var = _FakeVar("")
custom_popup_logic.discharge_date_var = _FakeVar("")
custom_popup_logic._popup_discharge_date_override = ""
custom_popup_logic._manual_discharge_date = False
custom_popup_logic.data = PatientData()
custom_popup_logic.output_vars = {"discharge": _FakeVar(False), "rvk": _FakeVar(False), DIARY_KIND: _FakeVar(False)}
custom_popup_logic._selected_custom_document_specs = lambda selected: (
    _SmokeDocumentTemplateSpec(
        id="custom_discharge",
        button_label="Выписной эпикриз хирурга",
        template="templates/discharge.docx",
        required_fields=("case.number", "diagnosis.main", "treatment.summary", "discharge.date"),
        role_id="discharge_epicrisis",
    ),
)
custom_popup_logic.selected_custom_docs = lambda: ["custom_discharge"]
custom_popup_logic._custom_requirement_flags = _main_module.CombinedMedicalDiaryApp._custom_requirement_flags.__get__(custom_popup_logic, _main_module.CombinedMedicalDiaryApp)
custom_popup_logic._selected_outputs_require_discharge_date = _main_module.CombinedMedicalDiaryApp._selected_outputs_require_discharge_date.__get__(custom_popup_logic, _main_module.CombinedMedicalDiaryApp)
custom_flags = custom_popup_logic._custom_requirement_flags(["custom_discharge"])
assert custom_flags["discharge"] is True
assert custom_flags["requires_discharge_date"] is True
custom_popup_logic._active_custom_requirement_flags = custom_flags
assert custom_popup_logic._selected_outputs_require_discharge_date() is True, "Custom discharge button must request discharge date"

custom_review_logic = _main_module.CombinedMedicalDiaryApp.__new__(_main_module.CombinedMedicalDiaryApp)
custom_review_logic.navigation_path_var = _FakeVar("")
custom_review_logic.case_number_var = _FakeVar("")
custom_review_logic.assigned_treatment_var = _FakeVar("")
custom_review_logic.diagnosis_var = _FakeVar("")
custom_review_logic.discharge_date_var = _FakeVar("")
custom_review_logic._popup_discharge_date_override = ""
custom_review_logic._manual_patient_name = False
custom_review_logic._manual_admission_date = False
custom_review_logic._manual_discharge_date = False
custom_review_logic._manual_diagnosis = False
custom_review_logic.data = PatientData(fio="Иванов И.И.", admission_date="10.06.2026")
custom_review_logic._base_output_dir = lambda: _SmokePath.cwd()
custom_review_logic._patient_output_dir_for_data = lambda data, base_dir: _SmokePath(base_dir) / "Иванов"
custom_review_logic._case_number_popup_default = lambda: ""
custom_review_logic._current_discharge_date_value = lambda: ""
custom_review_logic._custom_requirement_flags = lambda selected: custom_flags
custom_review_logic._build_patient_case_review_for_selection = _main_module.CombinedMedicalDiaryApp._build_patient_case_review_for_selection.__get__(custom_review_logic, _main_module.CombinedMedicalDiaryApp)
custom_review = custom_review_logic._build_patient_case_review_for_selection([], False, ["custom_discharge"])
missing_custom_keys = {field.key for field in custom_review.critical_missing()}
assert {"case_number", "diagnosis", "treatment", "discharge_date"}.issubset(missing_custom_keys), missing_custom_keys
