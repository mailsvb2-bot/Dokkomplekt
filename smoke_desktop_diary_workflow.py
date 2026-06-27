"""Regression smoke for desktop intake and configurable diary schedules."""

from __future__ import annotations

import json
import os
import tempfile
import time
from unittest.mock import patch
from datetime import date
from pathlib import Path

from docx import Document

from medical_field_line_pairs import assert_field_line_pairs_lock
from medical_service import MedicalDocumentService
from desktop_intake import (
    DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
    DesktopCandidate,
    desktop_path,
    assert_desktop_intake_lock,
    is_likely_primary_document,
    mark_seen,
    normalize_intake_settings,
    prepare_patient_work_folder,
    primary_document_score,
    safe_patient_subfolder,
    scan_primary_candidates,
    prompt_intake_folder,
    should_prompt_intake_setup,
    signature_key,
)
from desktop_patient_folder import assert_desktop_patient_folder_lock, build_patient_folder_info, patient_folder_name
from desktop_intake_mixin import DesktopIntakeMixin
from settings_mixin import SettingsMixin
from medical_admission_resolver import assert_admission_resolver_lock, extract_admission_date_from_primary_docx, extract_admission_date_from_primary_text
from diary_batch import fill_diary_batch
from diary_dates import parse_admission_month_year, parse_full_datetime
from diary_schedule import (
    DiaryScheduleSpec,
    assert_diary_schedule_lock,
    describe_schedule,
    expand_day_offsets,
    expand_hour_intervals,
    infer_diary_schedule_from_docx,
    parse_day_offsets,
    parse_hour_offsets,
    planned_diary_dates,
    planned_diary_time_labels,
)
from universal_diary_generation import assert_universal_diary_generation_lock, render_diary_documents_from_pack
from universal_diary_templates import assert_universal_diary_template_lock, looks_like_diary_template
from universal_fields import PatientCase
from universal_profiles import DocumentPack, DocumentTemplateSpec, load_document_pack, save_document_pack
from universal_template_engine import render_template_to_docx


def _make_primary_doc(path: Path, *, fio: str = "Иванов Иван Иванович", admission: str = "05.05.2026", title_date: bool = True) -> Path:
    doc = Document()
    if title_date:
        doc.add_paragraph(f"{admission} Первичный осмотр")
    else:
        doc.add_paragraph("Направление на госпитализацию")
    doc.add_paragraph(f"Ф.И.О.: {fio}")
    doc.add_paragraph("Дата рождения: 04.01.2000")
    doc.add_paragraph(f"Дата поступления: {admission}")
    doc.add_paragraph("Диагноз: K35.8")
    doc.save(path)
    return path




def _make_table_primary_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Направление на госпитализацию")
    table = doc.add_table(rows=6, cols=2)
    pairs = [
        ("Ф.И.О.", "Сидоров Сергей Петрович"),
        ("Дата рождения", "04.01.2000"),
        ("Дата поступления", "23.06.2026"),
        ("Место работы", "ООО Ромашка"),
        ("Должность", "врач"),
        ("Диагноз", "K35.8 Острый аппендицит"),
    ]
    for row_index, (label, value) in enumerate(pairs):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
    doc.add_paragraph("Жалобы: тревога")
    doc.add_paragraph("План лечения: терапия")
    doc.save(path)
    return path

def _make_status_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Пациент спокоен, жалоб активно не предъявляет, состояние стабильное, лечение переносит удовлетворительно.")
    doc.add_paragraph("Пациент контактен, ориентирован, режим соблюдает, назначения выполняет, динамика без отрицательной тенденции.")
    doc.save(path)
    return path


def _make_diary_template(path: Path, *, rows: int = 12, embedded_texts: bool = False) -> Path:
    doc = Document()
    if embedded_texts:
        doc.add_paragraph("Пациент спокоен, жалоб активно не предъявляет, состояние стабильное, лечение переносит удовлетворительно.")
        doc.add_paragraph("Пациент контактен, ориентирован, режим соблюдает, назначения выполняет, динамика без отрицательной тенденции.")
    table = doc.add_table(rows=rows + 1, cols=4)
    headers = ["День госпитализации", "Число", "Месяц/год", "Дневник наблюдения"]
    for index, label in enumerate(headers):
        table.cell(0, index).text = label
    for day in range(1, rows + 1):
        table.cell(day, 0).text = str(day)
        table.cell(day, 1).text = ""
        table.cell(day, 2).text = ""
        table.cell(day, 3).text = ""
    doc.save(path)
    return path


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class _DesktopIntakePromptHarness(SettingsMixin, DesktopIntakeMixin):
    def __init__(self, settings_path: Path) -> None:
        self._settings_path = settings_path
        self._settings = self._load_settings()
        self.root = object()
        self.log_messages: list[str] = []
        self.started = False

    def _log(self, message: str) -> None:
        self.log_messages.append(message)

    def _ensure_background_intake_agent_installed(self, *, start_now: bool = True) -> bool:
        self.log_messages.append("background watcher install skipped by smoke harness")
        return True

    def _start_desktop_intake_watcher(self) -> None:
        self.started = True


class _FakeTkRoot:
    def __init__(self) -> None:
        self.after_calls: list[tuple[int, object]] = []

    def after(self, delay_ms: int, callback):
        self.after_calls.append((delay_ms, callback))
        return f"after-{len(self.after_calls)}"


class _DesktopPollHarness(SettingsMixin, DesktopIntakeMixin):
    def __init__(self, settings_path: Path, folder: Path, *, enabled: bool = True) -> None:
        self._settings_path = settings_path
        self._settings = {}
        self.root = _FakeTkRoot()
        self._desktop_intake_enabled = enabled
        self._desktop_intake_asked = True
        self._desktop_intake_folder = str(folder)
        self._desktop_intake_prompt_version = "v2"
        self._desktop_intake_seen_signatures: set[str] = set()
        self._desktop_intake_poll_job = None
        self._desktop_intake_popup_open = False
        self._desktop_intake_last_popup_opened = False
        self.opened_paths: list[Path] = []

    def _open_desktop_intake_popup(self, primary_path: str | Path) -> bool:
        self.opened_paths.append(Path(primary_path))
        self._desktop_intake_last_popup_opened = True
        return False


def _assert_desktop_path_and_polling_hardening(tmp: Path) -> None:
    one_drive = tmp / "OneDrive"
    one_drive_desktop = one_drive / "Desktop"
    one_drive_desktop.mkdir(parents=True)
    fake_home = tmp / "home"
    fake_home_desktop = fake_home / "Desktop"
    fake_home_desktop.mkdir(parents=True)
    # A stale local USERPROFILE/Desktop may exist even when Explorer is actually
    # redirected to OneDrive/Desktop.  The first-launch intake folder must follow
    # the redirected desktop, not the stale profile folder.
    with (
        patch.dict(os.environ, {"USERPROFILE": str(fake_home), "OneDrive": str(one_drive)}, clear=True),
        patch("pathlib.Path.home", return_value=fake_home),
        patch("desktop_intake._desktop_from_windows_registry", return_value=None),
    ):
        assert desktop_path() == one_drive_desktop

    legacy_bool_folder = tmp / "legacy_bool_folder"
    legacy_bool_folder.mkdir()
    legacy_false = normalize_intake_settings({
        "asked": "нет",
        "enabled": "false",
        "folder": str(legacy_bool_folder),
        "prompt_version": "old",
    })
    assert legacy_false["asked"] is False
    assert legacy_false["enabled"] is False
    assert should_prompt_intake_setup({
        "asked": "да",
        "enabled": "false",
        "folder": str(legacy_bool_folder),
        "prompt_version": "old",
    }) is True, "String 'false' must not suppress the upgraded desktop-intake prompt"
    legacy_true = normalize_intake_settings({"asked": "да", "enabled": "true", "folder": str(legacy_bool_folder)})
    assert legacy_true["asked"] is True and legacy_true["enabled"] is True

    intake = tmp / "poll_intake"
    intake.mkdir()
    old = time.time() - 3
    primary = _make_primary_doc(intake / "Первичный осмотр для отмены.docx")
    os.utime(primary, (old, old))
    poller = _DesktopPollHarness(tmp / "poll_settings" / "settings.json", intake, enabled=True)
    poller._poll_desktop_intake_folder()
    assert poller.opened_paths == [primary], poller.opened_paths
    assert poller._desktop_intake_seen_signatures, "Closed/cancelled popup must not reopen forever for the same file"
    assert scan_primary_candidates(intake, poller._desktop_intake_seen_signatures) == ()
    assert poller.root.after_calls, "Enabled poller must reschedule itself"

    disabled = _DesktopPollHarness(tmp / "disabled_settings" / "settings.json", intake, enabled=False)
    disabled._poll_desktop_intake_folder()
    assert disabled.root.after_calls == [], "Disabled desktop intake must not keep scheduling background polling"

    fresh = intake / "fresh.docx"
    fresh.write_bytes(b"not a complete docx yet")
    with patch("desktop_intake.is_likely_primary_document", side_effect=AssertionError("fresh file should not be opened yet")):
        assert scan_primary_candidates(intake, set(poller._desktop_intake_seen_signatures)) == ()

    missing = intake / "missing.docx"
    before_children = {child.name for child in intake.iterdir()}
    missing_failed = False
    try:
        prepare_patient_work_folder(intake, missing, folder_name="Missing Patient")
    except FileNotFoundError:
        missing_failed = True
    assert missing_failed, "Missing primary must fail before creating an empty patient folder"
    after_children = {child.name for child in intake.iterdir()}
    assert before_children == after_children, "Missing primary must not leave an empty patient folder"

    failing = _make_primary_doc(intake / "Первичный осмотр failure.docx", fio="Фролов Фёдор Фёдорович", admission="08.05.2026")
    os.utime(failing, (old, old))
    failing_before_children = {child.name for child in intake.iterdir()}
    move_copy_failed = False
    with patch("desktop_intake.shutil.move", side_effect=PermissionError("locked")), patch("desktop_intake.shutil.copy2", side_effect=OSError("copy failed")):
        try:
            prepare_patient_work_folder(intake, failing, folder_name="Failure Patient")
        except RuntimeError as exc:
            move_copy_failed = True
            assert "Не удалось перенести первичный документ" in str(exc)
            assert "copy failed" in str(exc)
    assert move_copy_failed, "Move+copy failure must be visible to the doctor"
    assert failing.exists(), "Original primary should remain when the move/copy attempt fails"
    assert not (intake / "Failure Patient").exists(), "Failed move/copy must clean the empty patient folder"
    assert {child.name for child in intake.iterdir()} == failing_before_children


def _assert_desktop_intake_startup_prompt_persistence(tmp: Path) -> None:
    desktop = tmp / "Desktop"
    desktop.mkdir()
    settings_path = tmp / "settings" / "settings.json"

    with patch("desktop_intake.desktop_path", return_value=desktop):
        app = _DesktopIntakePromptHarness(settings_path)
        app._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", return_value=True) as ask_yes:
            app._bootstrap_desktop_intake_watcher()
        assert ask_yes.call_count == 1, "Clean first launch must ask about creating the desktop intake folder"
        intake_folder = desktop / "Выписанные пациенты"
        assert intake_folder.exists() and intake_folder.is_dir(), "Yes answer must create the desktop intake folder"
        assert app.started, "Enabled desktop intake must start the watcher after the first-run Yes answer"

        saved = json.loads(settings_path.read_text(encoding="utf-8"))
        intake_settings = saved.get("desktop_intake", {})
        assert intake_settings.get("asked") is True
        assert intake_settings.get("enabled") is True
        assert intake_settings.get("folder") == str(intake_folder)
        assert intake_settings.get("prompt_version") == DESKTOP_INTAKE_SETUP_PROMPT_VERSION

        reopened = _DesktopIntakePromptHarness(settings_path)
        reopened._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", side_effect=AssertionError("Prompt repeated after saved Yes answer")):
            reopened._bootstrap_desktop_intake_watcher()
        assert reopened.started, "Saved Yes answer must enable the watcher without asking again"

    decline_settings_path = tmp / "decline" / "settings.json"
    with patch("desktop_intake.desktop_path", return_value=desktop):
        declined = _DesktopIntakePromptHarness(decline_settings_path)
        declined._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", return_value=False) as ask_no:
            declined._bootstrap_desktop_intake_watcher()
        assert ask_no.call_count == 1, "Clean first launch must ask even when the answer is No"
        saved_decline = json.loads(decline_settings_path.read_text(encoding="utf-8"))
        assert saved_decline["desktop_intake"]["asked"] is True
        assert saved_decline["desktop_intake"]["enabled"] is False
        assert saved_decline["desktop_intake"]["prompt_version"] == DESKTOP_INTAKE_SETUP_PROMPT_VERSION

        declined_reopened = _DesktopIntakePromptHarness(decline_settings_path)
        declined_reopened._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", side_effect=AssertionError("Prompt repeated after saved No answer")):
            declined_reopened._bootstrap_desktop_intake_watcher()
        assert not declined_reopened.started, "Saved No answer must stay disabled without asking again"

    old_v2_settings_path = tmp / "old_v2" / "settings.json"
    old_v2_settings_path.parent.mkdir(parents=True)
    old_v2_settings_path.write_text(
        json.dumps({
            "desktop_intake": {
                "asked": True,
                "enabled": False,
                "folder": str(tmp / "stale" / "Выписанные пациенты"),
                "prompt_version": "v2",
                "seen_signatures": [],
            }
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    with patch("desktop_intake.desktop_path", return_value=desktop):
        upgraded = _DesktopIntakePromptHarness(old_v2_settings_path)
        upgraded._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", return_value=True) as ask_upgrade:
            upgraded._bootstrap_desktop_intake_watcher()
        assert ask_upgrade.call_count == 1, "Old v2 prompt settings must be re-asked after the fixed build"
        upgraded_saved = json.loads(old_v2_settings_path.read_text(encoding="utf-8"))
        assert upgraded_saved["desktop_intake"]["prompt_version"] == DESKTOP_INTAKE_SETUP_PROMPT_VERSION
        assert upgraded_saved["desktop_intake"]["folder"] == str(desktop / "Выписанные пациенты")
        assert (desktop / "Выписанные пациенты").exists(), "Upgraded prompt must create the real Desktop intake folder"

    missing_enabled_path = tmp / "missing_enabled" / "settings.json"
    missing_enabled_path.parent.mkdir(parents=True)
    missing_enabled_path.write_text(
        json.dumps({
            "desktop_intake": {
                "asked": True,
                "enabled": True,
                "folder": str(tmp / "missing" / "Выписанные пациенты"),
                "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
                "seen_signatures": [],
            }
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    with patch("desktop_intake.desktop_path", return_value=desktop):
        missing_enabled = _DesktopIntakePromptHarness(missing_enabled_path)
        missing_enabled._init_desktop_intake_state()
        with patch("desktop_intake_mixin.messagebox.askyesno", return_value=False) as ask_missing:
            missing_enabled._bootstrap_desktop_intake_watcher()
        assert ask_missing.call_count == 1, "Enabled settings with a missing folder must ask instead of silently recreating"
        missing_saved = json.loads(missing_enabled_path.read_text(encoding="utf-8"))
        assert missing_saved["desktop_intake"]["enabled"] is False


def _assert_desktop_intake(tmp: Path) -> None:
    assert_desktop_intake_lock()
    assert_desktop_patient_folder_lock()
    assert_admission_resolver_lock()
    normalized_old = normalize_intake_settings({"asked": True, "enabled": False, "folder": str(tmp / "old")})
    assert should_prompt_intake_setup(normalized_old), "Old asked=True settings without prompt_version must ask again after the intake feature upgrade"
    normalized_v2 = normalize_intake_settings({"asked": True, "enabled": False, "folder": str(tmp / "old"), "prompt_version": "v2", "seen_signatures": ["x", "a" * 64]})
    assert should_prompt_intake_setup(normalized_v2), "Broken v2 prompt settings must be re-asked in v3"
    normalized_current = normalize_intake_settings({"asked": True, "enabled": False, "folder": str(tmp / "old"), "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION, "seen_signatures": ["x", "a" * 64]})
    assert not should_prompt_intake_setup(normalized_current), "A doctor decline must be respected for the current prompt version"
    assert normalized_current["seen_signatures"] == ("a" * 64,)
    stale_folder = tmp / "stale" / "Выписанные пациенты"
    with patch("desktop_intake.desktop_path", return_value=tmp / "DesktopForPrompt"):
        (tmp / "DesktopForPrompt").mkdir()
        assert prompt_intake_folder(stale_folder) == tmp / "DesktopForPrompt" / "Выписанные пациенты"
    assert len(signature_key(tmp / "Иванов.docx", 1, 2)) == 64
    assert patient_folder_name("Иванов Иван Иванович", "05.05.2026") == "Иванов И.И. май 2026"
    assert extract_admission_date_from_primary_text("Направление на госпитализацию\nФИО: Иванов Иван Иванович\nДата рождения: 04.01.2000\nДата поступления: 05.05.2026") == "05.05.2026"
    assert extract_admission_date_from_primary_text("Дата рождения\n04.01.2000\nДата поступления\n05.05.2026") == "05.05.2026"
    intake = tmp / "Выписанные пациенты"
    intake.mkdir()
    old = time.time() - 3
    primary = _make_primary_doc(intake / "Первичный осмотр Иванов.docx")
    assert extract_admission_date_from_primary_docx(primary) == "05.05.2026"
    folder_info = build_patient_folder_info(primary)
    assert folder_info.folder_name == "Иванов И.И. май 2026", folder_info
    no_title = _make_primary_doc(intake / "Направление без даты в заголовке.docx", fio="Сидоров Сергей Сергеевич", admission="23.06.2026", title_date=False)
    os.utime(no_title, (old, old))
    no_title_info = build_patient_folder_info(no_title)
    assert no_title_info.admission_date == "23.06.2026", no_title_info
    assert no_title_info.folder_name == "Сидоров С.С. июнь 2026", no_title_info
    no_title.unlink()

    birth_only = intake / "Первичный без даты поступления.docx"
    birth_doc = Document()
    birth_doc.add_paragraph("Первичный осмотр")
    birth_doc.add_paragraph("Ф.И.О.: Петров Пётр Петрович")
    birth_doc.add_paragraph("Дата рождения: 04.01.2000")
    birth_doc.add_paragraph("Диагноз: K35.8")
    birth_doc.save(birth_only)
    birth_only_info = build_patient_folder_info(birth_only)
    assert birth_only_info.fio == "Петров Пётр Петрович", birth_only_info
    assert birth_only_info.admission_date == "", birth_only_info
    assert birth_only_info.folder_name == "Петров П.П", birth_only_info
    birth_only.unlink()

    assert_field_line_pairs_lock()
    table_primary = _make_table_primary_doc(intake / "Табличное направление.docx")
    os.utime(table_primary, (old, old))
    parsed_table = MedicalDocumentService().parse_primary_document(table_primary)
    assert parsed_table.fio == "Сидоров Сергей Петрович", parsed_table
    assert parsed_table.birth == "04.01.2000", parsed_table.birth
    assert parsed_table.admission_date == "23.06.2026", parsed_table.admission_date
    assert parsed_table.work_org == "ООО Ромашка", parsed_table.work_org
    assert parsed_table.position == "врач", parsed_table.position
    table_info = build_patient_folder_info(table_primary)
    assert table_info.folder_name == "Сидоров С.П. июнь 2026", table_info
    table_primary.unlink()
    os.utime(primary, (old, old))
    seen: set[str] = set()
    temp_word = _make_primary_doc(intake / "~$Первичный осмотр.docx")
    os.utime(temp_word, (old, old))
    assert not is_likely_primary_document(temp_word), "Temporary Word lock files must not trigger desktop intake"
    discharge = Document()
    discharge.add_paragraph("Выписной эпикриз. Пациент Иванов И.И. Рекомендации.")
    discharge_path = intake / "Выписной эпикриз Иванов.docx"
    discharge.save(discharge_path)
    os.utime(discharge_path, (old, old))
    assert not is_likely_primary_document(discharge_path), "Desktop intake must not treat discharge as primary source"
    assert primary_document_score("Выписка после госпитализации. Рекомендации.") < 7
    assert primary_document_score("Направление на госпитализацию\nФ.И.О. Иванов Иван Иванович\nДата поступления 05.05.2026\nДиагноз K35.8") >= 7
    candidates = scan_primary_candidates(intake, seen)
    assert len(candidates) == 1, candidates
    mark_seen(seen, candidates[0])
    assert scan_primary_candidates(intake, seen) == ()
    first = safe_patient_subfolder(intake, primary)
    assert first.name == "Иванов И.И. май 2026", first
    first.mkdir()
    second = safe_patient_subfolder(intake, primary)
    assert second.name.endswith("(2)"), second
    movable = _make_primary_doc(intake / "Первичный осмотр Петров.docx", fio="Петров Пётр Петрович", admission="06.05.2026")
    os.utime(movable, (old, old))
    patient_dir, moved_primary = prepare_patient_work_folder(intake, movable)
    assert patient_dir.name == "Петров П.П. май 2026", patient_dir
    assert patient_dir.exists() and moved_primary.exists() and moved_primary.parent == patient_dir
    assert not movable.exists(), "Processed primary DOCX should leave top-level intake folder when move is possible"
    locked = _make_primary_doc(intake / "Первичный осмотр locked.docx", fio="Кузнецов Кирилл Кириллович", admission="07.05.2026")
    os.utime(locked, (old, old))
    with patch("desktop_intake.shutil.move", side_effect=PermissionError("locked")):
        locked_dir, locked_effective = prepare_patient_work_folder(intake, locked)
    assert locked_effective.parent == locked_dir and locked_effective.exists(), locked_effective
    assert not locked.exists(), "Copy fallback should remove the top-level source when unlink is possible"
    nested = intake / "already_created"
    nested.mkdir()
    _make_primary_doc(nested / "Вложенный первичный.docx")
    assert scan_primary_candidates(intake, seen) == (), "Desktop intake must scan top-level files only and ignore output subfolders"


def _assert_diary_schedule(tmp: Path) -> None:
    assert_diary_schedule_lock()
    assert parse_day_offsets("+1, +2, +3, +5, +7, +14, +21, +28, +35, +42", require_minimum=True) == (1, 2, 3, 5, 7, 14, 21, 28, 35, 42)
    got_minimum_error = False
    try:
        parse_day_offsets("1,2,3", require_minimum=True)
    except ValueError:
        got_minimum_error = True
    if not got_minimum_error:
        raise AssertionError("Manual day schedule must require at least 10 numbers")
    assert parse_hour_offsets("1,2,3,4") == (1, 2, 3, 4)
    for bad in ("-1, 2, 3", "+1, -2, 3"):
        got_negative_error = False
        try:
            parse_day_offsets(bad)
        except ValueError:
            got_negative_error = True
        assert got_negative_error, f"Negative day schedule must be rejected: {bad}"
    assert expand_day_offsets((1, 2, 5), 5) == (1, 2, 5, 8, 11)
    assert expand_hour_intervals((1,), 4) == (1, 2, 3, 4)
    template = _make_diary_template(tmp / "diary.docx")
    assert looks_like_diary_template(template)
    inferred = infer_diary_schedule_from_docx([template])
    assert inferred.day_offsets[:5] == (1, 2, 3, 4, 5), inferred
    assert "от даты поступления" in describe_schedule(inferred)
    dates = planned_diary_dates(date(2026, 6, 10), DiaryScheduleSpec("daily", (1, 2, 5), (), 1.0, "test"), limit=4)
    assert dates == (date(2026, 6, 11), date(2026, 6, 12), date(2026, 6, 15), date(2026, 6, 18))
    labels = planned_diary_time_labels(DiaryScheduleSpec("hourly", (), (1, 2, 25), 1.0, "test"), limit=3)
    assert labels == ("01:00", "03:00", "04:00")
    assert parse_full_datetime("10.06.2026 14:30").hour == 14
    assert parse_admission_month_year("10.06.2026 14:30") == (6, 2026)


def _assert_diary_batch_generation(tmp: Path) -> None:
    status = _make_status_doc(tmp / "texts.docx")
    template = _make_diary_template(tmp / "table.docx", rows=12)
    out = tmp / "out"
    result = fill_diary_batch(
        status_files=[status],
        diary_files=[template],
        output_dir=out,
        patient_name="Иванов И.И.",
        admission_value="10.06.2026",
        gender_source_name="Иванов И.И.",
        diary_day_offsets=(1, 2, 3, 5, 7, 14, 21, 28, 35, 42),
    )
    assert result.created_files and result.created_files[0].exists()
    text = _doc_text(result.created_files[0])
    assert "11" in text and "06.2026" in text and "Пациент спокоен" in text

    hourly_out = tmp / "hourly"
    hourly_result = fill_diary_batch(
        status_files=[status],
        diary_files=[template],
        output_dir=hourly_out,
        patient_name="Иванов И.И.",
        admission_value="10.06.2026 14:00",
        gender_source_name="Иванов И.И.",
        diary_hour_offsets=(1,),
        diary_frequency_mode="hourly",
    )
    assert hourly_result.created_files
    hourly_text = _doc_text(hourly_result.created_files[0])
    assert "15:00" in hourly_text and "16:00" in hourly_text, hourly_text


def _assert_custom_diary_pack(tmp: Path) -> None:
    assert_universal_diary_template_lock()
    assert_universal_diary_generation_lock()
    profile_dir = tmp / "profile"
    templates_dir = profile_dir / "templates"
    templates_dir.mkdir(parents=True)
    template = _make_diary_template(templates_dir / "doctor_diary.docx", rows=10, embedded_texts=True)
    status = _make_status_doc(tmp / "texts_for_pack.docx")
    schedule = DiaryScheduleSpec("hourly", (1, 2, 3, 5, 7, 14, 21, 28, 35, 42), (1, 2, 3), 1.0, "test")
    pack = DocumentPack(
        pack_id="test.diary",
        name="Test diary profile",
        documents=(
            DocumentTemplateSpec(
                id="doctor_diary",
                button_label="Дневники врача",
                template="templates/doctor_diary.docx",
                category="diaries",
                role_id="daily_diary",
                diary_schedule=schedule.to_dict(),
            ),
        ),
    )
    save_document_pack(pack, profile_dir / "pack.json")
    loaded = load_document_pack(profile_dir / "pack.json")
    assert loaded.documents[0].diary_schedule["hour_offsets"] == [1, 2, 3]
    case = PatientCase()
    case.update_from_pairs({"patient.fio": "Иванов И.И.", "admission.date": "10.06.2026 08:00"})
    result = render_diary_documents_from_pack(
        pack=loaded,
        case=case,
        document_ids=["doctor_diary"],
        output_dir=tmp / "pack_out",
        base_dir=profile_dir,
        status_files=[],
        patient_name="Иванов И.И.",
        admission_value="10.06.2026 08:00",
        gender_source_name="Иванов И.И.",
        frequency_mode="hourly",
    )
    assert result.created_files and "09:00" in _doc_text(result.created_files[0])
    assert "Пациент спокоен" in _doc_text(result.created_files[0])

    placeholder_doc = Document()
    placeholder_doc.add_paragraph("График: {{diary.dates}} / {{diary.frequency}}")
    placeholder_path = tmp / "placeholder_diary.docx"
    placeholder_doc.save(placeholder_path)
    out_path = tmp / "placeholder_out.docx"
    render_template_to_docx(template_path=placeholder_path, output_path=out_path, case=case, document=loaded.documents[0], strict=False)
    rendered = _doc_text(out_path)
    assert "10.06.2026 09:00" in rendered and "ежечасно" in rendered


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="desktop_diary_workflow_") as raw:
        tmp = Path(raw)
        _assert_desktop_intake(tmp)
        _assert_desktop_path_and_polling_hardening(tmp)
        _assert_desktop_intake_startup_prompt_persistence(tmp)
        _assert_diary_schedule(tmp)
        _assert_diary_batch_generation(tmp)
        _assert_custom_diary_pack(tmp)
    print("DESKTOP DIARY WORKFLOW SMOKE OK")


if __name__ == "__main__":
    main()
