"""Targeted smoke locks for the post-v1.4.43 follow-up regressions.

These are intentionally small behavioral checks for bugs that the broad smoke
suite did not catch: table diagnosis boundaries, true manual diary selections,
and the optional background intake agent running safely from source/EXE builds.
"""

from __future__ import annotations

from pathlib import Path
import tempfile

from docx import Document

from app_config import DIARY_KIND
from medical_parser import MedicalTextParser
from universal_scanner import scan_docx
from universal_fields import PatientCase
from universal_profiles import DocumentTemplateSpec
from universal_template_engine import render_template_to_docx
from files_mixin import FilesMixin
import files_mixin as files_mixin_module


class _Var:
    def __init__(self, value=""):
        self.value = value

    def get(self):
        return self.value

    def set(self, value):
        self.value = value


class _FakeFiles(FilesMixin):
    def __init__(self, root: Path):
        self.root_dir = root
        self.status_files = []
        self.diary_files = []
        self.diary_texts_dir = ""
        self.diary_template_dir = ""
        self._diary_text_files_auto_selected = False
        self._diary_files_auto_selected = False
        self.diagnosis_var = _Var("F20.0 Параноидная шизофрения")
        self.output_dir_var = _Var("")
        self.output_vars = {DIARY_KIND: _Var(False)}
        self.logged = []

    def _dialog_initial_dir(self, *_args):
        return str(self.root_dir)

    def _get_saved_directory(self, *_args):
        return ""

    def _remember_dialog_directory(self, *_args, **_kwargs):
        return None

    def _remember_numbered_diary_template_dir(self, *_args, **_kwargs):
        return None

    def _folder_contains_numbered_diary_templates(self, folder):
        return any(Path(folder).glob("*.docx"))

    def _update_diary_text_label(self, **_kwargs):
        return None

    def _update_diary_template_label(self, **_kwargs):
        return None

    def _redraw_selection_controls(self):
        return None

    def _set_output_dir_auto(self, path):
        self.output_dir_var.set(str(path))

    def _log(self, text):
        self.logged.append(text)


def _docx_with_table(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("Ф.И.О.", "Петров Петр Петрович"),
        ("Диагноз", "F20.0 Параноидная шизофрения"),
        ("Дата поступления", "12.01.2026"),
        ("Лечение", "терапия"),
    ]
    for row, (left, right) in zip(table.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
    doc.save(path)


def _empty_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("test")
    doc.save(path)


def test_table_diagnosis_boundaries(tmp: Path) -> None:
    source = tmp / "primary_table.docx"
    _docx_with_table(source)
    parsed = MedicalTextParser().parse_docx(source)
    assert parsed.diagnosis == "F20.0 Параноидная шизофрения", parsed.diagnosis
    assert "Дата поступления" not in parsed.diagnosis
    assert "Лечение" not in parsed.diagnosis
    scan = scan_docx(source)
    best = scan.best_matches()
    assert best["diagnosis.icd10"].value == "F20.0"
    assert best["diagnosis.main"].value == "F20.0 Параноидная шизофрения", best["diagnosis.main"].value
    assert "Дата поступления" not in best["diagnosis.main"].value
    assert "Лечение" not in best["diagnosis.main"].value


def test_manual_diary_text_selection_is_not_overridden(tmp: Path) -> None:
    texts = tmp / "texts"
    texts.mkdir()
    manual = texts / "manual_selected.docx"
    auto = texts / "F20.0 Параноидная шизофрения.docx"
    _empty_docx(manual)
    _empty_docx(auto)
    fake = _FakeFiles(tmp)
    old = files_mixin_module.filedialog.askopenfilenames
    files_mixin_module.filedialog.askopenfilenames = lambda **_kwargs: (str(manual),)
    try:
        fake.choose_status_files()
    finally:
        files_mixin_module.filedialog.askopenfilenames = old
    assert fake.status_files == [str(manual)]
    assert fake._diary_text_files_auto_selected is False
    assert fake.diary_texts_dir == str(texts)
    assert fake.output_dir_var.get() == str(texts)
    # A later automatic pass must respect the manual override and not switch to auto.
    assert fake._auto_select_diary_text_by_diagnosis(ask_folder=False) is True
    assert fake.status_files == [str(manual)]


def test_manual_diary_date_template_selection_is_not_auto_replaced(tmp: Path) -> None:
    dates = tmp / "dates"
    dates.mkdir()
    selected = dates / "10.docx"
    other = dates / "12.docx"
    _empty_docx(selected)
    _empty_docx(other)
    fake = _FakeFiles(tmp)
    old_file = files_mixin_module.filedialog.askopenfilename
    old_dir = files_mixin_module.filedialog.askdirectory
    files_mixin_module.filedialog.askopenfilename = lambda **_kwargs: str(selected)
    files_mixin_module.filedialog.askdirectory = lambda **_kwargs: ""
    try:
        fake.choose_diary_files()
    finally:
        files_mixin_module.filedialog.askopenfilename = old_file
        files_mixin_module.filedialog.askdirectory = old_dir
    assert fake.diary_files == [str(selected)]
    assert fake._diary_files_auto_selected is False
    assert fake.diary_template_dir == str(dates)
    assert fake.output_vars[DIARY_KIND].get() is True
    assert fake.output_dir_var.get() == str(dates)


def test_primary_parser_reads_header_footer(tmp: Path) -> None:
    source = tmp / "header_footer_primary.docx"
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "ФИО: Иванов Иван Иванович\nДиагноз: K29.3 Хронический гастрит"
    footer_table = doc.sections[0].footer.add_table(rows=1, cols=2, width=1000000)
    footer_table.rows[0].cells[0].text = "История болезни №"
    footer_table.rows[0].cells[1].text = "ИБ-77"
    doc.add_paragraph("Дата поступления: 01.02.2026")
    doc.save(source)
    parsed = MedicalTextParser().parse_docx(source)
    assert parsed.fio.strip() == "Иванов Иван Иванович", parsed.fio
    assert parsed.diagnosis == "K29.3 Хронический гастрит", parsed.diagnosis
    assert parsed.case_number == "77", parsed.case_number


def test_custom_renderer_preserves_run_formatting(tmp: Path) -> None:
    template = tmp / "formatting_template.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("До ")
    bold_run = paragraph.add_run("{{patient.fio}}")
    bold_run.bold = True
    paragraph.add_run(" после")
    doc.save(template)
    case = PatientCase()
    case.set("patient.fio", "Иванов И.И.")
    spec = DocumentTemplateSpec(id="fmt", button_label="Формат", template=template.name, required_fields=("patient.fio",))
    output = tmp / "formatting_output.docx"
    render_template_to_docx(template_path=template, output_path=output, case=case, document=spec)
    rendered = Document(output)
    runs = rendered.paragraphs[0].runs
    assert rendered.paragraphs[0].text == "До Иванов И. И. после"
    assert runs[0].text == "До " and runs[0].bold is None
    assert runs[1].text == "Иванов И. И." and runs[1].bold is True
    assert runs[-1].text == " после" and runs[-1].bold is None


def test_background_agent_pending_handshake(tmp: Path) -> None:
    import desktop_intake_agent

    seen: set[str] = set()
    pending_path = tmp / "pending.docx"
    _empty_docx(pending_path)
    signature = "c" * 64
    old_log = desktop_intake_agent._write_log
    desktop_intake_agent._write_log = lambda _message: None
    try:
        active, changed = desktop_intake_agent._resolve_pending_state({"pending": {"path": str(pending_path), "signature": signature, "launched_at": 0}}, seen)
        assert active == {} and changed is True and signature not in seen, "Expired pending launch must retry without marking seen"
        pending_path.unlink()
        active, changed = desktop_intake_agent._resolve_pending_state({"pending": {"path": str(pending_path), "signature": signature, "launched_at": 0}}, seen)
        assert active == {} and changed is True and signature in seen, "Moved/removed pending file confirms processing"
    finally:
        desktop_intake_agent._write_log = old_log


def test_background_agent_contracts() -> None:
    import desktop_intake_agent
    from desktop_intake import DESKTOP_INTAKE_SETUP_PROMPT_VERSION

    desktop_intake_agent.assert_desktop_intake_agent_lock()
    assert desktop_intake_agent._safe_float("bad", 3.5) == 3.5
    assert desktop_intake_agent._safe_float("4.25", 0.0) == 4.25
    seen = desktop_intake_agent._state_seen_signatures({"seen_signatures": ["a" * 64, "bad", "B" * 64]})
    assert seen == {"a" * 64, "b" * 64}
    disabled = {"desktop_intake": {"asked": True, "enabled": False, "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION}}
    assert desktop_intake_agent._setting_is_current_explicit_no(disabled) is True
    assert desktop_intake_agent._launch_command()
    assert "--intake-agent" in Path("main.py").read_text(encoding="utf-8")



def test_background_agent_respects_active_gui_runtime_lock(tmp: Path) -> None:
    import time
    import desktop_intake_agent

    old_data_root = desktop_intake_agent._data_root
    try:
        desktop_intake_agent._data_root = lambda: tmp  # type: ignore[assignment]
        assert desktop_intake_agent.is_gui_runtime_active() is False
        desktop_intake_agent.write_gui_runtime_lock()
        assert desktop_intake_agent.is_gui_runtime_active() is True
        stale = {"version": desktop_intake_agent.AGENT_VERSION, "pid": 123456, "updated_at": time.time() - desktop_intake_agent.GUI_ACTIVE_SECONDS - 10}
        desktop_intake_agent._save_json(desktop_intake_agent._gui_lock_path(), stale)
        assert desktop_intake_agent.is_gui_runtime_active() is False
    finally:
        desktop_intake_agent._data_root = old_data_root  # type: ignore[assignment]


def test_desktop_intake_bootstrap_installs_background_agent_once(tmp: Path) -> None:
    from pathlib import Path
    from unittest.mock import patch
    import desktop_intake_mixin
    from desktop_intake import DESKTOP_INTAKE_FOLDER_NAME, DESKTOP_INTAKE_SETUP_PROMPT_VERSION

    calls: list[bool] = []

    class Root:
        def after(self, _delay, callback):
            return "job"
        def protocol(self, *_args):
            return None

    class App(desktop_intake_mixin.DesktopIntakeMixin):
        def __init__(self):
            self.root = Root()
            self._folder = tmp / DESKTOP_INTAKE_FOLDER_NAME
            self._folder.mkdir()
            self._settings = {}
            self._desktop_intake_enabled = True
            self._desktop_intake_asked = True
            self._desktop_intake_folder = str(self._folder)
            self._desktop_intake_prompt_version = DESKTOP_INTAKE_SETUP_PROMPT_VERSION
            self._desktop_intake_seen_signatures = set()
            self._desktop_intake_poll_job = "already-running"
            self._desktop_intake_popup_open = False
            self._desktop_intake_last_popup_opened = False
            self._desktop_intake_popup_outcome = ""
            self._desktop_intake_gui_lock_job = None
        def _settings_payload_for_disk(self):
            return {"asked": True, "enabled": True, "folder": self._desktop_intake_folder, "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION}
        def _save_settings(self):
            return None
        def _log(self, _message):
            return None
        def _ask_create_desktop_intake_folder(self):
            raise AssertionError("prompt should not run for current enabled settings")
        def _ensure_background_intake_agent_installed(self, *, start_now=True):
            calls.append(start_now)
            return True

    with patch("desktop_intake_agent.write_gui_runtime_lock", lambda: None):
        App()._bootstrap_desktop_intake_watcher()
    assert calls == [True]


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="followup_regressions_") as tmp_dir:
        tmp = Path(tmp_dir)
        test_table_diagnosis_boundaries(tmp)
        test_manual_diary_text_selection_is_not_overridden(tmp)
        test_manual_diary_date_template_selection_is_not_auto_replaced(tmp)
        test_primary_parser_reads_header_footer(tmp)
        test_custom_renderer_preserves_run_formatting(tmp)
        test_background_agent_pending_handshake(tmp)
        test_background_agent_contracts()
        test_background_agent_respects_active_gui_runtime_lock(tmp)
        test_desktop_intake_bootstrap_installs_background_agent_once(tmp)
    print("FOLLOWUP REGRESSIONS SMOKE OK")


if __name__ == "__main__":
    main()
