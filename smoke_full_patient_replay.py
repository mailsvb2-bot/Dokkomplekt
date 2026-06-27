"""Headless replay of the full patient path that caused repeated regressions."""

from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from desktop_intake import scan_primary_candidates
from diary_batch import fill_diary_batch
from doctor_action_journal import _history_dir
from installation_diagnostics import assert_installation_diagnostics_lock, collect_installation_diagnostics
from diary_creation_wizard import assert_diary_creation_wizard_lock
from medical_docx_reader import extract_docx_text
from medical_parser import MedicalTextParser


def _doc(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def main() -> None:
    assert_installation_diagnostics_lock()
    assert_diary_creation_wizard_lock()
    rows = collect_installation_diagnostics(None)
    assert rows and any("Watcher" in row.name for row in rows)

    risky_files = (
        "desktop_intake.py",
        "dnd_mixin.py",
        "files_mixin.py",
        "dialog_fields_core.py",
        "window_document_mapper.py",
        "diary_text_selection.py",
        "actions_creation_execution.py",
    )
    for filename in risky_files:
        source = Path(filename).read_text(encoding="utf-8")
        assert "except Exception:\n" not in source, f"silent bare Exception fallback in {filename}"

    with TemporaryDirectory() as tmp:
        root = Path(tmp)
        intake = root / "Выписанные пациенты"
        patient_dir = intake / "Иванов И.И. февраль 2026"
        patient_dir.mkdir(parents=True)
        primary = intake / "первичный.docx"
        _doc(
            primary,
            [
                "История болезни № 123",
                "Пациент: Иванов Иван Иванович",
                "Дата поступления: 10.02.2026",
                "Дата выписки: 12.02.2026",
                "Диагноз: K35.8 Острый аппендицит",
                "Лечение: режим, наблюдение, терапия по назначению врача",
            ],
        )
        parsed = MedicalTextParser().parse_docx(primary)
        assert parsed.fio
        assert parsed.case_number == "123"
        assert "K35.8" in parsed.diagnosis
        assert parsed.admission_date == "10.02.2026"

        # Top-level dropped DOCX must be discoverable after Explorer finishes copy.
        # Force an old mtime so the smoke does not sleep.
        import os, time
        old = time.time() - 5
        os.utime(primary, (old, old))
        candidates = scan_primary_candidates(intake, set())
        assert any(item.path == primary for item in candidates), [str(item.path) for item in candidates]

        template = root / "10.docx"
        template_doc = Document()
        table = template_doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = "Число"
        table.cell(0, 1).text = "Месяц/год"
        table.cell(0, 2).text = "Дневник"
        template_doc.save(template)

        texts = root / "дневники аппендицит.docx"
        _doc(texts, ["Состояние стабильное, жалоб активно не предъявляет."])
        result = fill_diary_batch(
            status_files=[texts],
            diary_files=[template],
            output_dir=patient_dir,
            patient_name="Иванов И.И.",
            admission_value="10.02.2026",
            discharge_value="12.02.2026",
            remove_holiday_rows=False,
            force_final_diary=False,
        )
        assert result.created_files
        output_text = extract_docx_text(result.created_files[0])
        assert "10" in output_text
        assert "02.2026" in output_text
        assert "Состояние стабильное" in output_text
        assert _history_dir(patient_dir) != patient_dir / "_medical_autofill_history"
        assert not (patient_dir / "_medical_autofill_history").exists()

    print("FULL PATIENT REPLAY SMOKE OK")


if __name__ == "__main__":
    main()
