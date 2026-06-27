"""Deterministic stress smoke for project auditor and universal generation.

The goal is not to benchmark a real clinic workload.  It protects against
classes of regressions that ordinary happy-path smoke tests missed: generated
output pollution, duplicate symbols, large import graphs, large medpacks and
batch output-name allocation.
"""

from __future__ import annotations

from pathlib import Path
import os
import sys
import shutil
import tempfile
import time

from docx import Document

from project_auditor import audit_project
from universal_fields import PatientCase
from universal_generation import _available_batch_path, analyze_pack_readiness, render_documents_from_pack
from universal_profiles import DocumentPack, DocumentTemplateSpec
from app_initialization import AppInitializationMixin
from files_mixin import FilesMixin
from desktop_intake import (
    DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
    DesktopCandidate,
    is_likely_primary_document,
    mark_seen,
    prepare_patient_work_folder,
    scan_primary_candidates,
    should_prompt_intake_setup,
    signature_key,
)
from diary_schedule import DiaryScheduleSpec
from language_preferences import LanguagePreferences
from medical_models import PatientData, build_patient_case_review, expected_medical_filenames
from medical_renderer import MedicalDocumentRenderer
from actions_reports import append_generation_history
from actions_creation_orchestrator import ActionsCreationOrchestratorMixin
from medical_service import create_documents_batch, discover_primary_documents, save_batch_generation_report
from settings_mixin import SettingsMixin

ROOT = Path(__file__).resolve().parent
SELF_TEST_NAME = "smoke_stress_hardening.py"
STRESS_HARDENING_LOCK_VERSION = "v1.3"
STRESS_PROJECT_AUDITOR_IGNORES_GENERATED_OUTPUTS = True
STRESS_BATCH_ALLOCATOR_HAS_NO_999_LIMIT = True
STRESS_RENDER_BUDGET_SECONDS = 4.0
STRESS_AUDITOR_BUDGET_SECONDS = 8.0
STRESS_V1428_REGRESSION_GUARDS = True




class _UpdateCheckFakeApp(ActionsCreationOrchestratorMixin):
    def __init__(self, manifest_path: Path) -> None:
        self._settings = {"update_manifest": str(manifest_path)}

class _FakeStringVar:
    def __init__(self, value: str = "") -> None:
        self._value = value

    def get(self) -> str:
        return self._value

    def set(self, value: str) -> None:
        self._value = value


class _OutputDirFakeApp(FilesMixin):
    def __init__(self) -> None:
        self.output_dir_var = _FakeStringVar()
        self._suspend_output_dir_tracking = False
        self._manual_output_dir = False
        self._output_dir_auto_locked_to_patient = False



class _V1427FakeService:
    def __init__(self) -> None:
        self.counter = 0

    def parse_primary_document(self, _path: Path) -> PatientData:
        self.counter += 1
        return PatientData(fio=f"Пациент {self.counter}", case_number=str(self.counter))


class _V1427FakeApp(AppInitializationMixin):
    def __init__(self) -> None:
        self.service = _V1427FakeService()
        self._primary_parse_cache: dict[str, tuple[tuple[int, int, str], PatientData]] = {}


class _V1427FakeSettings(SettingsMixin):
    def __init__(self) -> None:
        self._settings = {
            "language": {"ui_language": "ru", "spellcheck_enabled": "false"},
            "desktop_intake": {
                "asked": "да",
                "enabled": "false",
                "folder": "C:/tmp",
                "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION,
                "seen_signatures": ["a" * 64, "not-a-hash"],
            },
        }


def _assert_v1427_regression_guards() -> None:
    with tempfile.TemporaryDirectory(prefix="medical_autofill_v1427_") as raw:
        tmp = Path(raw)
        path = tmp / "primary.docx"
        path.write_bytes(b"same-size-1")
        base = 1_700_000_000_000_000_000
        os.utime(path, ns=(base, base))
        app = _V1427FakeApp()
        first = app._parse_primary_document(path)
        path.write_bytes(b"same-size-2")
        os.utime(path, ns=(base + 1, base + 1))
        second = app._parse_primary_document(path)
        if first.case_number == second.case_number:
            raise AssertionError("Primary DOCX cache must invalidate by st_mtime_ns")

        hidden = tmp / ".hidden.docx"
        hidden.write_bytes(b"hidden temp")
        os.utime(hidden, ns=(base, base))
        if is_likely_primary_document(hidden) or scan_primary_candidates(tmp, set()):
            raise AssertionError("Hidden dot DOCX files must not trigger desktop intake")

        watch = tmp / "watch"
        watch.mkdir()
        source = watch / "primary.docx"
        source.write_bytes(b"payload")
        original_move = shutil.move

        def broken_move(_src: str, _dst: str) -> str:
            raise PermissionError("simulated Explorer lock")

        shutil.move = broken_move  # type: ignore[assignment]
        try:
            patient_dir, moved = prepare_patient_work_folder(watch, source, folder_name="Пациент")
        finally:
            shutil.move = original_move  # type: ignore[assignment]
        if not patient_dir.exists() or not moved.exists() or source.exists():
            raise AssertionError("Copy fallback must use patient copy and try to remove watched source")

        prefs = LanguagePreferences.from_settings({"spellcheck_enabled": "false"})
        payload = _V1427FakeSettings()._settings_payload_for_disk()
        if prefs.spellcheck_enabled or payload["language"]["spellcheck_enabled"] is not False:
            raise AssertionError("String false must remain false in language settings")
        if payload["desktop_intake"]["enabled"] is not False or payload["desktop_intake"]["seen_signatures"] != ["a" * 64]:
            raise AssertionError("Desktop intake settings payload normalization regressed")

        spec = DiaryScheduleSpec.from_dict({
            "mode": "hourly",
            "day_offsets": [True, "2", 0, -1],
            "hour_offsets": [False, "3", "bad"],
            "confidence": "broken",
        })
        if spec.day_offsets != (2, 0) or spec.hour_offsets != (3,) or spec.confidence != 0.0:
            raise AssertionError("Diary schedule corrupted-settings guard regressed")

        try:
            MedicalDocumentRenderer().render("unknown_kind", tmp / "in.docx", tmp / "out.docx", PatientData())
        except ValueError as exc:
            if "Неизвестный тип медицинского документа" not in str(exc):
                raise AssertionError("Unknown renderer error text regressed") from exc
        else:
            raise AssertionError("Unknown renderer kind must raise ValueError")

        key = signature_key(tmp / "Иванов.docx", 1, 2)
        if len(key) != 64 or "Иванов" in key:
            raise AssertionError("Desktop intake signatures must remain hashed")
        if not should_prompt_intake_setup({}) or not should_prompt_intake_setup({"asked": True, "enabled": False, "prompt_version": "v2"}):
            raise AssertionError("First-launch/v2 prompt contract regressed")
        if should_prompt_intake_setup({"asked": True, "enabled": False, "prompt_version": DESKTOP_INTAKE_SETUP_PROMPT_VERSION}):
            raise AssertionError("Current explicit No must not nag every launch")
        seen: set[str] = set()
        mark_seen(seen, DesktopCandidate(tmp / "x.docx", (1, 2)))
        if len(next(iter(seen))) != 64:
            raise AssertionError("mark_seen must store hashed signatures")

        old_patient_dir = tmp / "Выписанные пациенты" / "Иванов"
        new_primary = tmp / "new_patient" / "primary.docx"
        new_primary.parent.mkdir(parents=True, exist_ok=True)
        new_primary.write_bytes(b"new")
        output_app = _OutputDirFakeApp()
        output_app._set_output_dir_auto_patient_scoped(old_patient_dir)
        output_app._release_patient_scoped_output_dir_before_new_primary()
        output_app._set_output_dir_from_primary_default(new_primary)
        if output_app.output_dir_var.get() != str(new_primary.parent.resolve()):
            raise AssertionError("Patient-scoped desktop-intake output folder must not leak to the next patient")

        manual_dir = tmp / "manual-output"
        output_app = _OutputDirFakeApp()
        output_app.output_dir_var.set(str(manual_dir))
        output_app._manual_output_dir = True
        output_app._output_dir_auto_locked_to_patient = False
        output_app._release_patient_scoped_output_dir_before_new_primary()
        output_app._set_output_dir_from_primary_default(new_primary)
        if output_app.output_dir_var.get() != str(manual_dir):
            raise AssertionError("Real manual output folder must not be overwritten by primary selection")


def _write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _make_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


def _case(**values: str) -> PatientCase:
    case = PatientCase()
    case.update_from_pairs(values)
    return case


def assert_stress_hardening_lock() -> None:
    if STRESS_HARDENING_LOCK_VERSION != "v1.3":
        raise AssertionError("Stress hardening lock changed unexpectedly")
    if not STRESS_PROJECT_AUDITOR_IGNORES_GENERATED_OUTPUTS:
        raise AssertionError("Project auditor must ignore generated smoke outputs")
    if not STRESS_BATCH_ALLOCATOR_HAS_NO_999_LIMIT:
        raise AssertionError("Batch output allocator must stay free of 999 duplicate limit")
    if not STRESS_V1428_REGRESSION_GUARDS:
        raise AssertionError("v1.4.28 regression guards must stay enabled")
    _assert_v1427_regression_guards()


def _assert_generated_outputs_do_not_break_auditor() -> None:
    generated = ROOT / "test_run_stress_hardening"
    generated.mkdir(exist_ok=True)
    try:
        (generated / "epi_cp1251.txt").write_bytes("ЭПИ".encode("cp1251"))
        report = audit_project(ROOT)
        if not report.ok:
            raise AssertionError("Project auditor must ignore generated test_run outputs:\n" + report.human_report())
    finally:
        shutil.rmtree(generated, ignore_errors=True)


def _assert_large_synthetic_project_audit_is_fast() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        for index in range(420):
            next_import = f"import m{index + 1}\n" if index < 419 else ""
            _write(root / f"m{index}.py", next_import + f"def f{index}():\n    return {index}\n")
        started = time.perf_counter()
        report = audit_project(root)
        elapsed = time.perf_counter() - started
        if not report.ok:
            raise AssertionError(report.human_report())
        if report.scanned_python_files != 420:
            raise AssertionError(f"Synthetic audit missed files: {report.scanned_python_files}")
        if elapsed > STRESS_AUDITOR_BUDGET_SECONDS:
            raise AssertionError(f"Project auditor stress too slow: {elapsed:.3f}s")


def _assert_duplicate_top_level_symbols_are_blocking() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        _write(root / "dup.py", "def x():\n    return 1\ndef x():\n    return 2\nclass A:\n    pass\nclass A:\n    pass\n")
        report = audit_project(root)
        codes = {item.rule_id for item in report.findings}
        if "PA014" not in codes or report.ok:
            raise AssertionError("Duplicate top-level symbols must be detected as blocking:\n" + report.human_report())


def _assert_batch_allocator_has_no_999_limit() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        reserved_names: dict[str, int] = {}
        reserved_paths: set[str] = set()
        paths = [_available_batch_path(root / "Документ.docx", reserved_names, reserved_paths) for _ in range(1105)]
        if len({path.name for path in paths}) != 1105:
            raise AssertionError("Batch allocator produced duplicate output names")
        if paths[-1].name != "Документ (1105).docx":
            raise AssertionError(f"Unexpected allocator tail: {paths[-1].name}")


def _assert_large_medpack_readiness_is_fast() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        _make_docx(root / "template.docx", "Пациент {{patient.fio}}")
        documents = tuple(
            DocumentTemplateSpec(id=f"doc_{index}", button_label=f"Документ {index}", template="template.docx", required_fields=("patient.fio",))
            for index in range(1500)
        )
        pack = DocumentPack(pack_id="stress", name="Stress", documents=documents)
        case = _case(**{"patient.fio": "Иванов Иван"})
        started = time.perf_counter()
        report = analyze_pack_readiness(pack, case, base_dir=root)
        elapsed = time.perf_counter() - started
        if report.ready_count != 1500 or report.blocked_count != 0:
            raise AssertionError(report.human_report())
        if elapsed > 2.0:
            raise AssertionError(f"Medpack readiness stress too slow: {elapsed:.3f}s")


def _assert_batch_render_is_reasonably_fast() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        _make_docx(root / "template.docx", "Пациент {{patient.fio}}. Диагноз {{diagnosis.main}}.")
        documents = tuple(
            DocumentTemplateSpec(
                id=f"doc_{index}",
                button_label="Одинаковый документ",
                template="template.docx",
                output_name="{{patient.fio}} Одинаковый документ.docx",
                required_fields=("patient.fio", "diagnosis.main"),
            )
            for index in range(36)
        )
        pack = DocumentPack(pack_id="render_stress", name="Render stress", documents=documents)
        case = _case(**{"patient.fio": "Иванов Иван", "diagnosis.main": "K35.8"})
        started = time.perf_counter()
        result = render_documents_from_pack(
            pack=pack,
            case=case,
            document_ids=[document.id for document in documents],
            output_dir=root / "out",
            base_dir=root,
            strict=True,
            output_language="ru",
            spellcheck_enabled=True,
        )
        elapsed = time.perf_counter() - started
        if not result.ok or len(result.created_files) != len(documents):
            raise AssertionError(result.human_report())
        if len(set(result.created_files)) != len(result.created_files):
            raise AssertionError("Batch render created duplicate output paths")
        if elapsed > STRESS_RENDER_BUDGET_SECONDS:
            raise AssertionError(f"Batch render stress too slow: {elapsed:.3f}s")




def _assert_patient_case_preflight_and_history() -> None:
    data = PatientData(
        fio="Иванов Иван Иванович",
        output_fio="Иванов Иван Иванович",
        case_number="12345",
        birth="01.01.1980",
        admission_date="10.02.2026",
        discharge_date="11.06.2026",
        diagnosis="K35.8 Острый аппендицит",
        treatment_plan="Терапия по назначению врача",
    )
    review = build_patient_case_review(
        data,
        selected_medical=["primary", "discharge"],
        selected_diaries=True,
        output_dir="/tmp/out",
        primary_path="primary.docx",
        manual_case_number=True,
        manual_treatment=True,
    )
    if review.critical_missing():
        raise AssertionError(review.as_text())
    names = expected_medical_filenames(review, ["primary", "discharge"])
    if "Иванов Иван Иванович Первичный осмотр.docx" not in names:
        raise AssertionError(str(names))
    incomplete = build_patient_case_review(PatientData(), selected_medical=["primary"])
    missing = {field.key for field in incomplete.critical_missing()}
    expected_missing = {"fio", "output_fio", "case_number", "admission_date", "diagnosis", "treatment"}
    if not expected_missing <= missing:
        raise AssertionError(f"Missing critical fields guard failed: {missing}")
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        patient_output_dir = tmp_path / "patient_output"
        patient_output_dir.mkdir()
        old_history_root = os.environ.get("MEDICAL_AUTOFILL_HISTORY_DIR")
        os.environ["MEDICAL_AUTOFILL_HISTORY_DIR"] = str(tmp_path / "history_root")
        try:
            log = append_generation_history(
                output_dir=patient_output_dir,
                review=review,
                selected_outputs=review.selected_outputs,
                created_files=[patient_output_dir / names[0]],
                errors=None,
            )
        finally:
            if old_history_root is None:
                os.environ.pop("MEDICAL_AUTOFILL_HISTORY_DIR", None)
            else:
                os.environ["MEDICAL_AUTOFILL_HISTORY_DIR"] = old_history_root
        if log is None or not log.exists():
            raise AssertionError("Generation history was not written")
        history_text = log.read_text(encoding="utf-8-sig")
        if "ref-" not in history_text or "created_file_count" not in history_text:
            raise AssertionError("Generation history misses privacy-safe support metadata")
        forbidden_history = ("Иванов", "12345", "10.02.2026", names[0])
        if any(item in history_text for item in forbidden_history):
            raise AssertionError("Generation history leaked patient identifiers")
        if patient_output_dir == log.parent or patient_output_dir in log.parents:
            raise AssertionError("Generation history must not be written into the patient/output folder")
        a_doc = tmp_path / "a.docx"
        _make_docx(a_doc, "ФИО: Иванов Иван Иванович")
        template_doc = tmp_path / "custom_template_named_like_patient.docx"
        _make_docx(template_doc, "Осмотр хирурга\nПациент {{patient.fio}}")
        (tmp_path / "~$temp.docx").write_text("fake", encoding="utf-8")
        discovered = discover_primary_documents(tmp_path)
        if discovered != (a_doc,):
            raise AssertionError(f"Unexpected batch discovery result: {discovered}")

        # Doctor-owned mode ships no built-in medical templates.  The legacy
        # fixed-document batch backend must therefore fail safely instead of
        # silently pretending that bundled templates exist.  Folder naming is
        # now configurable and shared by one-patient and batch flows.
        good = tmp_path / "batch_good.docx"
        doc = Document()
        for text in (
            "ФИО: Пакетный Пациент Один",
            "Дата рождения: 01.01.1980",
            "История болезни № 900",
            "Дата поступления: 10.06.2026",
            "Диагноз: K35.8 тестовый диагноз",
            "Лечение: терапия по назначению врача",
        ):
            doc.add_paragraph(text)
        doc.save(good)
        bad = tmp_path / "batch_bad.docx"
        bad_doc = Document()
        bad_doc.add_paragraph("Пустой файл без данных пациента")
        bad_doc.save(bad)
        batch = create_documents_batch(
            primary_documents=(good, good, bad),
            output_root=tmp_path / "batch_out",
            selected_docs=("primary",),
            folder_naming_settings={"parts": ["full_fio", "admission_date"], "date_format": "long"},
        )
        if batch.patient_count != 2 or batch.created_count != 0 or batch.error_count != 2:
            raise AssertionError(batch.human_report())
        if [Path(item.source).name for item in batch.items].count(good.name) != 1:
            raise AssertionError("Batch generation must deduplicate repeated source files")
        if "Не найдены" not in batch.human_report():
            raise AssertionError("Doctor-owned build must fail clearly when legacy fixed templates are requested")
        if "Пакетный Пациент Один 10.06.2026" not in batch.human_report():
            raise AssertionError("Batch patient folder naming settings were not applied")
        try:
            create_documents_batch(primary_documents=(), output_root=tmp_path / "empty_batch_out", selected_docs=("primary",))
        except ValueError as exc:
            if "нет ни одного" not in str(exc):
                raise AssertionError(str(exc)) from exc
        else:
            raise AssertionError("Empty batch generation must be rejected")
        report_file = save_batch_generation_report(batch, tmp_path / "batch_out" / "_medical_autofill_history" / "batch_generation_report.txt")
        report_text = report_file.read_text(encoding="utf-8")
        if "Пакетная обработка" not in report_text or "Ошибок: 2" not in report_text:
            raise AssertionError(report_text)
        second_report = save_batch_generation_report(batch, tmp_path / "batch_out" / "_medical_autofill_history" / "batch_generation_report.txt")
        if second_report == report_file:
            raise AssertionError("Batch report must not overwrite an existing report file")


def _assert_update_manifest_check_is_real(tmp_path: Path) -> None:
    manifest = tmp_path / "update_manifest.json"
    manifest.write_text("{\"version\": \"1.4.99\", \"download_url\": \"https://example.invalid/app.zip\", \"notes\": \"test\"}", encoding="utf-8")
    app = _UpdateCheckFakeApp(manifest)
    loaded, source = app._read_update_manifest()
    if not loaded or loaded.get("version") != "1.4.99" or source != str(manifest):
        raise AssertionError(f"Update manifest was not loaded: {loaded}, {source}")
    if app._version_tuple_for_compare("v1.4.35_name") >= app._version_tuple_for_compare("1.4.99"):
        raise AssertionError("Version comparison regressed")
    bad_manifest = tmp_path / "bad_manifest.json"
    bad_manifest.write_text("[]", encoding="utf-8")
    bad_app = _UpdateCheckFakeApp(bad_manifest)
    bad_loaded, bad_source = bad_app._read_update_manifest()
    if bad_loaded is not None or "JSON-объект" not in bad_source:
        raise AssertionError(f"Malformed update manifest must be rejected: {bad_loaded}, {bad_source}")


def main() -> None:
    assert_stress_hardening_lock()
    _assert_generated_outputs_do_not_break_auditor()
    _assert_large_synthetic_project_audit_is_fast()
    _assert_duplicate_top_level_symbols_are_blocking()
    _assert_batch_allocator_has_no_999_limit()
    _assert_large_medpack_readiness_is_fast()
    _assert_batch_render_is_reasonably_fast()
    _assert_patient_case_preflight_and_history()
    with tempfile.TemporaryDirectory(prefix="medical_autofill_update_manifest_") as raw:
        _assert_update_manifest_check_is_real(Path(raw))
    print("STRESS HARDENING SMOKE OK")


if __name__ == "__main__":
    main()
    # This is a disposable CI/smoke process. Heavy DOCX/Tk/platform probes can
    # occasionally leave helper handles alive after all assertions pass; exit
    # hard only after success so the smoke never reports OK and then hangs.
    sys.stdout.flush()
    sys.stderr.flush()
    os._exit(0)
