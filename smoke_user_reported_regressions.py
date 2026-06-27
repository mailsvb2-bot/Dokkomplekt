from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from desktop_intake import is_likely_primary_document, primary_document_score
from desktop_intake_agent import (
    DESKTOP_INTAKE_AGENT_AUTOSTART_INSTALL_SUPPORTED,
    DESKTOP_INTAKE_AGENT_HIDES_POWERSHELL_WINDOW,
    _shortcut_launch_target_and_args,
    _shortcut_launch_command,
)
from settings_mixin import SettingsMixin


class _Var:
    def __init__(self, value=""):
        self._value = value

    def get(self):
        return self._value

    def set(self, value):
        self._value = value


# 1) Folder watcher must not ignore minimal real primary templates.
minimal_primary = "Пациент Иванов И.И.\nДата поступления 01.02.2026\nДиагноз K35.8\nНомер истории 123"
assert primary_document_score(minimal_primary) >= 5
assert DESKTOP_INTAKE_AGENT_AUTOSTART_INSTALL_SUPPORTED is True
assert DESKTOP_INTAKE_AGENT_HIDES_POWERSHELL_WINDOW is True
_target, _arguments = _shortcut_launch_target_and_args()
assert _target
assert _shortcut_launch_command()[0]
agent_source = Path("desktop_intake_agent.py").read_text(encoding="utf-8")
assert "Start-Process" not in agent_source
assert "powershell.exe" not in agent_source
assert "MedicalDiaryAutofill Intake Agent.vbs" in agent_source
install_bat_source = Path("install_background_watcher.bat").read_text(encoding="utf-8")
assert "powershell" not in install_bat_source.lower()
assert "Start-Process" not in install_bat_source
assert "MedicalDiaryAutofill Intake Agent.vbs" in install_bat_source

# 2) RVK military commissariat default is a safe persisted setting.
settings_holder = SimpleNamespace(
    _settings={"defaults": {"rvk_military_commissariat": "Ленинский военкомат"}},
    _settings_path=Path("/tmp/settings.json"),
)
payload = SettingsMixin._settings_payload_for_disk(settings_holder)
assert payload["defaults"]["rvk_military_commissariat"] == "Ленинский военкомат"

# 3) New primary patient must clear stale manual diary files but keep folders for auto-selection.
from files_mixin import FilesMixin

fake = SimpleNamespace(
    assigned_treatment_var=_Var("old"),
    case_number_var=_Var("old"),
    expert_work_status_var=_Var(""),
    expert_work_org_var=_Var(""),
    expert_position_var=_Var(""),
    expert_sick_leave_needed_var=_Var("да"),
    expert_sick_leave_from_var=_Var(""),
    expert_sick_leave_number_var=_Var(""),
    vk_mse_work_org_var=_Var(""),
    vk_mse_position_var=_Var(""),
    sick_leave_vk_work_org_var=_Var(""),
    sick_leave_vk_position_var=_Var(""),
    sick_leave_vk_work_position_var=_Var(""),
    rvk_act_number_var=_Var("old"),
    rvk_military_commissariat_var=_Var("remembered"),
    rvk_work_position_var=_Var(""),
    sick_leave_vk_date_var=_Var(""),
    sick_leave_vk_protocol_number_var=_Var(""),
    sick_leave_vk_protocol_date_var=_Var(""),
    sick_leave_vk_commission_date_var=_Var(""),
    commission_date_var=_Var(""),
    commission_number_var=_Var(""),
    labs_text_var=_Var(""),
    labs_source_path_var=_Var(""),
    labs_explicit_date_var=_Var(""),
    labs_date_policy_var=_Var(""),
    labs_without_var=_Var(False),
    status_files=["old_text.docx"],
    diary_files=["10.docx"],
    diary_texts_dir="/texts",
    diary_template_dir="/dates",
    _diary_text_files_auto_selected=False,
    _diary_files_auto_selected=False,
    patient_name_var=_Var("old"),
    admission_date_var=_Var("old"),
    discharge_date_var=_Var("old"),
    diagnosis_var=_Var("old"),
    data=None,
)
fake._update_expert_sick_leave_display = lambda: None
fake._update_diary_text_label = lambda **_kwargs: None
fake._update_diary_template_label = lambda **_kwargs: None
fake._folder_contains_numbered_diary_templates = lambda _folder: True
fake._set_ui_var = lambda var, value: var.set(value)
fake._set_primary_drop_empty = lambda: None
FilesMixin._reset_primary_document_runtime_state(fake)
assert fake.status_files == []
assert fake.diary_files == []
assert fake._diary_text_files_auto_selected is True
assert fake._diary_files_auto_selected is True
assert fake.rvk_military_commissariat_var.get() == "remembered"

# 4) Custom block-03 buttons must have a first-click popup chain.
from layout_action_bar import LayoutActionBarMixin

class _ButtonFake:
    def __init__(self):
        self.called = []
        self.output_vars = {"custom_profile:rvk_doc": _Var(True)}

    def _custom_requirement_flags(self, ids):
        self.called.append(("flags", tuple(ids)))
        return {
            "diary": False,
            "regular": True,
            "discharge": False,
            "rvk": True,
            "commission": False,
            "vk_mse": False,
            "sick_leave_vk": False,
            "requires_case_number": True,
            "requires_diagnosis": True,
            "requires_treatment": True,
            "requires_discharge_date": True,
            "requires_labs": False,
        }

    def _prompt_rvk_details(self):
        self.called.append(("rvk",))
        return True

    def _update_selected_outputs_status(self):
        self.called.append(("status",))

    def _redraw_selection_controls(self):
        self.called.append(("redraw",))

button_fake = _ButtonFake()
button_fake._on_custom_output_toggle = LayoutActionBarMixin._on_custom_output_toggle.__get__(button_fake, _ButtonFake)
LayoutActionBarMixin._on_output_toggle(button_fake, "custom_profile:rvk_doc")
assert ("rvk",) in button_fake.called

# 5) Labs popup stays simple but the mouse scanner is visible again.
labs_source = Path("dialog_fields_core.py").read_text(encoding="utf-8")
visible_labs_block = labs_source.split("def open_labs_selection_scanner", 1)[0]
assert "Анализы — просто выберите один вариант" in visible_labs_block
assert "Вставить / ввести" in visible_labs_block
assert "Сканер мышкой" in visible_labs_block
assert "Пусть даты подставит программа" not in visible_labs_block

# 6) Discharge popup field mapping must not duplicate discharge_date.
dialog_expert = Path("dialog_expert.py").read_text(encoding="utf-8")
assert 'detail_fields.append("discharge_date")\n            detail_fields.append("discharge_date")' not in dialog_expert

# 7) Windows minimize/restore stays native-frame based.
chrome = Path("window_chrome_mixin.py").read_text(encoding="utf-8")
assert 'self.root.state("iconic")' in chrome
assert "def _raise_restored_window_safely" in chrome
assert "overrideredirect(False)" in chrome


# 8) Desktop intake patient folder is already final: do not nest another patient folder.
from actions_creation_foldering import ActionsCreationFolderingMixin

folder_fake = SimpleNamespace(
    _output_dir_auto_locked_to_patient=True,
    discharge_date_var=_Var(""),
)
folder_fake._base_output_dir = lambda: Path("/tmp/Выписанные пациенты/Иванов И.И. февраль 2026")
folder_fake._folder_naming_settings = lambda: {}
locked_dir = ActionsCreationFolderingMixin._patient_output_dir_for_data(
    folder_fake,
    SimpleNamespace(output_fio="Иванов И.И.", fio="Иванов Иван Иванович", admission_date="10.02.2026", discharge_date=""),
)
assert locked_dir == Path("/tmp/Выписанные пациенты/Иванов И.И. февраль 2026")

# 9) Doctor action journal must not create _medical_autofill_history inside patient folders.
from doctor_action_journal import _history_dir
patient_tmp = Path("/tmp/Выписанные пациенты/Иванов")
history_dir = _history_dir(patient_tmp)
assert history_dir != patient_tmp / "_medical_autofill_history"

# 10) Blank diary tables are valid doctor templates: date and text are filled.
from tempfile import TemporaryDirectory
from docx import Document
from diary_batch import fill_diary_batch
from medical_docx_reader import extract_docx_text

with TemporaryDirectory() as tmp:
    root = Path(tmp)
    template = root / "10.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=3)
    table.cell(0, 0).text = "Число"
    table.cell(0, 1).text = "Месяц/год"
    table.cell(0, 2).text = "Дневник"
    doc.save(template)
    texts = root / "тексты.docx"
    text_doc = Document()
    text_doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет.")
    text_doc.save(texts)
    out = root / "out"
    result = fill_diary_batch(
        status_files=[texts],
        diary_files=[template],
        output_dir=out,
        patient_name="Иванов И.И.",
        admission_value="10.02.2026",
        discharge_value="12.02.2026",
        remove_holiday_rows=False,
        force_final_diary=False,
    )
    assert result.filled_rows >= 1
    created_text = extract_docx_text(result.created_files[0])
    assert "10" in created_text
    assert "02.2026" in created_text
    assert "Состояние стабильное" in created_text

# 11) Preflight edit button must open an in-dialog editor instead of closing to main UI.
preflight_source = Path("actions_creation_preflight.py").read_text(encoding="utf-8")
assert "_edit_patient_case_inside_preflight" in preflight_source
assert "Окно проверки осталось открытым" in preflight_source


# 12) Generic field popup must not show analyses unless the caller explicitly asks for it.
fields_core = Path("dialog_fields_core.py").read_text(encoding="utf-8")
fields_mixin = Path("dialog_fields.py").read_text(encoding="utf-8")
assert "include_labs_block: bool = False" in fields_core
assert "if include_labs_block:" in fields_core
assert "include_labs_block=include_labs_block" in fields_mixin
assert "Сканер Word" in fields_core
assert "open_external_word_selection_scanner_dialog" in fields_core
assert "parent.wait_window(win)" in fields_core
assert "open_desktop_path" in fields_core
assert "subprocess.Popen" not in fields_core
printer_platform_source = Path("printer_platform.py").read_text(encoding="utf-8")
assert "stdin=subprocess.DEVNULL" in printer_platform_source
assert "creationflags=_creationflags_no_window()" in printer_platform_source

# 13) Creation requirements must pass labs block only for documents that actually need analyses.
actions_batch_source = Path("actions_creation_batch.py").read_text(encoding="utf-8")
layout_source = Path("layout_action_bar.py").read_text(encoding="utf-8")
dialog_expert_source = Path("dialog_expert.py").read_text(encoding="utf-8")
assert 'include_labs_block=bool(custom.get("requires_labs"))' in actions_batch_source
assert 'include_labs_block=bool(flags.get("requires_labs"))' in layout_source
assert "include_labs_block: bool = False" in dialog_expert_source

# 14) Service history and batch reports must never be written inside patient output folders.
reports_source = Path("actions_reports.py").read_text(encoding="utf-8")
assert "def _technical_history_root" in reports_source
assert "return root / _HISTORY_DIR_NAME" not in reports_source
assert "patient_ref" in reports_source
assert "diagnosis_code" not in reports_source
assert '"diagnosis": review.value("diagnosis")' not in reports_source
assert "errors_redacted" in reports_source
assert 'Path(output_root) / "_medical_autofill_history"' not in actions_batch_source
assert 'output_root / "_medical_autofill_history"' not in actions_batch_source

# 15) Release check must execute the user-reported smoke, not just check it is wired.
release_source = Path("release_check.py").read_text(encoding="utf-8")
assert 'smoke_user_reported_regressions.py' in release_source and 'runpy.run_path' in release_source
assert 'smoke_test_combined.py' in release_source
assert '.ruff_cache"}' not in release_source.split("IGNORED_DIR_NAMES", 1)[1].split("\n", 1)[0]

# 16) UI diary creation must refuse empty-text diaries by default.
diary_batch_source = Path("diary_batch.py").read_text(encoding="utf-8")
assert "allow_empty_statuses: bool = False" in diary_batch_source
assert "не будет создавать пустые дневники без текстов" in diary_batch_source

# 17) Diary text auto-discovery must recognize realistic doctor folder names and be deterministic.
files_source = Path("files_mixin.py").read_text(encoding="utf-8")
assert "sorted(root.iterdir()" in files_source
for token in ("статус", "наблюден", "status", "statuses"):
    assert token in files_source
assert "self.output_vars[DIARY_KIND].set(True)" in files_source

# 18) Medical flow must not silently swallow labs failures.
medical_flow_source = Path("actions_medical_flow.py").read_text(encoding="utf-8")
assert "actions_medical_flow.labs_override" in medical_flow_source

# 19) Drag-and-drop registration/parser failures must be logged.
dnd_source = Path("dnd_mixin.py").read_text(encoding="utf-8")
assert "dnd_mixin.import_tkinterdnd2" in dnd_source
assert "dnd_mixin.parse_drop_event_tcl" in dnd_source
assert "dnd_mixin.read_text_snippet" in dnd_source

# 20) No exact silent Exception swallows remain in runtime code except hard last-resort modules.
allowed_silent = {"diagnostic_logging.py", "medical_gender.py"}
allowed_gate_files = {
    "architecture_contracts.py",
    "performance_check.py",
    "prod_audit.py",
    "release_check.py",
}
for path in Path(".").glob("*.py"):
    if (
        path.name.startswith("smoke")
        or path.name.startswith("test_")
        or path.name.startswith("project_auditor")
        or path.name in allowed_silent
        or path.name in allowed_gate_files
    ):
        continue
    assert "except Exception:" not in path.read_text(encoding="utf-8"), f"silent Exception swallow in {path.name}"

# 21) Diary wizard must treat selected folders as usable state and require discharge date.
wizard_source = Path("diary_creation_wizard.py").read_text(encoding="utf-8")
assert "diary_template_dir" in wizard_source
assert "diary_texts_dir" in wizard_source
assert "Дата выписки" in wizard_source and "на какой строке закончить" in wizard_source

# 22) Date-template autoselection may use parser data when title-date is absent.
template_source = Path("diary_template_selection.py").read_text(encoding="utf-8")
assert "data_admission_date" in template_source
assert "полный запрет на UI-дату ломал автоподбор" in template_source

# 23) Diary output fills the hospitalization-day service column too.
writer_apply_source = Path("diary_writer_apply.py").read_text(encoding="utf-8")
assert "hospitalization_day_col" in writer_apply_source
assert "День госпитализации" in writer_apply_source

# 24) Installation diagnostics recognizes the hidden VBS autostart script and stale locks.
diag_source = Path("installation_diagnostics.py").read_text(encoding="utf-8")
assert "MedicalDiaryAutofill Intake Agent.vbs" in diag_source
assert "stale-lock" in diag_source



# 25) Useful selection principles from v1.3.18 must remain available without
# adding built-in diary texts: exact template priority, content fallback and
# semantic matching of doctor-owned filename variants.
import tempfile
from datetime import datetime
from docx import Document
from diary_template_discovery import DiaryTemplateDiscoveryMixin
from diary_template_selection import DiaryTemplateSelectionMixin
from diary_text_selection import (
    diary_diagnosis_match_score,
    find_diary_text_file_for_diagnosis,
    normalize_diary_diagnosis_name,
)


class _DiarySelector(DiaryTemplateSelectionMixin, DiaryTemplateDiscoveryMixin):
    def __init__(self) -> None:
        self._diary_template_files_cache = {}
        self._diary_template_day_cache = {}


def _empty_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("template")
    doc.save(path)


def _content_numbered_docx(path: Path, day: int) -> None:
    doc = Document()
    doc.add_paragraph("Дневник наблюдения")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "№"
    table.rows[0].cells[1].text = "Число"
    table.rows[0].cells[2].text = "Дневник"
    table.rows[1].cells[0].text = "2"
    table.rows[1].cells[1].text = f"{day:02d}"
    table.rows[1].cells[2].text = ""
    doc.save(path)


with tempfile.TemporaryDirectory(prefix="selection_principles_") as tmp_dir:
    tmp = Path(tmp_dir)
    root = tmp / "шаблоны дневников"
    root.mkdir(parents=True)
    _empty_docx(root / "шаблон 15.docx")
    _empty_docx(root / "15(2).docx")
    _empty_docx(root / "16 дневник.docx")
    selector = _DiarySelector()
    found = selector._find_numbered_diary_template(root, 15)
    assert found is not None and found.name == "15(2).docx", found
    content_only = root / "странное название без числа.docx"
    _content_numbered_docx(content_only, 7)
    found_by_content = selector._find_numbered_diary_template(root, 7)
    assert found_by_content is not None and found_by_content.name == content_only.name
    assert [(day, reason) for day, reason, _dt in selector._diary_template_day_candidates(datetime(2026, 4, 2))] == [
        (2, "дате госпитализации"),
        (3, "первому дню дневника"),
    ]

    neutral = tmp / "тексты дневников"
    neutral.mkdir()
    _empty_docx(neutral / "дневники ВЭ острый аппендицит.docx")
    _empty_docx(neutral / "дневники ВЭ аппендицит после операции.docx")
    _empty_docx(neutral / "дневники ВЭ артериальная гипертензия.docx")
    assert normalize_diary_diagnosis_name("K35 Острый аппендицит") == "острый аппендицит"
    matched_neutral = find_diary_text_file_for_diagnosis(neutral, "K35 Острый аппендицит после операции")
    assert matched_neutral is not None and matched_neutral.name == "дневники ВЭ аппендицит после операции.docx"
    assert diary_diagnosis_match_score("I10 Артериальная гипертензия", "дневники ВЭ артериальная гипертензия.docx") >= 90

    legacy = tmp / "реальные имена текстов"
    legacy.mkdir()
    legacy_expected = {
        "дневники ВЭ олигофрены.docx": "F70.0 Легкая умственная отсталость",
        "дневники ВЭ олигофрены с астенией.docx": "F70.0 Легкая умственная отсталость с астеническим синдромом",
        "дневники ВЭ олигофрены с психопатизацией.docx": "F70 Умственная отсталость с психопатизацией",
        "дневники ВЭ легкая депрессия с датами.docx": "F32.0 Легкий депрессивный эпизод",
        "дневники ВЭ легкая органика.docx": "F06.6 Органическое эмоционально лабильное расстройство",
        "дневники ВЭ здоровые2.docx": "Психически здоров",
    }
    for filename in legacy_expected:
        _empty_docx(legacy / filename)
    for filename, diagnosis in legacy_expected.items():
        matched = find_diary_text_file_for_diagnosis(legacy, diagnosis)
        assert matched is not None, diagnosis
        assert matched.name == filename, (diagnosis, matched.name)

# 26) Existing hospitalization-day sequences from doctor templates are preserved;
# blank rows are filled, but non-empty 2/3/4/7/11 service schedules are not overwritten.
writer_apply_source = Path("diary_writer_apply.py").read_text(encoding="utf-8")
assert "Preserve doctor-owned values" in writer_apply_source
assert "if not str(hosp_cell.text or \"\").strip()" in writer_apply_source

# 27) v1.4.55 UI/program interaction closure: desktop-intake popup must expose
# explicit bulk controls, cancellation and a diary tile when diary inputs are ready.
desktop_popup_source = Path("desktop_intake_mixin.py").read_text(encoding="utf-8")
assert "Выбрать всё" in desktop_popup_source
assert "Снять всё" in desktop_popup_source
assert "Отмена" in desktop_popup_source
assert "DIARY_KIND" in desktop_popup_source and "DIARY_LABEL" in desktop_popup_source
assert "add_diary_entry" in desktop_popup_source

# 28) Preflight editor must allow intentional clearing and must not leave modal
# grabs after closing nested windows.
preflight_source = Path("actions_creation_preflight.py").read_text(encoding="utf-8")
assert "def _clear_required_review_value" in preflight_source
assert "close_required_popup" in preflight_source
assert "close_editor" in preflight_source
assert "close_preflight" in preflight_source
assert "labs=(self.labs_text_var.get().strip()" in preflight_source
assert preflight_source.count('("diagnosis", "Диагноз с МКБ-10"') == 1

# 29) Labs popup must stay simple but include both scanners, preserve modal
# close contracts and apply Word-scanner selection directly to labs_text_var.
fields_source = Path("dialog_fields_core.py").read_text(encoding="utf-8")
assert "Сканер мышкой" in fields_source and "Сканер Word" in fields_source
assert "close_manual_labs" in fields_source
assert "close_word_scanner" in fields_source
assert 'field_id == "labs.results"' in fields_source
assert "app.labs_text_var.set(normalize_labs_block" in fields_source
assert "В документе не найден текст" in fields_source
assert "Введите текст анализов" in fields_source

# 30) Custom documents must not report success when no file was actually created;
# custom diaries require discharge date exactly like ordinary diary creation.
universal_flow_source = Path("actions_universal_flow.py").read_text(encoding="utf-8")
assert "Custom-документы профиля не созданы" in universal_flow_source
assert "Custom-дневники не созданы" in universal_flow_source
assert "Не указана дата выписки для дневников" in universal_flow_source
execution_source = Path("actions_creation_execution.py").read_text(encoding="utf-8")
assert "Ничего не создано" in execution_source
assert "Создание остановлено без файлов" in execution_source

# 31) Diagnostics and journal must not mislead the doctor or leak full diagnosis prose.
diag_source = Path("installation_diagnostics.py").read_text(encoding="utf-8")
assert "custom_documents_for_main_ui" in diag_source
assert "_doctor_buttons_setup_completed(pack)" in diag_source
journal_source = Path("doctor_action_journal.py").read_text(encoding="utf-8")
assert "def _sanitize_details" in journal_source
assert "_diagnosis_code_for_journal(value)" in journal_source
assert 'DOCTOR_ACTION_JOURNAL_LOCK_VERSION = "v1.3"' in journal_source

# 32) Hidden watcher must have no PowerShell code path left.
agent_source = Path("desktop_intake_agent.py").read_text(encoding="utf-8")
assert "DESKTOP_INTAKE_AGENT_HAS_NO_POWERSHELL_CODE_PATH = True" in agent_source
assert "_powershell_single_quote" not in agent_source
assert "powershell.exe" not in agent_source.lower()

# 33) Medpack requirement flags must not classify ordinary phrase
# «дневной стационар» as a diary document and must not contain duplicated diary assignment.
main_docs_source = Path("universal_main_documents.py").read_text(encoding="utf-8")
assert main_docs_source.count('flags["diary"] = flags["diary"] or is_diary') == 1
from universal_profiles import DocumentTemplateSpec
from universal_main_documents import custom_requirement_flags_for_documents
non_diary_doc = DocumentTemplateSpec(id="d1", button_label="Справка о дневном стационаре", template="", category="regular", required_fields=())
assert custom_requirement_flags_for_documents([non_diary_doc])["diary"] is False

print("USER REPORTED REGRESSIONS SMOKE OK")
