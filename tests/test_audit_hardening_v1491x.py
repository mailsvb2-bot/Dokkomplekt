from __future__ import annotations

from datetime import datetime, timezone
import json
from pathlib import Path
import sys
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from desktop_patient_folder import _looks_like_human_fio
from medical_docx_blocks import extract_docx_text
from medical_docx_xml_fragments import ensure_docx_compatible
from medical_formatting import safe_filename
from medical_service import discover_primary_documents
from product_access import (
    ProductAccessManager,
    ProductAccessMixin,
    TEST_PRODUCT_ACCESS_DISABLED_ENV,
    WatermarkBatchResult,
    WatermarkResult,
    product_access_enforcement_enabled,
)


def _make_docx(path: Path, lines: list[str]) -> Path:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)
    return path


def _make_real_docm(path: Path) -> Path:
    source = path.with_suffix(".docx")
    _make_docx(source, ["01.09.2026 Первичный осмотр", "ФИО: ИВАНОВ ИВАН ИВАНОВИЧ"])
    with ZipFile(source, "r") as src, ZipFile(path, "w", ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    b"application/vnd.ms-word.document.macroEnabled.main+xml",
                )
            dst.writestr(info, data)
    source.unlink()
    return path


def test_docm_is_prepared_as_macro_free_docx_and_remains_readable(tmp_path: Path):
    docm = _make_real_docm(tmp_path / "primary.docm")

    compatible = ensure_docx_compatible(docm)

    assert compatible.suffix.lower() == ".docx"
    assert compatible != docm
    assert "Первичный осмотр" in extract_docx_text(docm)
    with ZipFile(compatible) as archive:
        names = {name.casefold() for name in archive.namelist()}
        assert not any("vbaproject" in name for name in names)


def test_safe_filename_never_truncates_known_document_suffix():
    result = safe_filename("Очень длинное имя документа " * 10 + ".docx")

    assert len(result) <= 80
    assert result.endswith(".docx")


def test_patient_folder_fio_accepts_uppercase_and_initials():
    assert _looks_like_human_fio("ИВАНОВ ИВАН ИВАНОВИЧ") is True
    assert _looks_like_human_fio("Иванов И.И.") is True
    assert _looks_like_human_fio("Иванов Иван Иванович") is True


def test_extract_docx_text_preserves_repeated_structural_labels(tmp_path: Path):
    path = tmp_path / "dates.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Дата"
    table.cell(0, 1).text = "01.01.1980"
    table.cell(1, 0).text = "Дата"
    table.cell(1, 1).text = "01.09.2026"
    doc.save(path)

    lines = extract_docx_text(path).splitlines()

    assert lines == ["Дата", "01.01.1980", "Дата", "01.09.2026"]


def test_batch_discovery_rejects_generated_discharge_documents(tmp_path: Path):
    primary = _make_docx(
        tmp_path / "Первичный.docx",
        ["01.09.2026 Первичный осмотр", "Жалобы", "Анамнез", "Диагноз"],
    )
    _make_docx(
        tmp_path / "Иванов Выписной эпикриз.docx",
        ["Выписной эпикриз", "Дата поступления 01.09.2026", "Дата выписки 10.09.2026"],
    )

    discovered = discover_primary_documents(tmp_path)

    assert discovered == (primary,)


def test_trial_state_recovers_from_corrupt_primary_copy(tmp_path: Path):
    now = datetime(2026, 9, 1, tzinfo=timezone.utc)
    manager = ProductAccessManager(tmp_path, now=now)
    manager.current_state()
    manager.record_created_documents(29)
    manager.state_path.write_text("{broken", encoding="utf-8")

    recovered = ProductAccessManager(tmp_path, now=now).current_state()

    assert recovered.plan == "trial"
    assert recovered.documents_used_total_trial == 29
    assert recovered.remaining_documents_month == 1


def test_trial_state_corruption_fails_closed_when_all_copies_are_invalid(tmp_path: Path):
    now = datetime(2026, 9, 1, tzinfo=timezone.utc)
    manager = ProductAccessManager(tmp_path, now=now)
    manager.current_state()
    manager.state_path.write_text("{broken", encoding="utf-8")
    manager.state_guard_path.write_text("{broken", encoding="utf-8")

    state = ProductAccessManager(tmp_path, now=now).current_state()

    assert state.active is False
    assert state.plan == "blocked"
    assert "повреждено" in state.warning.lower()


def test_packaged_runtime_cannot_disable_access_control_with_environment(monkeypatch):
    monkeypatch.setenv(TEST_PRODUCT_ACCESS_DISABLED_ENV, "1")
    monkeypatch.setenv("MEDICAL_AUTOFILL_DISABLE_PRODUCT_ACCESS", "1")
    monkeypatch.setattr(sys, "frozen", True, raising=False)

    assert product_access_enforcement_enabled() is True


class _CreatedFilesBase:
    def _created_files_from_results(self, created_medical, created_custom, diary_result):
        return list(created_medical) + list(created_custom)


class _Harness(ProductAccessMixin, _CreatedFilesBase):
    def __init__(self, manager):
        self.manager = manager

    def _product_access_manager(self):
        return self.manager


class _WatermarkManager:
    def current_watermark_text(self):
        return "TRIAL"

    def record_created_documents(self, count):
        raise AssertionError("counter must not be updated after watermark failure")


class _CounterFailureManager:
    def current_watermark_text(self):
        return ""

    def record_created_documents(self, count):
        raise OSError("disk full")


def test_watermark_failure_deletes_unlicensed_output(tmp_path: Path, monkeypatch):
    output = tmp_path / "generated.docx"
    output.write_bytes(b"not-a-real-docx")
    monkeypatch.setattr(
        "product_access.apply_watermark_to_files",
        lambda paths, text: WatermarkBatchResult((WatermarkResult(str(output), False, "cannot watermark"),)),
    )

    with pytest.raises(RuntimeError, match="водяной знак"):
        _Harness(_WatermarkManager())._created_files_from_results([output], [], None)

    assert not output.exists()


def test_usage_counter_failure_deletes_output_instead_of_failing_open(tmp_path: Path):
    output = tmp_path / "generated.docx"
    output.write_bytes(b"document")

    with pytest.raises(RuntimeError, match="счётчик лицензии"):
        _Harness(_CounterFailureManager())._created_files_from_results([output], [], None)

    assert not output.exists()


def test_packaged_runtime_ignores_public_key_environment_override(monkeypatch):
    from product_access import native

    monkeypatch.setenv(native.PUBLIC_KEY_ENV, "attacker-controlled-key")
    monkeypatch.setattr(sys, "frozen", True, raising=False)

    assert native._verification_key() != "attacker-controlled-key"


def test_packaged_runtime_rejects_legacy_flat_paid_license(tmp_path: Path, monkeypatch):
    from product_access import sign_license_payload
    from product_access.native import NativeProductAccessManager, NativeLicenseError

    monkeypatch.setenv("DOKKOMPLEKT_LICENSE_DIR", str(tmp_path))
    monkeypatch.setenv("DOKKOMPLEKT_LICENSE_VERIFY_SECRET", "legacy-secret")
    monkeypatch.setattr(sys, "frozen", True, raising=False)
    payload = {
        "license_id": "LEGACY-SELF-SIGNED",
        "plan": "doctor_pro",
        "valid_until": datetime(2099, 1, 1, tzinfo=timezone.utc).isoformat(),
        "allowed_machines": [],
    }
    signed = sign_license_payload(payload, "legacy-secret")

    with pytest.raises(NativeLicenseError, match="only native Ed25519"):
        NativeProductAccessManager(now=datetime(2026, 9, 1, tzinfo=timezone.utc)).install_license_text(
            json.dumps(signed, ensure_ascii=False)
        )
