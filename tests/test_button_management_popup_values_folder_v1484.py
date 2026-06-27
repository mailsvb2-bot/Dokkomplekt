from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from actions_universal_flow import ActionsUniversalFlowMixin
from desktop_patient_folder import PrimaryPatientFolderInfo, build_patient_folder_name_from_info, folder_naming_uses_discharge_date
from universal_case_adapter import merge_case_values
from universal_fields import PatientCase
from universal_profiles import DocumentPack, DocumentTemplateSpec, remove_document_button, rename_document_button
from universal_template_engine import render_template_to_docx


class _Var:
    def __init__(self, value="") -> None:
        self.value = value

    def get(self):
        return self.value

    def set(self, value):
        self.value = value


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _spec(doc_id: str, label: str, *, template: str = "templates/t.docx") -> DocumentTemplateSpec:
    return DocumentTemplateSpec(
        id=doc_id,
        button_label=label,
        template=template,
        required_fields=("patient.fio", "case.number"),
        role_id="discharge_epicrisis",
    )


def test_doctor_created_button_can_be_renamed_without_losing_generation_contract() -> None:
    pack = DocumentPack(pack_id="doctor.test", name="Тест", documents=(_spec("d1", "Старое название"), _spec("d2", "Другая кнопка")))

    renamed = rename_document_button(pack, "d1", "  Выписной эпикриз новый  ")
    pack.documents = tuple(renamed if old.id == renamed.id else old for old in pack.documents)

    assert renamed.id == "d1"
    assert renamed.template == "templates/t.docx"
    assert renamed.required_fields == ("patient.fio", "case.number")
    assert renamed.button_label == "Выписной эпикриз новый"
    assert renamed.button_label_source == "doctor_renamed"
    assert pack.document_by_id("d1").button_label == "Выписной эпикриз новый"


def test_rename_keeps_button_labels_unique_inside_profile() -> None:
    pack = DocumentPack(pack_id="doctor.test", name="Тест", documents=(_spec("d1", "Акт"), _spec("d2", "Эпикриз")))

    renamed = rename_document_button(pack, "d1", "Эпикриз")

    assert renamed.button_label == "Эпикриз (2)"


def test_doctor_created_button_delete_removes_only_profile_entry_not_template_file(tmp_path: Path) -> None:
    template = tmp_path / "templates" / "t.docx"
    template.parent.mkdir()
    _write_docx(template, "{{patient.fio}}")
    pack = DocumentPack(pack_id="doctor.test", name="Тест", documents=(_spec("d1", "Удалить"), _spec("d2", "Оставить")))

    removed, kept = remove_document_button(pack, "d1")
    pack.documents = kept

    assert removed.button_label == "Удалить"
    assert [doc.id for doc in pack.documents] == ["d2"]
    assert template.exists(), "Удаление кнопки не должно уничтожать DOCX-шаблон врача безвозвратно"


class _OverlayHarness(ActionsUniversalFlowMixin):
    def __init__(self) -> None:
        self.patient_name_var = _Var("Иванов Иван Иванович")
        self.case_number_var = _Var("777")
        self.diagnosis_var = _Var("I10 Гипертензия")
        self._popup_diagnosis_override = ""
        self.assigned_treatment_var = _Var("Терапия из popup")
        self.admission_date_var = _Var("01.06.2026")
        self.discharge_date_var = _Var("12.06.2026")
        self.commission_date_var = _Var("10.06.2026")
        self.vk_date_var = _Var("11.06.2026")
        self.vk_protocol_date_var = _Var("11.06.2026")
        self.sick_leave_vk_date_var = _Var("09.06.2026")
        self.sick_leave_vk_protocol_date_var = _Var("09.06.2026")
        self.sick_leave_vk_commission_date_var = _Var("09.06.2026")
        self.expert_sick_leave_from_var = _Var("02.06.2026")
        self.labs_explicit_date_var = _Var("03.06.2026")
        self.labs_without_var = _Var(False)
        self.labs_text_var = _Var("Hb 140")
        self.labs_source_path_var = _Var("popup")
        self.labs_date_policy_var = _Var("preserve_found_dates")
        self.expert_work_status_var = _Var("да")
        self.expert_work_org_var = _Var("ООО Ромашка")
        self.expert_position_var = _Var("инженер")
        self.expert_sick_leave_needed_var = _Var("да")
        self.expert_sick_leave_number_var = _Var("ЛН-555")
        self.rvk_act_number_var = _Var("РВК-42")
        self.rvk_military_commissariat_var = _Var("Военкомат № 1")
        self.rvk_work_position_var = _Var("ООО Ромашка / инженер")
        self.commission_number_var = _Var("КОМ-33")
        self.vk_protocol_number_var = _Var("ВК-88")
        self.vk_mse_work_org_var = _Var("АО Место")
        self.vk_mse_position_var = _Var("мастер")
        self.sick_leave_vk_protocol_number_var = _Var("БЛ-99")
        self.sick_leave_vk_work_org_var = _Var("ИП Работа")
        self.sick_leave_vk_position_var = _Var("оператор")
        self.sick_leave_vk_work_position_var = _Var("ИП Работа / оператор")
        self._semantic_date_state = {}

    def _normalize_yes_no(self, value: str) -> str:
        return str(value or "").strip().lower()


def test_popup_entered_numeric_requisites_are_overlaid_into_custom_docx_render(tmp_path: Path) -> None:
    app = _OverlayHarness()
    values = app._confirmed_universal_overlay_values()
    case = merge_case_values(PatientCase(), values, source_document="popup_state")
    template = tmp_path / "numbers.docx"
    output = tmp_path / "out.docx"
    _write_docx(
        template,
        "ИБ {{caseNo}} | ЛН {{expertSickLeaveNumber}} | РВК {{rvk.act_number}} | "
        "ВК {{vk_mse.protocol_number}} | БЛ-ВК {{sick_leave_vk.protocol_number}} | Комиссия {{commission.number}}",
    )
    document = DocumentTemplateSpec(
        id="numbers",
        button_label="Проверка цифр",
        template=template.name,
        required_fields=(
            "case.number",
            "expert.sick_leave_number",
            "rvk.act_number",
            "vk_mse.protocol_number",
            "sick_leave_vk.protocol_number",
            "commission.number",
        ),
        role_id="vk_mse",
    )

    result = render_template_to_docx(template_path=template, output_path=output, case=case, document=document)

    assert result.ok
    text = _read_docx(output)
    assert "ИБ 777" in text
    assert "ЛН ЛН-555" in text
    assert "РВК РВК-42" in text
    assert "ВК ВК-88" in text
    assert "БЛ-ВК БЛ-99" in text
    assert "Комиссия КОМ-33" in text


def test_patient_subfolder_name_follows_confirmed_popup_principle_with_discharge_date() -> None:
    info = PrimaryPatientFolderInfo(
        fio="Петров Пётр Петрович",
        admission_date="01.06.2026",
        folder_name="legacy-default",
    )
    settings = {
        "parts": ["surname_initials", "admission_discharge_dates"],
        "date_format": "short",
        "doctor_confirmed": True,
    }

    assert folder_naming_uses_discharge_date(settings)
    assert build_patient_folder_name_from_info(info, settings=settings, discharge_date="12.06.2026") == "Петров П.П. 01.06.26-12.06.26"


def test_patient_subfolder_name_does_not_fall_back_to_old_default_when_custom_principle_exists() -> None:
    info = PrimaryPatientFolderInfo(
        fio="Сидоров Сергей Сергеевич",
        admission_date="05.05.2026",
        folder_name="Сидоров С.С. май 2026",
    )
    settings = {
        "parts": ["full_fio", "discharge_month"],
        "date_format": "long",
        "doctor_confirmed": True,
    }

    assert build_patient_folder_name_from_info(info, settings=settings, discharge_date="20.06.2026") == "Сидоров Сергей Сергеевич июнь 2026"
