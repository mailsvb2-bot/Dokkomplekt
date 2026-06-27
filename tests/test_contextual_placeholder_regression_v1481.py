from __future__ import annotations

from pathlib import Path

from docx import Document

from universal_case_adapter import patient_data_to_case
from universal_fields import normalize_field_id, normalize_field_id_for_context
from universal_profiles import DocumentPack, DocumentTemplateSpec
from universal_template_engine import (
    extract_template_placeholders,
    missing_required_fields,
    render_template_to_docx,
    validate_document_pack,
    validate_template,
)
from universal_main_documents import custom_requirement_flags_for_documents
from medical_models import PatientData


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def _read_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def test_global_ambiguous_aliases_stay_backward_compatible_but_context_routes_role_fields() -> None:
    assert normalize_field_id("Место работы") == "patient.work"
    assert normalize_field_id("Должность") == "patient.position"

    assert normalize_field_id_for_context("Место работы", role_id="vk_mse") == "vk_mse.work"
    assert normalize_field_id_for_context("Должность", role_id="vk_mse") == "vk_mse.position"
    assert normalize_field_id_for_context("Номер протокола", role_id="vk_mse") == "vk_mse.protocol_number"
    assert normalize_field_id_for_context("Протокол N", role_id="vk_mse") == "vk_mse.protocol_number"
    assert normalize_field_id_for_context("Номер", role_id="vk_mse") == "vk_mse.protocol_number"
    assert normalize_field_id_for_context("Дата протокола", role_id="vk_mse") == "vk_mse.protocol_date"
    assert normalize_field_id_for_context("От", role_id="vk_mse") == "vk_mse.protocol_date"
    assert normalize_field_id_for_context("Дата проведения ВК", role_id="vk_mse") == "vk_mse.date"

    assert normalize_field_id_for_context("Место работы", role_id="sickLeaveVk") == "sick_leave_vk.work"
    assert normalize_field_id_for_context("Должность", role_id="sickLeaveVk") == "sick_leave_vk.position"
    assert normalize_field_id_for_context("Номер протокола", role_id="sickLeaveVk") == "sick_leave_vk.protocol_number"
    assert normalize_field_id_for_context("Дата протокола", role_id="sickLeaveVk") == "sick_leave_vk.protocol_date"
    assert normalize_field_id_for_context("Дата проведения ВК", role_id="sickLeaveVk") == "sick_leave_vk.date"

    assert normalize_field_id_for_context("Номер", role_id="jointMedicalExam") == "commission.number"
    assert normalize_field_id_for_context("Дата комиссии", role_id="jointMedicalExam") == "commission.date"


def test_contextual_docx_placeholder_extraction_for_vk_and_sick_leave_roles(tmp_path: Path) -> None:
    vk_template = tmp_path / "vk.docx"
    _write_docx(vk_template, "{{Место работы}} {{Должность}} {{Номер протокола}} {{Дата протокола}} {{Дата проведения ВК}}")
    fields = [item.field_id for item in extract_template_placeholders(vk_template, role_id="vk_mse", button_label="ВК на МСЭ")]
    assert fields == [
        "vk_mse.work",
        "vk_mse.position",
        "vk_mse.protocol_number",
        "vk_mse.protocol_date",
        "vk_mse.date",
    ]

    sick_template = tmp_path / "sick.docx"
    _write_docx(sick_template, "{{Место работы}} {{Должность}} {{Номер протокола}} {{Дата протокола}} {{Дата проведения ВК}}")
    fields = [item.field_id for item in extract_template_placeholders(sick_template, role_id="sickLeaveVk", button_label="ВК больничный")]
    assert fields == [
        "sick_leave_vk.work",
        "sick_leave_vk.position",
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.protocol_date",
        "sick_leave_vk.date",
    ]


def test_contextual_render_uses_document_role_not_global_patient_work(tmp_path: Path) -> None:
    template = tmp_path / "vk_render.docx"
    out = tmp_path / "rendered.docx"
    _write_docx(template, "{{Место работы}} | {{Должность}} | {{Номер протокола}} | {{Дата протокола}}")

    data = PatientData()
    data.output_fio = "Иванов Иван Иванович"
    data.work_org = "Пациентская работа"
    data.position = "Пациентская должность"
    data.vk_mse_work_org = "Организация ВК"
    data.vk_mse_position = "Должность ВК"
    data.vk_protocol_number = "П-77"
    data.vk_protocol_date = "14.06.2026"
    case = patient_data_to_case(data)
    document = DocumentTemplateSpec(
        id="vk_doc",
        button_label="ВК на МСЭ",
        template=template.name,
        required_fields=("Место работы", "Должность", "Номер протокола", "Дата протокола"),
        role_id="vk_mse",
    )

    assert missing_required_fields(case, document) == ()
    result = render_template_to_docx(template_path=template, output_path=out, case=case, document=document)

    assert result.ok
    text = _read_text(out)
    assert "Организация ВК | Должность ВК | П-77 | 14.06.2026" in text
    assert "Пациентская работа" not in text
    assert "Пациентская должность" not in text


def test_profile_loading_and_validation_preserve_contextual_required_fields(tmp_path: Path) -> None:
    template = tmp_path / "commission.docx"
    _write_docx(template, "{{Номер}} {{Дата комиссии}}")

    document = DocumentTemplateSpec.from_dict({
        "id": "commission_doc",
        "button_label": "Совместный осмотр",
        "template": template.name,
        "required_fields": ["Номер", "Дата комиссии"],
        "role_id": "jointMedicalExam",
        "category": "medical",
    })
    assert document.required_fields == ("commission.number", "commission.date")

    validation = validate_template(
        template,
        required_fields=document.required_fields,
        role_id=document.role_id,
        category=document.category,
        button_label=document.button_label,
    )
    assert validation.ok
    assert validation.missing_required_placeholders == ()

    pack = DocumentPack(pack_id="doctor.test", name="test", documents=(document,))
    pack_validation = validate_document_pack(pack, base_dir=tmp_path)
    assert pack_validation.errors == ()
    assert pack_validation.ok


def test_custom_requirement_flags_use_contextual_fields_for_regulatory_roles() -> None:
    vk = DocumentTemplateSpec(
        id="vk_doc",
        button_label="ВК на МСЭ",
        template="templates/vk.docx",
        required_fields=("Место работы", "Должность", "Номер протокола", "Дата протокола"),
        role_id="vk_mse",
    )
    sick = DocumentTemplateSpec(
        id="sick_doc",
        button_label="ВК больничный",
        template="templates/sick.docx",
        required_fields=("Место работы", "Должность", "Номер протокола", "Дата протокола"),
        role_id="sickLeaveVk",
    )

    flags = custom_requirement_flags_for_documents([vk, sick])

    assert flags["vk_mse"] is True
    assert flags["sick_leave_vk"] is True
    assert flags["regular"] is True
