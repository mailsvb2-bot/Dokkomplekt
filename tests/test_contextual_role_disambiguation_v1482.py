from __future__ import annotations

from pathlib import Path

from docx import Document

from universal_fields import normalize_field_id, normalize_field_id_for_context
from universal_profiles import DocumentTemplateSpec
from universal_template_engine import extract_template_placeholders, missing_required_fields, render_template_to_docx, validate_template
from universal_fields import PatientCase


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def test_context_role_wins_for_ambiguous_vk_protocol_labels() -> None:
    # Globally this label is ambiguous and may be associated with VK/MSE by the
    # registry, but inside a sick-leave VK document it must not fill the MSE field.
    assert normalize_field_id("Номер протокола ВК") == "vk_mse.protocol_number"
    assert normalize_field_id_for_context("Номер протокола ВК", role_id="sickLeaveVk") == "sick_leave_vk.protocol_number"
    assert normalize_field_id_for_context("Протокол No", role_id="sickLeaveVk") == "sick_leave_vk.protocol_number"
    assert normalize_field_id_for_context("Дата от протокола", role_id="sickLeaveVk") == "sick_leave_vk.protocol_date"
    assert normalize_field_id_for_context("От протокола", role_id="vk_mse") == "vk_mse.protocol_date"


def test_vk_mse_role_accepts_generic_commission_date_without_becoming_joint_exam() -> None:
    assert normalize_field_id_for_context("Дата комиссии", role_id="vk_mse") == "vk_mse.date"
    assert normalize_field_id_for_context("Дата проведения комиссии", role_id="vk_mse") == "vk_mse.date"
    assert normalize_field_id_for_context("Дата комиссии", role_id="jointMedicalExam") == "commission.date"
    assert normalize_field_id_for_context("Дата проведения комиссии", role_id="sickLeaveVk") == "sick_leave_vk.commission_date"


def test_rvk_context_owns_work_position_slash_label() -> None:
    # The slash spelling is common in hand-written templates.  It used to be
    # captured by the sick-leave VK alias even when the document role was RVK.
    assert normalize_field_id_for_context("Место работы / должность", role_id="rvk") == "rvk.work_position"
    assert normalize_field_id_for_context("Работа и должность", role_id="rvk") == "rvk.work_position"


def test_docx_placeholder_extraction_uses_contextual_role_before_global_registry(tmp_path: Path) -> None:
    sick_template = tmp_path / "sick_vk.docx"
    _write_docx(sick_template, "{{Номер протокола ВК}} {{Протокол No}} {{Дата от протокола}} {{Дата проведения комиссии}}")
    sick_fields = [item.field_id for item in extract_template_placeholders(sick_template, role_id="sickLeaveVk", button_label="ВК больничный")]
    assert sick_fields == [
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.protocol_date",
        "sick_leave_vk.commission_date",
    ]

    rvk_template = tmp_path / "rvk.docx"
    _write_docx(rvk_template, "{{Место работы / должность}}")
    rvk_fields = [item.field_id for item in extract_template_placeholders(rvk_template, role_id="rvk", button_label="Акт для РВК")]
    assert rvk_fields == ["rvk.work_position"]


def test_contextual_required_fields_render_without_wrong_role_values(tmp_path: Path) -> None:
    template = tmp_path / "sick_vk_render.docx"
    output = tmp_path / "out.docx"
    _write_docx(template, "{{Номер протокола ВК}} | {{Протокол No}} | {{Дата от протокола}} | {{Дата проведения комиссии}}")

    document = DocumentTemplateSpec.from_dict({
        "id": "sick_vk",
        "button_label": "ВК больничный",
        "template": template.name,
        "role_id": "sickLeaveVk",
        "required_fields": ["Номер протокола ВК", "Протокол No", "Дата от протокола", "Дата проведения комиссии"],
    })
    assert document.required_fields == (
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.protocol_date",
        "sick_leave_vk.commission_date",
    )
    validation = validate_template(template, required_fields=document.required_fields, role_id=document.role_id, button_label=document.button_label)
    assert validation.ok

    case = PatientCase()
    case.set("vk_mse.protocol_number", "НЕ ТОТ ПРОТОКОЛ")
    case.set("sick_leave_vk.protocol_number", "БЛ-12")
    case.set("sick_leave_vk.protocol_date", "20.06.2026")
    case.set("sick_leave_vk.commission_date", "21.06.2026")

    assert missing_required_fields(case, document) == ()
    result = render_template_to_docx(template_path=template, output_path=output, case=case, document=document)
    assert result.ok
    text = _read_docx(output)
    assert "БЛ-12 | БЛ-12 | 20.06.2026 | 21.06.2026" in text
    assert "НЕ ТОТ ПРОТОКОЛ" not in text

from types import SimpleNamespace

from actions_universal_flow import ActionsUniversalFlowMixin
from universal_case_adapter import merge_case_values, merge_patient_cases


class _Var:
    def __init__(self, value: str = "") -> None:
        self.value = value

    def get(self) -> str:
        return self.value


class _UniversalFake(ActionsUniversalFlowMixin):
    def __init__(self) -> None:
        self.patient_name_var = _Var("Иванов Иван Иванович")
        self.case_number_var = _Var("UI-777")
        self.diagnosis_var = _Var("I10 Гипертензия")
        self.assigned_treatment_var = _Var("Лечение из UI")
        self.labs_text_var = _Var("ОАК без особенностей")
        self.labs_source_path_var = _Var("manual")
        self.labs_date_policy_var = _Var("preserve_found_dates")
        self.labs_without_var = _Var("")
        self.expert_work_status_var = _Var("да")
        self.expert_work_org_var = _Var("Организация из UI")
        self.expert_position_var = _Var("Должность из UI")
        self.expert_sick_leave_needed_var = _Var("да")
        self.expert_sick_leave_number_var = _Var("ЛН-UI")
        self.rvk_act_number_var = _Var("РВК-UI")
        self.rvk_military_commissariat_var = _Var("Военкомат UI")
        self.rvk_work_position_var = _Var("Работа/должность UI")
        self.commission_number_var = _Var("К-UI")
        self.vk_protocol_number_var = _Var("ВК-UI")
        self.vk_mse_work_org_var = _Var("ВК организация UI")
        self.vk_mse_position_var = _Var("ВК должность UI")
        self.sick_leave_vk_protocol_number_var = _Var("БЛ-UI")
        self.sick_leave_vk_work_org_var = _Var("БЛ организация UI")
        self.sick_leave_vk_position_var = _Var("БЛ должность UI")
        self.sick_leave_vk_work_position_var = _Var("БЛ работа/должность UI")
        self._popup_diagnosis_override = ""
        self._dates = {
            "admission_date": "10.06.2026",
            "discharge_date": "20.06.2026",
            "labs_explicit_date": "11.06.2026",
            "expert_sick_leave_from": "12.06.2026",
            "vk_date": "13.06.2026",
            "vk_protocol_date": "14.06.2026",
            "sick_leave_vk_date": "15.06.2026",
            "sick_leave_vk_protocol_date": "16.06.2026",
            "sick_leave_vk_commission_date": "17.06.2026",
            "commission_date": "18.06.2026",
        }

    def _normalize_yes_no(self, value: str) -> str:
        return "да" if str(value).strip().lower() in {"да", "yes", "true"} else "нет"


def test_doctor_confirmed_ui_values_are_final_overlay_for_custom_case(monkeypatch) -> None:
    fake = _UniversalFake()

    def fake_current_semantic_date(_app, key):
        return fake._dates.get(key, "")

    import actions_universal_flow
    monkeypatch.setattr(actions_universal_flow, "current_semantic_date", fake_current_semantic_date)

    source_case = PatientCase()
    source_case.set("case.number", "OLD-001", confidence=0.93, source_document="source_scan")
    source_case.set("diagnosis.main", "F00 Старый диагноз", confidence=0.99, source_document="source_scan")
    source_case.set("treatment.plan", "Старое лечение", confidence=0.99, source_document="source_scan")

    parsed_case = PatientCase()
    parsed_case.set("case.number", "UI-777", confidence=0.90, source_document="override_data")
    parsed_case.set("diagnosis.main", "I10 Гипертензия", confidence=0.90, source_document="override_data")
    parsed_case.set("treatment.plan", "Лечение из UI", confidence=0.90, source_document="override_data")

    # This reproduces the regression: high-confidence profile scanning could win
    # over the same values already corrected in UI/popup state.
    regressed = merge_patient_cases(parsed_case, source_case)
    assert regressed.get("case.number") == "OLD-001"
    assert regressed.get("diagnosis.main") == "F00 Старый диагноз"

    fixed = merge_case_values(regressed, fake._confirmed_universal_overlay_values(), source_document="doctor_confirmed_ui_state")
    assert fixed.get("case.number") == "UI-777"
    assert fixed.get("diagnosis.main") == "I10 Гипертензия"
    assert fixed.get("treatment.plan") == "Лечение из UI"
    assert fixed.get("admission.date") == "10.06.2026"
    assert fixed.get("discharge.date") == "20.06.2026"
    assert fixed.get("labs.results") == "ОАК без особенностей"
    assert fixed.get("expert.sick_leave_number") == "ЛН-UI"
    assert fixed.get("rvk.act_number") == "РВК-UI"
    assert fixed.get("vk_mse.protocol_number") == "ВК-UI"
    assert fixed.get("sick_leave_vk.protocol_number") == "БЛ-UI"
