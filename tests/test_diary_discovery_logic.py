from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from diary_template_discovery import DiaryTemplateDiscoveryMixin


class Discovery(DiaryTemplateDiscoveryMixin):
    def __init__(self) -> None:
        self._diary_template_files_cache = {}
        self._diary_template_day_cache = {}
        self._diary_template_folder_contains_cache = {}


def _docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Дневник наблюдения")
    doc.add_paragraph("2 15 состояние стабильное")
    doc.save(path)
    return path


def test_template_filename_day_accepts_numbered_docx(tmp_path: Path) -> None:
    path = _docx(tmp_path / "15.docx")
    assert Discovery._template_filename_day(path) == 15


def test_folder_contains_numbered_templates_by_content(tmp_path: Path) -> None:
    path = _docx(tmp_path / "нестандартное имя.docx")
    assert path.exists()
    assert Discovery()._folder_contains_numbered_diary_templates(tmp_path) is True
