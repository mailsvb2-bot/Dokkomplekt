from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from diary_batch import default_observation_diary_dates, fill_diary_batch, is_non_working_day
from diary_text_parser import extract_statuses_from_docx
from medical_docx_xml_fragments import SUPPORTED_WORD_SUFFIXES, is_supported_word_file


def _paragraph_text(path: Path) -> str:
    return "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs if paragraph.text.strip())


def test_diary_single_route_ignores_calendar_table_file(tmp_path: Path) -> None:
    status_docx = tmp_path / "texts.docx"
    source = Document()
    source.add_paragraph("01.06.2026 First diary text is long enough for extraction and output.")
    source.add_paragraph("02.06.2026 Second diary text is long enough for extraction and output.")
    source.add_paragraph("02.06.2026 Second diary text is long enough for extraction and output.")
    source.save(status_docx)
    assert len(extract_statuses_from_docx(status_docx)) == 2

    table_docx = tmp_path / "calendar_table.docx"
    template = Document()
    table = template.add_table(rows=1, cols=4)
    for index, header in enumerate(["n", "day", "month", "text"]):
        table.rows[0].cells[index].text = header
    for day in [10, 11, 12, 13, 14, 15]:
        row = table.add_row()
        row.cells[0].text = str(day)
        row.cells[1].text = str(day)
        row.cells[2].text = ""
        row.cells[3].text = "OLD_TABLE_CONTENT_SHOULD_NOT_BE_USED"
    template.save(table_docx)

    result = fill_diary_batch(
        status_files=[status_docx],
        diary_files=[table_docx],
        output_dir=tmp_path / "out",
        patient_name="Ivanova Irina",
        gender_source_name="Ivanova Irina",
        admission_value="10.06.2026",
        discharge_value="12.06.2026",
        repeat_statuses=True,
        force_final_diary=True,
    )

    assert result.processed_files == 1
    assert result.final_rows_filled == 1
    assert result.removed_after_discharge_rows == 3
    assert result.month_cells_filled == 0
    text = _paragraph_text(Path(result.created_files[0]))
    assert "11.06.26 First diary text" in text
    assert "12.06.26 Состояние улучшилось" in text
    assert "OLD_TABLE_CONTENT_SHOULD_NOT_BE_USED" not in text
    assert "13.06.26" not in text


def test_diary_filler_supports_merged_status_cells_once(tmp_path: Path) -> None:
    merged_status = tmp_path / "merged_status.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    merged_cell = table.cell(0, 0).merge(table.cell(0, 1))
    merged_cell.text = "Merged diary text is long enough for strict extraction and appears once."
    doc.save(merged_status)
    assert extract_statuses_from_docx(merged_status) == ["Merged diary text is long enough for strict extraction and appears once."]


def test_word_format_contract_accepts_doc_docx_docm() -> None:
    assert {".doc", ".docx", ".docm"}.issubset(SUPPORTED_WORD_SUFFIXES)
    assert is_supported_word_file("source.doc")
    assert is_supported_word_file("source.docx")
    assert is_supported_word_file("source.docm")
    assert not is_supported_word_file("source.pdf")


def test_default_diary_calendar_skips_weekends_and_fixed_holidays() -> None:
    dates = default_observation_diary_dates(date(2026, 1, 1), limit=8)
    assert len(dates) == 8
    assert dates == tuple(dict.fromkeys(dates))
    assert all(not is_non_working_day(item) for item in dates)
    assert dates[0] >= date(2026, 1, 12)
