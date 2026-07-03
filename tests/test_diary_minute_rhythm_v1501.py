from __future__ import annotations

from pathlib import Path

from docx import Document

from diary_batch import fill_diary_batch
from diary_schedule import diary_calendar_schedule_from_choice, diary_minute_schedule_from_choice, expand_minute_intervals


def _paragraph_text(path: Path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs if p.text.strip())


def test_minute_rhythm_choices_are_parsed() -> None:
    assert diary_minute_schedule_from_choice("каждые 4 часа").minute_offsets == (240,)
    assert diary_minute_schedule_from_choice("каждый час").minute_offsets == (60,)
    assert diary_minute_schedule_from_choice("30 минут").minute_offsets == (30,)
    assert diary_minute_schedule_from_choice("15 минут").minute_offsets == (15,)
    assert diary_minute_schedule_from_choice("5 минут").minute_offsets == (5,)
    assert diary_minute_schedule_from_choice("45 минут").minute_offsets == (45,)
    assert expand_minute_intervals((30,), 4) == (30, 60, 90, 120)


def test_daily_style_can_generate_intraday_entries(tmp_path: Path) -> None:
    status_docx = tmp_path / "texts.docx"
    doc = Document()
    doc.add_paragraph("Neutral diary text one for generated timed output.")
    doc.add_paragraph("Neutral diary text two for generated timed output.")
    doc.save(status_docx)

    day_spec = diary_calendar_schedule_from_choice("1")
    rhythm = diary_minute_schedule_from_choice("каждые 4 часа")
    result = fill_diary_batch(
        status_files=[status_docx],
        diary_files=[],
        output_dir=tmp_path / "out",
        patient_name="Ivanova Irina",
        gender_source_name="Ivanova Irina",
        admission_value="10.06.2026 08:00",
        discharge_value="11.06.2026",
        repeat_statuses=True,
        force_final_diary=False,
        diary_day_offsets=day_spec.day_offsets,
        diary_minute_offsets=rhythm.minute_offsets,
        diary_frequency_mode="hourly",
    )
    text = _paragraph_text(Path(result.created_files[0]))
    assert "11.06.26 08:00" in text
    assert "11.06.26 12:00" in text
    assert "11.06.26 16:00" in text
