from __future__ import annotations

from pathlib import Path

from docx import Document

from diary_batch import fill_diary_batch
from app import RegressionStateOverlayMixin


def _make_status_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет.")
    doc.save(path)
    return path


def _make_diary_template(path: Path) -> Path:
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Число"
    table.cell(0, 1).text = "Месяц"
    table.cell(0, 2).text = "Дневник"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = ""
    table.cell(2, 2).text = ""
    doc.save(path)
    return path


def test_selected_diary_template_is_filled_as_table_not_replaced_by_text_docx(tmp_path):
    status = _make_status_doc(tmp_path / "texts.docx")
    template = _make_diary_template(tmp_path / "01.docx")
    out = tmp_path / "out"

    result = fill_diary_batch(
        status_files=[status],
        diary_files=[template],
        output_dir=out,
        patient_name="Иванов И.И.",
        admission_value="01.06.2026",
        discharge_value="03.06.2026",
        repeat_statuses=True,
        text_output=False,
        write_report=False,
    )

    assert result.processed_files == 1
    assert len(result.created_files) == 1
    created = Document(str(result.created_files[0]))
    assert created.tables, "The Dates template table must survive; text fallback creates no tables."
    text = "\n".join(cell.text for table in created.tables for row in table.rows for cell in row.cells)
    assert "Состояние стабильное" in text
    assert "01" in text or "1" in text


class _AccessEstimateProbe(RegressionStateOverlayMixin):
    pass


def test_product_access_estimate_counts_selected_diary_templates():
    probe = _AccessEstimateProbe()
    probe.diary_files = ["01.docx", "02.docx", "03.docx"]

    assert probe._estimate_selected_document_count([], True, []) == 3
    assert probe._estimate_selected_document_count(["primary"], True, ["custom"]) == 5
