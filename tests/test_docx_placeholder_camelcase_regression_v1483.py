from __future__ import annotations

from pathlib import Path

from docx import Document

from universal_template_engine import extract_template_placeholders, normalize_placeholder_id, render_template_to_docx, validate_template
from universal_profiles import DocumentTemplateSpec
from universal_fields import PatientCase


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def test_docx_placeholder_normalizer_preserves_camelcase_semantics() -> None:
    aliases = {
        "patientName": "patient.fio",
        "fullName": "patient.fio",
        "caseNo": "case.number",
        "medicalRecordNo": "case.number",
        "mainDiagnosis": "diagnosis.main",
        "sickLeaveFrom": "expert.sick_leave_from",
        "expertSickLeaveNumber": "expert.sick_leave_number",
        "workPlace": "patient.work",
        "jobTitle": "patient.position",
    }
    for raw, expected in aliases.items():
        assert normalize_placeholder_id(raw) == expected


def test_extract_template_placeholders_accepts_export_camelcase_docx_fields(tmp_path: Path) -> None:
    template = tmp_path / "camelcase_template.docx"
    _write_docx(
        template,
        "{{patientName}} {{caseNo}} {{medicalRecordNo}} {{mainDiagnosis}} "
        "{{sickLeaveFrom}} {{expertSickLeaveNumber}} {{workPlace}} {{jobTitle}}",
    )

    fields = [item.field_id for item in extract_template_placeholders(template)]

    assert fields == [
        "patient.fio",
        "case.number",
        "case.number",
        "diagnosis.main",
        "expert.sick_leave_from",
        "expert.sick_leave_number",
        "patient.work",
        "patient.position",
    ]


def test_camelcase_docx_template_validates_and_renders_from_canonical_state(tmp_path: Path) -> None:
    template = tmp_path / "render_camelcase.docx"
    output = tmp_path / "out.docx"
    _write_docx(template, "{{patientName}} | {{caseNo}} | {{mainDiagnosis}} | {{expertSickLeaveNumber}}")

    document = DocumentTemplateSpec.from_dict({
        "id": "camelcase_doc",
        "button_label": "Custom export template",
        "template": template.name,
        "required_fields": ["patientName", "caseNo", "mainDiagnosis", "expertSickLeaveNumber"],
    })

    validation = validate_template(template, required_fields=document.required_fields, button_label=document.button_label)
    assert validation.ok

    case = PatientCase()
    case.set("patient.fio", "Иванов Иван Иванович")
    case.set("case.number", "ИБ-42")
    case.set("diagnosis.main", "I10 Гипертензия")
    case.set("expert.sick_leave_number", "ЛН-77")

    result = render_template_to_docx(template_path=template, output_path=output, case=case, document=document)
    assert result.ok
    assert _read_docx(output) == "Иванов Иван Иванович | ИБ-42 | I10 Гипертензия | ЛН-77"

from universal_template_engine import attach_template_to_pack, infer_document_spec_from_template
from universal_profiles import DocumentPack


def test_explicit_role_context_is_used_while_attaching_template(tmp_path: Path) -> None:
    template = tmp_path / "ambiguous_sick_vk.docx"
    profile_dir = tmp_path / "profile"
    _write_docx(template, "{{Дата комиссии}} {{Номер протокола ВК}}")

    # Reproduces the regression: the UI may know the role from button mapping,
    # even when title auto-detection is weak. Placeholder extraction must use
    # that explicit role before saving required_fields into pack.json.
    pack = DocumentPack(pack_id="doctor.test", name="Test")
    spec, _copied = attach_template_to_pack(
        pack,
        template,
        profile_dir,
        button_label="ВК больничный",
        role_id="sickLeaveVk",
    )

    assert spec.role_id == "sickLeaveVk"
    assert spec.required_fields == ("sick_leave_vk.commission_date", "sick_leave_vk.protocol_number")
    assert pack.documents[0].required_fields == spec.required_fields


def test_infer_document_spec_accepts_explicit_role_before_placeholder_extraction(tmp_path: Path) -> None:
    template = tmp_path / "ambiguous_rvk.docx"
    _write_docx(template, "{{Место работы / должность}}")

    spec = infer_document_spec_from_template(
        template,
        button_label="Акт для РВК",
        role_id="rvk",
    )

    assert spec.role_id == "rvk"
    assert spec.required_fields == ("rvk.work_position",)
