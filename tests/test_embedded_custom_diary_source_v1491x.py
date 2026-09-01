from pathlib import Path

from docx import Document

from actions_universal_flow import ActionsUniversalFlowMixin
from universal_profiles import default_document_pack
from universal_template_engine import attach_template_to_pack


class _Harness(ActionsUniversalFlowMixin):
    def __init__(self, profile_path: Path):
        self.status_files = []
        self.profile_path = profile_path
        self.chooser_calls = 0

    def _universal_profile_path(self):
        return self.profile_path

    def _auto_select_diary_text_by_diagnosis(self, *, ask_folder=False):
        return False

    def choose_status_files(self):
        self.chooser_calls += 1


def test_embedded_doctor_diary_does_not_force_duplicate_text_chooser(tmp_path: Path):
    profile_path = tmp_path / "profiles" / "doctor.medpack.json"
    profile_path.parent.mkdir(parents=True)
    source = tmp_path / "Дневники.docx"
    doc = Document()
    doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет.")
    doc.add_paragraph("Контактен, ориентирован, назначения выполняет.")
    doc.save(source)

    pack = default_document_pack()
    attach_template_to_pack(
        pack,
        source,
        profile_path.parent,
        button_label="Дневники наблюдения",
        document_id="doctor_diary",
        category="diaries",
        role_id="daily_diary",
    )

    harness = _Harness(profile_path)
    harness._ensure_diary_text_files_for_creation(pack, ["doctor_diary"])

    assert harness.status_files == []
    assert harness.chooser_calls == 0
