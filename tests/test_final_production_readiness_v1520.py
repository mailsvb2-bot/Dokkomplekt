from __future__ import annotations

from datetime import datetime, timezone
import json
import os
from pathlib import Path
import subprocess
import sys
import types
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pytest

from medical_paths import atomic_write_json
from diary_schedule import diary_minute_schedule_from_choice
from document_intelligence.form_fill import fill_docx_visible_fields, visible_fill_field_ids
from document_output_format import export_docx_to_pdf
from medical_paths import interprocess_file_lock
from output_transaction import OutputTransaction
from product_access import ProductAccessManager
from universal_fields import PatientCase
from universal_generation import render_documents_from_pack
from universal_profiles import DocumentPack, DocumentTemplateSpec, default_document_pack, save_document_pack
from universal_template_engine import export_document_pack_zip, import_document_pack_zip, infer_document_spec_from_template, render_template_to_docx


def _docx(path: Path, text: str = "Шаблон") -> Path:
    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return path


def _case(**values: str) -> PatientCase:
    case = PatientCase()
    for field_id, value in values.items():
        case.set(field_id.replace("__", "."), value)
    return case


def test_visible_fullwidth_colon_and_no_colon_are_fillable(tmp_path: Path) -> None:
    for name, text in (("wide.docx", "ФИО： ______"), ("space.docx", "ФИО ______")):
        path = _docx(tmp_path / name, text)
        assert visible_fill_field_ids(path) == ("patient.fio",)
        assert fill_docx_visible_fields(path, {"patient.fio": "Иванов Иван"}) == ("patient.fio",)
        assert "Иванов Иван" in Document(path).paragraphs[0].text


def test_visible_suffix_is_preserved(tmp_path: Path) -> None:
    path = _docx(tmp_path / "amount.docx", "Сумма: ______ руб.")
    (field_id,) = visible_fill_field_ids(path)
    fill_docx_visible_fields(path, {field_id: "1250"})
    assert Document(path).paragraphs[0].text == "Сумма: 1250 руб."


def test_header_footer_and_same_cell_table_are_fillable(tmp_path: Path) -> None:
    path = tmp_path / "stories.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "ФИО: ______"
    document.sections[0].footer.paragraphs[0].text = "Диагноз: ______"
    document.add_table(rows=1, cols=1).cell(0, 0).text = "ФИО: ______"
    document.save(path)
    assert set(visible_fill_field_ids(path)) == {"patient.fio", "diagnosis.main"}
    fill_docx_visible_fields(path, {"patient.fio": "Иванов", "diagnosis.main": "F20.0"})
    rendered = Document(path)
    assert "Иванов" in rendered.sections[0].header.paragraphs[0].text
    assert "F20.0" in rendered.sections[0].footer.paragraphs[0].text
    assert rendered.tables[0].cell(0, 0).text == "ФИО: Иванов"


def test_nested_table_is_fillable(tmp_path: Path) -> None:
    path = tmp_path / "nested.docx"
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "ФИО"
    nested.cell(0, 1).text = "______"
    document.save(path)
    assert visible_fill_field_ids(path) == ("patient.fio",)
    fill_docx_visible_fields(path, {"patient.fio": "Иванов"})
    rendered = Document(path)
    assert rendered.tables[0].cell(0, 0).tables[0].cell(0, 1).text == "Иванов"


def test_raw_textbox_paragraph_is_fillable(tmp_path: Path) -> None:
    path = tmp_path / "textbox.docx"
    document = Document()
    host = document.add_paragraph("Host")
    content = OxmlElement("w:txbxContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "ФИО: ______"
    run.append(text); paragraph.append(run); content.append(paragraph); host._p.append(content)
    document.save(path)
    assert visible_fill_field_ids(path) == ("patient.fio",)
    fill_docx_visible_fields(path, {"patient.fio": "Иванов"})
    values = [node.text for node in Document(path).element.body.iter(qn("w:t"))]
    assert "ФИО: Иванов" in values


def test_run_formatting_survives_visible_fill(tmp_path: Path) -> None:
    path = tmp_path / "format.docx"
    document = Document()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("ФИО: "); label.bold = True
    blank = paragraph.add_run("______"); blank.italic = True
    document.save(path)
    fill_docx_visible_fields(path, {"patient.fio": "Иванов"})
    runs = Document(path).paragraphs[0].runs
    assert runs[0].text == "ФИО: " and runs[0].bold is True
    assert runs[1].text == "Иванов" and runs[1].italic is True


def test_inferred_visible_fields_are_strict_required_fields(tmp_path: Path) -> None:
    template = tmp_path / "header.docx"
    document = Document(); document.sections[0].header.paragraphs[0].text = "ФИО: ______"; document.save(template)
    spec = infer_document_spec_from_template(template, button_label="Осмотр")
    assert "patient.fio" in spec.required_fields
    output = tmp_path / "out.docx"
    with pytest.raises(ValueError):
        render_template_to_docx(template_path=template, output_path=output, case=PatientCase(), document=spec, strict=True)
    assert not output.exists()
    render_template_to_docx(template_path=template, output_path=output, case=_case(patient__fio="Иванов"), document=spec, strict=True)
    assert "Иванов" in Document(output).sections[0].header.paragraphs[0].text


def test_empty_or_unknown_selection_does_not_create_output_directory(tmp_path: Path) -> None:
    template = _docx(tmp_path / "template.docx", "ФИО: ______")
    spec = infer_document_spec_from_template(template, button_label="Осмотр")
    pack = DocumentPack("p", "P", documents=(spec,))
    for ids in ([], ["unknown"]):
        out = tmp_path / ("empty" if not ids else "unknown")
        result = render_documents_from_pack(pack=pack, case=_case(patient__fio="Иванов"), document_ids=ids, output_dir=out, base_dir=tmp_path, strict=True)
        assert not result.created_files
        assert result.skipped_documents
        assert not out.exists()


def test_export_preserves_distinct_same_basename_templates(tmp_path: Path) -> None:
    left = tmp_path / "a"; right = tmp_path / "b"; left.mkdir(); right.mkdir()
    p1 = _docx(left / "same.docx", "ФИО: ______")
    p2 = _docx(right / "same.docx", "Диагноз: ______")
    pack = DocumentPack("p", "P", documents=(
        DocumentTemplateSpec("a", "A", str(p1), required_fields=("patient.fio",)),
        DocumentTemplateSpec("b", "B", str(p2), required_fields=("diagnosis.main",)),
    ))
    archive = export_document_pack_zip(pack, tmp_path / "profile.medpack.zip")
    with zipfile.ZipFile(archive) as zf:
        manifest = json.loads(zf.read("pack.json"))
        template_refs = [item["template"] for item in manifest["documents"]]
        members = [name for name in zf.namelist() if name.startswith("templates/")]
    assert len(set(template_refs)) == 2
    assert len(members) == 2


def test_failed_import_rolls_back_copied_templates_and_manifest(tmp_path: Path) -> None:
    good = _docx(tmp_path / "good.docx", "ФИО: ______")
    pack = DocumentPack("p", "P", documents=(
        DocumentTemplateSpec("good", "Good", "templates/good.docx", required_fields=("patient.fio",)),
        DocumentTemplateSpec("missing", "Missing", "templates/missing.docx"),
    ))
    archive = tmp_path / "broken.medpack.zip"
    with zipfile.ZipFile(archive, "w") as zf:
        zf.writestr("pack.json", json.dumps(pack.to_dict(), ensure_ascii=False))
        zf.write(good, "templates/good.docx")
    target = tmp_path / "imported"
    with pytest.raises(ValueError):
        import_document_pack_zip(archive, target)
    assert not any(path.is_file() for path in target.rglob("*"))


def test_pdf_rejects_missing_source_before_word_automation(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        export_docx_to_pdf(tmp_path / "missing.docx")
    assert not (tmp_path / "missing.pdf").exists()


def test_pdf_is_readonly_no_recent_and_partial_file_is_removed(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = _docx(tmp_path / "source.docx")
    target = tmp_path / "source.pdf"
    calls: dict[str, object] = {}
    class FakeDocument:
        def ExportAsFixedFormat(self, **kwargs):
            calls["export"] = kwargs; target.write_bytes(b"partial"); raise RuntimeError("Word failed")
        def Close(self, save): calls["close"] = save
    class FakeDocuments:
        def Open(self, **kwargs): calls["open"] = kwargs; return FakeDocument()
    class FakeWord:
        def __init__(self): self.Documents = FakeDocuments(); self.Visible = True; self.DisplayAlerts = 99
        def Quit(self): calls["quit"] = True
    word = FakeWord()
    client = types.ModuleType("win32com.client"); client.DispatchEx = lambda _name: word
    package = types.ModuleType("win32com"); package.client = client
    monkeypatch.setitem(sys.modules, "win32com", package); monkeypatch.setitem(sys.modules, "win32com.client", client)
    with pytest.raises(RuntimeError, match="Word failed"):
        export_docx_to_pdf(source)
    assert not target.exists()
    assert calls["open"] == {"FileName": str(source.resolve()), "ReadOnly": True, "AddToRecentFiles": False, "ConfirmConversions": False}
    assert word.Visible is False and word.DisplayAlerts == 0 and calls["close"] is False and calls["quit"] is True


def test_mixed_hour_minute_rhythm_is_total_duration() -> None:
    assert diary_minute_schedule_from_choice("2 часа 30 минут").minute_offsets == (150,)
    assert diary_minute_schedule_from_choice("1 час 45 минут").minute_offsets == (105,)


def test_trial_reservation_cannot_cross_total_limit(tmp_path: Path) -> None:
    manager = ProductAccessManager(tmp_path, now=datetime(2026, 9, 3, tzinfo=timezone.utc))
    manager.record_created_documents(29)
    with pytest.raises(PermissionError):
        manager.reserve_created_documents(2)
    assert manager.current_state().documents_used_total_trial == 29


def test_trial_configuration_limits_are_real_gates(tmp_path: Path) -> None:
    manager = ProductAccessManager(tmp_path, now=datetime(2026, 9, 3, tzinfo=timezone.utc))
    assert not manager.check_configuration_limits(template_count=6).allowed
    assert not manager.check_configuration_limits(profile_count=2).allowed
    assert manager.check_configuration_limits(template_count=5, profile_count=1).allowed


def test_concurrent_product_usage_has_no_lost_updates_or_shared_tmp_race(tmp_path: Path) -> None:
    storage = tmp_path / "license"
    manager = ProductAccessManager(storage, now=datetime(2026, 9, 3, tzinfo=timezone.utc)); manager.current_state()
    code = "from product_access import ProductAccessManager; ProductAccessManager().record_created_documents(1)"
    env = dict(os.environ); env["DOKKOMPLEKT_LICENSE_DIR"] = str(storage); env["PYTHONPATH"] = str(Path.cwd())
    processes = [subprocess.Popen([sys.executable, "-c", code], cwd=Path.cwd(), env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True) for _ in range(8)]
    failures = []
    for process in processes:
        stdout, stderr = process.communicate(timeout=30)
        if process.returncode:
            failures.append((process.returncode, stdout, stderr))
    assert not failures
    assert ProductAccessManager(storage, now=datetime(2026, 9, 3, tzinfo=timezone.utc)).current_state().documents_used_total_trial == 8
    assert not list(storage.glob("*.tmp"))


def test_atomic_json_concurrent_writers_never_share_fixed_temp(tmp_path: Path) -> None:
    target = tmp_path / "state.json"
    code = "import os; from medical_paths import atomic_write_json; atomic_write_json(os.environ['TARGET'], {'writer': os.environ['WRITER']})"
    base = dict(os.environ); base["TARGET"] = str(target); base["PYTHONPATH"] = str(Path.cwd())
    processes = []
    for index in range(12):
        env = dict(base); env["WRITER"] = str(index)
        processes.append(subprocess.Popen([sys.executable, "-c", code], cwd=Path.cwd(), env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True))
    for process in processes:
        stdout, stderr = process.communicate(timeout=30)
        assert process.returncode == 0, (stdout, stderr)
    assert json.loads(target.read_text())["writer"] in {str(i) for i in range(12)}
    assert not list(tmp_path.glob("*.tmp"))


def test_interprocess_lock_recovers_dead_stale_owner(tmp_path: Path) -> None:
    lock = tmp_path / "state.lock"
    lock.write_text(json.dumps({"pid": 99999999, "token": "dead", "created_at": 0}))
    old = datetime(2000, 1, 1, tzinfo=timezone.utc).timestamp(); os.utime(lock, (old, old))
    with interprocess_file_lock(lock, timeout_seconds=1, stale_seconds=0.01):
        assert lock.exists()
    assert not lock.exists()


def test_settings_and_profile_backups_are_bounded(tmp_path: Path) -> None:
    from settings_mixin import SettingsMixin
    settings = SettingsMixin(); settings._settings_path = tmp_path / "settings.json"; settings._settings = {}
    for index in range(40):
        settings._settings["printer"] = f"p{index}"; assert settings._save_settings()
    assert len(list((tmp_path / "_settings_backups").glob("*.json"))) <= 24
    profile = tmp_path / "profile.medpack.json"; pack = default_document_pack()
    for index in range(45):
        pack.name = f"P{index}"; save_document_pack(pack, profile, backup_reason="audit")
    assert len(list((tmp_path / "_profile_backups").glob("*.json"))) <= 32


def test_desktop_intake_seen_state_is_bounded_consistently(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import desktop_intake
    import desktop_intake_agent as agent
    values = [f"{index:064x}" for index in range(1500)]
    normalized = desktop_intake.normalize_intake_settings({"seen_signatures": values})
    assert len(normalized["seen_signatures"]) == desktop_intake.DESKTOP_INTAKE_MAX_SEEN_SIGNATURES
    state_path = tmp_path / "agent.json"; monkeypatch.setattr(agent, "_state_path", lambda: state_path)
    agent._save_state(set(values), last_launch=0.0)
    assert len(json.loads(state_path.read_text())["seen_signatures"]) == desktop_intake.DESKTOP_INTAKE_MAX_SEEN_SIGNATURES


def test_output_transaction_rejects_reported_file_outside_staging(tmp_path: Path) -> None:
    transaction = OutputTransaction(tmp_path / "final"); stage = transaction.begin()
    outside = tmp_path / "outside.docx"; outside.write_bytes(b"x")
    with pytest.raises(RuntimeError, match="вне staging"):
        transaction.validate_reported_files([outside])
    transaction.rollback(); assert not stage.exists()


def test_output_transaction_rejects_unexpected_binary_file(tmp_path: Path) -> None:
    final = tmp_path / "final"; transaction = OutputTransaction(final); stage = transaction.begin()
    expected = stage / "expected.docx"; expected.write_bytes(b"expected")
    rogue = stage / "rogue.pdf"; rogue.write_bytes(b"rogue")
    with pytest.raises(RuntimeError, match="неучтённые файлы"):
        transaction.commit(expected_files=[expected])
    assert not final.exists()
    transaction.rollback()


def test_concurrent_native_license_installs_use_unique_atomic_temp(tmp_path: Path) -> None:
    storage = tmp_path / "native-license"
    code = r'''
import json, os
import product_access.native as native
native._entitlement_payload = lambda _text: {
    "license_id": "native-test",
    "plan": "doctor_start",
    "valid_until": "2099-01-01T00:00:00+00:00",
    "features": [native.RUST_NATIVE_VERIFIED_FEATURE],
    "signature": "rust-ed25519",
}
payload = {"schema": native.RUST_LICENSE_SCHEMA, "license": {"writer": os.environ["WRITER"]}}
native.NativeProductAccessManager().install_license_text(json.dumps(payload))
'''
    base = dict(os.environ)
    base["DOKKOMPLEKT_LICENSE_DIR"] = str(storage)
    base["PYTHONPATH"] = str(Path.cwd())
    processes = []
    for index in range(8):
        env = dict(base)
        env["WRITER"] = str(index)
        processes.append(
            subprocess.Popen(
                [sys.executable, "-c", code],
                cwd=Path.cwd(),
                env=env,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
        )
    failures = []
    for process in processes:
        stdout, stderr = process.communicate(timeout=30)
        if process.returncode:
            failures.append((process.returncode, stdout, stderr))
    assert not failures
    payload = json.loads((storage / "license.json").read_text())
    assert payload["schema"] == "dokkomplekt.license.v1"
    assert payload["license"]["writer"] in {str(index) for index in range(8)}
    assert not list(storage.glob("*.tmp"))


def test_concurrent_docm_conversion_has_no_shared_temp_race(tmp_path: Path) -> None:
    source_docx = _docx(tmp_path / "macro-source.docx", "ФИО: ______")
    docm = tmp_path / "macro-source.docm"
    with zipfile.ZipFile(source_docx, "r") as src, zipfile.ZipFile(docm, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    b"application/vnd.ms-word.document.macroEnabled.main+xml",
                )
            dst.writestr(info, data)
    source_docx.unlink()

    from medical_docx_xml_fragments import _conversion_target

    target = _conversion_target(docm)
    target.unlink(missing_ok=True)
    code = r'''
import os, zipfile
from medical_docx_xml_fragments import convert_docm_to_docx
result = convert_docm_to_docx(os.environ["DOCM"])
with zipfile.ZipFile(result) as archive:
    assert "word/document.xml" in archive.namelist()
'''
    base = dict(os.environ)
    base["DOCM"] = str(docm)
    base["PYTHONPATH"] = str(Path.cwd())
    processes = [
        subprocess.Popen(
            [sys.executable, "-c", code],
            cwd=Path.cwd(),
            env=base,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        for _ in range(8)
    ]
    failures = []
    for process in processes:
        stdout, stderr = process.communicate(timeout=30)
        if process.returncode:
            failures.append((process.returncode, stdout, stderr))
    assert not failures
    assert target.exists() and target.stat().st_size > 0
    assert not list(target.parent.glob(f".{target.name}.*"))
