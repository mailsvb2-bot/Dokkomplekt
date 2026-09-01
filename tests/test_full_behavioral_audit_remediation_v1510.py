from __future__ import annotations

import os
import time
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document


class Var:
    def __init__(self, value=""):
        self.value = value
    def get(self):
        return self.value
    def set(self, value):
        self.value = value


def _make_primary(path: Path, *, fio: str = "Иванов Иван Иванович", admission: str = "01.06.2026") -> Path:
    doc = Document()
    doc.add_paragraph("Первичный осмотр")
    doc.add_paragraph(f"ФИО: {fio}")
    doc.add_paragraph(f"Дата поступления: {admission}")
    doc.add_paragraph("Диагноз: K35.8 Острый аппендицит")
    doc.add_paragraph("Жалобы: боли в животе")
    doc.save(path)
    old = time.time() - 5
    os.utime(path, (old, old))
    return path


def _make_docm(path: Path) -> Path:
    source = path.with_suffix(".docx")
    doc = Document()
    doc.add_paragraph("Пациент {{patient.fio}}")
    doc.save(source)
    with ZipFile(source, "r") as src, ZipFile(path, "w", ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    b"application/vnd.ms-word.document.macroEnabled.main+xml",
                )
            dst.writestr(info, data)
    source.unlink()
    return path


def test_icd_subrubric_conflict_is_hard_mismatch():
    from diary_text_selection import diary_diagnosis_match_score
    assert diary_diagnosis_match_score("K35.8 Острый аппендицит", "K35.2 Острый аппендицит") == 0
    assert diary_diagnosis_match_score("F20.0 Шизофрения", "F20.1 Шизофрения") == 0
    assert diary_diagnosis_match_score("K35.8 Острый аппендицит", "K35.8 Острый аппендицит") >= 100


def test_surname_initials_preserve_preabbreviated_name():
    from desktop_patient_folder import _fio_to_surname_initials
    assert _fio_to_surname_initials("Иванов Иван Иванович") == "Иванов И.И."
    assert _fio_to_surname_initials("Иванов И.И.") == "Иванов И.И."


def test_intake_rejects_arbitrary_docx_and_accepts_primary(tmp_path: Path):
    from desktop_intake import scan_primary_candidates
    random = tmp_path / "random.docx"
    doc = Document(); doc.add_paragraph("Совершенно произвольный офисный документ без данных пациента."); doc.save(random)
    old = time.time() - 5; os.utime(random, (old, old))
    assert scan_primary_candidates(tmp_path, set()) == ()
    primary = _make_primary(tmp_path / "primary.docx")
    found = scan_primary_candidates(tmp_path, set())
    assert [item.path for item in found] == [primary]


def test_prepare_patient_work_folder_transaction_keeps_source(tmp_path: Path):
    from desktop_intake import prepare_patient_work_folder
    source = _make_primary(tmp_path / "primary.docx")
    patient_dir, staged = prepare_patient_work_folder(tmp_path, source, folder_name="Иванов И.И. июнь 2026", keep_source=True)
    assert source.exists()
    assert staged.exists()
    assert staged.parent == patient_dir
    assert staged.read_bytes() == source.read_bytes()


def test_same_patient_reuses_existing_folder_but_different_case_does_not(tmp_path: Path):
    from desktop_intake import safe_patient_subfolder
    existing = tmp_path / "Иванов И.И. июнь 2026"
    existing.mkdir()
    _make_primary(existing / "old.docx", fio="Иванов Иван Иванович", admission="01.06.2026")
    same = _make_primary(tmp_path / "same.docx", fio="Иванов Иван Иванович", admission="01.06.2026")
    other = _make_primary(tmp_path / "other.docx", fio="Иванов Иван Иванович", admission="15.06.2026")
    assert safe_patient_subfolder(tmp_path, same, folder_name=existing.name) == existing
    assert safe_patient_subfolder(tmp_path, other, folder_name=existing.name) != existing


def test_date_template_changes_text_diary_calendar_without_restoring_table_writer(tmp_path: Path):
    from diary_batch import fill_diary_batch
    statuses = tmp_path / "texts.docx"
    doc = Document(); doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет."); doc.add_paragraph("Состояние без отрицательной динамики, лечение переносит."); doc.save(statuses)
    dates = tmp_path / "01.docx"
    d = Document(); table = d.add_table(rows=3, cols=2)
    table.cell(0,0).text="Число"; table.cell(0,1).text="Месяц/год"
    table.cell(1,0).text="03"; table.cell(1,1).text="06.2026"
    table.cell(2,0).text="05"; table.cell(2,1).text="06.2026"
    d.save(dates)
    result = fill_diary_batch(status_files=[statuses], diary_files=[dates], output_dir=tmp_path / "out", patient_name="Иванов Иван Иванович", admission_value="01.06.2026", discharge_value="06.06.2026", diary_day_offsets=(1,2,3,4), force_final_diary=False, text_output=True)
    out = Document(result.created_files[0])
    text = "\n".join(p.text for p in out.paragraphs)
    assert "03.06.26" in text
    assert "05.06.26" in text
    assert "02.06.26" not in text
    assert len(out.tables) == 0


def test_diary_parser_does_not_turn_heading_into_status(tmp_path: Path):
    from diary_text_parser import extract_statuses_from_docx
    path = tmp_path / "texts.docx"
    doc = Document(); doc.add_heading("УНИКАЛЬНЫЙ ЗАГОЛОВОК ВРАЧА", level=1); doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет."); doc.save(path)
    statuses = extract_statuses_from_docx(path)
    assert statuses == ["Состояние стабильное, жалоб активно не предъявляет."]


def test_nested_table_placeholder_is_rendered_and_reported(tmp_path: Path):
    from universal_fields import PatientCase
    from universal_profiles import DocumentPack
    from universal_template_engine import attach_template_to_pack
    from universal_generation import render_documents_from_pack
    template = tmp_path / "nested.docx"
    doc = Document(); outer=doc.add_table(rows=1, cols=1); inner=outer.cell(0,0).add_table(rows=1, cols=1); inner.cell(0,0).text="Пациент {{patient.fio}}"; doc.save(template)
    pack=DocumentPack(pack_id="p",name="p")
    spec,_=attach_template_to_pack(pack,template,tmp_path/"profile",button_label="Nested")
    case=PatientCase(); case.set("patient.fio","Петров Пётр Петрович",source_document="test",confidence=1.0)
    result=render_documents_from_pack(pack=pack,case=case,document_ids=[spec.id],output_dir=tmp_path/"out",base_dir=tmp_path/"profile",strict=True)
    assert len(result.created_files)==1
    rendered=Document(result.created_files[0])
    assert "Петров Пётр Петрович" in rendered.tables[0].cell(0,0).tables[0].cell(0,0).text
    assert result.render_results[0].missing_fields == ()


def test_strict_custom_render_refuses_missing_placeholder(tmp_path: Path):
    from universal_fields import PatientCase
    from universal_profiles import DocumentPack
    from universal_template_engine import attach_template_to_pack
    from universal_generation import render_documents_from_pack
    template=tmp_path/"strict.docx"; doc=Document(); doc.add_paragraph("История № {{case.number}} Пациент {{patient.fio}}"); doc.save(template)
    pack=DocumentPack(pack_id="p",name="p"); spec,_=attach_template_to_pack(pack,template,tmp_path/"profile",button_label="Strict")
    case=PatientCase(); case.set("patient.fio","Иванов Иван Иванович",source_document="test",confidence=1.0)
    result=render_documents_from_pack(pack=pack,case=case,document_ids=[spec.id],output_dir=tmp_path/"out",base_dir=tmp_path/"profile",strict=True)
    assert result.created_files == ()
    assert any("case.number" in item for item in result.skipped_documents)


def test_docm_template_is_owned_as_macro_free_docx(tmp_path: Path):
    from universal_profiles import DocumentPack
    from universal_template_engine import attach_template_to_pack
    source=_make_docm(tmp_path/"template.docm")
    pack=DocumentPack(pack_id="p",name="p")
    spec,copied=attach_template_to_pack(pack,source,tmp_path/"profile",button_label="DOCM")
    assert copied.suffix.lower()==".docx"
    assert copied.exists()
    assert spec.template.endswith(".docx")


def test_visual_marker_preserves_runs_and_rejects_ambiguous_selection(tmp_path: Path):
    from universal_template_engine import replace_selection_with_placeholder
    path=tmp_path/"styled.docx"; doc=Document(); p=doc.add_paragraph(); r1=p.add_run("ФИО: "); r1.bold=True; r2=p.add_run("Иванов"); r2.italic=True; doc.save(path)
    replace_selection_with_placeholder(path,"Иванов","patient.fio",create_backup=False)
    out=Document(path); assert out.paragraphs[0].runs[0].bold is True; assert out.paragraphs[0].runs[1].italic is True; assert "{{patient.fio}}" in out.paragraphs[0].text
    dup=tmp_path/"dup.docx"; doc=Document(); doc.add_paragraph("Диагноз"); doc.add_paragraph("Диагноз"); doc.save(dup)
    with pytest.raises(ValueError, match="встречается"):
        replace_selection_with_placeholder(dup,"Диагноз","diagnosis.main",create_backup=False)


def test_patient_reset_clears_external_sources_and_undo_history():
    from files_mixin import FilesMixin
    names = [
        "assigned_treatment_var","case_number_var","expert_work_status_var","expert_work_org_var","expert_position_var","expert_sick_leave_needed_var","expert_sick_leave_from_var","expert_sick_leave_number_var","vk_mse_work_org_var","vk_mse_position_var","sick_leave_vk_work_org_var","sick_leave_vk_position_var","sick_leave_vk_work_position_var","rvk_act_number_var","rvk_work_position_var","sick_leave_vk_date_var","sick_leave_vk_protocol_number_var","sick_leave_vk_protocol_date_var","sick_leave_vk_commission_date_var","commission_date_var","commission_number_var","vk_date_var","vk_protocol_number_var","vk_protocol_date_var","labs_text_var","labs_source_path_var","labs_explicit_date_var","labs_date_policy_var","labs_without_var","patient_name_var","admission_date_var","discharge_date_var","diagnosis_var","epi_path_var","additional_info_text_var","additional_info_source_path_var","diary_treatment_correction_var"
    ]
    h=SimpleNamespace(**{name:Var("OLD") for name in names})
    h.status_files=[]; h.diary_files=[]; h.diary_texts_dir=""; h.diary_template_dir=""; h._diary_text_files_auto_selected=False; h._diary_files_auto_selected=False
    h._field_undo_suspended=False; h._field_undo_stack={"patient_name":["PATIENT_A"]}; h._field_undo_vars={"patient_name":h.patient_name_var}; h._field_undo_last={"patient_name":"OLD"}
    h._semantic_date_state={"x":1}; h._update_expert_sick_leave_display=lambda:None; h._update_diary_text_label=lambda **kw:None; h._update_diary_template_label=lambda **kw:None; h._set_ui_var=lambda var,value:var.set(value); h._set_primary_drop_empty=lambda:None
    FilesMixin._reset_primary_document_runtime_state(h)
    assert h.epi_path_var.get()==""
    assert h.additional_info_text_var.get()==""
    assert h.additional_info_source_path_var.get()==""
    assert h.diary_treatment_correction_var.get()==""
    assert h._field_undo_stack["patient_name"] == []


def test_dead_agent_pid_makes_lock_stale_immediately(monkeypatch, tmp_path: Path):
    import desktop_intake_agent as agent
    lock=tmp_path/"agent.lock"; lock.write_text("pid=424242\nversion=x\n",encoding="utf-8"); os.utime(lock,(time.time(),time.time()))
    monkeypatch.setattr(agent,"_pid_is_running",lambda pid:False)
    assert agent._lock_is_stale(lock) is True


def test_target_facility_is_exact_required_organization():
    from medical_constants import TARGET_MEDICAL_FACILITY
    assert TARGET_MEDICAL_FACILITY == "ГБУЗ НО «НКЦПЗ» диспансер №2"
