from __future__ import annotations

from datetime import date
import json
import os
from pathlib import Path
from types import SimpleNamespace
import zipfile

import pytest
from docx import Document
from docx.oxml import OxmlElement


def _docx(path: Path, *paragraphs: str) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


def test_ambiguous_same_fio_does_not_merge_distinct_admission(monkeypatch, tmp_path: Path) -> None:
    import desktop_intake
    import desktop_patient_folder

    source = tmp_path / "source.docx"; source.write_bytes(b"new-admission")
    folder = tmp_path / "patient"; folder.mkdir()
    old = folder / "old.docx"; old.write_bytes(b"old-admission")

    def info(path):
        p = Path(path)
        return SimpleNamespace(fio="Иванов Иван Иванович", admission_date="" if p == source else "01.08.2026")

    monkeypatch.setattr(desktop_patient_folder, "build_patient_folder_info", info)
    assert desktop_intake._same_patient_existing_folder(folder, source) is False


def test_ambiguous_retry_reuses_only_byte_identical_primary(monkeypatch, tmp_path: Path) -> None:
    import desktop_intake
    import desktop_patient_folder

    source = tmp_path / "source.docx"; source.write_bytes(b"same")
    folder = tmp_path / "patient"; folder.mkdir()
    old = folder / "old.docx"; old.write_bytes(b"same")
    monkeypatch.setattr(desktop_patient_folder, "build_patient_folder_info", lambda _p: SimpleNamespace(fio="Иванов Иван Иванович", admission_date=""))
    assert desktop_intake._same_patient_existing_folder(folder, source) is True


def test_signature_uses_content_even_when_size_and_mtime_match(tmp_path: Path) -> None:
    from desktop_intake import signature_key
    path = tmp_path / "primary.docx"
    path.write_bytes(b"AAAA")
    first = signature_key(path, 123, 4)
    path.write_bytes(b"BBBB")
    second = signature_key(path, 123, 4)
    assert first != second


def test_seen_ledger_preserves_more_than_300_entries() -> None:
    from desktop_intake import normalize_intake_settings
    seen = [f"{i:064x}" for i in range(420)]
    normalized = normalize_intake_settings({"seen_signatures": seen})
    assert len(normalized["seen_signatures"]) == 420
    assert normalized["seen_signatures"][0] == seen[0]


def test_live_agent_pid_cannot_be_stolen_due_to_old_mtime(monkeypatch, tmp_path: Path) -> None:
    import desktop_intake_agent as agent
    lock = tmp_path / "agent.lock"
    lock.write_text("pid=4242\nversion=x\nprocess_started=proc:same\ntoken=" + "a" * 32 + "\n", encoding="utf-8")
    os.utime(lock, (1, 1))
    monkeypatch.setattr(agent, "_pid_is_running", lambda pid: pid == 4242)
    monkeypatch.setattr(agent, "_process_start_identity", lambda pid: "proc:same" if pid == 4242 else "")
    assert agent._lock_is_stale(lock) is False


def test_old_agent_cannot_delete_replacement_lock(tmp_path: Path) -> None:
    import desktop_intake_agent as agent
    lock = tmp_path / "agent.lock"
    lock.write_text(f"pid={os.getpid()}\nversion=x\ntoken={'b'*32}\n", encoding="utf-8")
    agent._release_agent_lock(None, lock, "a" * 32)
    assert lock.exists()


def test_live_gui_pid_beats_stale_heartbeat(monkeypatch, tmp_path: Path) -> None:
    import desktop_intake_agent as agent
    lock = tmp_path / "gui.json"
    lock.write_text(json.dumps({"pid": 5151, "process_started": "proc:same", "updated_at": 1.0}), encoding="utf-8")
    monkeypatch.setattr(agent, "_gui_lock_path", lambda: lock)
    monkeypatch.setattr(agent, "_pid_is_running", lambda pid: pid == 5151)
    monkeypatch.setattr(agent, "_process_start_identity", lambda pid: "proc:same" if pid == 5151 else "")
    assert agent.is_gui_runtime_active(now=999999.0) is True


def test_pending_absence_is_not_proof_of_processing(monkeypatch, tmp_path: Path) -> None:
    import desktop_intake_agent as agent
    signature = "c" * 64
    state = {"pending": {"signature": signature, "launched_at": 1.0}}
    seen: set[str] = set()
    monkeypatch.setattr(agent, "_settings_seen_signatures", lambda: set())
    monkeypatch.setattr(agent, "_signature_present_in_folder", lambda folder, sig: False)
    monkeypatch.setattr(agent.time, "time", lambda: 1.0 + agent.PENDING_RETRY_SECONDS + 5)
    pending, changed = agent._resolve_pending_state(state, seen, tmp_path)
    assert pending == {}
    assert changed is True
    assert signature not in seen


def test_custom_requirements_ignore_label_path_description_and_optional_fields() -> None:
    from universal_main_documents import custom_requirement_flags_for_documents
    base = SimpleNamespace(
        id="stable-id", role_id="", category="medical", button_label="Обычная справка",
        template="templates/neutral.docx", description="", required_fields=("patient.fio",), optional_fields=("case.number",),
    )
    renamed = SimpleNamespace(**{**base.__dict__, "button_label": "Выписной эпикриз комиссия РВК", "template": "/tmp/discharge-rvk.docx", "description": "МCЭ больничный"})
    assert custom_requirement_flags_for_documents([base]) == custom_requirement_flags_for_documents([renamed])
    flags = custom_requirement_flags_for_documents([renamed])
    assert flags["requires_fio"] is True
    assert flags["requires_case_number"] is False
    assert flags["discharge"] is False
    assert flags["rvk"] is False


def test_missing_profile_template_remains_visible_as_broken_button(tmp_path: Path) -> None:
    from universal_main_documents import custom_documents_for_main_ui
    from universal_profiles import DocumentPack, DocumentTemplateSpec
    pack = DocumentPack(pack_id="x", name="x", documents=(DocumentTemplateSpec(id="mine", button_label="Мой документ", template="templates/missing.docx"),))
    docs = custom_documents_for_main_ui(pack, base_dir=tmp_path)
    assert len(docs) == 1
    assert docs[0].available is False
    assert "не найден" in docs[0].problem.lower()


def test_profile_validation_and_export_fail_on_dangling_template(tmp_path: Path) -> None:
    from universal_profiles import DocumentPack, DocumentTemplateSpec
    from universal_template_engine import export_document_pack_zip, validate_document_pack
    pack = DocumentPack(pack_id="x", name="x", documents=(DocumentTemplateSpec(id="mine", button_label="Мой документ", template="templates/missing.docx"),))
    validation = validate_document_pack(pack, base_dir=tmp_path)
    assert validation.ok is False
    with pytest.raises(ValueError, match="нельзя экспортировать"):
        export_document_pack_zip(pack, tmp_path / "broken.medpack.zip", template_base_dir=tmp_path)
    assert not (tmp_path / "broken.medpack.zip").exists()


def test_import_fails_if_manifest_references_missing_template(tmp_path: Path) -> None:
    from universal_profiles import DocumentPack, DocumentTemplateSpec
    from universal_template_engine import import_document_pack_zip
    pack = DocumentPack(pack_id="x", name="x", documents=(DocumentTemplateSpec(id="mine", button_label="Мой документ", template="templates/missing.docx"),))
    archive = tmp_path / "broken.medpack.zip"
    with zipfile.ZipFile(archive, "w") as zf:
        zf.writestr("pack.json", json.dumps(pack.to_dict(), ensure_ascii=False))
    with pytest.raises(ValueError, match="нет шаблона"):
        import_document_pack_zip(archive, tmp_path / "imported")


def test_attach_invalid_template_is_atomic(tmp_path: Path) -> None:
    from universal_profiles import DocumentPack
    from universal_template_engine import attach_template_to_pack
    source = _docx(tmp_path / "bad.docx", "{{definitely.unknown.field}}")
    profile = tmp_path / "profile"
    pack = DocumentPack(pack_id="x", name="x")
    with pytest.raises(ValueError):
        attach_template_to_pack(pack, source, profile, button_label="Bad")
    assert pack.documents == ()
    assert not list((profile / "templates").glob("*.docx"))


def test_strict_multi_render_does_not_emit_good_subset_when_one_template_missing(tmp_path: Path) -> None:
    from universal_fields import PatientCase
    from universal_generation import render_documents_from_pack
    from universal_profiles import DocumentPack, DocumentTemplateSpec
    profile = tmp_path / "profile"; (profile / "templates").mkdir(parents=True)
    _docx(profile / "templates" / "good.docx", "{{patient.fio}}")
    good = DocumentTemplateSpec(id="good", button_label="Good", template="templates/good.docx", required_fields=("patient.fio",))
    bad = DocumentTemplateSpec(id="bad", button_label="Bad", template="templates/missing.docx", required_fields=("patient.fio",))
    pack = DocumentPack(pack_id="x", name="x", documents=(good, bad))
    case = PatientCase(); case.set("patient.fio", "Иванов Иван")
    result = render_documents_from_pack(pack=pack, case=case, document_ids=["good", "bad"], output_dir=tmp_path / "out", base_dir=profile, strict=True)
    assert result.created_files == ()
    assert result.skipped_documents
    assert not list((tmp_path / "out").glob("*.docx"))


def test_output_transaction_rolls_back_overwrite_on_mid_commit_failure(monkeypatch, tmp_path: Path) -> None:
    import output_transaction as module
    from output_transaction import OutputTransaction
    final = tmp_path / "patient"; final.mkdir()
    old = final / "a.docx"; old.write_text("OLD", encoding="utf-8")
    tx = OutputTransaction(final, overwrite_paths=(old,))
    stage = tx.begin()
    (stage / "a.docx").write_text("NEW-A", encoding="utf-8")
    (stage / "b.docx").write_text("NEW-B", encoding="utf-8")
    real_move = module.OutputTransaction._move_no_replace
    calls = {"n": 0}
    def fail_third(src, dst):
        calls["n"] += 1
        # 1=backup old a.docx, 2=commit new a.docx, 3=commit b.docx.
        if calls["n"] == 3:
            raise OSError("simulated commit failure")
        return real_move(src, dst)
    monkeypatch.setattr(module.OutputTransaction, "_move_no_replace", staticmethod(fail_third))
    with pytest.raises(OSError, match="simulated"):
        tx.commit()
    assert old.read_text(encoding="utf-8") == "OLD"
    assert not (final / "b.docx").exists()


def test_hourly_inpatient_diaries_keep_weekend_dates() -> None:
    from diary_batch import _hourly_text_diary_datetimes
    moments = _hourly_text_diary_datetimes("05.09.2026 10:00", date(2026, 9, 6), limit=3, hour_offsets=(1,))
    assert [m.date() for m in moments] == [date(2026, 9, 5), date(2026, 9, 5), date(2026, 9, 5)]


def test_repeated_diary_statuses_are_preserved(tmp_path: Path) -> None:
    from diary_batch import read_statuses_from_files
    text = "Состояние стабильное, контакт продуктивный, жалоб активно не предъявляет."
    source_a = _docx(tmp_path / "statuses_a.docx", text)
    source_b = _docx(tmp_path / "statuses_b.docx", text)
    statuses = read_statuses_from_files([source_a, source_b])
    assert statuses == [text, text]


def test_unreadable_selected_date_template_is_hard_error(tmp_path: Path) -> None:
    from diary_batch import _day_offsets_from_date_templates
    broken = tmp_path / "01.docx"; broken.write_bytes(b"not a zip")
    with pytest.raises(ValueError, match="Не удалось прочитать выбранный файл дат"):
        _day_offsets_from_date_templates([broken], admission_date_value=date(2026, 9, 1), discharge_date_value=None)


def test_yearless_diary_month_rolls_into_next_year(tmp_path: Path) -> None:
    from diary_batch import _day_offsets_from_date_templates
    source = _docx(tmp_path / "dates.docx", "15.02")
    offsets = _day_offsets_from_date_templates([source], admission_date_value=date(2026, 12, 10), discharge_date_value=None)
    assert offsets == ((date(2027, 2, 15) - date(2026, 12, 10)).days,)


def test_initials_first_patient_name_uses_surname_for_gender() -> None:
    from diary_gender import detect_gender_from_patient_name
    assert detect_gender_from_patient_name("И.И. Иванов") == "male"
    assert detect_gender_from_patient_name("А.А. Петрова") == "female"


def test_confirmed_folder_rule_never_silently_drops_missing_component() -> None:
    from desktop_patient_folder import build_patient_folder_name
    settings = {"parts": ["surname_initials", "discharge_date"], "date_format": "short", "doctor_confirmed": True}
    with pytest.raises(ValueError, match="Дата выписки"):
        build_patient_folder_name(fio="Иванов Иван Иванович", admission_date="01.09.2026", discharge_date="", settings=settings, fallback="Иванов", strict=True)


def test_textbox_placeholder_is_validated_and_rendered(tmp_path: Path) -> None:
    from universal_fields import PatientCase
    from universal_profiles import DocumentTemplateSpec
    from universal_template_engine import extract_template_placeholders, render_template_to_docx
    template = tmp_path / "textbox.docx"
    doc = Document(); outer = doc.add_paragraph(); run = outer.add_run()
    box = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p"); r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "Пациент: {{patient.fio}}"
    r.append(t); p.append(r); box.append(p); run._r.append(box); doc.save(template)
    assert [item.field_id for item in extract_template_placeholders(template)] == ["patient.fio"]
    case = PatientCase(); case.set("patient.fio", "Петров Петр")
    spec = DocumentTemplateSpec(id="box", button_label="Box", template=str(template), required_fields=("patient.fio",))
    out = tmp_path / "out.docx"
    render_template_to_docx(template_path=template, output_path=out, case=case, document=spec, strict=True)
    with zipfile.ZipFile(out) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "Петров Петр" in xml
    assert "{{patient.fio}}" not in xml
