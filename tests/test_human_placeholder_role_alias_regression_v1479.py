from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from actions_required_fields_popup import (
    _is_admission_date_field,
    _is_case_number_field,
    _is_diagnosis_field,
    _is_discharge_date_field,
    _is_labs_field,
    _is_treatment_field,
    _store_key_for_field,
)
from universal_fields import PatientCase, normalize_field_id
from universal_main_documents import custom_requirement_flags_for_documents
from universal_template_engine import extract_template_placeholders, missing_required_fields


def _field(key: str, label: str = "", value: str = ""):
    return SimpleNamespace(key=key, field_id=key, label=label, placeholder="", reason="", value=value)


def _doc(*fields: str, label: str = "Custom", role_id: str = ""):
    return SimpleNamespace(
        id="custom_doc",
        document_id="custom_doc",
        role_id=role_id,
        category="documents",
        button_label=label,
        template="templates/custom.docx",
        description="",
        required_fields=tuple(fields),
        optional_fields=(),
    )


def test_human_docx_placeholder_aliases_resolve_to_canonical_fields(tmp_path: Path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{ФИО}} {{История болезни №}} {{Дата поступления}} {{Дата выписки}}")
    doc.add_paragraph("{{Диагноз}} {{Код МКБ}} {{Назначенное лечение}} {{Лабораторные исследования}}")
    doc.save(template)

    fields = [item.field_id for item in extract_template_placeholders(template)]

    assert fields == [
        "patient.fio",
        "case.number",
        "admission.date",
        "discharge.date",
        "diagnosis.main",
        "diagnosis.icd10",
        "treatment.plan",
        "labs.results",
    ]


def test_common_short_and_human_aliases_share_patient_case_state():
    case = PatientCase()
    case.set("patientName", "Иванов Иван Иванович")
    case.set("case #", "ИБ-77")
    case.set("Дата выписки", "12.06.2026")
    case.set("Назначенное лечение", "Терапия")

    assert case.get("patient.fio") == "Иванов Иван Иванович"
    assert case.get("case.number") == "ИБ-77"
    assert case.get("discharge.date") == "12.06.2026"
    assert case.get("treatment.plan") == "Терапия"


def test_required_popup_detects_human_placeholder_labels_as_same_semantics():
    samples = [
        (_field("patientName", "ФИО"), None, "fio"),
        (_field("case #", "История болезни №"), _is_case_number_field, "case_number"),
        (_field("Дата поступления", "Поступил"), _is_admission_date_field, "admission_date"),
        (_field("Дата выписки", "Выписан"), _is_discharge_date_field, "discharge_date"),
        (_field("Диагноз", "Клинический диагноз"), _is_diagnosis_field, "diagnosis"),
        (_field("Назначенное лечение", "Назначения"), _is_treatment_field, "treatment"),
        (_field("Лабораторные исследования", "Анализы"), _is_labs_field, "labs"),
    ]
    for field, detector, store_key in samples:
        if detector is not None:
            assert detector(field), field
        assert _store_key_for_field(field) == store_key


def test_custom_requirement_flags_cover_human_fields_and_camelcase_roles():
    flags = custom_requirement_flags_for_documents([
        _doc("case #", "Диагноз", "Назначенное лечение", "Дата выписки", "Лабораторные исследования", role_id="dischargeEpicrisis"),
        _doc(role_id="sickLeaveVk"),
        _doc(role_id="militaryCommissariatAct"),
    ])

    assert flags["discharge"] is True
    assert flags["sick_leave_vk"] is True
    assert flags["rvk"] is True
    assert flags["requires_case_number"] is True
    assert flags["requires_diagnosis"] is True
    assert flags["requires_treatment"] is True
    assert flags["requires_discharge_date"] is True
    assert flags["requires_labs"] is True


def test_missing_required_fields_normalizes_human_required_field_ids():
    document = SimpleNamespace(required_fields=("ФИО", "История болезни №", "Дата выписки"))
    case = PatientCase()
    case.set("patient.fio", "Иванов Иван Иванович")
    case.set("caseNo", "ИБ-77")

    assert missing_required_fields(case, document) == ("discharge.date",)


def test_representative_alias_normalizer_samples():
    assert normalize_field_id("medicalRecordNo") == "case.number"
    assert normalize_field_id("fullName") == "patient.fio"
    assert normalize_field_id("patientFio") == "patient.fio"
    assert normalize_field_id("jobTitle") == "patient.position"
    assert normalize_field_id("workPlace") == "patient.work"
