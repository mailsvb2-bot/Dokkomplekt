from __future__ import annotations

from docx import Document

from medical_docx_editor import DocxBlockEditor
from medical_models import PatientData
from medical_renderer_labs import (
    MedicalRendererLabsMixin,
    canonical_labs_placeholder,
    extract_labs_from_file,
    normalize_date,
)
from universal_case_adapter import patient_data_to_case
from universal_fields import normalize_field_id
from universal_profiles import DocumentTemplateSpec
from universal_template_engine import missing_required_fields


def test_labs_case_adapter_never_uses_epi_text_as_labs() -> None:
    data = PatientData(
        fio="Иванов Иван Иванович",
        epi_text="ЭПИ не является анализами",
        labs_text="",
        labs_without=False,
    )
    case = patient_data_to_case(data)
    assert case.get("labs.results") == ""


def test_without_labs_is_explicit_value_for_required_custom_template() -> None:
    data = PatientData(fio="Иванов Иван Иванович", labs_without=True)
    case = patient_data_to_case(data)
    document = DocumentTemplateSpec(
        id="custom_labs",
        button_label="Шаблон с анализами",
        template="templates/labs.docx",
        required_fields=("labs.results",),
    )
    assert case.get("labs.results") == "Нет анализов"
    assert missing_required_fields(case, document) == ()


def test_labs_date_normalization_uses_strict_episode_date_parser() -> None:
    assert normalize_date("1126") == "01.01.2026"
    assert normalize_date("110626") == "11.06.2026"
    assert normalize_date("99.99.2026") == ""
    assert normalize_date("31.02.2026") == ""


def test_extract_labs_file_rejects_unrelated_document_text(tmp_path) -> None:
    path = tmp_path / "not_labs.txt"
    path.write_text("ФИО: Иванов Иван Иванович\nДиагноз: K35 Острый аппендицит", encoding="utf-8")
    try:
        extract_labs_from_file(path)
    except ValueError as exc:
        assert "не распознан блок анализов" in str(exc)
    else:
        raise AssertionError("unrelated text file must not be accepted as analyses")


def test_legacy_labs_markers_are_removed_not_auto_filled() -> None:
    doc = Document()
    doc.add_paragraph("ОАК")
    doc.add_paragraph("ОАМ")
    doc.add_paragraph("Диагноз")
    editor = DocxBlockEditor(doc)
    MedicalRendererLabsMixin._replace_lab_lines(
        editor,
        {"day1": "02.01.2026", "day2": "03.01.2026", "flg": "01.12.2025"},
        data=PatientData(),
    )
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "в норме" not in text
    assert "ОАК" not in text
    assert "ОАМ" not in text
    assert "Диагноз" in text


def test_labs_aliases_are_consistent_across_template_and_field_registry() -> None:
    assert canonical_labs_placeholder("laboratory.results") == "labs.results"
    assert canonical_labs_placeholder("LAB_BLOCK") == "labs.results"
    assert normalize_field_id("laboratory.results") == "labs.results"
    assert normalize_field_id("labs_block") == "labs.results"
