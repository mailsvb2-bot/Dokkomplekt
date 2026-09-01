from __future__ import annotations

from pathlib import Path

from docx import Document

from universal_fields import PatientCase
from universal_generation import render_documents_from_pack
from universal_profiles import DocumentPack, DocumentTemplateSpec


def test_parallel_batch_render_preserves_order_unique_names_and_content(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Пациент {{patient.fio}}. Диагноз {{diagnosis.main}}.")
    doc.save(template)

    documents = tuple(
        DocumentTemplateSpec(
            id=f"parallel_{index}",
            button_label="Одинаковый документ",
            template=template.name,
            output_name="{{patient.fio}} Одинаковый документ.docx",
            required_fields=("patient.fio", "diagnosis.main"),
        )
        for index in range(8)
    )
    pack = DocumentPack(pack_id="parallel_regression", name="Parallel regression", documents=documents)
    case = PatientCase()
    case.update_from_pairs({"patient.fio": "Иванов Иван", "diagnosis.main": "K35.8"})

    result = render_documents_from_pack(
        pack=pack,
        case=case,
        document_ids=[item.id for item in documents],
        output_dir=tmp_path / "out",
        base_dir=tmp_path,
        strict=True,
        output_language="ru",
        spellcheck_enabled=True,
    )

    assert result.ok
    assert len(result.created_files) == 8
    assert [Path(path).name for path in result.created_files] == [
        "Иванов Иван Одинаковый документ.docx",
        *[f"Иванов Иван Одинаковый документ ({index}).docx" for index in range(2, 9)],
    ]
    assert len(set(result.created_files)) == 8

    for path in result.created_files:
        rendered = Document(path)
        text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
        assert "Иванов Иван" in text
        assert "K35.8" in text
        assert "{{" not in text
