from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from icd10_f_search import normalize_diagnosis_with_icd10
from medical_parser import MedicalTextParser
from universal_scanner import scan_docx


def _write_docx(path: Path, lines: list[str]) -> Path:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)
    return path


def test_parser_reads_diagnosis_from_plain_docx(tmp_path: Path) -> None:
    path = _write_docx(
        tmp_path / "primary.docx",
        [
            "Первичный осмотр",
            "Ф.И.О.: Иванов Иван Иванович",
            "Дата поступления: 10.02.2026",
            "Диагноз: K35.8 Острый аппендицит",
        ],
    )
    data = MedicalTextParser().parse_docx(path)
    assert data.fio == "Иванов Иван Иванович"
    assert data.admission_date == "10.02.2026"
    assert data.diagnosis.startswith("K35.8")


def test_parser_reads_header_footer_docx(tmp_path: Path) -> None:
    path = tmp_path / "header_footer.docx"
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Ф.И.О.: Петров Петр Петрович"
    section.footer.paragraphs[0].text = "Диагноз: I10 Гипертензивная болезнь"
    doc.add_paragraph("Дата поступления: 11.03.2026")
    doc.save(path)

    data = MedicalTextParser().parse_docx(path)
    assert data.fio == "Петров Петр Петрович"
    assert data.diagnosis.startswith("I10")


def test_universal_scanner_keeps_icd10_code_separate(tmp_path: Path) -> None:
    path = _write_docx(tmp_path / "scanner.docx", ["Диагноз: F06.7 лёгкое когнитивное расстройство"])
    scan = scan_docx(path)
    case = scan.patient_case()
    assert case.get("diagnosis.icd10") == "F06.7"
    assert "Дата" not in case.get("diagnosis.main")


def test_icd10_normalizer_accepts_manual_code() -> None:
    normalized = normalize_diagnosis_with_icd10("K35.8 Острый аппендицит")
    assert normalized.startswith("K35.8")
