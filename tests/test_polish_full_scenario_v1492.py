from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from actions_required_fields_popup import _is_admission_date_field, _is_case_number_field, _is_diagnosis_field
from dialog_fields_popup import DialogDiagnosisPopup
from diary_batch import fill_diary_batch
from diary_gender import detect_gender_from_patient_name
from i18n_strings import tr
from icd10_f_search import format_diagnosis, normalize_diagnosis_with_icd10, normalize_required_diagnosis_with_icd10, search_icd10_f
from medical_admission_resolver import extract_admission_date_from_primary_text
from medical_language_catalog import language_profile, normalize_language_id
from medical_language_detector import detect_text_language
from medical_parser import MedicalTextParser
from personal_document_buttons import (
    recognize_document_title_from_template,
    regular_document_role_choices,
    safe_profile_filename,
    suggest_button_label_for_template,
)
from regulatory_document_classifier import classify_docx
from universal_scanner import scan_docx


def _docx(path: Path, lines: list[str]) -> Path:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)
    return path


def _polish_primary_docx(path: Path) -> Path:
    return _docx(path, [
        "Karta informacyjna leczenia szpitalnego",
        "Pacjent: Anna Kowalska",
        "Nr historii choroby: 123/PL",
        "Data urodzenia: 04.01.1980",
        "Data przyjęcia: 2 czerwca 2026",
        "Data wypisu: 12.06.2026",
        "Rozpoznanie: Ostre zapalenie wyrostka robaczkowego K35.8",
        "Leczenie: appendektomia, antybiotykoterapia",
        "Zalecenia: kontrola w poradni",
    ])


def test_polish_language_is_first_class_runtime_language() -> None:
    detection = detect_text_language("Historia choroby. Rozpoznanie i leczenie. Data przyjęcia. Pacjent.")

    assert normalize_language_id("Polski") == "pl"
    assert language_profile("pl").native_name == "Polski"
    assert tr("button.language", "pl") == "Język"
    assert detection.language_id == "pl"
    assert detection.confidence >= 0.9


def test_polish_primary_document_parser_extracts_case_dates_diagnosis_and_treatment(tmp_path: Path) -> None:
    path = _polish_primary_docx(tmp_path / "karta_informacyjna.docx")
    text = "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
    parsed = MedicalTextParser().parse_docx(path)

    assert extract_admission_date_from_primary_text(text) == "02.06.2026"
    assert parsed.fio == "Anna Kowalska"
    assert parsed.case_number == "123/PL"
    assert parsed.birth == "04.01.1980"
    assert parsed.admission_date == "02.06.2026"
    assert parsed.diagnosis == "K35.8 Ostre zapalenie wyrostka robaczkowego"
    assert parsed.treatment_plan == "appendektomia, antybiotykoterapia"
    assert parsed.has_treatment_section is True


def test_polish_universal_scanner_extracts_constructor_fields(tmp_path: Path) -> None:
    path = _polish_primary_docx(tmp_path / "polish_source.docx")
    scan = scan_docx(path)
    best = scan.best_matches()

    assert scan.detected_language == "pl"
    assert scan.language_confidence >= 0.9
    assert best["document.title"].value == "Karta informacyjna leczenia szpitalnego"
    assert best["patient.fio"].value == "Anna Kowalska"
    assert best["case.number"].value == "123/PL"
    assert best["patient.birth_date"].value == "04.01.1980"
    assert best["admission.date"].value == "02.06.2026"
    assert best["discharge.date"].value == "12.06.2026"
    assert best["diagnosis.icd10"].value == "K35.8"
    assert best["diagnosis.main"].value == "Ostre zapalenie wyrostka robaczkowego K35.8"
    assert best["treatment.plan"].value == "appendektomia, antybiotykoterapia"
    assert best["recommendations"].value == "kontrola w poradni"
    assert scan.missing_field_ids() == ()


def test_polish_block03_button_uses_visible_word_title_and_polish_role(tmp_path: Path) -> None:
    path = _polish_primary_docx(tmp_path / "karta.docx")
    classification = classify_docx(path)
    suggestion = suggest_button_label_for_template(path, preferred_language="pl", ui_language="pl")

    assert recognize_document_title_from_template(path) == "Karta informacyjna leczenia szpitalnego"
    assert classification.role_id == "discharge_epicrisis"
    assert suggestion.language_id == "pl"
    assert suggestion.source_language == "pl"
    assert suggestion.role_id == "discharge_epicrisis"
    assert suggestion.label == "Karta informacyjna leczenia szpitalnego"
    assert suggestion.source == "template_top_title"
    assert safe_profile_filename("Chirurgia Łódź") == "Chirurgia_Łódź.medpack.json"
    assert "Karta informacyjna leczenia szpitalnego [discharge_epicrisis]" in regular_document_role_choices("pl")


def test_polish_required_popup_field_detection_and_icd_selector() -> None:
    assert DialogDiagnosisPopup.is_diagnosis_label("Rozpoznanie") is True
    assert _is_diagnosis_field(SimpleNamespace(key="Rozpoznanie")) is True
    assert _is_admission_date_field(SimpleNamespace(key="Data przyjęcia")) is True
    assert _is_case_number_field(SimpleNamespace(key="Nr historii choroby")) is True

    appendicitis = search_icd10_f("Ostre zapalenie wyrostka", language_id="pl", limit=3)
    assert appendicitis
    assert format_diagnosis(appendicitis[0], language_id="pl").startswith("K35 Ostre zapalenie")
    assert normalize_diagnosis_with_icd10("Ostre zapalenie wyrostka robaczkowego", language_id="pl").startswith("K35")
    assert normalize_required_diagnosis_with_icd10("Nadciśnienie tętnicze", language_id="pl").startswith("I10")
    assert normalize_required_diagnosis_with_icd10("same niekliniczne słowa", language_id="pl") == ""


def test_polish_diary_text_output_keeps_texts_and_adapts_gender(tmp_path: Path) -> None:
    status_path = _docx(tmp_path / "teksty_dziennikow.docx", ["Pacjent przyjęty. Hospitalizowany. Leczenie kontynuowane."])
    out_dir = tmp_path / "wynik"

    result = fill_diary_batch(
        status_files=[status_path],
        diary_files=[],
        output_dir=out_dir,
        patient_name="Anna Kowalska",
        gender_source_name="Anna Kowalska",
        admission_value="02.06.2026",
        discharge_value="06.06.2026",
        repeat_statuses=True,
        text_output=True,
    )

    assert detect_gender_from_patient_name("Anna Kowalska") == "female"
    assert result.created_files
    paragraphs = [paragraph.text for paragraph in Document(str(result.created_files[0])).paragraphs if paragraph.text.strip()]
    assert paragraphs[0].startswith("02.06.26 ")
    assert "Pacjentka przyjęta. Hospitalizowana." in paragraphs[0]
