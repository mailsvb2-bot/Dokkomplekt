from __future__ import annotations

from pathlib import Path

from docx import Document

from actions_universal_flow import ActionsUniversalFlowMixin
from desktop_patient_folder import PrimaryPatientFolderInfo, build_patient_folder_name_from_info
from universal_case_adapter import merge_case_values
from universal_fields import PatientCase, normalize_field_id, normalize_field_id_for_context
from universal_profiles import DocumentPack
from universal_template_engine import attach_template_to_pack, normalize_placeholder_id, render_template_to_docx


GLOBAL_ALIAS_CASES = [
    ("patientName", "patient.fio"),
    ("patient_name", "patient.fio"),
    ("fullName", "patient.fio"),
    ("ФИО", "patient.fio"),
    ("ф.и.о.", "patient.fio"),
    ("Пациент", "patient.fio"),
    ("caseNo", "case.number"),
    ("caseNumber", "case.number"),
    ("medicalRecordNo", "case.number"),
    ("История болезни №", "case.number"),
    ("№ истории болезни", "case.number"),
    ("admissionDate", "admission.date"),
    ("hospitalizationDate", "admission.date"),
    ("Дата поступления", "admission.date"),
    ("Дата госпитализации", "admission.date"),
    ("dischargeDate", "discharge.date"),
    ("Дата выписки", "discharge.date"),
    ("mainDiagnosis", "diagnosis.main"),
    ("clinicalDiagnosis", "diagnosis.main"),
    ("Диагноз", "diagnosis.main"),
    ("icdCode", "diagnosis.icd10"),
    ("Код МКБ-10", "diagnosis.icd10"),
    ("МКБ-10", "diagnosis.icd10"),
    ("treatmentPlan", "treatment.plan"),
    ("prescribedTreatment", "treatment.plan"),
    ("Назначенное лечение", "treatment.plan"),
    ("analysisResults", "labs.results"),
    ("labResults", "labs.results"),
    ("Лабораторные исследования", "labs.results"),
    ("LABS_BLOCK", "labs.results"),
    ("labsDate", "labs.date"),
    ("Дата анализов", "labs.date"),
    ("Дата лабораторных исследований", "labs.date"),
    ("expertSickLeaveNumber", "expert.sick_leave_number"),
    ("Номер больничного", "expert.sick_leave_number"),
    ("Номер ЛН", "expert.sick_leave_number"),
    ("sickLeaveFrom", "expert.sick_leave_from"),
    ("С какого числа больничный", "expert.sick_leave_from"),
    ("Нужен ЛН", "expert.sick_leave_needed"),
    ("workPlace", "patient.work"),
    ("Место работы", "patient.work"),
    ("jobTitle", "patient.position"),
    ("Должность", "patient.position"),
    ("rvkActNumber", "rvk.act_number"),
    ("Номер медицинского заключения РВК", "rvk.act_number"),
    ("Военкомат", "rvk.military_commissariat"),
    ("vkMseProtocolNumber", "vk_mse.protocol_number"),
    ("sickLeaveVkProtocolNumber", "sick_leave_vk.protocol_number"),
]


CONTEXT_ALIAS_CASES = [
    ("Место работы", "vk_mse", "vk_mse.work"),
    ("Должность", "vk_mse", "vk_mse.position"),
    ("Место работы / должность", "vk_mse", "vk_mse.work_position"),
    ("Дата комиссии", "vk_mse", "vk_mse.date"),
    ("Номер протокола ВК", "vk_mse", "vk_mse.protocol_number"),
    ("Дата протокола", "vk_mse", "vk_mse.protocol_date"),
    ("От", "vk_mse", "vk_mse.protocol_date"),
    ("Место работы", "sickLeaveVk", "sick_leave_vk.work"),
    ("Должность", "sickLeaveVk", "sick_leave_vk.position"),
    ("Место работы / должность", "sickLeaveVk", "sick_leave_vk.work_position"),
    ("Дата комиссии", "sickLeaveVk", "sick_leave_vk.commission_date"),
    ("Номер протокола ВК", "sickLeaveVk", "sick_leave_vk.protocol_number"),
    ("Дата протокола", "sickLeaveVk", "sick_leave_vk.protocol_date"),
    ("От", "sickLeaveVk", "sick_leave_vk.protocol_date"),
    ("Место работы / должность", "rvk", "rvk.work_position"),
    ("Работа и должность", "rvk", "rvk.work_position"),
    ("Номер медицинского заключения", "rvk", "rvk.act_number"),
    ("Военный комиссариат", "rvk", "rvk.military_commissariat"),
    ("Дата комиссии", "jointMedicalExam", "commission.date"),
    ("Номер", "jointMedicalExam", "commission.number"),
    ("Номер комиссии", "jointMedicalExam", "commission.number"),
]


class _Var:
    def __init__(self, value="") -> None:
        self.value = value

    def get(self):
        return self.value


class _OverlayHarness(ActionsUniversalFlowMixin):
    def __init__(self) -> None:
        self.patient_name_var = _Var("Иванов Иван Иванович")
        self.case_number_var = _Var("ИБ-100")
        self.diagnosis_var = _Var("I10 Гипертензия")
        self._popup_diagnosis_override = ""
        self.assigned_treatment_var = _Var("Лечение из popup")
        self.labs_without_var = _Var(False)
        self.labs_text_var = _Var("Hb 140")
        self.labs_source_path_var = _Var("popup")
        self.labs_date_policy_var = _Var("preserve_found_dates")
        self.expert_work_status_var = _Var("да")
        self.expert_work_org_var = _Var("ООО Ромашка")
        self.expert_position_var = _Var("инженер")
        self.expert_sick_leave_needed_var = _Var("да")
        self.expert_sick_leave_number_var = _Var("ЛН-555")
        self.rvk_act_number_var = _Var("РВК-42")
        self.rvk_military_commissariat_var = _Var("Военкомат № 1")
        self.rvk_work_position_var = _Var("ООО Ромашка / инженер")
        self.commission_number_var = _Var("КОМ-33")
        self.vk_protocol_number_var = _Var("ВК-88")
        self.vk_mse_work_org_var = _Var("АО Место")
        self.vk_mse_position_var = _Var("мастер")
        self.sick_leave_vk_protocol_number_var = _Var("БЛ-99")
        self.sick_leave_vk_work_org_var = _Var("ИП Работа")
        self.sick_leave_vk_position_var = _Var("оператор")
        self.sick_leave_vk_work_position_var = _Var("ИП Работа / оператор")
        self._semantic_date_state = {}

    def _normalize_yes_no(self, value: str) -> str:
        return str(value or "").strip().lower()


def _write_docx(path: Path, *paragraphs: str) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)




def _make_global_alias_test(raw: str, expected: str):
    def _test() -> None:
        assert normalize_field_id(raw) == expected
        assert normalize_placeholder_id(raw) == expected
    return _test


for _index, (_raw, _expected) in enumerate(GLOBAL_ALIAS_CASES, start=1):
    globals()[f"test_global_semantic_alias_matrix_{_index:02d}"] = _make_global_alias_test(_raw, _expected)


def _make_context_alias_test(raw: str, role_id: str, expected: str):
    def _test() -> None:
        assert normalize_field_id_for_context(raw, role_id=role_id) == expected
        assert normalize_placeholder_id(raw, role_id=role_id) == expected
    return _test


for _index, (_raw, _role_id, _expected) in enumerate(CONTEXT_ALIAS_CASES, start=1):
    globals()[f"test_context_semantic_alias_matrix_{_index:02d}"] = _make_context_alias_test(_raw, _role_id, _expected)

def test_ambiguous_work_position_label_requires_document_context() -> None:
    try:
        normalize_field_id("Место работы / должность")
    except ValueError:
        return
    raise AssertionError("Ambiguous human label must require document context")


def test_vkmse_combined_work_position_is_rendered_from_popup_state(tmp_path: Path) -> None:
    app = _OverlayHarness()
    values = app._confirmed_universal_overlay_values()
    assert values["vk_mse.work"] == "АО Место"
    assert values["vk_mse.position"] == "мастер"
    assert values["vk_mse.work_position"] == "АО Место, мастер"

    case = merge_case_values(PatientCase(), values, source_document="popup_state")
    template = tmp_path / "vk_mse.docx"
    output = tmp_path / "out.docx"
    _write_docx(template, "ВК {{Место работы / должность}} | Протокол {{Номер протокола ВК}}")

    pack = DocumentPack(pack_id="doctor.vkmse", name="Профиль")
    spec, copied = attach_template_to_pack(pack, template, tmp_path / "profile", button_label="ВК на МСЭ", role_id="vk_mse")
    assert spec.required_fields == ("vk_mse.work_position", "vk_mse.protocol_number")
    result = render_template_to_docx(template_path=copied, output_path=output, case=case, document=spec)
    assert result.ok
    text = _read_docx(output)
    assert "ВК АО Место, мастер" in text
    assert "Протокол ВК-88" in text


def test_legacy_pending_path_is_used_once_without_reintroducing_pathful_state(tmp_path: Path) -> None:
    import desktop_intake_agent

    pending_path = tmp_path / "pending.docx"
    _write_docx(pending_path, "primary")
    signature = "d" * 64
    seen: set[str] = set()
    active, changed = desktop_intake_agent._resolve_pending_state(
        {"pending": {"path": str(pending_path), "signature": signature, "launched_at": 0}},
        seen,
    )
    assert active == {}
    assert changed is True
    assert signature not in seen

    pending_path.unlink()
    active, changed = desktop_intake_agent._resolve_pending_state(
        {"pending": {"path": str(pending_path), "signature": signature, "launched_at": 0}},
        seen,
    )
    assert active == {}
    assert changed is True
    assert signature in seen
    assert "path" not in desktop_intake_agent._pending_from_state(
        {"pending": {"path": str(pending_path), "signature": signature, "launched_at": 0}}
    )



FOLDER_NAMING_CASES = [
    ({"parts": ["surname_initials", "admission_discharge_dates"], "date_format": "short", "doctor_confirmed": True}, "Петров Пётр Петрович", "01.06.2026", "12.06.2026", "Петров П.П. 01.06.26-12.06.26"),
    ({"parts": ["full_fio", "discharge_month"], "date_format": "long", "doctor_confirmed": True}, "Сидоров Сергей Сергеевич", "05.05.2026", "20.06.2026", "Сидоров Сергей Сергеевич июнь 2026"),
    ({"parts": ["surname_initials", "admission_date"], "date_format": "short", "doctor_confirmed": True}, "Иванова Мария Петровна", "03.04.2026", "", "Иванова М.П. 03.04.26"),
]


def _make_folder_naming_test(settings, fio: str, admission: str, discharge: str, expected: str):
    def _test() -> None:
        info = PrimaryPatientFolderInfo(fio=fio, admission_date=admission, folder_name="legacy-default")
        assert build_patient_folder_name_from_info(info, settings=settings, discharge_date=discharge) == expected
    return _test


for _index, (_settings, _fio, _admission, _discharge, _expected) in enumerate(FOLDER_NAMING_CASES, start=1):
    globals()[f"test_folder_naming_matrix_keeps_confirmed_principle_{_index:02d}"] = _make_folder_naming_test(
        _settings, _fio, _admission, _discharge, _expected
    )
