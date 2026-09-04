from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import hmac
import json
from pathlib import Path

from docx import Document

import desktop_intake_agent as agent
from medical_constants import RVK_COMMISSARIAT_PRESETS
from product_access import ProductAccessManager, machine_fingerprint, stable_json
from universal_fields import PatientCase
from universal_main_documents import custom_requirement_flags_for_documents
from universal_profiles import DocumentPack
from universal_template_engine import attach_template_to_pack
from universal_generation import render_documents_from_pack


def _legacy_v2_key() -> bytes:
    seed = f"dokkomplekt-product-access-v2|{machine_fingerprint()}"
    return hashlib.sha256(seed.encode("utf-8", errors="replace")).digest()


def _write_v2_trial(path: Path, *, started: str, used: int) -> None:
    payload = {
        "state_version": 2,
        "trial_started_at": started,
        "usage_by_month": {"2026-08": used},
        "trial_created_total": used,
    }
    payload["_state_mac"] = hmac.new(
        _legacy_v2_key(), stable_json(payload).encode("utf-8"), hashlib.sha256
    ).hexdigest()
    (path / "product_access_state.json").write_text(json.dumps(payload), encoding="utf-8")
    (path / "product_access_guard.json").write_text(json.dumps(payload), encoding="utf-8")


def test_first_public_release_resets_only_pre_release_trial_once(tmp_path: Path) -> None:
    _write_v2_trial(tmp_path, started="2026-06-20T00:00:00+00:00", used=30)
    first_now = datetime(2026, 9, 4, 9, 0, tzinfo=timezone.utc)
    manager = ProductAccessManager(tmp_path, now=first_now)
    state = manager.current_state()

    assert state.active is True
    assert state.documents_used_total_trial == 0
    assert state.trial_started_at.startswith("2026-09-04T09:00:00")
    stored = json.loads(manager.state_path.read_text("utf-8"))
    assert stored["trial_epoch"] == "v1.4.91-public"
    assert stored["trial_created_total"] == 0

    # The integrity-protected epoch prevents reinstall/update from refreshing
    # the public trial again after its genuine 14-day window has elapsed.
    later = ProductAccessManager(tmp_path, now=datetime(2026, 9, 25, 9, 0, tzinfo=timezone.utc)).current_state()
    assert later.active is False
    assert later.reason == "trial_expired"
    assert later.trial_started_at == state.trial_started_at


def test_pre_release_v3_trial_is_also_reset_once(tmp_path: Path) -> None:
    old = ProductAccessManager(tmp_path, now=datetime(2026, 6, 20, 8, 0, tzinfo=timezone.utc))
    old._save_state_payload({
        "trial_started_at": "2026-06-20T08:00:00+00:00",
        "usage_by_month": {"2026-06": 30},
        "trial_created_total": 30,
    })

    current = ProductAccessManager(tmp_path, now=datetime(2026, 9, 4, 10, 0, tzinfo=timezone.utc)).current_state()

    assert current.active is True
    assert current.documents_used_total_trial == 0
    assert current.trial_started_at.startswith("2026-09-04T10:00:00")
    payload = json.loads((tmp_path / "product_access_state.json").read_text("utf-8"))
    assert payload["trial_epoch"] == "v1.4.91-public"


def test_rvk_doctor_template_generates_complete_docx(tmp_path: Path) -> None:
    template = tmp_path / "rvk.docx"
    doc = Document()
    doc.add_paragraph(
        "ФИО {{patientName}} | ИБ {{caseNo}} | Акт № {{rvk.act_number}} | "
        "Военкомат {{rvk.military_commissariat}} | Выписка {{discharge.date}} | "
        "Диагноз {{mainDiagnosis}} | Лечение {{treatment.plan}}"
    )
    doc.save(template)

    profile_dir = tmp_path / "profile"
    pack = DocumentPack(pack_id="doctor.rvk", name="РВК")
    spec, _copied = attach_template_to_pack(
        pack, template, profile_dir, button_label="Акт для РВК", role_id="rvk"
    )
    flags = custom_requirement_flags_for_documents((spec,))
    assert flags["rvk"] is True
    assert flags["requires_discharge_date"] is True
    assert flags["requires_case_number"] is True
    assert flags["requires_diagnosis"] is True
    assert flags["requires_treatment"] is True

    case = PatientCase()
    for field_id, value in {
        "patient.fio": "Иванов Иван Иванович",
        "case.number": "777",
        "rvk.act_number": "РВК-42",
        "rvk.military_commissariat": "Ленинский",
        "discharge.date": "20.06.2026",
        "diagnosis.main": "F20.0",
        "treatment.plan": "Терапия по назначению врача",
    }.items():
        case.set(field_id, value, source_document="test")

    result = render_documents_from_pack(
        pack=pack,
        case=case,
        document_ids=[spec.id],
        output_dir=tmp_path / "out",
        base_dir=profile_dir,
        strict=True,
    )
    assert result.ok
    assert len(result.created_files) == 1
    rendered = Document(result.created_files[0])
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Акт № РВК-42" in text
    assert "Военкомат Ленинский" in text
    assert "Выписка 20.06.2026" in text


def test_rvk_popup_keeps_required_commissariat_presets_and_manual_entry() -> None:
    assert RVK_COMMISSARIAT_PRESETS == ("Ленинский", "Канавинский", "Сормовский", "Московский")
    source = Path("dialog_document_details.py").read_text(encoding="utf-8")
    assert 'add_entry("Военкомат / организация направления", military_var' in source


def test_uninstall_requests_live_agent_shutdown_before_files_are_removed(tmp_path: Path, monkeypatch) -> None:
    startup = tmp_path / "agent.vbs"
    shortcut = tmp_path / "agent.lnk"
    handoff = tmp_path / "handoff.json"
    lock = tmp_path / "agent.lock"
    startup.write_text("x", encoding="utf-8")
    shortcut.write_text("x", encoding="utf-8")
    lock.write_text("pid=1", encoding="utf-8")

    monkeypatch.setattr(agent, "_startup_agent_script_path", lambda: startup)
    monkeypatch.setattr(agent, "_legacy_startup_shortcut_path", lambda: shortcut)
    monkeypatch.setattr(agent, "_handoff_path", lambda: handoff)
    monkeypatch.setattr(agent, "_lock_path", lambda: lock)
    monkeypatch.setattr(agent, "_gui_lock_path", lambda: tmp_path / "gui.lock")
    monkeypatch.setattr(agent, "is_gui_runtime_active", lambda: False)
    states = iter((True, False, False))
    monkeypatch.setattr(agent, "_agent_lock_has_live_owner", lambda: next(states, False))
    monkeypatch.setattr(agent.time, "sleep", lambda _seconds: None)

    ok, message = agent.uninstall_agent_autostart(wait_seconds=1.0)

    assert ok is True, message
    assert not startup.exists()
    assert not shortcut.exists()
    assert not lock.exists()
    assert not handoff.exists()


def test_agent_obeys_explicit_uninstall_shutdown_handoff(monkeypatch) -> None:
    monkeypatch.setattr(agent, "_read_agent_handoff", lambda: {"shutdown_requested": True})
    assert agent._agent_is_retired() is True


def test_foreground_gui_heartbeat_obeys_uninstall_shutdown_contract() -> None:
    source = Path("desktop_intake_mixin.py").read_text(encoding="utf-8")
    refresh = source[source.index("def _refresh_gui_runtime_lock"):source.index("def _close_app_with_runtime_lock_release")]
    assert "uninstall_shutdown_requested" in refresh
    assert "self._close_app_with_runtime_lock_release()" in refresh


def test_installer_aborts_before_delete_when_runtime_shutdown_fails() -> None:
    script = Path("installer/Dokkomplekt.iss").read_text(encoding="utf-8")
    assert "function InitializeUninstall(): Boolean;" in script
    assert "--uninstall-intake-agent" in script
    assert "Result := False;" in script
    assert "SuppressibleMsgBox" in script
    assert "ArchitecturesAllowed=x64compatible" in script


def test_current_public_trial_epoch_ignores_and_heals_stale_legacy_guard(tmp_path: Path) -> None:
    """A stale pre-public redundant copy must not re-expire a reset public trial."""
    current_now = datetime(2026, 9, 4, 11, 0, tzinfo=timezone.utc)
    manager = ProductAccessManager(tmp_path, now=current_now)
    manager._save_state_payload({
        "trial_started_at": "2026-09-04T11:00:00+00:00",
        "usage_by_month": {},
        "trial_created_total": 0,
        "trial_epoch": "v1.4.91-public",
        "trial_public_reset_at": "2026-09-04T11:00:00+00:00",
    })

    # Simulate the exact Windows failure mode: files were updated, but one old
    # redundant owner (Registry in production; guard file here) survived the
    # public-reset write and still carries the months-old pre-release trial.
    legacy = {
        "state_version": 2,
        "trial_started_at": "2026-06-20T00:00:00+00:00",
        "usage_by_month": {"2026-06": 30},
        "trial_created_total": 30,
    }
    legacy["_state_mac"] = hmac.new(
        _legacy_v2_key(), stable_json(legacy).encode("utf-8"), hashlib.sha256
    ).hexdigest()
    manager.state_guard_path.write_text(json.dumps(legacy), encoding="utf-8")

    state = ProductAccessManager(tmp_path, now=current_now).current_state()

    assert state.active is True
    assert state.reason == "trial_active"
    assert state.documents_used_total_trial == 0
    assert state.trial_started_at == "2026-09-04T11:00:00+00:00"
    healed_guard = json.loads((tmp_path / "product_access_guard.json").read_text("utf-8"))
    assert healed_guard["trial_epoch"] == "v1.4.91-public"
    assert healed_guard["trial_started_at"] == "2026-09-04T11:00:00+00:00"
    assert healed_guard["trial_created_total"] == 0


def test_trial_allows_normal_multi_document_patient_set_up_to_total_allowance(tmp_path: Path) -> None:
    manager = ProductAccessManager(tmp_path, now=datetime(2026, 9, 4, 12, 0, tzinfo=timezone.utc))
    decision = manager.check_document_creation(8)
    assert decision.allowed is True
    assert decision.code == "ok_trial"
    assert decision.state.documents_limit_month == 30


def test_fresh_trial_reports_real_days_left(tmp_path: Path) -> None:
    manager = ProductAccessManager(tmp_path, now=datetime(2026, 9, 4, 12, 0, tzinfo=timezone.utc))
    state = manager.current_state()
    assert state.active is True
    assert state.days_left == 14
