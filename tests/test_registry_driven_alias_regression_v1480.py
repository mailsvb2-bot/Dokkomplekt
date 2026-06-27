from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from actions_creation_preflight import ActionsCreationReviewMixin
from actions_required_fields_popup import (
    _RequiredFieldsDialog,
    _is_semantic_date_field,
    _store_key_for_field,
)
from universal_fields import PatientCase, normalize_field_id
from universal_template_engine import extract_template_placeholders


class _Var:
    def __init__(self, value: str = "") -> None:
        self.value = value

    def get(self) -> str:
        return self.value

    def set(self, value: str) -> None:
        self.value = value


def _field(key: str, label: str = "", value: str = ""):
    return SimpleNamespace(key=key, field_id=key, label=label, placeholder="", reason="", value=value)


class _PopupApp:
    def __init__(self) -> None:
        self.stored: list[tuple[str, str]] = []
        self.dates: list[tuple[str, str]] = []

    def _store_required_review_value(self, key: str, value: str) -> None:
        self.stored.append((key, value))

    def _store_popup_date_value(self, key: str, value: str, **_kwargs) -> bool:
        self.dates.append((key, value))
        return True


class _PreflightApp(ActionsCreationReviewMixin):
    def __init__(self) -> None:
        self.data = SimpleNamespace(
            expert_sick_leave_number="",
            expert_sick_leave_needed="",
            expert_sick_leave_from="",
            commission_number="",
            rvk_act_number="",
        )
        self.expert_sick_leave_number_var = _Var()
        self.expert_sick_leave_needed_var = _Var()
        self.expert_sick_leave_from_var = _Var()
        self.commission_number_var = _Var()
        self.rvk_act_number_var = _Var()
        self.dates: list[tuple[str, str]] = []
        self.display_updates = 0

    def _normalize_yes_no(self, value: str) -> str:
        return "да" if str(value).strip().lower() in {"да", "yes", "нужен"} else "нет"

    def _store_popup_date_value(self, key: str, value: str, **_kwargs) -> bool:
        self.dates.append((key, value))
        getattr(self, f"{key}_var").set(value)
        return True

    def _update_expert_sick_leave_display(self) -> None:
        self.display_updates += 1


def test_registry_definitions_drive_human_placeholder_aliases(tmp_path: Path):
    template = tmp_path / "human_aliases.docx"
    doc = Document()
    doc.add_paragraph("{{МКБ-10}} {{Код МКБ-10}} {{Дата анализов}} {{Дата лабораторных исследований}}")
    doc.add_paragraph("{{Номер больничного}} {{С какого числа больничный}} {{Дата совместного осмотра}}")
    doc.add_paragraph("{{Номер медицинского заключения РВК}}")
    doc.save(template)

    fields = [item.field_id for item in extract_template_placeholders(template)]

    assert fields == [
        "diagnosis.icd10",
        "diagnosis.icd10",
        "labs.date",
        "labs.date",
        "expert.sick_leave_number",
        "expert.sick_leave_from",
        "commission.date",
        "rvk.act_number",
    ]


def test_field_normalizer_uses_registry_labels_and_export_forms():
    samples = [
        ("ICD-10", "diagnosis.icd10"),
        ("MKB-10", "diagnosis.icd10"),
        ("МКБ-10", "diagnosis.icd10"),
        ("код МКБ-10", "diagnosis.icd10"),
        ("labsDate", "labs.date"),
        ("Дата анализов", "labs.date"),
        ("Дата лабораторных исследований", "labs.date"),
        ("sickLeaveNumber", "expert.sick_leave_number"),
        ("Номер больничного", "expert.sick_leave_number"),
        ("С какого числа больничный", "expert.sick_leave_from"),
        ("Нужен ЛН", "expert.sick_leave_needed"),
        ("Дата совместного осмотра", "commission.date"),
        ("Номер медицинского заключения РВК", "rvk.act_number"),
    ]
    for raw, expected in samples:
        assert normalize_field_id(raw) == expected


def test_patient_case_accepts_new_registry_aliases_without_second_state():
    case = PatientCase()
    case.set("код МКБ-10", "I10")
    case.set("Дата анализов", "10.06.2026")
    case.set("Номер больничного", "ЛН-55")
    case.set("С какого числа больничный", "12.06.2026")
    case.set("Дата совместного осмотра", "13.06.2026")

    assert case.get("diagnosis.icd10") == "I10"
    assert case.get("labs.date") == "10.06.2026"
    assert case.get("expert.sick_leave_number") == "ЛН-55"
    assert case.get("expert.sick_leave_from") == "12.06.2026"
    assert case.get("commission.date") == "13.06.2026"


def test_required_popup_store_keys_cover_registry_aliases():
    assert _store_key_for_field(_field("Дата анализов", "Дата анализов")) == "labs_explicit_date"
    assert _is_semantic_date_field(_field("Дата анализов", "Дата анализов"))
    assert _store_key_for_field(_field("Номер больничного", "Номер больничного")) == "expert_sick_leave_number"
    assert _store_key_for_field(_field("С какого числа больничный", "С какого числа больничный")) == "expert_sick_leave_from"
    assert _is_semantic_date_field(_field("С какого числа больничный", "С какого числа больничный"))
    assert _store_key_for_field(_field("Нужен ЛН", "Нужен ЛН")) == "expert_sick_leave_needed"
    assert _store_key_for_field(_field("Номер медицинского заключения РВК", "Акт РВК №")) == "rvk_act_number"


def test_required_popup_routes_semantic_dates_and_text_targets_to_existing_state():
    app = _PopupApp()
    fields = [
        _field("Дата анализов", "Дата анализов"),
        _field("С какого числа больничный", "С какого числа больничный"),
        _field("Номер больничного", "Номер больничного"),
    ]
    dialog = _RequiredFieldsDialog(app, SimpleNamespace(), fields)
    dialog.win = SimpleNamespace()

    assert dialog._store_value("Дата анализов", "10.06.2026")
    assert dialog._store_value("С какого числа больничный", "12.06.2026")
    assert dialog._store_value("Номер больничного", "ЛН-55")

    assert app.dates == [("labs_explicit_date", "10.06.2026"), ("expert_sick_leave_from", "12.06.2026")]
    assert app.stored == [("expert_sick_leave_number", "ЛН-55")]


def test_required_review_storage_covers_expert_and_regulatory_fields():
    app = _PreflightApp()

    app._store_required_review_value("expert_sick_leave_needed", "да")
    app._store_required_review_value("expert_sick_leave_number", "ЛН-55")
    app._store_required_review_value("expert_sick_leave_from", "12.06.2026")
    app._store_required_review_value("commission_number", "К-7")
    app._store_required_review_value("rvk_act_number", "РВК-9")

    assert app.expert_sick_leave_needed_var.get() == "да"
    assert app.expert_sick_leave_number_var.get() == "ЛН-55"
    assert app.expert_sick_leave_from_var.get() == "12.06.2026"
    assert app.commission_number_var.get() == "К-7"
    assert app.rvk_act_number_var.get() == "РВК-9"
    assert app.data.expert_sick_leave_needed == "да"
    assert app.data.expert_sick_leave_number == "ЛН-55"
    assert app.data.commission_number == "К-7"
    assert app.data.rvk_act_number == "РВК-9"
    assert app.dates == [("expert_sick_leave_from", "12.06.2026")]
    assert app.display_updates >= 2
