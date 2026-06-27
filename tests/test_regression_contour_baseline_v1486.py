from __future__ import annotations

from pathlib import Path

from docx import Document

from actions_universal_flow import ActionsUniversalFlowMixin
from desktop_patient_folder import PrimaryPatientFolderInfo, build_patient_folder_name_from_info, folder_naming_uses_discharge_date
from universal_case_adapter import merge_case_values, merge_patient_cases
from universal_fields import PatientCase
from universal_generation import analyze_pack_readiness, render_documents_from_pack
from universal_profiles import DocumentPack, current_builtin_documents
from universal_template_engine import attach_template_to_pack, import_document_pack_zip, export_document_pack_zip

ROOT = Path(__file__).resolve().parents[1]


class _Var:
    def __init__(self, value="") -> None:
        self.value = value

    def get(self):
        return self.value


class _DoctorPopupHarness(ActionsUniversalFlowMixin):
    def __init__(self) -> None:
        self.patient_name_var = _Var("Иванов Иван Иванович")
        self.case_number_var = _Var("UI-777")
        self.diagnosis_var = _Var("I10 Гипертензия")
        self._popup_diagnosis_override = ""
        self.assigned_treatment_var = _Var("Лечение из popup")
        self.labs_without_var = _Var(False)
        self.labs_text_var = _Var("Hb 140")
        self.labs_source_path_var = _Var("popup")
        self.labs_date_policy_var = _Var("preserve_found_dates")
        self.expert_work_status_var = _Var("да")
        self.expert_work_org_var = _Var("ООО Ромашка")
        self.expert_position_var = _Var("инженер")
        self.expert_sick_leave_needed_var = _Var("да")
        self.expert_sick_leave_number_var = _Var("ЛН-555")
        self.rvk_act_number_var = _Var("РВК-42")
        self.rvk_military_commissariat_var = _Var("Военкомат № 1")
        self.rvk_work_position_var = _Var("ООО Ромашка / инженер")
        self.commission_number_var = _Var("КОМ-33")
        self.vk_protocol_number_var = _Var("ВК-88")
        self.vk_mse_work_org_var = _Var("АО Место")
        self.vk_mse_position_var = _Var("мастер")
        self.sick_leave_vk_protocol_number_var = _Var("БЛ-99")
        self.sick_leave_vk_work_org_var = _Var("ИП Работа")
        self.sick_leave_vk_position_var = _Var("оператор")
        self.sick_leave_vk_work_position_var = _Var("ИП Работа / оператор")
        self._semantic_date_state = {}
        self._dates = {
            "admission_date": "01.06.2026",
            "discharge_date": "12.06.2026",
            "labs_explicit_date": "03.06.2026",
            "expert_sick_leave_from": "02.06.2026",
            "vk_date": "11.06.2026",
            "vk_protocol_date": "11.06.2026",
            "sick_leave_vk_date": "09.06.2026",
            "sick_leave_vk_protocol_date": "09.06.2026",
            "sick_leave_vk_commission_date": "10.06.2026",
            "commission_date": "10.06.2026",
        }

    def _normalize_yes_no(self, value: str) -> str:
        return "да" if str(value or "").strip().lower() in {"да", "yes", "true", "1"} else "нет"


def _write_docx(path: Path, *paragraphs: str) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _case_from_popup(monkeypatch) -> PatientCase:
    app = _DoctorPopupHarness()

    def fake_current_semantic_date(_app, key):
        return app._dates.get(key, "")

    import actions_universal_flow

    monkeypatch.setattr(actions_universal_flow, "current_semantic_date", fake_current_semantic_date)

    scanner_case = PatientCase()
    scanner_case.set("case.number", "OLD-001", confidence=0.99, source_document="source_scan")
    scanner_case.set("diagnosis.main", "OLD diagnosis", confidence=0.99, source_document="source_scan")
    scanner_case.set("expert.sick_leave_number", "OLD-LN", confidence=0.99, source_document="source_scan")
    ui_case = PatientCase()
    ui_case.set("case.number", "UI-777", confidence=0.90, source_document="ui")
    regressed_without_final_overlay = merge_patient_cases(ui_case, scanner_case)
    assert regressed_without_final_overlay.get("case.number") == "OLD-001"

    return merge_case_values(regressed_without_final_overlay, app._confirmed_universal_overlay_values(), source_document="doctor_confirmed_ui_state")


def test_regression_contour_is_wired_into_docs_build_and_ci() -> None:
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    contour = (ROOT / "REGRESSION_CONTOUR.md").read_text(encoding="utf-8")
    matrix = (ROOT / "REGRESSION_MATRIX.md").read_text(encoding="utf-8")
    workflow = (ROOT / ".github/workflows/windows-build.yml").read_text(encoding="utf-8")
    build = (ROOT / "build_exe_windows.bat").read_text(encoding="utf-8", errors="replace")
    release_check = (ROOT / "release_check.py").read_text(encoding="utf-8")

    required = "python tools/run_regression_contour.py"
    assert required in readme
    assert required in contour
    assert required in workflow
    assert required in build
    assert "REGRESSION_CONTOUR.md" in release_check
    assert "REGRESSION_MATRIX.md" in release_check
    for phrase in [
        "Doctor-owned constructor",
        "Popup values",
        "Folder naming",
        "Archive/CI hygiene",
    ]:
        assert phrase in matrix


def test_no_bundled_templates_or_builtin_documents_returned() -> None:
    embedded = (ROOT / "embedded_templates.py").read_text(encoding="utf-8")
    assert "TEMPLATE_B64: dict[str, str] = {}" in embedded
    assert current_builtin_documents() == ()


def test_full_doctor_regression_replay_from_template_to_output(tmp_path: Path, monkeypatch) -> None:
    source_templates = tmp_path / "source_templates"
    source_templates.mkdir()
    sick_template = source_templates / "sick_vk.docx"
    rvk_template = source_templates / "rvk.docx"
    removable_template = source_templates / "removable.docx"
    _write_docx(
        sick_template,
        "ФИО {{patientName}} | ИБ {{caseNo}} | Диагноз {{mainDiagnosis}}",
        "ЛН {{expertSickLeaveNumber}} | Протокол {{Номер протокола ВК}} | Комиссия {{Дата проведения комиссии}}",
    )
    _write_docx(rvk_template, "РВК {{rvk.act_number}} | Работа {{Место работы / должность}}")
    _write_docx(removable_template, "Удаляемая кнопка {{patientName}}")

    profile_dir = tmp_path / "profile"
    pack = DocumentPack(pack_id="doctor.regression", name="Профиль врача")
    sick_spec, sick_copy = attach_template_to_pack(pack, sick_template, profile_dir, button_label="ВК больничный", role_id="sickLeaveVk")
    rvk_spec, rvk_copy = attach_template_to_pack(pack, rvk_template, profile_dir, button_label="Акт для РВК", role_id="rvk")
    removable_spec, removable_copy = attach_template_to_pack(pack, removable_template, profile_dir, button_label="Лишняя кнопка", role_id="dischargeEpicrisis")

    renamed = pack.rename_document(sick_spec.id, "ВК больничный контроль")
    removed = pack.remove_document(removable_spec.id)

    assert renamed.id == sick_spec.id
    assert renamed.template == sick_spec.template
    assert renamed.required_fields == sick_spec.required_fields
    assert renamed.button_label == "ВК больничный контроль"
    assert removed.id == removable_spec.id
    assert removable_copy.exists(), "Удаление кнопки не должно удалять скопированный DOCX врача"
    assert pack.document_by_id(removable_spec.id) is None
    assert pack.document_by_id(rvk_spec.id) is not None
    assert sick_copy.exists() and rvk_copy.exists()
    assert renamed.required_fields == (
        "patient.fio",
        "case.number",
        "diagnosis.main",
        "expert.sick_leave_number",
        "sick_leave_vk.protocol_number",
        "sick_leave_vk.commission_date",
    )
    assert rvk_spec.required_fields == ("rvk.act_number", "rvk.work_position")

    case = _case_from_popup(monkeypatch)
    assert case.get("case.number") == "UI-777"
    assert case.get("diagnosis.main") == "I10 Гипертензия"
    assert case.get("expert.sick_leave_number") == "ЛН-555"
    assert case.get("sick_leave_vk.protocol_number") == "БЛ-99"
    assert case.get("sick_leave_vk.commission_date") == "10.06.2026"
    assert case.get("rvk.act_number") == "РВК-42"
    assert case.get("rvk.work_position") == "ООО Ромашка / инженер"

    readiness = analyze_pack_readiness(pack, case, base_dir=profile_dir)
    assert set(readiness.ready_document_ids) == {sick_spec.id, rvk_spec.id}
    assert not readiness.blocked_document_ids

    out_dir = tmp_path / "generated"
    result = render_documents_from_pack(
        pack=pack,
        case=case,
        document_ids=[sick_spec.id, rvk_spec.id],
        output_dir=out_dir,
        base_dir=profile_dir,
        strict=True,
    )

    assert result.ok
    assert len(result.created_files) == 2
    combined = "\n".join(_read_docx(Path(path)) for path in result.created_files)
    assert "ФИО Иванов Иван Иванович" in combined
    assert "ИБ UI-777" in combined
    assert "Диагноз I10 Гипертензия" in combined
    assert "ЛН ЛН-555" in combined
    assert "Протокол БЛ-99" in combined
    assert "Комиссия 10.06.2026" in combined
    assert "РВК РВК-42" in combined
    assert "Работа ООО Ромашка / инженер" in combined
    assert "OLD-001" not in combined
    assert "OLD-LN" not in combined


def test_folder_naming_contract_is_part_of_regression_contour() -> None:
    info = PrimaryPatientFolderInfo(
        fio="Петров Пётр Петрович",
        admission_date="01.06.2026",
        folder_name="legacy-default",
    )
    settings = {
        "parts": ["surname_initials", "admission_discharge_dates"],
        "date_format": "short",
        "doctor_confirmed": True,
    }
    assert folder_naming_uses_discharge_date(settings)
    assert build_patient_folder_name_from_info(info, settings=settings, discharge_date="12.06.2026") == "Петров П.П. 01.06.26-12.06.26"


def test_medpack_export_import_keeps_template_paths_portable_and_buttons_profile_owned(tmp_path: Path) -> None:
    template = tmp_path / "source.docx"
    _write_docx(template, "{{patientName}} {{caseNo}}")
    profile_dir = tmp_path / "profile"
    pack = DocumentPack(pack_id="doctor.portable", name="Профиль врача")
    spec, copied = attach_template_to_pack(pack, template, profile_dir, button_label="Первичная кнопка", role_id="dischargeEpicrisis")
    pack.rename_document(spec.id, "Переименованная кнопка")

    archive = tmp_path / "profile.medpack.zip"
    export_document_pack_zip(pack, archive, template_base_dir=profile_dir)
    imported_pack, imported_manifest = import_document_pack_zip(archive, tmp_path / "imported")

    assert imported_manifest.exists()
    imported_doc = imported_pack.documents[0]
    assert imported_doc.button_label == "Переименованная кнопка"
    assert imported_doc.button_label_source == "doctor_renamed"
    assert imported_doc.required_fields == ("patient.fio", "case.number")
    assert imported_doc.template.startswith("templates/")
    assert str(tmp_path) not in imported_doc.template
    assert copied.exists()
