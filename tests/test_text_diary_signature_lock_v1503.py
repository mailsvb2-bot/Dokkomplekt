from __future__ import annotations

from pathlib import Path

from docx import Document

from diary_batch import TEXT_DIARY_SIGNATURE_LINES, fill_diary_batch


def _status_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Состояние стабильное, жалоб активно не предъявляет, назначения выполняет.")
    doc.add_paragraph("Состояние спокойное, сон и аппетит без существенной отрицательной динамики.")
    doc.save(path)


def _paragraphs(path: Path) -> list[str]:
    return [paragraph.text.strip() for paragraph in Document(str(path)).paragraphs if paragraph.text.strip()]


def test_text_diary_has_doctor_and_head_signature_after_each_daily_entry(tmp_path):
    texts = tmp_path / "texts.docx"
    _status_docx(texts)

    result = fill_diary_batch(
        status_files=[texts],
        diary_files=[],
        output_dir=tmp_path / "out",
        patient_name="Иванов И.И.",
        admission_value="10.06.2026",
        discharge_value="12.06.2026",
        gender_source_name="Иванов И.И.",
        diary_day_offsets=(1, 2),
        force_final_diary=True,
    )

    lines = _paragraphs(result.created_files[0])
    diary_lines = [line for line in lines if line.startswith(("11.06.26", "12.06.26"))]
    assert len(diary_lines) == 2
    for diary_line in diary_lines:
        index = lines.index(diary_line)
        assert lines[index + 1:index + 3] == list(TEXT_DIARY_SIGNATURE_LINES)


def test_text_diary_has_doctor_and_head_signature_after_each_hourly_entry(tmp_path):
    texts = tmp_path / "texts.docx"
    _status_docx(texts)

    result = fill_diary_batch(
        status_files=[texts],
        diary_files=[],
        output_dir=tmp_path / "out",
        patient_name="Иванов И.И.",
        admission_value="10.06.2026 14:00",
        gender_source_name="Иванов И.И.",
        diary_hour_offsets=(1,),
        diary_frequency_mode="hourly",
        force_final_diary=False,
    )

    lines = _paragraphs(result.created_files[0])
    for prefix in ("10.06.26 15:00", "10.06.26 16:00"):
        index = next(index for index, line in enumerate(lines) if line.startswith(prefix))
        assert lines[index + 1:index + 3] == list(TEXT_DIARY_SIGNATURE_LINES)
