from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from universal_fields import PatientCase
from universal_profiles import default_document_pack


def _ordinary_word_form(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("ВЫПИСНОЙ ЭПИКРИЗ")
    doc.add_paragraph("ФИО: ______")
    doc.add_paragraph("Дата выписки: ______")
    doc.add_paragraph("Диагноз: ______")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Номер истории болезни"
    table.cell(0, 1).text = ""
    doc.save(path)


def test_active_doctor_profile_is_persisted_as_safe_setting(tmp_path: Path):
    from settings_mixin import SettingsMixin

    class Harness(SettingsMixin):
        pass

    app = Harness()
    app._settings_path = tmp_path / "settings.json"
    profile = tmp_path / "profiles" / "doctor.medpack.json"
    app._settings = {"active_universal_profile": str(profile)}

    app._save_settings()
    stored = json.loads(app._settings_path.read_text("utf-8"))

    assert stored["active_universal_profile"] == str(profile)
    app._settings = {}
    assert app._load_settings()["active_universal_profile"] == str(profile)


def test_clean_profile_read_does_not_rewrite_or_backup(monkeypatch, tmp_path: Path):
    import universal_profiles

    path = tmp_path / "doctor.medpack.json"
    universal_profiles.save_document_pack(default_document_pack(), path)
    before = path.read_bytes()
    save_calls: list[str] = []
    real_save = universal_profiles.save_document_pack

    def spy_save(*args, **kwargs):
        save_calls.append(str(kwargs.get("backup_reason", "")))
        return real_save(*args, **kwargs)

    monkeypatch.setattr(universal_profiles, "save_document_pack", spy_save)
    pack = universal_profiles.ensure_default_pack(path)

    assert pack.pack_id
    assert save_calls == []
    assert path.read_bytes() == before


def test_doctor_confirmation_marker_has_one_canonical_owner():
    from layout_checklist import _doctor_buttons_setup_completed, mark_doctor_buttons_setup_completed

    pack = default_document_pack()
    assert not _doctor_buttons_setup_completed(pack)

    mark_doctor_buttons_setup_completed(pack)

    assert _doctor_buttons_setup_completed(pack)


def test_ordinary_word_form_without_placeholders_is_valid_and_rendered(tmp_path: Path):
    from medical_docx_reader import extract_docx_text
    from universal_template_engine import attach_template_to_pack, render_template_to_docx, validate_template

    source = tmp_path / "ordinary.docx"
    _ordinary_word_form(source)

    validation = validate_template(source, role_id="discharge", category="medical", button_label="Выписной эпикриз")
    assert validation.ok
    assert validation.placeholders == ()
    assert {"patient.fio", "case.number", "discharge.date", "diagnosis.main"} <= set(validation.visible_fields)

    profile = tmp_path / "profile"
    pack = default_document_pack()
    spec, copied = attach_template_to_pack(
        pack,
        source,
        profile,
        button_label="Выписной эпикриз",
        role_id="discharge",
    )
    assert spec.template.startswith("templates/")
    assert {"patient.fio", "case.number", "discharge.date", "diagnosis.main"} <= set(spec.required_fields)

    case = PatientCase()
    case.update_from_pairs(
        {
            "patient.fio": "Орлова Мария Ивановна",
            "case.number": "314/26",
            "discharge.date": "19.05.2026",
            "diagnosis.main": "F32.1 Депрессивный эпизод",
        }
    )
    output = tmp_path / "result.docx"
    result = render_template_to_docx(template_path=copied, output_path=output, case=case, document=spec)
    text = extract_docx_text(output)

    assert result.ok
    assert "ФИО: Орлова Мария Ивановна" in text
    assert "Дата выписки: 19.05.2026" in text
    assert "Диагноз: F32.1 Депрессивный эпизод" in text
    assert "314/26" in text
    assert "______" not in text


def test_pdf_import_rolls_back_all_copies_when_later_source_fails(monkeypatch, tmp_path: Path):
    import pdf_template_importer
    from document_intelligence.pdf_reader import PdfReadResult

    first = tmp_path / "first.pdf"
    second = tmp_path / "second.pdf"
    first.write_bytes(b"%PDF-1.4 first")
    second.write_bytes(b"%PDF-1.4 second")

    monkeypatch.setattr(
        pdf_template_importer,
        "read_pdf_text",
        lambda path: PdfReadResult(str(path), "ФИО: ______", page_count=1),
    )
    calls = 0

    def fake_write(_source: Path, target: Path, _text: str) -> str:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise RuntimeError("second conversion failed")
        doc = Document()
        doc.add_paragraph("ФИО: ______")
        doc.save(target)
        return "text"

    monkeypatch.setattr(pdf_template_importer, "_write_pdf_template", fake_write)
    pack = default_document_pack()

    with pytest.raises(RuntimeError, match="second conversion failed"):
        pdf_template_importer.import_pdf_templates_to_pack(pack, [first, second], tmp_path / "profile")

    assert pack.documents == ()
    assert list((tmp_path / "profile" / "templates").glob("*.docx")) == []


def test_usage_reservation_refunds_known_commit_failure_and_finalize_keeps_charge(tmp_path: Path):
    from product_access import ProductAccessManager

    manager = ProductAccessManager(storage_dir=tmp_path / "license")
    before = manager.current_state().documents_used_total_trial

    reservation = manager.reserve_created_documents(2)
    assert manager.current_state().documents_used_total_trial == before + 2
    manager.release_created_documents(reservation)
    assert manager.current_state().documents_used_total_trial == before
    manager.release_created_documents(reservation)
    assert manager.current_state().documents_used_total_trial == before

    reservation = manager.reserve_created_documents(1)
    manager.finalize_created_documents(reservation)
    assert manager.current_state().documents_used_total_trial == before + 1
    # Finalized reservations cannot be refunded by a duplicate late cleanup.
    manager.release_created_documents(reservation)
    assert manager.current_state().documents_used_total_trial == before + 1


def test_successful_pdf_export_removes_intermediate_docx(monkeypatch, tmp_path: Path):
    import actions_document_intelligence_flow

    source = tmp_path / "result.docx"
    source.write_bytes(b"docx")
    target = tmp_path / "result.pdf"

    def fake_export(path: Path) -> Path:
        assert Path(path) == source
        target.write_bytes(b"pdf")
        return target

    monkeypatch.setattr("document_output_format.export_docx_to_pdf", fake_export)
    mixin = actions_document_intelligence_flow.ActionsDocumentIntelligenceFlowMixin()
    exported = mixin._export_custom_documents_to_pdf([source])

    assert exported == [target]
    assert target.exists()
    assert not source.exists()


def test_visible_word_fields_share_one_canonical_semantic_owner(tmp_path: Path):
    from universal_document_principles import missing_fields_from_principles
    from universal_template_engine import attach_template_to_pack

    source = tmp_path / "ordinary.docx"
    _ordinary_word_form(source)
    pack = default_document_pack()
    spec, _copied = attach_template_to_pack(
        pack,
        source,
        tmp_path / "profile",
        button_label="Выписной эпикриз",
        role_id="discharge",
    )
    case = PatientCase()
    case.update_from_pairs(
        {
            "patient.fio": "Орлова Мария Ивановна",
            "case.number": "314/26",
            "discharge.date": "19.05.2026",
            "diagnosis.main": "F32.1 Депрессивный эпизод",
        }
    )

    assert missing_fields_from_principles(case, spec, base_dir=tmp_path / "profile") == ()
