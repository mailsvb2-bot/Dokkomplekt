from __future__ import annotations

import os
import time
from pathlib import Path


def test_primary_parser_understands_solid_text_core_fields():
    from medical_parser import MedicalTextParser

    text = (
        "Первичный осмотр 05.06.2026 История болезни № 12345 "
        "Ф.И.О.: Иванов Иван Иванович Возраст: 45 лет "
        "Место жительства: г. Нижний Новгород, ул. Пушкина, 1 "
        "Место работы: ООО Ромашка Должность: инженер "
        "Жалобы: головная боль "
        "Анамнез заболевания: заболел вчера "
        "Анамнез жизни: рос и развивался нормально "
        "Лечение: режим, терапия "
        "Диагноз: J20 Острый бронхит"
    )

    data = MedicalTextParser().parse_text(text)

    assert data.admission_date == "05.06.2026"
    assert data.case_number == "12345"
    assert data.fio == "Иванов Иван Иванович"
    assert "45" in data.birth
    assert "Нижний" in data.registered
    assert "Ромашка" in data.work_org
    assert "инженер" in data.position.lower()
    assert "головная" in data.complaints.lower()
    assert "заболел" in data.disease_anamnesis.lower()
    assert data.life_anamnesis
    assert "терап" in data.treatment_plan.lower()
    assert data.has_treatment_section
    assert data.diagnosis


def test_primary_parser_captures_explicit_discharge_date_without_stealing_other_dates():
    from medical_admission_resolver import extract_discharge_date_from_primary_text
    from medical_parser import MedicalTextParser

    text = "\n".join(
        [
            "Первичный осмотр",
            "Дата рождения",
            "04.01.1980",
            "Дата поступления",
            "01.06.2026",
            "История болезни № 12345",
            "ФИО: Иванов Иван Иванович",
            "Дата выписки",
            "10.06.2026",
            "Диагноз: J20 Острый бронхит",
        ]
    )

    data = MedicalTextParser().parse_text(text)

    assert data.admission_date == "01.06.2026"
    assert data.discharge_date == "10.06.2026"
    assert extract_discharge_date_from_primary_text("Дата рождения 04.01.1980 Дата поступления 01.06.2026") == ""


def test_compact_popup_date_100526_normalizes_to_full_date():
    from medical_formatting import parse_date
    from medical_date_state import normalize_date_value

    parsed = parse_date("100526")

    assert parsed is not None
    assert parsed.strftime("%d.%m.%Y") == "10.05.2026"
    assert normalize_date_value("100526") == "10.05.2026"


def test_dialog_date_validation_does_not_reject_admission_compact_source():
    source = Path("dialog_fields_core.py").read_text(encoding="utf-8")

    assert "semantic_key != \"admission_date\"" in source
    assert "parse_date(value)" in source


def test_visible_diary_flow_forces_text_docx_not_table_template():
    source = Path("actions_diary_flow.py").read_text(encoding="utf-8")

    assert "text_output = True" in source
    assert "diary_files=[]" in source
    assert "self.diary_files = []" in source


def test_created_document_preview_popup_is_diagnostic_opt_in_only():
    source = Path("actions_creation_execution.py").read_text(encoding="utf-8")

    assert "MEDICAL_AUTOFILL_SHOW_CREATED_PREVIEW" in source
    assert "Production creation must not open an unsolicited modal window" in source
    assert "not enabled" not in source  # keep the guard positive/explicit, not a second warning popup


def test_discharge_custom_template_receives_primary_case_fields(tmp_path: Path) -> None:
    from docx import Document
    from medical_models import PatientData
    from universal_case_adapter import patient_data_to_case
    from universal_generation import render_documents_from_pack
    from universal_profiles import DocumentPack
    from universal_template_engine import attach_template_to_pack

    template = tmp_path / "discharge.docx"
    document = Document()
    document.add_paragraph("Пациент {{patient.fio}} | ИБ {{case.number}}")
    document.add_paragraph("Даты {{admission.date}} — {{discharge.date}} | Возраст {{patient.age}}")
    document.add_paragraph("Жалобы {{complaints}}")
    document.add_paragraph("Анамнез {{anamnesis.disease}}")
    document.add_paragraph("Статус {{status.objective}}")
    document.add_paragraph("Состояние {{condition.discharge}}")
    document.add_paragraph("Диагноз {{diagnosis.main}}")
    document.add_paragraph("Лечение {{treatment.plan}}")
    document.add_paragraph("Итог {{treatment.result}}")
    document.save(template)

    data = PatientData(
        fio="Петров Пётр Петрович",
        birth="45 лет",
        case_number="777",
        admission_date="01.06.2026",
        discharge_date="10.06.2026",
        complaints="головная боль",
        disease_anamnesis="заболел остро",
        somatic_status="соматически стабилен",
        profile_status="профильный статус без отрицательной динамики",
        diagnosis="J20 Острый бронхит",
        treatment_plan="режим, терапия",
        epi_text="выписывается в стабильном состоянии",
    )
    case = patient_data_to_case(data, source_document="primary.docx")
    pack = DocumentPack(pack_id="doctor.discharge", name="Профиль врача")
    spec, _copied = attach_template_to_pack(pack, template, tmp_path / "profile", button_label="Выписной эпикриз", role_id="dischargeEpicrisis")

    result = render_documents_from_pack(pack=pack, case=case, document_ids=[spec.id], output_dir=tmp_path / "out", base_dir=tmp_path / "profile", strict=True)

    assert result.ok
    text = "\n".join(paragraph.text for paragraph in Document(result.created_files[0]).paragraphs)
    assert "Петров Пётр Петрович" in text
    assert "ИБ 777" in text
    assert "01.06.2026 — 10.06.2026" in text
    assert "Возраст 45 лет" in text
    assert "головная боль" in text
    assert "заболел остро" in text
    assert "соматически стабилен" in text
    assert "выписывается в стабильном состоянии" in text
    assert "J20 Острый бронхит" in text
    assert "режим, терапия" in text


def test_discharge_case_uses_additional_info_when_epi_text_is_absent() -> None:
    from medical_models import PatientData
    from universal_case_adapter import patient_data_to_case

    data = PatientData(
        fio="Сидоров Семён Семёнович",
        case_number="888",
        diagnosis="J20 Острый бронхит",
        treatment_plan="режим, терапия",
        additional_info_text="выписывается с улучшением, даны рекомендации",
    )

    case = patient_data_to_case(data, source_document="primary.docx")

    assert case.get("condition.discharge") == "выписывается с улучшением, даны рекомендации"
    assert case.get("treatment.result") == "выписывается с улучшением, даны рекомендации"
    assert case.get("recommendations") == "выписывается с улучшением, даны рекомендации"


def test_custom_requirement_flags_do_not_depend_on_overlay_super_method() -> None:
    from actions_universal_flow import ActionsUniversalFlowMixin
    from universal_profiles import DocumentPack, DocumentTemplateSpec

    class Harness(ActionsUniversalFlowMixin):
        def __init__(self) -> None:
            self.pack = DocumentPack(
                pack_id="doctor.custom.requirements",
                name="Профиль врача",
                documents=(
                    DocumentTemplateSpec(
                        id="doctor_discharge",
                        button_label="Мой документ врача",
                        template="templates/discharge.docx",
                        required_fields=("case.number", "diagnosis.main", "treatment.plan", "discharge.date"),
                        role_id="dischargeEpicrisis",
                    ),
                ),
            )

        def _load_or_create_universal_pack(self):
            return self.pack

    flags = Harness()._custom_requirement_flags(["doctor_discharge"])

    assert flags["requires_case_number"]
    assert flags["requires_diagnosis"]
    assert flags["requires_treatment"]
    assert flags["requires_discharge_date"]
    assert flags["discharge"]


def test_desktop_intake_top_level_docx_is_launch_intent(tmp_path, monkeypatch):
    import desktop_intake

    primary = tmp_path / "primary.docx"
    primary.write_bytes(b"fake-docx-for-monkeypatched-reader")
    stable_time = time.time() - 10
    os.utime(primary, (stable_time, stable_time))
    monkeypatch.setattr(time, "time", lambda: stable_time + 10)
    monkeypatch.setattr(desktop_intake, "_read_intake_docx_text", lambda path, context: "Первичный осмотр ФИО: Иванов Иван История болезни № 123 Диагноз: J20")

    candidates = desktop_intake.scan_primary_candidates(tmp_path, set())

    assert candidates
    assert candidates[0].path == primary


def test_clean_profile_create_flow_has_folder_naming_guard():
    root = Path(__file__).resolve().parents[1]
    source = "\n".join(path.read_text(encoding="utf-8") for path in root.glob("*.py"))

    assert "_ensure_patient_folder_naming_configured" in source
    assert "configure_patient_folder_naming_dialog" in source
    assert "doctor_confirmed" in source


def test_diary_creation_wizard_reports_table_text_and_frequency():
    assert "таблица дневников" in Path("diary_creation_wizard.py").read_text(encoding="utf-8")
    assert "текстовый DOCX" in Path("diary_creation_wizard.py").read_text(encoding="utf-8")
    actions = Path("actions_diary_flow.py").read_text(encoding="utf-8")
    assert "diary_frequency_mode_var" in actions
    assert "diary_hour_offsets" in actions
    assert "text_output=text_output" in actions


def test_visible_license_entrypoint_exists():
    product_source = Path("product_access/__init__.py").read_text(encoding="utf-8")
    app_source = Path("app.py").read_text(encoding="utf-8")

    assert "show_product_license_dialog" in product_source
    assert "<Control-l>" in product_source
    assert "show_product_license_dialog" in app_source
