from pathlib import Path

from docx import Document

from medical_models import PatientData
from universal_case_adapter import patient_data_to_case
from universal_generation import analyze_pack_readiness, render_documents_from_pack
from universal_profiles import DocumentPack
from universal_template_engine import attach_template_to_pack


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_doctor_owned_template_receives_icd10_code_from_primary_diagnosis(tmp_path: Path) -> None:
    """Primary document has a diagnosis with ICD-10 code; doctor template asks for the code separately."""

    template = tmp_path / "doctor_discharge_template.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Discharge epicrisis"
    doc.add_paragraph("Patient: {{patient.fio}}")
    doc.add_paragraph("Case number: {{case.number}}")
    doc.add_paragraph("Diagnosis: {{diagnosis.main}}")
    doc.add_paragraph("ICD-10 code: {{diagnosis.icd10}}")
    doc.add_paragraph("Admission date: {{admission.date}}")
    doc.add_paragraph("Discharge date: {{discharge.date}}")
    doc.save(template)

    pack = DocumentPack(pack_id="doctor.scenario", name="Doctor scenario")
    spec, _copied = attach_template_to_pack(
        pack,
        template,
        tmp_path / "profile",
        button_label="Discharge epicrisis",
        role_id="discharge_epicrisis",
    )
    assert "diagnosis.icd10" in spec.required_fields

    patient = PatientData(
        fio="Ivanov Ivan",
        output_fio="Ivanov Ivan",
        case_number="42",
        admission_date="01.06.2026",
        discharge_date="10.06.2026",
        diagnosis="K35 Acute appendicitis",
    )
    case = patient_data_to_case(patient, source_document="primary.docx")

    assert case.get("diagnosis.main") == "K35 Acute appendicitis"
    assert case.get("diagnosis.icd10") == "K35"

    readiness = analyze_pack_readiness(pack, case, base_dir=tmp_path / "profile")
    assert readiness.ready_document_ids == (spec.id,), readiness.human_report()
    assert readiness.blocked_document_ids == ()

    result = render_documents_from_pack(
        pack=pack,
        case=case,
        document_ids=[spec.id],
        output_dir=tmp_path / "out",
        base_dir=tmp_path / "profile",
        strict=True,
    )

    assert result.ok, result.human_report()
    assert len(result.created_files) == 1
    rendered_text = _docx_text(Path(result.created_files[0]))
    assert "K35 Acute appendicitis" in rendered_text
    assert "ICD-10 code: K35" in rendered_text


def test_doctor_owned_template_does_not_invent_icd10_for_unresolved_text() -> None:
    patient = PatientData(
        fio="Petrov Petr",
        output_fio="Petrov Petr",
        diagnosis="unmapped local wording",
    )
    case = patient_data_to_case(patient)

    assert case.get("diagnosis.main") == "unmapped local wording"
    assert case.get("diagnosis.icd10") == ""
