from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from diary_batch import fill_diary_batch
from diary_calendar import default_observation_diary_dates, is_non_working_day
from diary_text_output import DynamicEpicrisisInput, build_dynamic_epicrisis_text
from medical_word_format import SUPPORTED_WORD_SUFFIXES, is_supported_word_file


def _read_docx_paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(str(path)).paragraphs if paragraph.text.strip()]


def test_word_format_contract_accepts_doc_docx_docm() -> None:
    assert {".doc", ".docx", ".docm"}.issubset(SUPPORTED_WORD_SUFFIXES)
    assert is_supported_word_file("source.doc")
    assert is_supported_word_file("source.docx")
    assert is_supported_word_file("source.docm")
    assert not is_supported_word_file("source.pdf")


def test_default_diary_calendar_skips_weekends_and_fixed_holidays() -> None:
    dates = default_observation_diary_dates(date(2026, 1, 1), limit=8)
    assert len(dates) == 8
    assert dates == tuple(dict.fromkeys(dates))
    assert all(not is_non_working_day(item) for item in dates)
    assert dates[0] >= date(2026, 1, 12)


def test_text_diary_output_uses_selected_texts_gender_and_periodic_summary(tmp_path: Path) -> None:
    texts = tmp_path / "texts.docx"
    source = Document()
    source.add_paragraph("Пациент был спокоен, жалоб активно не предъявлял, инструкции выполнял.")
    source.add_paragraph("Пациент сообщил об улучшении сна, фон настроения ровный, поведение упорядоченное.")
    source.save(texts)

    result = fill_diary_batch(
        status_files=[texts],
        diary_files=[],
        output_dir=tmp_path / "out",
        patient_name="Иванова Ирина Ивановна",
        gender_source_name="Иванова Ирина Ивановна",
        admission_value="10.06.2026",
        discharge_value="30.06.2026",
        repeat_statuses=True,
        text_output=True,
        sick_leave_dynamic_epicrisis=True,
        sick_leave_from="10.06.2026",
        birth_date="01.01.1980",
        complaints="жалобы уменьшились",
        treatment="терапия по назначению",
        profile_status="контакт доступен",
        treatment_correction="Коррекция терапии не требуется.",
    )

    assert result.processed_files == 1
    assert result.created_files[0].exists()
    joined = "\n".join(_read_docx_paragraphs(Path(result.created_files[0])))
    assert "Пациентка была спокойна" in joined
    assert "не предъявляла" in joined
    assert "Динамический эпикриз" in joined
    assert "Продолжение лечения" in joined
    assert "Заведующий отделением" in joined
    assert "Лечащий врач" in joined


def test_dynamic_summary_fallback_correction_text() -> None:
    text = build_dynamic_epicrisis_text(DynamicEpicrisisInput(patient_name="Петров Пётр", sick_leave_from="01.06.2026"))
    assert "Лекарства принимает согласно назначениям." in text
    assert "Продолжение лечения" in text
