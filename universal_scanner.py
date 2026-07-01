"""DOCX scanner/mapper for configurable medical profiles.

The scanner is deliberately deterministic.  It proposes what it found and gives
confidence levels, but the doctor must be able to confirm or correct the result
before values are reused across generated documents.
"""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from pathlib import Path
import re
from typing import Iterable, Mapping, Sequence, cast

from docx import Document

from universal_fields import FieldRegistry, PatientCase, default_field_registry, normalize_field_id
from universal_profiles import ExtractionRule
from medical_language_detector import detect_text_language
from medical_formatting import parse_date

_DATE_RE = re.compile(r"(?<!\d)([0-3]?\d[.\-/][01]?\d[.\-/](?:19|20)?\d{2}|[0-3]?\d[01]\d(?:19|20)?\d{2})(?!\d)")
_CASE_RE = re.compile(r"(?i)(?:история\s+болезни|иб|nr\s+historii\s+choroby|numer\s+historii\s+choroby|historia\s+choroby\s*(?:nr|n|no)?|nr\s+dokumentacji|numer\s+dokumentacji|nr\s+karty)\s*(?:№|n|no|nr|numer|номер)?\s*[:\-–—]?\s*([A-Za-zА-Яа-яŁłŃńÓóŚśŹźŻż0-9/\\\-]+)\b")
_SNILS_RE = re.compile(r"(?<!\d)(\d{3}[- ]?\d{3}[- ]?\d{3}[- ]?\d{2})(?!\d)")
_PASSPORT_RE = re.compile(r"(?i)(?:паспорт[^0-9]{0,24})?(\d{2}\s?\d{2}\s?\d{6})")
_ICD10_RE = re.compile(r"(?i)\b([A-ZА-Я]\s?\d{2}(?:\.\d+)?)\b")


def _normalize_icd10_code(value: str) -> str:
    """Return a clean ICD-10 code or an empty string.

    A very old scanner path treated the alias ``F`` itself as a label and then
    captured the rest of the diagnosis as the ICD value (``06.7 diagnosis``).
    Codes must therefore be extracted by regex, not by generic label-after logic.
    """

    match = _ICD10_RE.search(str(value or ""))
    if not match:
        return ""
    code = match.group(1).replace(" ", "").upper()
    # Common Cyrillic look-alikes from Word templates.
    translation = cast(dict[str | int, str | int | None], {"А": "A", "В": "B", "Е": "E", "К": "K", "М": "M", "Н": "H", "О": "O", "Р": "P", "С": "C", "Т": "T", "Х": "X"})
    code = code.translate(str.maketrans(translation))
    return code
_LABEL_VALUE_RE_TEMPLATE = r"(?is)(?:^|[\n\r]|\b){label}\s*(?:[:\-–—№Nn]+)?\s*(.+?)(?=$|[\n\r])"
_BLOCK_GUARD_WORDS = (
    "Ф.И.О.", "ФИО", "Дата рождения", "Год рождения", "Возраст", "История болезни",
    "Дата поступления", "Дата госпитализации", "Дата приема", "Дата приёма", "Дата осмотра", "Дата выписки", "Жалобы",
    "Анамнез жизни", "Анамнез заболевания", "Психический статус", "Соматический статус",
    "Объективный статус", "Диагноз", "Лечение", "План лечения", "Назначенное лечение",
    "Рекомендовано", "Рекомендации", "Врач", "Зав. отделением", "Начмед",
    "Pacjent", "Pacjentka", "Imię i nazwisko", "Imie i nazwisko", "Nazwisko i imię", "Nazwisko i imie",
    "Data urodzenia", "PESEL", "Nr historii choroby", "Numer historii choroby", "Historia choroby",
    "Data przyjęcia", "Data przyjecia", "Data hospitalizacji", "Data wypisu",
    "Rozpoznanie", "Diagnoza", "Leczenie", "Plan leczenia", "Zalecone leczenie", "Zastosowane leczenie",
    "Skargi", "Dolegliwości", "Dolegliwosci", "Wywiad", "Stan psychiczny", "Stan przedmiotowy",
    "Wyniki badań", "Wyniki badan", "Zalecenia", "Lekarz", "Ordynator", "Podpis",
)


@dataclass(frozen=True)
class DocumentBlock:
    index: int
    kind: str
    text: str
    path_hint: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class FieldMatch:
    field_id: str
    label: str
    value: str
    confidence: float
    strategy: str
    block_index: int
    start: int = -1
    end: int = -1
    context: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class DocumentScanResult:
    source_path: str
    blocks: tuple[DocumentBlock, ...]
    matches: tuple[FieldMatch, ...] = ()
    warnings: tuple[str, ...] = ()
    detected_language: str = "auto"
    language_confidence: float = 0.0

    def patient_case(self) -> PatientCase:
        case = PatientCase()
        for match in sorted(self.matches, key=lambda item: item.confidence):
            case.set(match.field_id, match.value, confidence=match.confidence, source_document=self.source_path, source_hint=match.context)
        return case

    def best_matches(self) -> dict[str, FieldMatch]:
        result: dict[str, FieldMatch] = {}
        for match in self.matches:
            old = result.get(match.field_id)
            if old is None or match.confidence > old.confidence:
                result[match.field_id] = match
        return result

    def review_rows(self, *, medium_threshold: float = 0.78, high_threshold: float = 0.9) -> tuple[dict[str, object], ...]:
        """Return UI-friendly rows for doctor confirmation.

        The product rule is explicit: low/medium confidence values should be
        visible to the doctor before they are reused in generated documents.
        """

        rows: list[dict[str, object]] = []
        for field_id, match in sorted(self.best_matches().items()):
            if match.confidence >= high_threshold:
                level = "high"
            elif match.confidence >= medium_threshold:
                level = "medium"
            else:
                level = "low"
            rows.append({
                "field_id": field_id,
                "label": match.label,
                "value": match.value,
                "confidence": match.confidence,
                "level": level,
                "strategy": match.strategy,
                "context": match.context,
            })
        return tuple(rows)

    def missing_field_ids(self, registry: FieldRegistry | None = None) -> tuple[str, ...]:
        registry = registry or default_field_registry()
        found = set(self.best_matches())
        return tuple(definition.id for definition in registry.definitions() if definition.required_by_default and definition.id not in found)

    def to_dict(self) -> dict:
        return {
            "source_path": self.source_path,
            "blocks": [block.to_dict() for block in self.blocks],
            "matches": [match.to_dict() for match in self.matches],
            "warnings": list(self.warnings),
            "detected_language": self.detected_language,
            "language_confidence": self.language_confidence,
        }

    def human_report(self) -> str:
        best = self.best_matches()
        lines = ["Разметчик документа", f"Файл: {Path(self.source_path).name}", f"Язык документа: {self.detected_language} ({int(round(self.language_confidence * 100))}%)", ""]
        if not best:
            lines.append("Поля не найдены автоматически. Выделите нужный фрагмент мышкой и назначьте смысловое поле.")
            return "\n".join(lines)
        for field_id in sorted(best):
            match = best[field_id]
            percent = int(round(match.confidence * 100))
            lines.append(f"{match.label}: {match.value}  ({percent}%, {match.strategy})")
        if self.warnings:
            lines.append("")
            lines.append("Предупреждения:")
            lines.extend(f"• {warning}" for warning in self.warnings)
        return "\n".join(lines)


def extract_docx_blocks(path: str | Path) -> tuple[DocumentBlock, ...]:
    candidate = Path(path).expanduser()
    if not candidate.exists() or not candidate.is_file():
        raise FileNotFoundError(f"DOCX не найден: {candidate}")
    if candidate.suffix.lower() not in {".docx", ".docm"}:
        raise ValueError(f"Разметчик принимает только DOCX/DOCM, получено: {candidate.suffix or 'без расширения'}")
    doc = Document(str(candidate))
    blocks: list[DocumentBlock] = []

    def add_paragraphs(paragraphs, prefix: str) -> None:
        for paragraph_index, paragraph in enumerate(paragraphs):
            text = _clean(paragraph.text)
            if text:
                blocks.append(DocumentBlock(len(blocks), "paragraph", text, f"{prefix}.paragraph[{paragraph_index}]"))

    def add_tables(tables, prefix: str) -> None:
        for table_index, table in enumerate(tables):
            for row_index, row in enumerate(table.rows):
                cells = [_clean(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    blocks.append(DocumentBlock(len(blocks), "table_row", row_text, f"{prefix}.table[{table_index}].row[{row_index}]"))

    add_paragraphs(doc.paragraphs, "body")
    add_tables(doc.tables, "body")
    # Hospital forms often keep document title, department, patient identity or
    # signatures in Word headers/footers. The mapper must show and scan those
    # regions too, otherwise a doctor sees a value in the DOCX but the profile
    # cannot learn it.
    for section_index, section in enumerate(doc.sections):
        for area_name, area in (("header", section.header), ("footer", section.footer)):
            add_paragraphs(area.paragraphs, f"section[{section_index}].{area_name}")
            add_tables(area.tables, f"section[{section_index}].{area_name}")
    return tuple(blocks)


def scan_many_docx(
    paths: Sequence[str | Path],
    *,
    registry: FieldRegistry | None = None,
    rules: Sequence[ExtractionRule] = (),
) -> DocumentScanResult:
    """Scan several source examples and merge them into one reviewable result."""

    if not paths:
        return DocumentScanResult("", (), (), ("Не выбраны исходные DOCX для обучения профиля.",), "auto", 0.0)
    scans = [scan_docx(path, registry=registry, rules=rules) for path in paths]
    return merge_scan_results(*scans)


def merge_scan_results(*results: DocumentScanResult) -> DocumentScanResult:
    """Merge several source documents into one reviewable case.

    This supports the real setup flow: a doctor can train a profile on several
    examples instead of trusting one randomly formatted document.  Higher
    confidence values win, while all original matches remain auditable.
    """

    if not results:
        return DocumentScanResult("", (), (), ("Нет документов для объединения.",), "auto", 0.0)
    blocks: list[DocumentBlock] = []
    matches: list[FieldMatch] = []
    warnings: list[str] = []
    block_offset = 0
    for result in results:
        for block in result.blocks:
            blocks.append(DocumentBlock(block.index + block_offset, block.kind, block.text, f"{Path(result.source_path).name}:{block.path_hint}"))
        for match in result.matches:
            matches.append(FieldMatch(match.field_id, match.label, match.value, match.confidence, match.strategy, match.block_index + block_offset, match.start, match.end, f"{Path(result.source_path).name}:{match.context}"))
        warnings.extend(result.warnings)
        block_offset += len(result.blocks)
    deduped = _dedupe_matches(matches)
    best_language = max(results, key=lambda item: item.language_confidence)
    return DocumentScanResult("; ".join(str(Path(item.source_path).name) for item in results), tuple(blocks), tuple(deduped), tuple(dict.fromkeys(warnings)), best_language.detected_language, best_language.language_confidence)


def scan_docx(
    path: str | Path,
    *,
    registry: FieldRegistry | None = None,
    rules: Sequence[ExtractionRule] = (),
) -> DocumentScanResult:
    registry = registry or default_field_registry()
    blocks = extract_docx_blocks(path)
    joined = "\n".join(block.text for block in blocks)
    language = detect_text_language(joined)
    matches: list[FieldMatch] = []
    warnings: list[str] = []

    matches.extend(_scan_with_registry_aliases(blocks, registry))
    matches.extend(_scan_known_regexes(blocks, registry))
    matches.extend(_scan_block_fields(joined, blocks, registry))
    matches.extend(_scan_with_saved_rules(blocks, registry, rules))

    deduped = _dedupe_matches(matches)
    required_missing = [definition.label for definition in registry.definitions() if definition.required_by_default and definition.id not in {m.field_id for m in deduped}]
    if required_missing:
        warnings.append("Не найдены обязательные поля: " + ", ".join(required_missing))
    return DocumentScanResult(str(Path(path).expanduser()), blocks, tuple(deduped), tuple(warnings), language.language_id, language.confidence)


def learn_rule_from_selection(
    blocks: Sequence[DocumentBlock],
    *,
    field_id: str,
    selected_text: str,
    registry: FieldRegistry | None = None,
) -> ExtractionRule:
    """Create a first extraction rule from a doctor's mouse selection."""

    registry = registry or default_field_registry()
    normalized_field = normalize_field_id(field_id)
    selection = _clean(selected_text)
    if not selection:
        raise ValueError("Сначала выделите фрагмент документа.")
    best_block = _find_block_containing(blocks, selection)
    if best_block is None:
        return ExtractionRule(normalized_field, "exact_selection", selected_text=selection, confidence=0.55, created_from="manual_selection")
    prefix = best_block.text[: best_block.text.find(selection)]
    label = _nearest_label(prefix, registry.aliases_for(normalized_field))
    if label:
        return ExtractionRule(normalized_field, "label_after", label=label, block_hint=best_block.path_hint, selected_text=selection, confidence=0.86, created_from="manual_selection")
    return ExtractionRule(normalized_field, "exact_selection", block_hint=best_block.path_hint, selected_text=selection, confidence=0.62, created_from="manual_selection")


def _scan_with_registry_aliases(blocks: Sequence[DocumentBlock], registry: FieldRegistry) -> list[FieldMatch]:
    matches: list[FieldMatch] = []
    for block in blocks:
        for definition in registry.definitions():
            for alias in registry.aliases_for(definition.id):
                value_span = _value_after_label(block.text, alias)
                if not value_span:
                    continue
                value, start, end = value_span
                if definition.value_kind == "date":
                    value = _first_date(value) or value
                if definition.id == "diagnosis.icd10":
                    code = _normalize_icd10_code(block.text) or _normalize_icd10_code(value)
                    if not code:
                        continue
                    value = code
                elif definition.value_kind == "identifier" and definition.id == "case.number":
                    case = _CASE_RE.search(block.text)
                    if case:
                        value = case.group(1)
                value = _trim_value(value)
                if not value or _looks_like_only_label(value, alias):
                    continue
                matches.append(
                    FieldMatch(definition.id, definition.label, value, 0.78, "label_after", block.index, start, end, block.path_hint)
                )
    return matches


def _scan_known_regexes(blocks: Sequence[DocumentBlock], registry: FieldRegistry) -> list[FieldMatch]:
    result: list[FieldMatch] = []
    for block in blocks:
        case = _CASE_RE.search(block.text)
        if case:
            definition = registry.require("case.number")
            result.append(FieldMatch(definition.id, definition.label, case.group(1).strip(), 0.93, "regex_case_number", block.index, case.start(1), case.end(1), block.path_hint))
        snils = _SNILS_RE.search(block.text)
        if snils:
            definition = registry.require("patient.snils")
            result.append(FieldMatch(definition.id, definition.label, snils.group(1).strip(), 0.86, "regex_snils", block.index, snils.start(1), snils.end(1), block.path_hint))
        passport = _PASSPORT_RE.search(block.text)
        if passport and "паспорт" in block.text.lower():
            definition = registry.require("patient.passport")
            result.append(FieldMatch(definition.id, definition.label, passport.group(1).strip(), 0.78, "regex_passport", block.index, passport.start(1), passport.end(1), block.path_hint))
        icd = _ICD10_RE.search(block.text)
        if icd and any(word in block.text.lower() for word in ("диагноз", "мкб", "icd", "mkb", "rozpoznanie", "diagnoza", "kod rozpoznania", "f")):
            definition = registry.require("diagnosis.icd10")
            result.append(FieldMatch(definition.id, definition.label, _normalize_icd10_code(icd.group(1)), 0.84, "regex_icd10", block.index, icd.start(1), icd.end(1), block.path_hint))
        if any(word.lower() in block.text.lower() for word in ("первичный осмотр", "направление на госпитализацию", "выписной эпикриз", "badanie wstępne", "badanie wstepne", "skierowanie do szpitala", "skierowanie na hospitalizację", "skierowanie na hospitalizacje", "karta informacyjna", "epikryza", "historia choroby")):
            definition = registry.require("document.title")
            title = _trim_value(block.text)
            result.append(FieldMatch(definition.id, definition.label, title, 0.7, "title_phrase", block.index, 0, len(block.text), block.path_hint))
            date = _first_date(block.text)
            if date:
                admission = registry.require("admission.date")
                result.append(FieldMatch(admission.id, admission.label, date, 0.82, "title_date", block.index, block.text.find(date), block.text.find(date) + len(date), block.path_hint))
    return result


def _scan_block_fields(joined: str, blocks: Sequence[DocumentBlock], registry: FieldRegistry) -> list[FieldMatch]:
    matches: list[FieldMatch] = []
    block_field_ids = [definition.id for definition in registry.definitions() if definition.value_kind == "block"]
    guard_regex = _marker_regex(_all_aliases_for_fields(registry, block_field_ids))
    for field_id in block_field_ids:
        definition = registry.require(field_id)
        for alias in registry.aliases_for(field_id):
            pattern = re.compile(rf"(?is){re.escape(alias)}\s*[:\-–—]?\s*(.+?)(?=\n\s*(?:{guard_regex})\s*(?:[:\-–—]|\n|$)|$)")
            match = pattern.search(joined)
            if not match:
                continue
            value = _trim_value(match.group(1))
            if value and len(value) >= 3:
                block_index = _block_index_for_joined_offset(blocks, joined, match.start(1))
                matches.append(FieldMatch(field_id, definition.label, value, 0.81, "block_between_markers", block_index, match.start(1), match.end(1), alias))
                break
    return matches


def _scan_with_saved_rules(blocks: Sequence[DocumentBlock], registry: FieldRegistry, rules: Sequence[ExtractionRule]) -> list[FieldMatch]:
    matches: list[FieldMatch] = []
    for rule in rules:
        definition = registry.require(rule.field_id)
        if rule.strategy == "label_after" and rule.label:
            for block in blocks:
                value_span = _value_after_label(block.text, rule.label)
                if value_span:
                    value, start, end = value_span
                    value = _normalize_icd10_code(block.text) if definition.id == "diagnosis.icd10" else _trim_value(value)
                    if not value:
                        continue
                    matches.append(FieldMatch(definition.id, definition.label, value, max(0.0, min(1.0, rule.confidence)), "saved_label_after", block.index, start, end, block.path_hint))
                    break
        elif rule.strategy == "regex" and rule.regex:
            pattern = re.compile(rule.regex, re.IGNORECASE | re.DOTALL)
            for block in blocks:
                match = pattern.search(block.text)
                if match:
                    value = match.group(1) if match.groups() else match.group(0)
                    matches.append(FieldMatch(definition.id, definition.label, _trim_value(value), max(0.0, min(1.0, rule.confidence)), "saved_regex", block.index, match.start(), match.end(), block.path_hint))
                    break
        elif rule.strategy == "block_between_markers" and rule.label:
            joined = "\n".join(block.text for block in blocks)
            guard_regex = _marker_regex(_all_aliases_for_fields(registry, [definition.id]))
            pattern = re.compile(rf"(?is){re.escape(rule.label)}\s*[:\-–—]?\s*(.+?)(?=\n\s*(?:{guard_regex})\s*(?:[:\-–—]|\n|$)|$)")
            match = pattern.search(joined)
            if match:
                value = _trim_value(match.group(1))
                if value:
                    block_index = _block_index_for_joined_offset(blocks, joined, match.start(1))
                    matches.append(FieldMatch(definition.id, definition.label, value, max(0.0, min(1.0, rule.confidence)), "saved_block_between_markers", block_index, match.start(1), match.end(1), rule.label))
        elif rule.strategy == "table_cell" and rule.block_hint:
            for block in blocks:
                if block.path_hint == rule.block_hint and rule.selected_text:
                    pos = block.text.find(_clean(rule.selected_text))
                    if pos >= 0:
                        matches.append(FieldMatch(definition.id, definition.label, _clean(rule.selected_text), max(0.0, min(1.0, rule.confidence)), "saved_table_cell", block.index, pos, pos + len(_clean(rule.selected_text)), block.path_hint))
                        break
        elif rule.strategy == "exact_selection" and rule.selected_text:
            needle = _clean(rule.selected_text)
            for block in blocks:
                pos = block.text.find(needle)
                if pos >= 0:
                    matches.append(FieldMatch(definition.id, definition.label, needle, max(0.0, min(1.0, rule.confidence)), "saved_exact_selection", block.index, pos, pos + len(needle), block.path_hint))
                    break
    return matches


def _dedupe_matches(matches: Iterable[FieldMatch]) -> list[FieldMatch]:
    best: dict[tuple[str, str], FieldMatch] = {}
    for match in matches:
        if not match.value:
            continue
        key = (match.field_id, " ".join(match.value.lower().replace("ё", "е").split()))
        old = best.get(key)
        if old is None or match.confidence > old.confidence:
            best[key] = match
    # Keep the best match per field first, but preserve alternative values after it.
    return sorted(best.values(), key=lambda item: (item.field_id, -item.confidence, item.block_index))


def _value_after_label(text: str, label: str) -> tuple[str, int, int] | None:
    if not label.strip():
        return None
    pattern = re.compile(_LABEL_VALUE_RE_TEMPLATE.format(label=re.escape(label)), re.IGNORECASE)
    match = pattern.search(text)
    if not match:
        return None
    value = _trim_value(match.group(1))
    if not value:
        return None
    start = match.start(1)
    end = start + len(match.group(1))
    return value, start, end


def _first_date(text: str) -> str:
    parsed_textual = parse_date(text)
    if parsed_textual:
        return parsed_textual.strftime("%d.%m.%Y")
    match = _DATE_RE.search(text)
    if not match:
        return ""
    raw = match.group(1).strip().replace("-", ".").replace("/", ".")
    if "." not in raw and len(raw) in {4, 6, 8}:
        if len(raw) == 4:
            return f"01.{raw[:2]}.20{raw[2:]}"
        if len(raw) == 6:
            return f"{raw[:2]}.{raw[2:4]}.20{raw[4:]}"
        return f"{raw[:2]}.{raw[2:4]}.{raw[4:]}"
    parts = raw.split(".")
    if len(parts) == 3:
        day, month, year = parts
        if len(year) == 2:
            year = "20" + year if int(year) <= 50 else "19" + year
        return f"{int(day):02d}.{int(month):02d}.{year}"
    return raw


def _trim_value(value: str) -> str:
    text = _clean(value)
    for guard in _BLOCK_GUARD_WORDS:
        pos = text.lower().find(guard.lower())
        if pos > 0:
            text = text[:pos].strip(" ;,.|\n\t")
    return text.strip(" ;,.|\n\t")


def _clean(text: str) -> str:
    return re.sub(r"[ \t\u00a0]+", " ", str(text or "").replace("\r", "\n")).strip()


def _looks_like_only_label(value: str, label: str) -> bool:
    normalized_value = " ".join(value.lower().split()).strip(":-–— ")
    normalized_label = " ".join(label.lower().split()).strip(":-–— ")
    return normalized_value == normalized_label


def _marker_regex(markers: Sequence[str]) -> str:
    escaped = [re.escape(marker) for marker in markers if marker.strip()]
    return "|".join(sorted(set(escaped), key=len, reverse=True)) or r"$^"


def _all_aliases_for_fields(registry: FieldRegistry, field_ids: Sequence[str]) -> tuple[str, ...]:
    aliases: list[str] = []
    for field_id in field_ids:
        aliases.extend(registry.aliases_for(field_id))
    aliases.extend(_BLOCK_GUARD_WORDS)
    return tuple(dict.fromkeys(aliases))


def _block_index_for_joined_offset(blocks: Sequence[DocumentBlock], joined: str, offset: int) -> int:
    cursor = 0
    for block in blocks:
        block_len = len(block.text)
        if cursor <= offset <= cursor + block_len:
            return block.index
        cursor += block_len + 1
    return blocks[-1].index if blocks else -1


def _find_block_containing(blocks: Sequence[DocumentBlock], selected_text: str) -> DocumentBlock | None:
    needle = _clean(selected_text)
    for block in blocks:
        if needle and needle in block.text:
            return block
    normalized_needle = " ".join(needle.lower().replace("ё", "е").split())
    for block in blocks:
        normalized_block = " ".join(block.text.lower().replace("ё", "е").split())
        if normalized_needle and normalized_needle in normalized_block:
            return block
    return None


def _nearest_label(prefix: str, aliases: Sequence[str]) -> str:
    best_label = ""
    best_pos = -1
    prefix_low = prefix.lower().replace("ё", "е")
    for alias in aliases:
        pos = prefix_low.rfind(alias.lower().replace("ё", "е"))
        if pos > best_pos:
            best_pos = pos
            best_label = alias
    return best_label
