"""Universal DOCX template engine for configurable medical document packs.

This module is the next layer after ``universal_scanner``.  The scanner answers
"what did the source document mean?"; the template engine answers "which custom
DOCX templates can be created from that meaning, what fields do they require,
and how can a doctor move a pack between computers?".

The implementation is intentionally deterministic and local-only.  It renders
explicit ``{{field.id}}`` placeholders and refuses to silently ignore missing
required medical fields.
"""

from __future__ import annotations

from diagnostic_logging import record_soft_exception
from dataclasses import asdict, dataclass, replace
import json
from pathlib import Path, PurePosixPath
import re
import shutil
import zipfile
from typing import Iterable, Mapping, Sequence

from docx import Document
from medical_gender import remove_forbidden_hospitalization_phrase_from_document

from universal_fields import PatientCase, FieldRegistry, default_field_registry, normalize_field_id, normalize_field_id_for_context
from universal_profiles import DocumentPack, DocumentTemplateSpec, load_document_pack, save_document_pack

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
PACK_MANIFEST_NAME = "pack.json"
TEMPLATE_DIR_NAME = "templates"
_ALLOWED_PACK_SUFFIXES = {".json", ".medpack", ".zip"}


@dataclass(frozen=True)
class TemplatePlaceholder:
    """One ``{{field.id}}`` placeholder discovered in a DOCX template."""

    field_id: str
    block_hint: str
    raw: str

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateValidationResult:
    """Static validation for one custom DOCX template."""

    template_path: str
    placeholders: tuple[TemplatePlaceholder, ...]
    unknown_fields: tuple[str, ...] = ()
    missing_required_placeholders: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()

    @property
    def ok(self) -> bool:
        # A template with zero placeholders looks valid syntactically, but cannot
        # be filled by the universal renderer. Treat it as blocked, not merely a
        # warning, so doctors do not add inert templates to a paid profile.
        return bool(self.placeholders) and not self.unknown_fields and not self.missing_required_placeholders

    def field_ids(self) -> tuple[str, ...]:
        return tuple(dict.fromkeys(item.field_id for item in self.placeholders))

    def to_dict(self) -> dict:
        return {
            "template_path": self.template_path,
            "placeholders": [item.to_dict() for item in self.placeholders],
            "unknown_fields": list(self.unknown_fields),
            "missing_required_placeholders": list(self.missing_required_placeholders),
            "warnings": list(self.warnings),
            "ok": self.ok,
        }


@dataclass(frozen=True)
class PackValidationResult:
    """Validation summary for an entire configurable document pack."""

    pack_id: str
    errors: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()
    template_results: tuple[TemplateValidationResult, ...] = ()

    @property
    def ok(self) -> bool:
        return not self.errors and all(item.ok for item in self.template_results)

    def human_report(self) -> str:
        lines = [f"Проверка профиля: {self.pack_id}", f"Статус: {'OK' if self.ok else 'есть ошибки'}"]
        if self.errors:
            lines.append("")
            lines.append("Ошибки:")
            lines.extend("• " + error for error in self.errors)
        if self.warnings:
            lines.append("")
            lines.append("Предупреждения:")
            lines.extend("• " + warning for warning in self.warnings)
        if self.template_results:
            lines.append("")
            lines.append("Шаблоны:")
            for result in self.template_results:
                state = "OK" if result.ok else "ошибка"
                lines.append(f"• {Path(result.template_path).name}: {state}, placeholders={len(result.placeholders)}")
                for unknown in result.unknown_fields:
                    lines.append(f"  - неизвестное поле: {unknown}")
                for missing in result.missing_required_placeholders:
                    lines.append(f"  - нет обязательного placeholder: {missing}")
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "pack_id": self.pack_id,
            "ok": self.ok,
            "errors": list(self.errors),
            "warnings": list(self.warnings),
            "template_results": [item.to_dict() for item in self.template_results],
        }


@dataclass(frozen=True)
class RenderResult:
    """Result of rendering one custom DOCX document."""

    output_path: str
    replaced_fields: tuple[str, ...]
    missing_fields: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()

    @property
    def ok(self) -> bool:
        return not self.missing_fields

    def to_dict(self) -> dict:
        return asdict(self)


def normalize_placeholder_id(
    raw_field_id: str,
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> str:
    """Normalize template placeholder ids while allowing reserved aliases.

    Doctors may write either canonical placeholders (``{{labs.results}}``),
    export-style placeholders (``{{patientName}}`` / ``{{caseNo}}``) or
    human-friendly analysis aliases such as ``{{АНАЛИЗЫ}}`` /
    ``{{LABS_BLOCK}}``.  The important rule is that this DOCX layer must not
    pre-lowercase the raw id before the universal field normalizer sees it:
    camelCase splitting lives in :func:`normalize_field_id`, and losing the
    capitals here silently turns ``patientName`` into an unknown
    ``patientname`` field.
    """

    raw = str(raw_field_id or "").strip()
    from medical_renderer_labs import canonical_labs_placeholder

    # Preserve the historical labs shortcuts (LABS_BLOCK, БЛОК АНАЛИЗОВ, etc.)
    # without flattening camelCase export ids before semantic normalization.
    labs_key = "_".join(str(raw or "").strip().lower().replace("-", "_").split())
    labs_alias = canonical_labs_placeholder(raw)
    text = labs_alias if labs_alias != labs_key else raw

    if str(text).strip().lower().startswith("document."):
        return str(text).strip().lower()
    if role_id or category or button_label:
        return normalize_field_id_for_context(text, role_id=role_id, category=category, document_label=button_label)
    return normalize_field_id(text)


def _document_context_kwargs(document: object | None = None, **overrides: str) -> dict[str, str]:
    if document is None:
        return {
            "role_id": str(overrides.get("role_id", "") or ""),
            "category": str(overrides.get("category", "") or ""),
            "button_label": str(overrides.get("button_label", "") or ""),
        }
    return {
        "role_id": str(overrides.get("role_id", "") or getattr(document, "role_id", "") or ""),
        "category": str(overrides.get("category", "") or getattr(document, "category", "") or ""),
        "button_label": str(overrides.get("button_label", "") or getattr(document, "button_label", "") or ""),
    }


def extract_template_placeholders(
    path: str | Path,
    *,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> tuple[TemplatePlaceholder, ...]:
    """Read a DOCX/DOCM template and return all semantic placeholders."""

    candidate = _existing_docx(path, "шаблон документа")
    doc = Document(str(candidate))
    found: list[TemplatePlaceholder] = []

    for paragraph, hint in _iter_docx_paragraphs(doc):
        for match in PLACEHOLDER_RE.finditer(paragraph.text or ""):
            found.append(TemplatePlaceholder(normalize_placeholder_id(match.group(1), role_id=role_id, category=category, button_label=button_label), hint, match.group(0)))
    return tuple(found)


def validate_template(
    template_path: str | Path,
    *,
    required_fields: Sequence[str] = (),
    registry: FieldRegistry | None = None,
    role_id: str = "",
    category: str = "",
    button_label: str = "",
) -> TemplateValidationResult:
    """Validate placeholders in a user-supplied DOCX template."""

    registry = registry or default_field_registry()
    placeholders = extract_template_placeholders(template_path, role_id=role_id, category=category, button_label=button_label)
    known = set(registry.ids()) | {"document.id", "document.label", "document.category", "document.description"}
    unknown = sorted({item.field_id for item in placeholders if item.field_id not in known and not item.field_id.startswith("custom.")})
    placeholder_fields = {item.field_id for item in placeholders}
    missing_required = sorted({
        normalize_field_id_for_context(field_id, role_id=role_id, category=category, document_label=button_label)
        for field_id in required_fields
        if normalize_field_id_for_context(field_id, role_id=role_id, category=category, document_label=button_label) not in placeholder_fields
    })
    warnings: list[str] = []
    if not placeholders:
        warnings.append("В шаблоне нет placeholders вида {{patient.fio}} — такой шаблон нельзя заполнить универсальным движком.")
    duplicate_fields = [field_id for field_id in placeholder_fields if sum(1 for item in placeholders if item.field_id == field_id) > 1]
    if duplicate_fields:
        warnings.append("Повторяющиеся поля в шаблоне: " + ", ".join(sorted(set(duplicate_fields))))
    return TemplateValidationResult(str(Path(template_path).expanduser()), placeholders, tuple(unknown), tuple(missing_required), tuple(warnings))


def infer_document_spec_from_template(
    template_path: str | Path,
    *,
    button_label: str | None = None,
    document_id: str | None = None,
    category: str = "medical",
    registry: FieldRegistry | None = None,
    role_id: str = "",
) -> DocumentTemplateSpec:
    """Create a dynamic document-button spec from a DOCX template."""

    path = _existing_docx(template_path, "шаблон документа")
    registry = registry or default_field_registry()
    explicit_role_id = str(role_id or "").strip()
    try:
        from personal_document_buttons import suggest_button_label_for_template
        suggestion = suggest_button_label_for_template(
            path,
            preferred_language="auto",
            ui_language="ru",
            explicit_specialty="",
            fallback_label=button_label or path.stem,
        )
        label = (button_label or suggestion.label or path.stem).strip()
        suggested_role_id = suggestion.role_id if suggestion.role_id != "unknown" else ""
        role_id = explicit_role_id or suggested_role_id
        button_language = suggestion.language_id
        source_language = suggestion.source_language
        label_source = "manual" if button_label else suggestion.source
    except Exception as exc:
        record_soft_exception("universal_template_engine.infer_spec", exc, detail=str(path))
        label = (button_label or path.stem).strip()
        role_id = explicit_role_id
        button_language = "auto"
        source_language = "auto"
        label_source = "manual" if button_label else "template_title"
    placeholders = extract_template_placeholders(path, role_id=role_id, category=category, button_label=label)
    semantic_fields = tuple(
        dict.fromkeys(
            item.field_id
            for item in placeholders
            if not item.field_id.startswith("document.") and (item.field_id in registry or item.field_id.startswith("custom."))
        )
    )
    doc_id = _safe_document_id(document_id or path.stem or label)
    return DocumentTemplateSpec(
        id=doc_id,
        button_label=label,
        template=path.name,
        output_name="{{patient.fio}} " + label + ".docx",
        required_fields=semantic_fields,
        optional_fields=(),
        category=category or "medical",
        description="Создано автоматически по placeholders в пользовательском DOCX-шаблоне.",
        role_id=role_id,
        button_language=button_language,
        source_language=source_language,
        button_label_source=label_source,
    )


def build_document_pack_from_templates(
    *,
    pack_id: str,
    name: str,
    specialty: str = "",
    template_paths: Sequence[str | Path],
    registry: FieldRegistry | None = None,
) -> DocumentPack:
    """Build a custom DocumentPack from a list of DOCX templates."""

    registry = registry or default_field_registry()
    documents = tuple(infer_document_spec_from_template(path, registry=registry) for path in template_paths)
    return DocumentPack(pack_id=pack_id, name=name, specialty=specialty, documents=documents)


def validate_document_pack(pack: DocumentPack, *, base_dir: str | Path | None = None) -> PackValidationResult:
    """Validate document ids, field ids and template files for a profile."""

    registry = pack.registry()
    errors: list[str] = []
    warnings: list[str] = []
    template_results: list[TemplateValidationResult] = []

    if not pack.documents:
        warnings.append("В профиле нет документов: блок 03 не сможет построить пользовательские кнопки.")
    seen_ids: set[str] = set()
    for document in pack.documents:
        if not document.id.strip():
            errors.append("Документ без id.")
        if document.id in seen_ids:
            errors.append(f"Дублируется id документа: {document.id}")
        seen_ids.add(document.id)
        context_kwargs = _document_context_kwargs(document)
        for field_id in [*document.required_fields, *document.optional_fields]:
            try:
                normalized = normalize_field_id_for_context(field_id, **context_kwargs)
            except ValueError as exc:
                errors.append(f"{document.button_label}: некорректное поле {field_id!r}: {exc}")
                continue
            if normalized not in registry and not normalized.startswith("custom."):
                errors.append(f"{document.button_label}: поле не найдено в реестре: {normalized}")
        if not document.template:
            errors.append(f"{document.button_label}: не указан template.")
            continue
        template_path = _resolve_pack_template_path(document.template, base_dir)
        if template_path.exists() and template_path.suffix.lower() in {".docx", ".docm"}:
            template_results.append(validate_template(template_path, required_fields=document.required_fields, registry=registry, role_id=document.role_id, category=document.category, button_label=document.button_label))
        elif not any(ch in document.template for ch in "*?"):
            warnings.append(f"{document.button_label}: шаблон пока не найден рядом с профилем: {document.template}")
    return PackValidationResult(pack.pack_id, tuple(errors), tuple(warnings), tuple(template_results))


def build_render_context(case: PatientCase, document: DocumentTemplateSpec, *, output_language: str = "auto", spellcheck_enabled: bool = True) -> dict[str, str]:
    """Create the placeholder replacement map for one patient case and document."""

    raw_context = {field_id: value.value for field_id, value in case.values.items()}
    try:
        from medical_orthography import correct_case_values
        context = correct_case_values(raw_context, language_id=output_language, enabled=spellcheck_enabled)
    except Exception as exc:
        record_soft_exception("universal_template_engine.orthography", exc)
        # Orthography is a safety net, not a render blocker. Rendering must never
        # fail only because a spelling dictionary/rule is unavailable.
        context = raw_context
    context.update(
        {
            "document.id": document.id,
            "document.label": document.button_label,
            "document.category": document.category,
            "document.description": document.description,
        }
    )
    if getattr(document, "category", "") == "diaries":
        context.update(_diary_context_values(case, document))
    return context


def missing_required_fields(case: PatientCase, document: DocumentTemplateSpec) -> tuple[str, ...]:
    """Return required fields that are absent before rendering."""

    missing: list[str] = []
    context_kwargs = _document_context_kwargs(document)
    for field_id in document.required_fields:
        normalized = normalize_field_id_for_context(field_id, **context_kwargs)
        if not case.get(normalized).strip():
            missing.append(normalized)
    return tuple(dict.fromkeys(missing))


def render_output_name(
    document: DocumentTemplateSpec,
    case: PatientCase,
    *,
    output_language: str = "auto",
    spellcheck_enabled: bool = True,
) -> str:
    """Render and sanitize the output DOCX name for a custom document."""

    context = build_render_context(case, document, output_language=output_language, spellcheck_enabled=spellcheck_enabled)
    raw = _replace_placeholders(document.output_name or "{{patient.fio}} {{document.label}}.docx", context, missing_value="", document=document)
    raw = re.sub(r"\s+", " ", raw).strip(" .") or (document.button_label or document.id or "Документ")
    if not raw.lower().endswith(".docx"):
        raw += ".docx"
    from medical_formatting import safe_filename

    return safe_filename(raw)


def render_template_to_docx(
    *,
    template_path: str | Path,
    output_path: str | Path,
    case: PatientCase,
    document: DocumentTemplateSpec,
    strict: bool = True,
    output_language: str = "auto",
    spellcheck_enabled: bool = True,
) -> RenderResult:
    """Render a custom DOCX template using explicit ``{{field.id}}`` placeholders."""

    template = _existing_docx(template_path, "шаблон документа")
    missing_required = missing_required_fields(case, document)
    if strict and missing_required:
        raise ValueError("Не заполнены обязательные поля для документа: " + ", ".join(missing_required))

    context = build_render_context(case, document, output_language=output_language, spellcheck_enabled=spellcheck_enabled)
    doc = Document(str(template))
    replaced: set[str] = set()
    missing_seen: set[str] = set()

    for paragraph, _hint in _iter_docx_paragraphs(doc):
        _replace_paragraph_placeholders(paragraph, context, replaced, missing_seen, document=document)

    remove_forbidden_hospitalization_phrase_from_document(doc)
    output = Path(output_path).expanduser()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return RenderResult(str(output), tuple(sorted(replaced)), tuple(sorted(missing_seen)), ())


def unique_document_id_for_pack(pack: DocumentPack, proposed_id: str, *, template_path: str | Path | None = None) -> str:
    """Return a document id that does not replace another document in this pack."""

    base = _safe_document_id(proposed_id or (Path(template_path).stem if template_path else "document"))
    used = {document.id for document in pack.documents}
    if base not in used:
        return base
    stem = _safe_document_id(Path(template_path).stem if template_path else base)
    for index in range(2, 1000):
        candidate = _safe_document_id(f"{base}_{stem}_{index}")
        if candidate not in used:
            return candidate
    raise ValueError(f"Не удалось создать уникальный id документа: {base}")


def attach_template_to_pack(
    pack: DocumentPack,
    template_path: str | Path,
    profile_dir: str | Path,
    *,
    button_label: str | None = None,
    document_id: str | None = None,
    category: str = "medical",
    registry: FieldRegistry | None = None,
    role_id: str = "",
    button_language: str = "auto",
    source_language: str = "auto",
    button_label_source: str = "manual",
) -> tuple[DocumentTemplateSpec, Path]:
    """Copy a user DOCX into the profile folder and attach it as a dynamic document.

    v1.4.1 stored only ``path.name`` in the JSON profile.  That made the UI look
    successful, but export/validation could later miss the physical DOCX when
    it lived in Downloads/Desktop.  The universal product contract is stricter:
    once a doctor adds a template, the profile owns a private copy under
    ``profiles/templates/`` and portable ``.medpack.zip`` can include it.
    """

    source = _existing_docx(template_path, "шаблон документа")
    profile_root = Path(profile_dir).expanduser()
    templates_dir = profile_root / TEMPLATE_DIR_NAME
    templates_dir.mkdir(parents=True, exist_ok=True)
    try:
        source_already_owned = source.parent.resolve() == templates_dir.resolve()
    except Exception as exc:
        record_soft_exception("universal_template_engine.template_dir_resolve", exc, detail=str(source))
        source_already_owned = False
    target = source if source_already_owned else _available_template_copy_path(templates_dir / source.name)
    if source.resolve() != target.resolve():
        shutil.copy2(source, target)
    spec = infer_document_spec_from_template(target, button_label=button_label, document_id=document_id, category=category, registry=registry, role_id=role_id)
    spec = replace(
        spec,
        id=unique_document_id_for_pack(pack, spec.id, template_path=target),
        template=(PurePosixPath(TEMPLATE_DIR_NAME) / target.name).as_posix(),
        role_id=role_id or spec.role_id,
        button_language=button_language or spec.button_language,
        source_language=source_language or spec.source_language,
        button_label_source=button_label_source or spec.button_label_source,
    )
    pack.add_document(spec)
    return spec, target

def export_document_pack_zip(pack: DocumentPack, target_zip: str | Path, *, template_base_dir: str | Path | None = None) -> Path:
    """Export a profile as a portable ``.medpack.zip`` with manifest and templates."""

    target = Path(target_zip).expanduser()
    target.parent.mkdir(parents=True, exist_ok=True)
    base = Path(template_base_dir).expanduser() if template_base_dir else None
    manifest = pack.to_dict()
    portable_documents: list[dict] = []
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        seen: set[str] = set()
        for document in pack.documents:
            doc_data = document.to_dict()
            source = _resolve_pack_template_path(document.template, base)
            if source.exists() and source.is_file() and source.suffix.lower() in {".docx", ".docm"}:
                arcname = PurePosixPath(TEMPLATE_DIR_NAME) / source.name
                if arcname.as_posix() not in seen:
                    zf.write(source, arcname.as_posix())
                    seen.add(arcname.as_posix())
                # A portable medpack must never keep an absolute path from the
                # exporting doctor's desktop/downloads in pack.json.
                doc_data["template"] = arcname.as_posix()
            portable_documents.append(doc_data)
        manifest["documents"] = portable_documents
        zf.writestr(PACK_MANIFEST_NAME, json.dumps(manifest, ensure_ascii=False, indent=2) + "\n")
    return target


def import_document_pack_zip(source_zip: str | Path, target_dir: str | Path) -> tuple[DocumentPack, Path]:
    """Import a portable medpack/json profile safely into ``target_dir``.

    Imported templates are copied with collision-safe names.  This prevents one
    doctor's ``templates/Осмотр.docx`` from overwriting another profile's file
    when several medpack archives are imported into the same local profile
    directory.
    """

    source = Path(source_zip).expanduser()
    if source.suffix.lower() not in _ALLOWED_PACK_SUFFIXES:
        raise ValueError(f"Неподдерживаемый формат профиля: {source.suffix}")
    target = Path(target_dir).expanduser()
    target.mkdir(parents=True, exist_ok=True)

    if source.suffix.lower() == ".json" or source.name.endswith(".medpack.json"):
        pack = load_document_pack(source)
        pack = _copy_json_profile_templates(pack, source.parent, target)
        out = _available_profile_manifest_path(target / source.name)
        save_document_pack(pack, out)
        return pack, out

    with zipfile.ZipFile(source, "r") as zf:
        infos = zf.infolist()
        names = [info.filename for info in infos]
        if PACK_MANIFEST_NAME not in names:
            raise ValueError("В medpack-архиве нет pack.json.")
        if len(infos) > 250:
            raise ValueError("Слишком много файлов внутри medpack-архива.")
        total_size = sum(max(0, info.file_size) for info in infos)
        if total_size > 100 * 1024 * 1024:
            raise ValueError("medpack-архив слишком большой для безопасного импорта.")
        for name in names:
            _assert_safe_zip_name(name)
        manifest_data = json.loads(zf.read(PACK_MANIFEST_NAME).decode("utf-8"))
        if not isinstance(manifest_data, dict):
            raise ValueError("pack.json внутри medpack должен содержать JSON-объект.")
        pack = DocumentPack.from_dict(manifest_data)
        pack = _copy_zip_profile_templates(pack, zf, set(names), target)

    pack_path = _available_profile_manifest_path(target / PACK_MANIFEST_NAME)
    save_document_pack(pack, pack_path)
    return pack, pack_path


def save_pack_report(report: PackValidationResult, path: str | Path) -> Path:
    """Write a human-readable validation report for support/QA."""

    target = Path(path).expanduser()
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(report.human_report() + "\n", encoding="utf-8")
    return target






def _diary_context_values(case: PatientCase, document: DocumentTemplateSpec) -> dict[str, str]:
    try:
        from diary_dates import parse_full_datetime
        from diary_schedule import DiaryScheduleSpec, describe_schedule, planned_diary_datetimes
        admission = parse_full_datetime(case.get("admission.date"))
        spec = DiaryScheduleSpec.from_dict(getattr(document, "diary_schedule", None))
        moments = planned_diary_datetimes(admission, spec, limit=20)
        if spec.mode == "hourly" and spec.hour_offsets:
            formatted = [item.strftime("%d.%m.%Y %H:%M") for item in moments]
        else:
            formatted = [item.strftime("%d.%m.%Y") for item in moments]
        return {
            "diary.schedule": describe_schedule(spec),
            "diary.dates": "\n".join(formatted),
            "diary.entries": "\n".join(f"{index + 1}. {value}" for index, value in enumerate(formatted)),
            "diary.frequency": "ежечасно" if spec.mode == "hourly" else "ежедневно",
        }
    except Exception as exc:
        record_soft_exception("universal_template_engine.diary_context", exc)
        return {
            "diary.schedule": "",
            "diary.dates": "",
            "diary.entries": "",
            "diary.frequency": "",
        }


def _iter_docx_paragraphs(doc: Document):
    """Yield paragraphs from body, tables, headers and footers with stable hints."""

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        yield paragraph, f"paragraph[{paragraph_index}]"
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield paragraph, f"table[{table_index}].row[{row_index}].cell[{cell_index}].paragraph[{paragraph_index}]"
    for section_index, section in enumerate(doc.sections):
        for area_name, area in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(area.paragraphs):
                yield paragraph, f"section[{section_index}].{area_name}.paragraph[{paragraph_index}]"
            for table_index, table in enumerate(area.tables):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        for paragraph_index, paragraph in enumerate(cell.paragraphs):
                            yield paragraph, f"section[{section_index}].{area_name}.table[{table_index}].row[{row_index}].cell[{cell_index}].paragraph[{paragraph_index}]"


def _resolve_pack_template_path(template_value: str, base_dir: str | Path | None) -> Path:
    template = Path(template_value).expanduser()
    if template.is_absolute():
        return template
    if base_dir:
        base = Path(base_dir).expanduser()
        direct = base / template
        if direct.exists():
            return direct
        in_templates = base / TEMPLATE_DIR_NAME / template.name
        if in_templates.exists():
            return in_templates
        return direct
    return template

def _available_template_copy_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem} ({index}){suffix}")
        if not candidate.exists():
            return candidate
    raise FileExistsError(f"Не удалось сохранить копию шаблона: {path}")


def _available_profile_manifest_path(path: Path) -> Path:
    candidate = Path(path).expanduser()
    if not candidate.exists():
        return candidate
    stem = candidate.name[:-len(".medpack.json")] if candidate.name.endswith(".medpack.json") else candidate.stem
    suffix = ".medpack.json" if candidate.name.endswith(".medpack.json") else candidate.suffix
    for index in range(2, 1000):
        next_candidate = candidate.with_name(f"{stem} ({index}){suffix}")
        if not next_candidate.exists():
            return next_candidate
    raise FileExistsError(f"Не удалось сохранить импортированный профиль: {candidate}")


def _copy_json_profile_templates(pack: DocumentPack, source_base: Path, target_base: Path) -> DocumentPack:
    updated_documents: list[DocumentTemplateSpec] = []
    templates_dir = target_base / TEMPLATE_DIR_NAME
    templates_dir.mkdir(parents=True, exist_ok=True)
    for document in pack.documents:
        source = _resolve_pack_template_path(document.template, source_base)
        if source.exists() and source.is_file() and source.suffix.lower() in {".docx", ".docm"}:
            target = _available_template_copy_path(templates_dir / source.name)
            if source.resolve() != target.resolve():
                shutil.copy2(source, target)
            document = replace(document, template=(PurePosixPath(TEMPLATE_DIR_NAME) / target.name).as_posix())
        updated_documents.append(document)
    pack.documents = tuple(updated_documents)
    return pack


def _copy_zip_profile_templates(pack: DocumentPack, zf: zipfile.ZipFile, names: set[str], target_base: Path) -> DocumentPack:
    updated_documents: list[DocumentTemplateSpec] = []
    templates_dir = target_base / TEMPLATE_DIR_NAME
    templates_dir.mkdir(parents=True, exist_ok=True)
    for document in pack.documents:
        template_value = str(document.template or "").replace("\\", "/").strip()
        candidates = [template_value]
        if template_value:
            candidates.append((PurePosixPath(TEMPLATE_DIR_NAME) / PurePosixPath(template_value).name).as_posix())
        archive_name = next((item for item in candidates if item in names and PurePosixPath(item).suffix.lower() in {".docx", ".docm"}), "")
        if archive_name:
            target = _available_template_copy_path(templates_dir / PurePosixPath(archive_name).name)
            with zf.open(archive_name, "r") as source_file, target.open("wb") as target_file:
                shutil.copyfileobj(source_file, target_file)
            document = replace(document, template=(PurePosixPath(TEMPLATE_DIR_NAME) / target.name).as_posix())
        updated_documents.append(document)
    pack.documents = tuple(updated_documents)
    return pack

def _replace_placeholders(text: str, context: Mapping[str, str], *, missing_value: str, document: DocumentTemplateSpec | None = None) -> str:
    def repl(match: re.Match[str]) -> str:
        field_id = normalize_placeholder_id(match.group(1), **_document_context_kwargs(document))
        return str(context.get(field_id, missing_value))

    return PLACEHOLDER_RE.sub(repl, text or "")


def _replace_placeholders_with_report(text: str, context: Mapping[str, str], *, document: DocumentTemplateSpec | None = None) -> tuple[str, tuple[str, ...], tuple[str, ...]]:
    replaced: list[str] = []
    missing: list[str] = []

    def repl(match: re.Match[str]) -> str:
        field_id = normalize_placeholder_id(match.group(1), **_document_context_kwargs(document))
        value = str(context.get(field_id, ""))
        if value:
            replaced.append(field_id)
            return value
        missing.append(field_id)
        return ""

    return PLACEHOLDER_RE.sub(repl, text or ""), tuple(dict.fromkeys(replaced)), tuple(dict.fromkeys(missing))


def _replace_paragraph_placeholders(
    paragraph,
    context: Mapping[str, str],
    replaced: set[str],
    missing_seen: set[str],
    *,
    document: DocumentTemplateSpec | None = None,
) -> None:
    """Replace placeholders while preserving surrounding run formatting.

    Older builds used ``paragraph.text = ...`` for every replacement.  That was
    reliable for split placeholders but it flattened all runs in the paragraph.
    This routine edits only the runs that contain the placeholder text.  If a
    placeholder is split across runs, the replacement inherits the first run's
    formatting while text before/after the placeholder keeps its own runs.
    """

    source = "".join(run.text for run in paragraph.runs) if paragraph.runs else (paragraph.text or "")
    if "{{" not in source:
        return
    matches = list(PLACEHOLDER_RE.finditer(source))
    if not matches:
        return

    replacements: list[tuple[int, int, str]] = []
    for match in matches:
        field_id = normalize_placeholder_id(match.group(1), **_document_context_kwargs(document))
        value = str(context.get(field_id, ""))
        if value:
            replaced.add(field_id)
        else:
            missing_seen.add(field_id)
        replacements.append((match.start(), match.end(), value))

    if not paragraph.runs:
        paragraph.text = _replace_placeholders(source, context, missing_value="", document=document)
        return

    run_spans: list[tuple[int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        end = cursor + len(run.text or "")
        run_spans.append((cursor, end))
        cursor = end
    if cursor != len(source):
        paragraph.text = _replace_placeholders(source, context, missing_value="", document=document)
        return

    def locate(offset: int) -> tuple[int, int]:
        for index, (start, end) in enumerate(run_spans):
            if start <= offset < end:
                return index, offset - start
            if offset == end and index + 1 < len(run_spans) and run_spans[index + 1][0] == end:
                continue
        last = max(0, len(run_spans) - 1)
        return last, max(0, min(offset - run_spans[last][0], len(paragraph.runs[last].text or "")))

    for start, end, value in reversed(replacements):
        start_run, start_offset = locate(start)
        end_run, end_offset = locate(max(start, end - 1))
        end_offset += 1
        if start_run == end_run:
            run = paragraph.runs[start_run]
            text = run.text or ""
            run.text = text[:start_offset] + value + text[end_offset:]
            continue
        first = paragraph.runs[start_run]
        first.text = (first.text or "")[:start_offset] + value
        for index in range(start_run + 1, end_run):
            paragraph.runs[index].text = ""
        last = paragraph.runs[end_run]
        last.text = (last.text or "")[end_offset:]


def _replace_paragraph_text(paragraph, replace_func) -> None:
    """Backward-compatible fallback used by older private imports."""

    old_text = paragraph.text
    if "{{" not in old_text:
        return
    new_text = replace_func(old_text)
    if new_text != old_text:
        paragraph.text = new_text


def _existing_docx(path: str | Path, label: str) -> Path:
    candidate = Path(path).expanduser()
    if not candidate.exists() or not candidate.is_file():
        raise FileNotFoundError(f"Не найден файл ({label}): {candidate}")
    if candidate.suffix.lower() not in {".docx", ".docm"}:
        raise ValueError(f"Неверный формат файла ({label}): {candidate.suffix or 'без расширения'}. Разрешено: .docx, .docm.")
    return candidate


def _safe_document_id(value: str) -> str:
    text = re.sub(r"[^a-zA-Zа-яА-Я0-9]+", "_", str(value or "").strip().lower()).strip("_")
    translit = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e", "ж": "zh", "з": "z", "и": "i", "й": "y",
        "к": "k", "л": "l", "м": "m", "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u", "ф": "f",
        "х": "h", "ц": "c", "ч": "ch", "ш": "sh", "щ": "sch", "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    }
    ascii_text = "".join(translit.get(ch, ch) for ch in text)
    ascii_text = re.sub(r"[^a-z0-9_]+", "_", ascii_text).strip("_")
    if not ascii_text or not ascii_text[0].isalpha():
        ascii_text = "document_" + (ascii_text or "custom")
    return ascii_text


def _assert_safe_zip_name(name: str) -> None:
    path = PurePosixPath(name)
    normalized = str(name or "")
    first_part = PurePosixPath(normalized).parts[0] if PurePosixPath(normalized).parts else ""
    if (
        path.is_absolute()
        or ".." in path.parts
        or normalized.startswith(("/", "\\"))
        or "\\" in normalized
        or ":" in first_part
    ):
        raise ValueError(f"Небезопасный путь внутри medpack: {name}")

# --- Safe template marking for the visual mouse/color scanner ---

def placeholder_for_field(field_id: str) -> str:
    """Return a canonical DOCX placeholder for a semantic field."""

    return "{{" + normalize_placeholder_id(field_id) + "}}"


@dataclass(frozen=True)
class TemplateMarkResult:
    template_path: str
    field_id: str
    placeholder: str
    strategy: str
    replacements: int
    backup_path: str = ""

    @property
    def ok(self) -> bool:
        return self.replacements > 0


def replace_selection_with_placeholder(
    template_path: str | Path,
    selected_text: str,
    field_id: str,
    *,
    create_backup: bool = True,
) -> TemplateMarkResult:
    """Replace the selected text in a DOCX template with ``{{field.id}}``."""

    return _apply_visual_placeholder(template_path, selected_text, field_id, mode="template_replace", create_backup=create_backup)


def insert_placeholder_after_selection(
    template_path: str | Path,
    selected_text: str,
    field_id: str,
    *,
    create_backup: bool = True,
) -> TemplateMarkResult:
    """Insert ``{{field.id}}`` after the selected anchor paragraph/cell."""

    return _apply_visual_placeholder(template_path, selected_text, field_id, mode="template_insert_after", create_backup=create_backup)


def _apply_visual_placeholder(
    template_path: str | Path,
    selected_text: str,
    field_id: str,
    *,
    mode: str,
    create_backup: bool,
) -> TemplateMarkResult:
    path = _existing_docx(template_path, "шаблон для цветной разметки")
    selected = " ".join(str(selected_text or "").replace("\r", "\n").split())
    if not selected:
        raise ValueError("Выделите текст/строку в шаблоне, куда нужно поставить поле.")
    placeholder = placeholder_for_field(field_id)
    backup = ""
    if create_backup:
        backup_path = _available_visual_backup_path(path)
        shutil.copy2(path, backup_path)
        backup = str(backup_path)
    doc = Document(str(path))
    replacements = 0
    for paragraph, _hint in _iter_docx_paragraphs(doc):
        current = paragraph.text or ""
        if mode == "template_replace":
            pos = current.find(selected)
            if pos < 0:
                continue
            paragraph.text = current[:pos] + placeholder + current[pos + len(selected):]
            replacements += 1
            break
        if mode == "template_insert_after":
            if current.find(selected) < 0 and selected not in " ".join(current.split()):
                continue
            if placeholder not in current:
                paragraph.text = current + ("\n" if current and not current.endswith("\n") else "") + placeholder
            replacements += 1
            break
        raise ValueError(f"Неизвестный режим цветной разметки: {mode}")
    if replacements:
        doc.save(str(path))
    return TemplateMarkResult(str(path), normalize_placeholder_id(field_id), placeholder, mode, replacements, backup)


def _available_visual_backup_path(path: Path) -> Path:
    base = path.with_name(path.stem + ".before_visual_marker" + path.suffix)
    if not base.exists():
        return base
    for index in range(2, 1000):
        candidate = path.with_name(f"{path.stem}.before_visual_marker_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise FileExistsError(f"Не удалось создать резервную копию шаблона: {path}")
